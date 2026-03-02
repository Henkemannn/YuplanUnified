from __future__ import annotations

import datetime
import re

from docx import Document  # type: ignore

from .base import ImportedMenuItem, MenuImporter, MenuImportResult, WeekImport, normalize_day

_WEEK_PATTERNS = [r"vecka\s*[:\.]?\s*(\d+)", r"v[\.:\s]*(\d+)", r"week\s+(\d+)"]
_NON_DISH_PHRASES = [
    "med reservation för ändringar",
    "ni når oss på telefon",
    "allt med röd text",
]
_PHONE_RE = re.compile(r"\b07\d{1,2}[-\s]?\d{2,3}[-\s]?\d{2,3}\b")


class DocxMenuImporter(MenuImporter):
    """Adapts legacy kommun Word parser to unified model.
    Focus: Alt1/Alt2/Dessert/Kväll mapping.
    """

    def can_handle(self, filename: str, mimetype: str | None, first_bytes: bytes) -> bool:
        return filename.lower().endswith(".docx")

    def parse(self, file_bytes: bytes, filename: str) -> MenuImportResult:
        from io import BytesIO

        doc = Document(BytesIO(file_bytes))
        content: list[str] = []
        for p in doc.paragraphs:
            txt = p.text.strip()
            if txt:
                for line in txt.split("\n"):
                    if line.strip():
                        content.append(line.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    ctext = cell.text.strip()
                    if ctext:
                        for line in ctext.split("\n"):
                            if line.strip():
                                content.append(line.strip())
        # Split into weeks (may be multi-week)
        week_sections = self._split_content_by_weeks(content)
        if not week_sections:
            return MenuImportResult(
                weeks=[],
                errors=["Could not detect any week number; supply week manually."],
                warnings=[],
            )
        results: list[WeekImport] = []
        current_year = datetime.date.today().year
        for week, lines in week_sections:
            items = self._extract_items(lines, current_year, week)
            results.append(WeekImport(year=current_year, week=week, items=items))
        return MenuImportResult(weeks=results)

    def _split_content_by_weeks(self, content: list[str]):
        out = []
        current_week = None
        current_lines: list[str] = []
        for line in content:
            lw = line.lower()
            matched = False
            for pattern in _WEEK_PATTERNS:
                m = re.search(pattern, lw)
                if m:
                    w = int(m.group(1))
                    if 1 <= w <= 52:
                        if current_week is not None and current_lines:
                            out.append((current_week, current_lines))
                        current_week = w
                        current_lines = []
                        matched = True
                        break
            if not matched and current_week is not None:
                current_lines.append(line)
        if current_week is not None and current_lines:
            out.append((current_week, current_lines))
        return out

    def _extract_items(self, lines: list[str], year: int, week: int) -> list[ImportedMenuItem]:
        items: list[ImportedMenuItem] = []
        current_day = None
        current_ctx: dict[str, str] | None = None
        # Patterns for variants
        alt_patterns = [
            ("alt1", re.compile(r"^(Alt\s*1:|Alternativ\s*1:|Lunch:)", re.IGNORECASE)),
            ("alt2", re.compile(r"^(Alt\s*2:|Alternativ\s*2:)", re.IGNORECASE)),
            ("dessert", re.compile(r"^(Dessert:)", re.IGNORECASE)),
            ("evening", re.compile(r"^(Kväll:)", re.IGNORECASE)),
        ]

        def norm_line(s: str) -> str:
            s = s.replace("\u00a0", " ").strip()
            return re.sub(r"\s+", " ", s)

        def is_non_dish_line(s: str) -> bool:
            low = s.lower().strip()
            if not low:
                return True
            for phrase in _NON_DISH_PHRASES:
                if phrase in low:
                    return True
            if _PHONE_RE.search(low):
                return True
            return False

        for raw in lines:
            for sub in raw.split("\n"):
                line = norm_line(sub)
                if not line:
                    continue
                if is_non_dish_line(line):
                    continue
                # Day detection: direct localized token detection via normalize_day map below
                # Try localized day tokens
                lower = line.lower()
                day_token = None
                for token, eng in list(
                    {k: v for k, v in normalize_day.__globals__["_day_map"].items()}.items()
                ):
                    if re.match(rf"^\b{re.escape(token)}\b", lower):
                        day_token = eng
                        break
                if day_token:
                    current_day = day_token
                    current_ctx = None
                    # Remove only the matched localized day token plus optional colon/punctuation and whitespace
                    # Strip first word and following punctuation
                    parts = line.split(":", 1)
                    if len(parts) == 2:
                        line_after = parts[1].strip()
                    else:
                        # fallback: remove first word
                        line_after = " ".join(line.split(" ")[1:]).strip()
                    if line_after:
                        current_ctx = self._handle_variant_line(items, current_day, line_after, alt_patterns, current_ctx)
                    continue
                if current_day:
                    current_ctx = self._handle_variant_line(items, current_day, line, alt_patterns, current_ctx)
        return items

    def _handle_variant_line(
        self,
        items: list[ImportedMenuItem],
        day: str,
        text: str,
        alt_patterns,
        current_ctx: dict[str, str] | None,
    ):
        matched = False
        for vtype, rgx in alt_patterns:
            m = rgx.match(text)
            if m:
                dish = text[m.end() :].strip()
                if not dish:
                    return current_ctx
                meal = "dinner" if vtype == "evening" else "lunch"
                canonical_variant = "main" if vtype == "evening" else vtype
                # Category inference: dessert -> dessert, evening -> evening, alts -> main
                if canonical_variant == "dessert":
                    category = "dessert"
                elif vtype == "evening":
                    category = "evening"
                else:
                    category = "main"
                current_ctx = {
                    "meal": meal,
                    "variant_type": canonical_variant,
                    "category": category,
                }
                items.append(
                    ImportedMenuItem(
                        day=day,
                        meal=meal,
                        variant_type=canonical_variant,
                        dish_name=dish,
                        category=category,
                        source_labels=[vtype],
                    )
                )
                matched = True
                break
        if not matched:
            if current_ctx is None:
                return current_ctx
            items.append(
                ImportedMenuItem(
                    day=day,
                    meal=current_ctx["meal"],
                    variant_type=current_ctx["variant_type"],
                    dish_name=text,
                    category=current_ctx["category"],
                    source_labels=["unlabeled"],
                )
            )
        return current_ctx
