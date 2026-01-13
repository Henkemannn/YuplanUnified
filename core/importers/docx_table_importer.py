from __future__ import annotations

from collections.abc import Sequence
from dataclasses import dataclass

try:  # pragma: no cover - optional dependency guard
    import docx
except Exception:  # noqa: BLE001
    docx = None  # type: ignore[assignment]

from .base_types import RawRow, UnsupportedFormatError

__all__ = ["DOCXImportResult", "parse_docx"]


@dataclass(slots=True)
class DOCXImportResult:
    rows: list[RawRow]
    headers: Sequence[str]


def parse_docx(data: bytes) -> DOCXImportResult:
    """Parse a simple DOCX table file into raw rows.

    Rules:
    - Use first table only. If none, return empty.
    - First row = headers (stripped).
    - Blank data rows (all empty cells) skipped.
    - Short rows padded with empty strings; long rows truncated.
    """
    if docx is None:  # pragma: no cover - environment branch
        raise UnsupportedFormatError("python-docx is not installed")

    # python-docx expects a path or a file-like; wrap bytes in BytesIO.
    from io import BytesIO

    document = docx.Document(BytesIO(data))  # type: ignore[attr-defined]
    if not document.tables:
        return DOCXImportResult(rows=[], headers=[])
    table = document.tables[0]
    if not table.rows:
        return DOCXImportResult(rows=[], headers=[])

    def _clean(s: str) -> str:
        return s.strip().strip("\ufeff")

    headers = [_clean(c.text) for c in table.rows[0].cells]

    rows: list[RawRow] = []
    for r in table.rows[1:]:
        cells = [_clean(c.text) for c in r.cells]
        if all(c == "" for c in cells):
            continue
        if len(cells) < len(headers):
            cells.extend(["" for _ in range(len(headers) - len(cells))])
        mapped: RawRow = {}
        for h, v in zip(headers, cells, strict=False):
            mapped[h] = v
        rows.append(mapped)

    return DOCXImportResult(rows=rows, headers=headers)
