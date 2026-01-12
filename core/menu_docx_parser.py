from __future__ import annotations

from io import BytesIO
from typing import Dict, Any, List

DAY_MAP = {
    # Swedish full
    "måndag": 1,
    "tisdag": 2,
    "onsdag": 3,
    "torsdag": 4,
    "fredag": 5,
    "lördag": 6,
    "söndag": 7,
    # Swedish short
    "mån": 1,
    "tis": 2,
    "ons": 3,
    "tors": 4,
    "fre": 5,
    "lör": 6,
    "sön": 7,
    # English
    "monday": 1,
    "tuesday": 2,
    "wednesday": 3,
    "thursday": 4,
    "friday": 5,
    "saturday": 6,
    "sunday": 7,
    "mon": 1,
    "tue": 2,
    "wed": 3,
    "thu": 4,
    "fri": 5,
    "sat": 6,
    "sun": 7,
}


class MenuDocxParseError(Exception):
    pass


def _collect_lines(document) -> List[str]:
    lines: List[str] = []
    # paragraphs
    try:
        for p in document.paragraphs:
            txt = (p.text or "").strip()
            if txt:
                lines.append(txt)
    except Exception:
        pass
    # tables
    try:
        for t in document.tables:
            for row in t.rows:
                for cell in row.cells:
                    txt = (cell.text or "").strip()
                    if txt:
                        lines.append(txt)
    except Exception:
        pass
    return lines


def parse_menu_docx(stream: BytesIO) -> Dict[str, Any]:
    try:
        from docx import Document  # type: ignore
    except Exception as e:  # pragma: no cover
        raise MenuDocxParseError(f"DOCX parsing not available: {e}")

    try:
        document = Document(stream)
    except Exception as e:
        raise MenuDocxParseError(f"Ogiltigt DOCX: {e}")

    lines = _collect_lines(document)
    text = "\n".join(lines)
    out: Dict[str, Any] = {
        "year": None,
        "weeks": {},  # map: week -> {days:{1..7:{lunch,dinner}}}
    }

    import re
    # First, detect explicit week headings as standalone lines: "v. 8"
    week_heading_re = re.compile(r"(?mi)^\s*v\.\s*(\d{1,2})\s*$")
    week_headings = list(week_heading_re.finditer(text))

    # Also detect range numbers (e.g., "v. 8-15") for fallback replication when no explicit headings
    weeks_set: set[int] = set()
    for m in re.finditer(r"\bv\.\s*(\d{1,2})(?:\s*[–\-]\s*(\d{1,2}))?", text, flags=re.IGNORECASE):
        a = int(m.group(1))
        b = m.group(2)
        if b:
            b_i = int(b)
            rng = range(min(a, b_i), max(a, b_i) + 1)
            weeks_set.update(rng)
        else:
            weeks_set.add(a)
    sorted_weeks = sorted(weeks_set) if weeks_set else []

    day_names = ["Måndag", "Tisdag", "Onsdag", "Torsdag", "Fredag", "Lördag", "Söndag"]
    day_pattern = r"\b(" + "|".join(day_names) + r")\b\s*:?"
    matches = list(re.finditer(day_pattern, text, flags=re.IGNORECASE))
    segments: List[tuple[int, str]] = []
    for i, m in enumerate(matches):
        start = m.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        seg = text[start:end]
        segments.append((DAY_MAP[m.group(1).lower()], seg))

    def _norm(s: str) -> str:
        s = (s or "").strip()
        s = re.sub(r"\s+", " ", s)
        return s

    def _extract(block: str, label: str) -> str:
        pat = r"\b" + label + r"\b\s*:\s*(.+?)(?=\bAlt\s*1\b|\bAlt\s*2\b|\bDessert\b|\bKväll\b|" + day_pattern + r"|$)"
        m = re.search(pat, block, flags=re.IGNORECASE | re.DOTALL)
        return _norm(m.group(1)) if m else ""

    def _extract_alt(block: str, n: int) -> str:
        pat = r"\bAlt\s*" + str(n) + r"\b\s*:\s*(.+?)(?=\bAlt\s*\d\b|\bDessert\b|\bKväll\b|" + day_pattern + r"|$)"
        m = re.search(pat, block, flags=re.IGNORECASE | re.DOTALL)
        return _norm(m.group(1)) if m else ""

    # Helper to init week days dict
    def _blank_days() -> Dict[int, Dict[str, Dict[str, str]]]:
        return {i: {"lunch": {"alt1_text": "", "alt2_text": "", "dessert": ""},
                    "dinner": {"alt1_text": "", "alt2_text": "", "dessert": ""}}
                for i in range(1, 8)}

    # If explicit week headings exist, split by them; else parse once and replicate across range
    if len(week_headings) >= 1:
        # Parse each explicit week block separately; no replication across range
        for i, wm in enumerate(week_headings):
            wnum = int(wm.group(1))
            start = wm.end()
            end = week_headings[i + 1].start() if i + 1 < len(week_headings) else len(text)
            wseg = text[start:end]
            # Extract day segments per week
            wday_matches = list(re.finditer(day_pattern, wseg, flags=re.IGNORECASE))
            out["weeks"][wnum] = {"days": _blank_days()}
            for j, dm in enumerate(wday_matches):
                d_idx = DAY_MAP[dm.group(1).lower()]
                d_start = dm.start()
                d_end = wday_matches[j + 1].start() if j + 1 < len(wday_matches) else len(wseg)
                dseg = wseg[d_start:d_end]
                parts = re.split(r"(?i)\bKväll\b\s*:", dseg, maxsplit=1)
                lunch_part = parts[0]
                dinner_part = parts[1] if len(parts) > 1 else ""
                l_alt1 = _extract_alt(lunch_part, 1) or _extract(lunch_part, "Lunch")
                l_alt2 = _extract_alt(lunch_part, 2)
                l_des = _extract(lunch_part, "Dessert")
                d_alt1 = _extract_alt(dinner_part, 1)
                if not d_alt1:
                    m0 = re.search(r"^(.+?)(?=\bAlt\s*\d\b|\bDessert\b|" + day_pattern + r"|$)", dinner_part, flags=re.IGNORECASE | re.DOTALL)
                    d_alt1 = _norm(m0.group(1)) if m0 else ""
                d_des = _extract(dinner_part, "Dessert")
                out["weeks"][wnum]["days"][d_idx]["lunch"]["alt1_text"] = l_alt1
                out["weeks"][wnum]["days"][d_idx]["lunch"]["alt2_text"] = l_alt2
                out["weeks"][wnum]["days"][d_idx]["lunch"]["dessert"] = l_des
                out["weeks"][wnum]["days"][d_idx]["dinner"]["alt1_text"] = d_alt1
                out["weeks"][wnum]["days"][d_idx]["dinner"]["alt2_text"] = ""
                out["weeks"][wnum]["days"][d_idx]["dinner"]["dessert"] = d_des
    else:
        # Parse once and replicate across sorted_weeks
        temp_days = _blank_days()
        for idx, seg in segments:
            parts = re.split(r"(?i)\bKväll\b\s*:", seg, maxsplit=1)
            lunch_part = parts[0]
            dinner_part = parts[1] if len(parts) > 1 else ""
            l_alt1 = _extract_alt(lunch_part, 1) or _extract(lunch_part, "Lunch")
            l_alt2 = _extract_alt(lunch_part, 2)
            l_des = _extract(lunch_part, "Dessert")
            d_alt1 = _extract_alt(dinner_part, 1)
            if not d_alt1:
                m0 = re.search(r"^(.+?)(?=\bAlt\s*\d\b|\bDessert\b|" + day_pattern + r"|$)", dinner_part, flags=re.IGNORECASE | re.DOTALL)
                d_alt1 = _norm(m0.group(1)) if m0 else ""
            d_des = _extract(dinner_part, "Dessert")
            temp_days[idx]["lunch"]["alt1_text"] = l_alt1
            temp_days[idx]["lunch"]["alt2_text"] = l_alt2
            temp_days[idx]["lunch"]["dessert"] = l_des
            temp_days[idx]["dinner"]["alt1_text"] = d_alt1
            temp_days[idx]["dinner"]["alt2_text"] = ""
            temp_days[idx]["dinner"]["dessert"] = d_des
        for w in sorted_weeks:
            out["weeks"][w] = {"days": temp_days}

    return out
