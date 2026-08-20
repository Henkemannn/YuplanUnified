from __future__ import annotations

import csv
import io
import re
from dataclasses import dataclass

from docx import Document
from openpyxl import load_workbook

from werkzeug.datastructures import FileStorage

_SUPPORTED_EXTENSIONS = {".txt", ".csv", ".xlsx", ".docx"}
_TEXT_HEADER_TOKENS = {
    "text",
    "line",
    "name",
    "dish",
    "dish_name",
    "composition",
    "composition_name",
    "title",
    "ratt",
    "ratt_namn",
}
_HEADING_PREFIXES = (
    "week",
    "vecka",
    "monday",
    "tuesday",
    "wednesday",
    "thursday",
    "friday",
    "saturday",
    "sunday",
    "mandag",
    "tisdag",
    "onsdag",
    "torsdag",
    "fredag",
    "lordag",
    "sondag",
)
_LABEL_TOKENS = {
    "menu",
    "meny",
    "lunch",
    "middag",
    "dinner",
    "breakfast",
    "frukost",
    "kvallsmat",
    "special",
    "specialkost",
}
_REVIEW_IGNORE_TOKENS = {
    "monday",
    "tuesday",
    "wednesday",
    "thursday",
    "friday",
    "saturday",
    "sunday",
    "mon",
    "tue",
    "wed",
    "thu",
    "fri",
    "sat",
    "sun",
    "mandag",
    "tisdag",
    "onsdag",
    "torsdag",
    "fredag",
    "lordag",
    "sondag",
    "january",
    "february",
    "march",
    "april",
    "may",
    "june",
    "july",
    "august",
    "september",
    "october",
    "november",
    "december",
    "januari",
    "februari",
    "mars",
    "maj",
    "juni",
    "juli",
    "augusti",
    "oktober",
}
_MEAL_LABEL_TOKENS = {
    "lunch",
    "middag",
    "dinner",
    "breakfast",
    "frukost",
    "kvallsmat",
    "special",
    "specialkost",
}

_WEEKDAY_TOKENS = {
    "monday",
    "tuesday",
    "wednesday",
    "thursday",
    "friday",
    "saturday",
    "sunday",
    "mon",
    "tue",
    "wed",
    "thu",
    "fri",
    "sat",
    "sun",
    "mandag",
    "tisdag",
    "onsdag",
    "torsdag",
    "fredag",
    "lordag",
    "sondag",
}
_LEADING_LABEL_PREFIX_RE = re.compile(r"^([A-Za-zÅÄÖåäö\s]+):\s*")
_LEADING_LABEL_PREFIX_TOKENS = {
    "dessert",
    "kvall",
    "lunch",
    "middag",
    "frukost",
    "supper",
    "dinner",
}
_PREFIX_STRIP_RE = re.compile(
    r"^(?:"
    r"menyval\s*\d+|menu\s*choice\s*\d+|menu\s*val\s*\d+|"
    r"alt\s*[12]|alternativ\s*[12]|"
    r"lunch|middag|kvall|kvallsmat|dinner|breakfast|frokost|"
    r"monday|tuesday|wednesday|thursday|friday|saturday|sunday|"
    r"mandag|tisdag|onsdag|torsdag|fredag|lordag|sondag"
    r")\s*[:\-]\s*",
    flags=re.IGNORECASE,
)
_STANDALONE_SERVING_WORD_RE = re.compile(r"\b(?:serveras|serveres|servert)\b(?!\s+med\b)", flags=re.IGNORECASE)
_STRONG_COMPONENT_CONNECTOR_RE = re.compile(
    r"\b(?:serveras\s+med|serveres\s+med|servert\s+med|served\s+with|with|med)\b",
    flags=re.IGNORECASE,
)
_TOP_LEVEL_COMPONENT_CONNECTOR_RE = re.compile(r"\s+(?:och|og|samt)\s+", flags=re.IGNORECASE)
_TAIL_COMPONENT_CONNECTOR_RE = re.compile(r"\s+(?:och|og|samt|and)\s+", flags=re.IGNORECASE)
_COMPONENT_NOISE_RE = re.compile(r"^(?:serveras(?:\s+med)?|served\s+with|with|med)$", flags=re.IGNORECASE)
_LEADING_COMPONENT_FILLER_RE = re.compile(r"^(?:serveras(?:\s+med)?|served\s+with|with|med)\b\s*", flags=re.IGNORECASE)
_TRAILING_COMPONENT_FILLER_RE = re.compile(r"\s*\b(?:serveras(?:\s+med)?|served\s+with|with|med)\b$", flags=re.IGNORECASE)
_DESCRIPTIVE_COMPONENT_FRAGMENT_RE = re.compile(
    r"^(?:med\s+smak\s+av|smak\s+av|med\s+inslag\s+av|inslag\s+av|smaksatt(?:\s+med)?)\b",
    flags=re.IGNORECASE,
)
DEFAULT_COMPONENT_CATEGORIES = ("main", "side", "sauce", "dessert")


@dataclass(frozen=True)
class BuilderFileImportLine:
    raw_text: str
    normalized_text: str
    classification: str
    reason: str | None = None


@dataclass(frozen=True)
class BuilderFileImportPreview:
    file_type: str
    lines: list[str]
    importable_lines: list[str]
    ignored_lines: list[BuilderFileImportLine]
    classified_lines: list[BuilderFileImportLine]
    csv_column: str | None = None
    csv_column_index: int | None = None


def parse_builder_import_file(
    file_storage: FileStorage,
    *,
    csv_column: str | None = None,
) -> BuilderFileImportPreview:
    filename = str(getattr(file_storage, "filename", "") or "").strip()
    if not filename:
        raise ValueError("file name is required")

    lower_name = filename.lower()
    if not any(lower_name.endswith(ext) for ext in _SUPPORTED_EXTENSIONS):
        raise ValueError("unsupported file type; use .txt, .csv, .xlsx, or .docx")

    raw_bytes = file_storage.read()
    if not raw_bytes:
        raise ValueError("file is empty")

    if lower_name.endswith(".txt"):
        return _build_preview(file_type="txt", lines=_parse_txt_lines(raw_bytes))

    if lower_name.endswith(".xlsx"):
        lines, used_column, used_index = _parse_xlsx_lines(raw_bytes, csv_column=csv_column)
        return _build_preview(
            file_type="xlsx",
            lines=lines,
            csv_column=used_column,
            csv_column_index=used_index,
        )

    if lower_name.endswith(".docx"):
        return _build_preview(file_type="docx", lines=_parse_docx_lines(raw_bytes))

    lines, used_column, used_index = _parse_csv_lines(raw_bytes, csv_column=csv_column)
    return _build_preview(
        file_type="csv",
        lines=lines,
        csv_column=used_column,
        csv_column_index=used_index,
    )


def _build_preview(
    *,
    file_type: str,
    lines: list[str],
    csv_column: str | None = None,
    csv_column_index: int | None = None,
) -> BuilderFileImportPreview:
    classified = classify_builder_import_lines(lines)
    importable = [item.normalized_text for item in classified if item.classification == "importable_dish"]
    ignored = [item for item in classified if item.classification == "ignored_noise"]
    if not importable:
        raise ValueError("file contains no importable dish lines")

    return BuilderFileImportPreview(
        file_type=file_type,
        lines=importable,
        importable_lines=importable,
        ignored_lines=ignored,
        classified_lines=classified,
        csv_column=csv_column,
        csv_column_index=csv_column_index,
    )


def _decode_utf8_text(raw_bytes: bytes) -> str:
    try:
        return raw_bytes.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise ValueError("file must be UTF-8 encoded") from exc


def _normalize_non_empty_lines(values: list[str]) -> list[str]:
    normalized = [str(item or "").strip() for item in values]
    return [item for item in normalized if item]


def _parse_txt_lines(raw_bytes: bytes) -> list[str]:
    text = _decode_utf8_text(raw_bytes)
    lines = [str(item or "").strip() for item in text.splitlines()]
    if not lines:
        raise ValueError("file contains no importable lines")
    return lines


def _parse_csv_lines(
    raw_bytes: bytes,
    *,
    csv_column: str | None,
) -> tuple[list[str], str | None, int]:
    text = _decode_utf8_text(raw_bytes)
    reader = csv.reader(io.StringIO(text, newline=""))
    rows = [row for row in reader if any(str(cell or "").strip() for cell in row)]
    if not rows:
        raise ValueError("file contains no importable rows")

    extracted, used_column, used_index = _extract_lines_from_tabular_rows(rows, csv_column=csv_column)

    if not extracted:
        raise ValueError("file contains no importable lines")

    return extracted, used_column, used_index


def _parse_xlsx_lines(
    raw_bytes: bytes,
    *,
    csv_column: str | None,
) -> tuple[list[str], str | None, int]:
    workbook = None
    try:
        workbook = load_workbook(filename=io.BytesIO(raw_bytes), read_only=True, data_only=True)
    except Exception as exc:
        raise ValueError("invalid xlsx file") from exc

    if not workbook.worksheets:
        raise ValueError("file contains no importable rows")

    rows: list[list[str]] = []
    try:
        worksheet = workbook.worksheets[0]
        for row in worksheet.iter_rows(values_only=True):
            normalized_row = [str(cell or "").strip() for cell in row]
            if any(cell for cell in normalized_row):
                rows.append(normalized_row)
    finally:
        if workbook is not None:
            workbook.close()

    if not rows:
        raise ValueError("file contains no importable rows")

    extracted, used_column, used_index = _extract_lines_from_tabular_rows(rows, csv_column=csv_column)

    if not extracted:
        raise ValueError("file contains no importable lines")

    return extracted, used_column, used_index


def _parse_docx_lines(raw_bytes: bytes) -> list[str]:
    try:
        document = Document(io.BytesIO(raw_bytes))
    except Exception as exc:
        raise ValueError("invalid docx file") from exc

    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = str(paragraph.text or "").strip()
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = str(cell.text or "").replace("\n", " ").strip()
                if text:
                    lines.append(text)

    if not lines:
        raise ValueError("file contains no importable rows")

    return lines


def _extract_lines_from_tabular_rows(
    rows: list[list[str]],
    *,
    csv_column: str | None,
) -> tuple[list[str], str | None, int]:
    header_row = [str(cell or "").strip() for cell in rows[0]]
    header_tokens = [cell.lower() for cell in header_row]

    selected_index = 0
    selected_column = None
    skip_header = False

    csv_column_value = str(csv_column or "").strip()
    if csv_column_value:
        if csv_column_value.isdigit():
            selected_index = int(csv_column_value)
            if selected_index < 0:
                raise ValueError("csv_column index must be >= 0")
            selected_column = str(selected_index)
            skip_header = any(token in _TEXT_HEADER_TOKENS for token in header_tokens)
        else:
            match_index = next(
                (
                    idx
                    for idx, token in enumerate(header_tokens)
                    if token == csv_column_value.lower()
                ),
                None,
            )
            if match_index is None:
                raise ValueError("csv_column not found in header")
            selected_index = match_index
            selected_column = header_row[selected_index]
            skip_header = True
    else:
        detected_index = next(
            (
                idx
                for idx, token in enumerate(header_tokens)
                if token in _TEXT_HEADER_TOKENS
            ),
            None,
        )
        if detected_index is not None:
            selected_index = detected_index
            selected_column = header_row[selected_index]
            skip_header = True
        else:
            selected_index = 0
            selected_column = header_row[0] if header_row else "0"
            skip_header = False

    extracted: list[str] = []
    start_row = 1 if skip_header else 0
    for row in rows[start_row:]:
        value = row[selected_index] if selected_index < len(row) else ""
        normalized = str(value or "").strip()
        if normalized:
            extracted.append(normalized)

    return extracted, selected_column, selected_index


def classify_builder_import_lines(lines: list[str]) -> list[BuilderFileImportLine]:
    classified: list[BuilderFileImportLine] = []
    for line in lines:
        normalized = sanitize_builder_import_text(str(line or ""))
        classification, reason = _classify_single_line(normalized)
        classified.append(
            BuilderFileImportLine(
                raw_text=str(line or ""),
                normalized_text=normalized,
                classification=classification,
                reason=reason,
            )
        )
    return classified


def _classify_single_line(normalized: str) -> tuple[str, str | None]:
    if not normalized:
        return "ignored_noise", "blank"

    collapsed = re.sub(r"\s+", " ", normalized).strip()
    lower = collapsed.lower().strip(" :.-")
    folded_lower = _fold_ascii(lower)
    alnum = re.sub(r"[^a-z0-9]+", "", lower)

    if re.fullmatch(r"alt\s*[12]", lower):
        return "ignored_noise", "alt_marker"

    if lower in {"alt", "alt1", "alt2"}:
        return "ignored_noise", "alt_marker"

    if re.fullmatch(r"(?:week|vecka|v)\s*\d+", folded_lower):
        return "ignored_noise", "heading"

    if folded_lower in _WEEKDAY_TOKENS:
        return "ignored_noise", "weekday_or_date"

    if re.fullmatch(r"(?:alt|alternativ|menu\s*val|menyval)\s*\d+", folded_lower):
        return "ignored_noise", "alt_marker"

    if folded_lower in _MEAL_LABEL_TOKENS:
        return "ignored_noise", "label"

    if folded_lower in _REVIEW_IGNORE_TOKENS:
        return "ignored_noise", "weekday_or_date"

    if lower in _LABEL_TOKENS:
        return "ignored_noise", "label"

    if len(alnum) <= 1:
        return "ignored_noise", "near_blank"

    if re.fullmatch(r"[-–—_=+*/|.:,;\s]+", collapsed):
        return "ignored_noise", "separator"

    return "importable_dish", None


def sanitize_builder_import_text(value: str) -> str:
    text = str(value or "").strip()
    if not text:
        return ""

    collapsed = re.sub(r"\s+", " ", text)
    collapsed = collapsed.strip(" \t\r\n-|;:,.•")

    while True:
        updated = _strip_leading_label_prefix(collapsed)
        updated = _PREFIX_STRIP_RE.sub("", updated).strip()
        if updated == collapsed:
            break
        collapsed = updated

    collapsed = collapsed.strip(" \t\r\n-|;:,.•")
    return collapsed


def _strip_leading_label_prefix(value: str) -> str:
    text = str(value or "").strip()
    if not text:
        return ""

    matched = _LEADING_LABEL_PREFIX_RE.match(text)
    if not matched:
        return text

    token = _fold_ascii(matched.group(1)).strip()
    if token not in _LEADING_LABEL_PREFIX_TOKENS:
        return text

    return text[matched.end() :].strip()


def _normalize_import_component_candidate(candidate: str) -> str | None:
    text = str(candidate or "").strip(" \t\r\n-|;:,.•")
    while text:
        updated = _LEADING_COMPONENT_FILLER_RE.sub("", text).strip(" \t\r\n-|;:,.•")
        updated = _TRAILING_COMPONENT_FILLER_RE.sub("", updated).strip(" \t\r\n-|;:,.•")
        if updated == text:
            break
        text = updated
    if not text:
        return None
    if _COMPONENT_NOISE_RE.fullmatch(text):
        return None
    if _DESCRIPTIVE_COMPONENT_FRAGMENT_RE.match(text):
        return None
    return text[:1].upper() + text[1:]


def _split_import_component_candidates(text: str, *, tail_mode: bool) -> list[str]:
    source = re.sub(r"\s+", " ", str(text or "")).strip()
    if not source:
        return []

    source = _STANDALONE_SERVING_WORD_RE.sub(" ", source)
    source = re.sub(r"\s+", " ", source).strip()

    if not source:
        return []

    matches = list(_STRONG_COMPONENT_CONNECTOR_RE.finditer(source))
    if matches:
        suggestions: list[str] = []
        start = 0
        for match in matches:
            prefix = source[start:match.start()].strip()
            if prefix:
                suggestions.extend(_split_import_component_candidates(prefix, tail_mode=False))
            start = match.end()
        tail = source[start:].strip()
        if tail:
            suggestions.extend(_split_import_component_candidates(tail, tail_mode=True))
        return suggestions

    parts = _TAIL_COMPONENT_CONNECTOR_RE.split(source) if tail_mode else _TOP_LEVEL_COMPONENT_CONNECTOR_RE.split(source)
    suggestions: list[str] = []
    for part in parts:
        label = _normalize_import_component_candidate(part)
        if label and label not in suggestions:
            suggestions.append(label)
    return suggestions


def suggest_components_from_import_dish_name(name: str) -> list[str]:
    text = sanitize_builder_import_text(name)
    if not text:
        return []

    return _split_import_component_candidates(text, tail_mode=False)


def suggest_component_category(name: str) -> str:
    value = _fold_ascii(sanitize_builder_import_text(name))
    if not value:
        return "ovrigt"

    # Keep fish pudding in main unless a stronger sauce/side/dessert rule matches first.
    if "fiskpudding" in value:
        return "main"

    sauce_keywords = [
        "sas",
        "majonnas",
        "dressing",
        "sky",
        "vinaigrette",
        "graddsas",
    ]
    if value.endswith("sas") or any(keyword in value for keyword in sauce_keywords):
        return "sauce"

    side_keywords = [
        "potatis",
        "mos",
        "gratang",
        "ris",
        "pasta",
        "rotmos",
        "rodbetor",
        "gurka",
        "sallad",
        "stuvning",
    ]
    if any(keyword in value for keyword in side_keywords):
        return "side"

    dessert_keywords = [
        "pannacotta",
        "kaka",
        "paj",
        "mousse",
        "kram",
        "visp",
        "dessert",
        "tosca persikor",
        "persikor",
    ]
    if any(keyword in value for keyword in dessert_keywords):
        return "dessert"

    main_keywords = [
        "kyck",
        "flask",
        "kott",
        "fars",
        "isterband",
        "korv",
        "omelett",
        "fisk",
        "lax",
        "torsk",
        "sej",
        "sill",
        "tonfisk",
        "beef",
        "pork",
        "chicken",
        "meat",
    ]
    if any(keyword in value for keyword in main_keywords):
        return "main"
    return "ovrigt"


def suggest_component_tags(name: str) -> list[str]:
    value = _fold_ascii(sanitize_builder_import_text(name))
    if not value:
        return []

    ordered_rules = [
        ("fisk", ["fisk", "fish", "lax", "torsk", "sej", "sill", "tonfisk"]),
        ("kott", ["kott", "beef", "flask", "pork", "lamm", "veal", "isterband", "korv", "meat"]),
        ("kyckling", ["kyck", "chicken", "kalkon", "turkey", "anka", "duck"]),
        ("fars", ["fars", "mince"]),
        ("vegetariskt", ["vegetar", "veg", "tofu", "lins", "bon", "quorn", "halloumi", "falafel"]),
        ("italienskt", ["pasta", "lasagne", "risotto", "parmesan", "tomat", "bologn"]),
        ("asiatiskt", ["nudel", "soja", "ingefara", "curry", "wok", "teriyaki"]),
        ("husmanskost", ["husman", "kottbull", "pytt", "falukorv", "raggmunk", "kalops"]),
        ("dessert", ["dessert", "kaka", "pannacotta", "mousse", "kram", "paj", "glass", "pudding"]),
        ("sas", ["sas", "sauce", "dressing", "majonnas", "graddsas", "pesto"]),
    ]

    tags: list[str] = []
    for tag, keywords in ordered_rules:
        if any(keyword in value for keyword in keywords):
            tags.append(tag)
    return tags[:6]


def _fold_ascii(value: str) -> str:
    return (
        str(value or "")
        .replace("å", "a")
        .replace("ä", "a")
        .replace("ö", "o")
        .replace("Å", "a")
        .replace("Ä", "a")
        .replace("Ö", "o")
        .lower()
    )


__all__ = [
    "DEFAULT_COMPONENT_CATEGORIES",
    "BuilderFileImportLine",
    "BuilderFileImportPreview",
    "classify_builder_import_lines",
    "parse_builder_import_file",
    "sanitize_builder_import_text",
    "suggest_component_category",
    "suggest_component_tags",
    "suggest_components_from_import_dish_name",
]
