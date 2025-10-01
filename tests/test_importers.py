import io

from docx import Document

from core.importers.composite import CompositeMenuImporter
from core.importers.docx_importer import DocxMenuImporter
from core.importers.excel_importer import ExcelMenuImporter


def make_docx(lines):
    buf = io.BytesIO()
    doc = Document()
    for l in lines:
        p = doc.add_paragraph()
        p.add_run(l)
    doc.save(buf)
    return buf.getvalue()


def test_docx_basic_alt_detection():
    content = [
        "Vecka 12", "Mån: Alt1: Fiskgratäng", "Tis: Alt2: Köttfärssås", "Ons: Kväll: Soppa"
    ]
    data = make_docx(content)
    imp = DocxMenuImporter()
    res = imp.parse(data, "meny.docx")
    assert res.weeks, "Should parse at least one week"
    week = res.weeks[0]
    names = {i.dish_name for i in week.items}
    assert "Fiskgratäng" in names
    assert "Köttfärssås" in names
    assert "Soppa" in names
    # Evening variant should map to meal dinner and variant_type main
    evening = [i for i in week.items if i.meal=="dinner"]
    assert evening and evening[0].variant_type == "main"


def test_composite_unknown():
    comp = CompositeMenuImporter([DocxMenuImporter(), ExcelMenuImporter()])
    bogus = b"not a valid menu format"
    res = comp.parse(bogus, "foo.txt", "text/plain")
    assert res.errors and "No importer accepted file" in res.errors[0]
