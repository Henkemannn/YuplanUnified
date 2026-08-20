from __future__ import annotations

from io import BytesIO

from docx import Document
from openpyxl import Workbook
import pytest
from werkzeug.datastructures import FileStorage

from core.builder.file_import import (
    classify_builder_import_lines,
    parse_builder_import_file,
    suggest_component_category,
    suggest_component_tags,
    suggest_components_from_import_dish_name,
)


def _file(name: str, raw: bytes) -> FileStorage:
    return FileStorage(stream=BytesIO(raw), filename=name)


def _xlsx_bytes(rows: list[list[str]]) -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    for row in rows:
        sheet.append(row)

    stream = BytesIO()
    workbook.save(stream)
    workbook.close()
    return stream.getvalue()


def _docx_bytes(*, paragraphs: list[str] | None = None, table_rows: list[list[str]] | None = None) -> bytes:
    doc = Document()
    for value in paragraphs or []:
        doc.add_paragraph(str(value))

    rows = table_rows or []
    if rows:
        table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
        for row_index, row in enumerate(rows):
            for col_index, value in enumerate(row):
                table.rows[row_index].cells[col_index].text = str(value)

    stream = BytesIO()
    doc.save(stream)
    return stream.getvalue()


def test_parse_txt_file_to_normalized_lines() -> None:
    preview = parse_builder_import_file(
        _file("library.txt", b"Kottbullar med potatismos\n\n Fiskgratang  \n"),
    )

    assert preview.file_type == "txt"
    assert preview.lines == ["Kottbullar med potatismos", "Fiskgratang"]


def test_parse_csv_file_detects_likely_text_column() -> None:
    preview = parse_builder_import_file(
        _file(
            "library.csv",
            b"dish_name,category\nKottbullar med potatismos,main\nFiskgratang,main\n",
        ),
    )

    assert preview.file_type == "csv"
    assert preview.lines == ["Kottbullar med potatismos", "Fiskgratang"]
    assert preview.csv_column == "dish_name"
    assert preview.csv_column_index == 0


def test_parse_csv_file_supports_explicit_column_name() -> None:
    preview = parse_builder_import_file(
        _file(
            "library.csv",
            b"id,text,flag\n1,Kottbullar med potatismos,x\n2,Fiskgratang,y\n",
        ),
        csv_column="text",
    )

    assert preview.lines == ["Kottbullar med potatismos", "Fiskgratang"]
    assert preview.csv_column == "text"
    assert preview.csv_column_index == 1


def test_parse_csv_file_uses_first_column_when_no_text_header_found() -> None:
    preview = parse_builder_import_file(
        _file("library.csv", b"Kottbullar med potatismos,main\nFiskgratang,main\n"),
    )

    assert preview.lines == ["Kottbullar med potatismos", "Fiskgratang"]
    assert preview.csv_column_index == 0


def test_parse_xlsx_file_detects_likely_text_column() -> None:
    preview = parse_builder_import_file(
        _file(
            "library.xlsx",
            _xlsx_bytes(
                [
                    ["dish_name", "category"],
                    ["Kottbullar med potatismos", "main"],
                    ["Fiskgratang", "main"],
                ]
            ),
        ),
    )

    assert preview.file_type == "xlsx"
    assert preview.lines == ["Kottbullar med potatismos", "Fiskgratang"]
    assert preview.csv_column == "dish_name"
    assert preview.csv_column_index == 0


def test_parse_xlsx_file_supports_explicit_column_name() -> None:
    preview = parse_builder_import_file(
        _file(
            "library.xlsx",
            _xlsx_bytes(
                [
                    ["id", "text", "flag"],
                    ["1", "Kottbullar med potatismos", "x"],
                    ["2", "Fiskgratang", "y"],
                ]
            ),
        ),
        csv_column="text",
    )

    assert preview.file_type == "xlsx"
    assert preview.lines == ["Kottbullar med potatismos", "Fiskgratang"]
    assert preview.csv_column == "text"
    assert preview.csv_column_index == 1


def test_parse_rejects_unsupported_extension() -> None:
    with pytest.raises(ValueError, match="unsupported file type"):
        parse_builder_import_file(_file("library.pdf", b"x"))


def test_parse_txt_ignores_alt_markers_and_headings() -> None:
    preview = parse_builder_import_file(
        _file(
            "library.txt",
            b"Week 12\nAlt 1\nAlt 2\nMonday\nKottbullar med potatismos\n",
        ),
    )

    assert preview.importable_lines == ["Kottbullar med potatismos"]
    ignored_texts = [item.normalized_text for item in preview.ignored_lines]
    assert "Alt 1" in ignored_texts
    assert "Alt 2" in ignored_texts
    assert "Week 12" in ignored_texts
    assert "Monday" in ignored_texts


def test_parse_xlsx_ignores_markers_and_labels() -> None:
    preview = parse_builder_import_file(
        _file(
            "library.xlsx",
            _xlsx_bytes(
                [
                    ["text"],
                    ["Week 12"],
                    ["Alt 1"],
                    ["Lunch"],
                    ["Fiskgratang"],
                ]
            ),
        ),
    )

    assert preview.importable_lines == ["Fiskgratang"]
    ignored_reasons = {item.reason for item in preview.ignored_lines}
    assert "heading" in ignored_reasons
    assert "alt_marker" in ignored_reasons
    assert "label" in ignored_reasons


def test_parse_txt_keeps_valid_dishes_and_ignores_labels() -> None:
    preview = parse_builder_import_file(
        _file(
            "library.txt",
            b"Lunch\nFiskgratang\nMeny\nKottbullar med graddsas\n",
        ),
    )

    assert preview.importable_lines == ["Fiskgratang", "Kottbullar med graddsas"]
    ignored_reasons = {item.reason for item in preview.ignored_lines}
    assert "label" in ignored_reasons


def test_parse_docx_file_reads_paragraph_lines() -> None:
    preview = parse_builder_import_file(
        _file(
            "library.docx",
            _docx_bytes(paragraphs=["Week 12", "Alt 1", "Fiskgratang", "Kottbullar med potatismos"]),
        ),
    )

    assert preview.file_type == "docx"
    assert preview.importable_lines == ["Fiskgratang", "Kottbullar med potatismos"]


def test_parse_docx_file_reads_table_cells() -> None:
    preview = parse_builder_import_file(
        _file(
            "library.docx",
            _docx_bytes(
                table_rows=[
                    ["text", "tag"],
                    ["Kottbullar med potatismos", "main"],
                    ["Fiskgratang", "main"],
                ],
            ),
        ),
    )

    assert preview.file_type == "docx"
    assert "Kottbullar med potatismos" in preview.importable_lines
    assert "Fiskgratang" in preview.importable_lines


def test_classify_cleanup_strips_menu_prefixes() -> None:
    items = classify_builder_import_lines(["Menyval1:köttbullar", "Alt 1: Fiskgryta"])

    assert items[0].classification == "importable_dish"
    assert items[0].normalized_text == "köttbullar"
    assert items[1].classification == "importable_dish"
    assert items[1].normalized_text == "Fiskgryta"


def test_classify_cleanup_ignores_weekday_headers() -> None:
    items = classify_builder_import_lines(["Lördag", "Monday"])

    assert items[0].classification == "ignored_noise"
    assert items[0].reason in {"weekday_or_date", "heading"}
    assert items[1].classification == "ignored_noise"
    assert items[1].reason in {"weekday_or_date", "heading"}


def test_classify_strips_meal_prefix_dessert() -> None:
    items = classify_builder_import_lines(["Dessert: Chokladpudding"])

    assert items[0].classification == "importable_dish"
    assert items[0].normalized_text == "Chokladpudding"


def test_component_decomposition_strips_meal_prefix_kvall() -> None:
    components = suggest_components_from_import_dish_name("Kväll: Omelett med sparris")

    assert components == ["Omelett", "Sparris"]
    assert all("Dessert:" not in name for name in components)
    assert all("Kväll:" not in name for name in components)
    assert all("Lunch:" not in name for name in components)


def test_component_decomposition_strips_meal_prefix_lunch() -> None:
    components = suggest_components_from_import_dish_name("Lunch: Köttbullar med potatis")

    assert components == ["Köttbullar", "Potatis"]
    assert all("Dessert:" not in name for name in components)
    assert all("Kväll:" not in name for name in components)
    assert all("Lunch:" not in name for name in components)


def test_component_decomposition_ignores_serveras_only_suffix() -> None:
    components = suggest_components_from_import_dish_name("Soppa serveras")

    assert components == ["Soppa"]
    assert all(name.lower() != "serveras" for name in components)
    assert all(name.lower() != "serveras med" for name in components)


def test_component_decomposition_handles_serveras_med_connector() -> None:
    components = suggest_components_from_import_dish_name("Köttbullar serveras med potatis")

    assert components == ["Köttbullar", "Potatis"]
    assert all(name.lower() != "serveras" for name in components)
    assert all(name.lower() != "serveras med" for name in components)


def test_component_decomposition_handles_serveras_med_and_och() -> None:
    components = suggest_components_from_import_dish_name("Fisk serveras med kokt potatis och sås")

    assert components == ["Fisk", "Kokt potatis", "Sås"]
    assert all(name.lower() != "serveras" for name in components)
    assert all(name.lower() != "serveras med" for name in components)


@pytest.mark.parametrize(
    ("dish_name", "expected"),
    [
        ("Asiatisk svineribbe med hoisinsås og ris", ["Asiatisk svineribbe", "Hoisinsås", "Ris"]),
        ("Köttbullar med potatismos och gräddsås", ["Köttbullar", "Potatismos", "Gräddsås"]),
        ("Laks med poteter og saus", ["Laks", "Poteter", "Saus"]),
        ("Kylling serveres med ris og saus", ["Kylling", "Ris", "Saus"]),
        ("Kylling servert med ris og saus", ["Kylling", "Ris", "Saus"]),
        ("Köttbullar och potatismos", ["Köttbullar", "Potatismos"]),
        ("Laks og poteter", ["Laks", "Poteter"]),
        ("Köttbullar med potatismos samt en citrondoftande sås", ["Köttbullar", "Potatismos", "En citrondoftande sås"]),
        ("Fisk samt kokt potatis", ["Fisk", "Kokt potatis"]),
        ("Chicken with rice and sauce", ["Chicken", "Rice", "Sauce"]),
        ("Fish and chips", ["Fish and chips"]),
        ("Mac and cheese", ["Mac and cheese"]),
    ],
)
def test_component_decomposition_supports_scandinavian_and_cautious_english_connectors(
    dish_name: str,
    expected: list[str],
) -> None:
    assert suggest_components_from_import_dish_name(dish_name) == expected


def test_component_decomposition_ignores_descriptive_phrase_only_fragment() -> None:
    components = suggest_components_from_import_dish_name("Kycklinggryta med smak av dragon")

    assert components == ["Kycklinggryta"]
    assert all("smak av" not in name.lower() for name in components)


def test_suggest_component_category_rules_are_lightweight() -> None:
    assert suggest_component_category("Köttfärssås") == "sauce"
    assert suggest_component_category("Senapssås") == "sauce"
    assert suggest_component_category("Vitvinssås") == "sauce"
    assert suggest_component_category("Äppelcidersås") == "sauce"
    assert suggest_component_category("Gräddsås") == "sauce"
    assert suggest_component_category("Gurkmajonnäs") == "sauce"
    assert suggest_component_category("Hallonsås") == "sauce"
    assert suggest_component_category("Kokt potatis") == "side"
    assert suggest_component_category("Potatisgratäng") == "side"
    assert suggest_component_category("Potatismos") == "side"
    assert suggest_component_category("Rotmos") == "side"
    assert suggest_component_category("Stuvad potatis") == "side"
    assert suggest_component_category("Kycklingfilé") == "main"
    assert suggest_component_category("Laxfilé") == "main"
    assert suggest_component_category("Stekt fläsk") == "main"
    assert suggest_component_category("Omelett") == "main"
    assert suggest_component_category("Kycklinggryta") == "main"
    assert suggest_component_category("Fläskbog") == "main"
    assert suggest_component_category("Isterband") == "main"
    assert suggest_component_category("Rödbetor") == "side"
    assert suggest_component_category("Salt gurka") == "side"
    assert suggest_component_category("Majonnäs") == "sauce"
    assert suggest_component_category("Pannacotta") == "dessert"
    assert suggest_component_category("Vaniljvisp") == "dessert"
    assert suggest_component_category("Tosca persikor") == "dessert"
    assert suggest_component_category("Chokladmousse") == "dessert"
    assert suggest_component_category("Björnbärskräm") == "dessert"
    assert suggest_component_category("Chokladkaka") == "dessert"
    assert suggest_component_category("Fiskpudding") == "main"
    assert suggest_component_category("Vatten") == "ovrigt"


def test_suggest_component_tags_returns_keyword_matches() -> None:
    assert suggest_component_tags("Lax med kokt potatis") == ["fisk"]
    assert suggest_component_tags("Kycklinggryta") == ["kyckling"]
    assert suggest_component_tags("Vegetarisk soppa") == ["vegetariskt"]
    assert suggest_component_tags("Nötkött i sås") == ["kott", "sas"]
    assert suggest_component_tags("Potatissallad") == []
