from __future__ import annotations

import io
import tempfile
import os

from docx import Document
from openpyxl import Workbook
from core.app_factory import create_app

HEADERS = {"X-User-Role": "admin", "X-Tenant-Id": "1", "X-User-Id": "11"}


def _headers(*, role: str = "admin", tenant_id: int = 1, user_id: int = 11) -> dict[str, str]:
    return {
        "X-User-Role": role,
        "X-Tenant-Id": str(tenant_id),
        "X-User-Id": str(user_id),
    }


def _seed_session(client, *, role: str, tenant_id: int, user_id: int) -> None:
    with client.session_transaction() as sess:
        sess["role"] = role
        sess["tenant_id"] = tenant_id
        sess["user_id"] = user_id


def _app_with_builder_db(db_path: str):
    previous = os.environ.get("BUILDER_DB_PATH")
    os.environ["BUILDER_DB_PATH"] = db_path
    try:
        return create_app({"TESTING": True})
    finally:
        if previous is None:
            os.environ.pop("BUILDER_DB_PATH", None)
        else:
            os.environ["BUILDER_DB_PATH"] = previous


def _client():
    app = create_app({"TESTING": True})
    return app.test_client()


def _sqlite_client():
    fd, db_path = tempfile.mkstemp(prefix="builder_import_session_", suffix=".db")
    os.close(fd)
    app = create_app({"TESTING": True, "BUILDER_DB_PATH": db_path})
    return app.test_client()


def _xlsx_bytes(rows: list[list[str]]) -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    for row in rows:
        sheet.append(row)

    stream = io.BytesIO()
    workbook.save(stream)
    workbook.close()
    return stream.getvalue()


def _docx_bytes(*, paragraphs: list[str] | None = None, table_rows: list[list[str]] | None = None) -> bytes:
    doc = Document()
    for value in paragraphs or []:
        doc.add_paragraph(str(value))

    rows = table_rows or []
    if rows:
        table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
        for row_index, row in enumerate(rows):
            for col_index, value in enumerate(row):
                table.rows[row_index].cells[col_index].text = str(value)

    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()


def test_create_composition_endpoint() -> None:
    client = _client()

    rv = client.post(
        "/api/builder/compositions",
        json={
            "composition_id": "plate_1",
            "composition_name": "Fish Plate",
            "library_group": "weekly",
        },
        headers=HEADERS,
    )

    assert rv.status_code == 201
    body = rv.get_json() or {}
    assert body.get("ok") is True
    assert body.get("composition", {}).get("composition_id") == "plate_1"
    assert body.get("composition", {}).get("composition_name") == "Fish Plate"


def test_create_composition_endpoint_supports_generated_id_without_menu_context() -> None:
    client = _client()

    rv = client.post(
        "/api/builder/compositions",
        json={
            "composition_name": "Free Dish",
        },
        headers=HEADERS,
    )

    assert rv.status_code == 201
    body = rv.get_json() or {}
    assert body.get("ok") is True
    composition = body.get("composition") or {}
    generated_id = composition.get("composition_id")
    assert isinstance(generated_id, str)
    assert generated_id.startswith("cmp_")
    assert len(generated_id) == 10
    assert composition.get("composition_name") == "Free Dish"
    components = composition.get("components") or []
    assert len(components) == 1
    assert components[0].get("component_name") == "Free Dish"


def test_create_composition_endpoint_supports_empty_shell_when_seed_components_false() -> None:
    client = _client()

    rv = client.post(
        "/api/builder/compositions",
        json={
            "composition_name": "Open shell",
            "library_group": "kott",
            "seed_components": False,
        },
        headers=HEADERS,
    )

    assert rv.status_code == 201
    body = rv.get_json() or {}
    assert body.get("ok") is True
    composition = body.get("composition") or {}
    assert composition.get("composition_name") == "Open shell"
    assert composition.get("library_group") == "kott"
    assert composition.get("components") == []


def test_create_composition_endpoint_accepts_legacy_library_group() -> None:
    client = _client()

    rv = client.post(
        "/api/builder/compositions",
        json={
            "composition_name": "Open shell",
            "library_group": "weekly",
            "seed_components": False,
        },
        headers=HEADERS,
    )

    assert rv.status_code == 201
    data = rv.get_json()
    assert data["ok"] is True
    assert data["composition"]["library_group"] == "weekly"
    assert data["composition"]["components"] == []


def test_create_composition_endpoint_rejects_empty_name() -> None:
    client = _client()

    rv = client.post(
        "/api/builder/compositions",
        json={
            "composition_name": "",
            "seed_components": False,
        },
        headers=HEADERS,
    )

    assert rv.status_code == 400
    body = rv.get_json() or {}
    assert body.get("error") == "bad_request"


def test_patch_composition_metadata_endpoint_updates_name_and_category_without_changing_components() -> None:
    client = _client()
    created = client.post(
        "/api/builder/compositions",
        json={
            "composition_id": "plate_patch_meta",
            "composition_name": "Original Plate",
            "library_group": "ovrigt",
        },
        headers=HEADERS,
    )
    assert created.status_code == 201
    client.post(
        "/api/builder/compositions/plate_patch_meta/components",
        json={"component_name": "Fisk", "role": "main"},
        headers=HEADERS,
    )
    before = client.get("/api/builder/compositions", headers=HEADERS).get_json() or {}
    before_comp = next(item for item in (before.get("compositions") or []) if item.get("composition_id") == "plate_patch_meta")
    before_components = before_comp.get("components") or []

    rv = client.patch(
        "/api/builder/compositions/plate_patch_meta",
        json={
            "composition_name": "Updated Plate",
            "library_group": "fisk",
        },
        headers=HEADERS,
    )

    assert rv.status_code == 200
    body = rv.get_json() or {}
    assert body.get("ok") is True
    composition = body.get("composition") or {}
    assert composition.get("composition_name") == "Updated Plate"
    assert composition.get("library_group") == "fisk"
    assert (composition.get("components") or []) == before_components


def test_patch_composition_metadata_endpoint_rejects_empty_name() -> None:
    client = _client()
    client.post(
        "/api/builder/compositions",
        json={"composition_id": "plate_patch_empty", "composition_name": "Original"},
        headers=HEADERS,
    )

    rv = client.patch(
        "/api/builder/compositions/plate_patch_empty",
        json={"composition_name": "   "},
        headers=HEADERS,
    )

    assert rv.status_code == 400
    body = rv.get_json() or {}
    assert body.get("error") == "bad_request"


def test_patch_composition_metadata_endpoint_rejects_invalid_library_group() -> None:
    client = _client()
    client.post(
        "/api/builder/compositions",
        json={"composition_id": "plate_patch_bad_group", "composition_name": "Original"},
        headers=HEADERS,
    )

    rv = client.patch(
        "/api/builder/compositions/plate_patch_bad_group",
        json={"library_group": "weekly"},
        headers=HEADERS,
    )

    assert rv.status_code == 400
    body = rv.get_json() or {}
    assert body.get("error") == "bad_request"


def test_patch_composition_metadata_endpoint_supports_custom_menu_name_toggle() -> None:
    client = _client()
    created = client.post(
        "/api/builder/compositions",
        json={"composition_id": "plate_custom", "composition_name": "Lasagne"},
        headers=HEADERS,
    )
    assert created.status_code == 201

    enabled = client.patch(
        "/api/builder/compositions/plate_custom",
        json={
            "use_custom_menu_name": True,
            "menu_name": "  Hemlagad lasagne med tomat  ",
        },
        headers=HEADERS,
    )

    assert enabled.status_code == 200
    enabled_body = enabled.get_json() or {}
    enabled_composition = enabled_body.get("composition") or {}
    assert enabled_composition.get("use_custom_menu_name") is True
    assert enabled_composition.get("menu_name") == "Hemlagad lasagne med tomat"
    assert enabled_composition.get("effective_menu_name") == "Hemlagad lasagne med tomat"

    disabled = client.patch(
        "/api/builder/compositions/plate_custom",
        json={"use_custom_menu_name": False},
        headers=HEADERS,
    )

    assert disabled.status_code == 200
    disabled_body = disabled.get_json() or {}
    disabled_composition = disabled_body.get("composition") or {}
    assert disabled_composition.get("use_custom_menu_name") is False
    assert disabled_composition.get("menu_name") == "Hemlagad lasagne med tomat"
    assert disabled_composition.get("effective_menu_name") == "Lasagne"


def test_patch_composition_metadata_endpoint_returns_not_found_for_unknown_composition() -> None:
    client = _client()

    rv = client.patch(
        "/api/builder/compositions/missing_plate",
        json={"composition_name": "Updated"},
        headers=HEADERS,
    )

    assert rv.status_code == 404
    body = rv.get_json() or {}
    assert body.get("error") == "not_found"


def test_linked_component_edit_target_rewires_private_dish_and_copies_details_once() -> None:
    fd, db_path = tempfile.mkstemp(prefix="builder_api_linked_edit_target_", suffix=".db")
    os.close(fd)
    app = _app_with_builder_db(db_path)
    admin_client = app.test_client()
    cook_a_client = app.test_client()
    cook_b_client = app.test_client()
    admin_headers = _headers(role="admin", tenant_id=9001, user_id=3)
    cook_a_headers = _headers(role="cook", tenant_id=9001, user_id=4)
    cook_b_headers = _headers(role="cook", tenant_id=9001, user_id=5)
    _seed_session(admin_client, role="admin", tenant_id=9001, user_id=3)
    _seed_session(cook_a_client, role="cook", tenant_id=9001, user_id=4)
    _seed_session(cook_b_client, role="cook", tenant_id=9001, user_id=5)

    shared_component_rv = admin_client.post(
        "/api/builder/components",
        json={"component_name": "Linked dish fish"},
        headers=admin_headers,
    )
    assert shared_component_rv.status_code == 201
    shared_component_id = str(((shared_component_rv.get_json() or {}).get("component") or {}).get("component_id") or "")
    assert shared_component_id

    admin_client.patch(
        f"/api/builder/components/{shared_component_id}/details",
        json={
            "recipe_ingredient_rows": [{"ingredient_name": "Fish", "amount_value": "1", "amount_unit": "kg"}],
            "method_text": "Steam the fish",
            "calculation_cost": "123.45",
            "calculation_rows": [{"ingredient_name": "Fish", "amount_value": "1", "amount_unit": "kg"}],
            "allergens": ["fish"],
            "tags": ["seafood"],
        },
        headers=admin_headers,
    )

    shared_dish_rv = admin_client.post(
        "/api/builder/compositions",
        json={"composition_id": "cmp_5mr4rq", "composition_name": "Shared dish"},
        headers=admin_headers,
    )
    assert shared_dish_rv.status_code == 201
    admin_client.post(
        "/api/builder/compositions/cmp_5mr4rq/components",
        json={"component_name": "Linked dish fish", "role": "main"},
        headers=admin_headers,
    )

    cook_a_private_dish = cook_a_client.post(
        "/api/builder/compositions/cmp_5mr4rq/edit-target",
        headers=cook_a_headers,
    )
    assert cook_a_private_dish.status_code == 200
    cook_a_private_dish_id = ((cook_a_private_dish.get_json() or {}).get("composition") or {}).get("composition_id")
    assert cook_a_private_dish_id and cook_a_private_dish_id != "cmp_5mr4rq"

    linked_target = cook_a_client.post(
        f"/api/builder/compositions/{cook_a_private_dish_id}/components/{shared_component_id}/edit-target",
        headers=cook_a_headers,
    )
    assert linked_target.status_code == 200
    linked_body = linked_target.get_json() or {}
    assert linked_body.get("ok") is True
    assert linked_body.get("forked") is True
    private_component_id = ((linked_body.get("component") or {}).get("component_id") or "")
    assert private_component_id and private_component_id != shared_component_id
    linked_component_entry = ((linked_body.get("composition") or {}).get("components") or [])[0]
    assert ((linked_body.get("composition") or {}).get("components") or [])[0].get("component_id") == private_component_id
    assert linked_component_entry.get("role") == "main"
    source_component_entry = (admin_client.get("/api/builder/compositions/cmp_5mr4rq", headers=admin_headers).get_json() or {}).get("composition") or {}
    source_link_entry = (source_component_entry.get("components") or [])[0]
    assert linked_component_entry.get("sort_order") == source_link_entry.get("sort_order")

    repeated_target = cook_a_client.post(
        f"/api/builder/compositions/{cook_a_private_dish_id}/components/{shared_component_id}/edit-target",
        headers=cook_a_headers,
    )
    assert repeated_target.status_code == 200
    assert (((repeated_target.get_json() or {}).get("component") or {}).get("component_id")) == private_component_id

    copied_details = cook_a_client.get(f"/api/builder/components/{private_component_id}/details", headers=cook_a_headers)
    assert copied_details.status_code == 200
    copied_body = copied_details.get_json() or {}
    copied_details_payload = copied_body.get("details") or {}
    assert copied_details_payload.get("method_text") == "Steam the fish"
    assert copied_details_payload.get("recipe_ingredient_rows") == [{"ingredient_name": "Fish", "amount_value": "1", "amount_unit": "kg"}]
    assert copied_details_payload.get("calculation_cost") == "123.45"
    assert copied_details_payload.get("allergens") == ["fish"]

    cook_a_private_patch = cook_a_client.patch(
        f"/api/builder/components/{private_component_id}/details",
        json={"method_text": "Cook A method"},
        headers=cook_a_headers,
    )
    assert cook_a_private_patch.status_code == 200

    shared_after = admin_client.get(f"/api/builder/components/{shared_component_id}/details", headers=admin_headers)
    assert shared_after.status_code == 200
    assert ((shared_after.get_json() or {}).get("details") or {}).get("method_text") == "Steam the fish"

    cook_a_shared_patch = cook_a_client.patch(
        f"/api/builder/components/{shared_component_id}/details",
        json={"method_text": "Blocked"},
        headers=cook_a_headers,
    )
    assert cook_a_shared_patch.status_code == 400

    cook_b_private_dish = cook_b_client.post(
        "/api/builder/compositions/cmp_5mr4rq/edit-target",
        headers=cook_b_headers,
    )
    assert cook_b_private_dish.status_code == 200
    cook_b_private_dish_id = ((cook_b_private_dish.get_json() or {}).get("composition") or {}).get("composition_id")
    assert cook_b_private_dish_id and cook_b_private_dish_id != cook_a_private_dish_id

    cook_b_target = cook_b_client.post(
        f"/api/builder/compositions/{cook_b_private_dish_id}/components/{shared_component_id}/edit-target",
        headers=cook_b_headers,
    )
    assert cook_b_target.status_code == 200
    cook_b_private_component_id = ((cook_b_target.get_json() or {}).get("component") or {}).get("component_id")
    assert cook_b_private_component_id and cook_b_private_component_id != private_component_id

    cook_b_private_details = cook_b_client.get(f"/api/builder/components/{private_component_id}/details", headers=cook_b_headers)
    assert cook_b_private_details.status_code == 400

    cook_b_initial_details = cook_b_client.get(f"/api/builder/components/{cook_b_private_component_id}/details", headers=cook_b_headers)
    assert cook_b_initial_details.status_code == 200
    assert ((cook_b_initial_details.get_json() or {}).get("details") or {}).get("method_text") == "Steam the fish"


def test_composition_edit_target_endpoint_reuses_private_fork_and_preserves_source() -> None:
    fd, db_path = tempfile.mkstemp(prefix="builder_api_edit_target_", suffix=".db")
    os.close(fd)
    app = _app_with_builder_db(db_path)
    admin_client = app.test_client()
    cook_a_client = app.test_client()
    cook_b_client = app.test_client()
    admin_headers = _headers(role="admin", tenant_id=1, user_id=3)
    cook_a_headers = _headers(role="cook", tenant_id=1, user_id=4)
    cook_b_headers = _headers(role="cook", tenant_id=1, user_id=5)
    _seed_session(admin_client, role="admin", tenant_id=1, user_id=3)
    _seed_session(cook_a_client, role="cook", tenant_id=1, user_id=4)
    _seed_session(cook_b_client, role="cook", tenant_id=1, user_id=5)

    created = admin_client.post(
        "/api/builder/compositions",
        json={"composition_id": "cmp_5mr4rq", "composition_name": "Shared dish"},
        headers=admin_headers,
    )
    assert created.status_code == 201

    first = cook_a_client.post(
        "/api/builder/compositions/cmp_5mr4rq/edit-target",
        headers=cook_a_headers,
    )
    second = cook_a_client.post(
        "/api/builder/compositions/cmp_5mr4rq/edit-target",
        headers=cook_a_headers,
    )
    other = cook_b_client.post(
        "/api/builder/compositions/cmp_5mr4rq/edit-target",
        headers=cook_b_headers,
    )

    assert first.status_code == 200
    assert second.status_code == 200
    assert other.status_code == 200

    first_body = first.get_json() or {}
    second_body = second.get_json() or {}
    other_body = other.get_json() or {}
    assert first_body.get("ok") is True
    assert first_body.get("forked") is True
    assert first_body.get("source_composition_id") == "cmp_5mr4rq"
    assert (first_body.get("composition") or {}).get("composition_id") != "cmp_5mr4rq"
    assert (second_body.get("composition") or {}).get("composition_id") == (first_body.get("composition") or {}).get("composition_id")
    assert (other_body.get("composition") or {}).get("composition_id") != (first_body.get("composition") or {}).get("composition_id")

    source = admin_client.get("/api/builder/compositions/cmp_5mr4rq", headers=admin_headers)
    assert source.status_code == 200
    assert ((source.get_json() or {}).get("composition") or {}).get("composition_name") == "Shared dish"

    fork_read = cook_b_client.get(
        f"/api/builder/compositions/{(first_body.get('composition') or {}).get('composition_id')}",
        headers=cook_b_headers,
    )
    assert fork_read.status_code == 400
    assert (fork_read.get_json() or {}).get("error") == "bad_request"


def test_cook_can_patch_private_composition_but_not_shared_source() -> None:
    fd, db_path = tempfile.mkstemp(prefix="builder_api_private_patch_", suffix=".db")
    os.close(fd)
    app = _app_with_builder_db(db_path)
    admin_client = app.test_client()
    cook_client = app.test_client()
    admin_headers = _headers(role="admin", tenant_id=1, user_id=3)
    cook_headers = _headers(role="cook", tenant_id=1, user_id=4)
    _seed_session(admin_client, role="admin", tenant_id=1, user_id=3)
    _seed_session(cook_client, role="cook", tenant_id=1, user_id=4)

    created = admin_client.post(
        "/api/builder/compositions",
        json={"composition_id": "cmp_5mr4rq", "composition_name": "Shared dish"},
        headers=admin_headers,
    )
    assert created.status_code == 201

    shared_attempt = cook_client.patch(
        "/api/builder/compositions/cmp_5mr4rq",
        json={"composition_name": "Cook blocked"},
        headers=cook_headers,
    )
    assert shared_attempt.status_code == 404

    edit_target = cook_client.post(
        "/api/builder/compositions/cmp_5mr4rq/edit-target",
        headers=cook_headers,
    )
    private_id = ((edit_target.get_json() or {}).get("composition") or {}).get("composition_id")
    assert private_id and private_id != "cmp_5mr4rq"

    private_patch = cook_client.patch(
        f"/api/builder/compositions/{private_id}",
        json={"composition_name": "Cook editable dish", "library_group": "fisk"},
        headers=cook_headers,
    )
    assert private_patch.status_code == 200
    private_body = private_patch.get_json() or {}
    assert (private_body.get("composition") or {}).get("composition_name") == "Cook editable dish"
    assert (private_body.get("composition") or {}).get("library_group") == "fisk"

    source = admin_client.get("/api/builder/compositions/cmp_5mr4rq", headers=admin_headers)
    assert source.status_code == 200
    assert ((source.get_json() or {}).get("composition") or {}).get("composition_name") == "Shared dish"


def test_component_patch_passes_actor_for_scope_enforcement() -> None:
    fd, db_path = tempfile.mkstemp(prefix="builder_api_component_scope_", suffix=".db")
    os.close(fd)
    app = _app_with_builder_db(db_path)
    admin_client = app.test_client()
    cook_client = app.test_client()
    admin_headers = _headers(role="admin", tenant_id=1, user_id=3)
    cook_headers = _headers(role="cook", tenant_id=1, user_id=4)
    _seed_session(admin_client, role="admin", tenant_id=1, user_id=3)
    _seed_session(cook_client, role="cook", tenant_id=1, user_id=4)

    created = admin_client.post(
        "/api/builder/components",
        json={"component_name": "Shared component soup"},
        headers=admin_headers,
    )
    assert created.status_code == 201
    component_id = ((created.get_json() or {}).get("component") or {}).get("component_id")

    shared_patch = cook_client.patch(
        f"/api/builder/components/{component_id}",
        json={"component_name": "Cook blocked"},
        headers=cook_headers,
    )
    assert shared_patch.status_code == 400

    forked = cook_client.post(f"/api/builder/components/{component_id}/fork", headers=cook_headers)
    forked_id = ((forked.get_json() or {}).get("component") or {}).get("component_id")
    assert forked_id and forked_id != component_id

    private_patch = cook_client.patch(
        f"/api/builder/components/{forked_id}",
        json={"component_name": "Cook allowed", "category": "side"},
        headers=cook_headers,
    )
    assert private_patch.status_code == 200
    private_body = private_patch.get_json() or {}
    assert (private_body.get("component") or {}).get("component_name") == "Cook allowed"
    assert (private_body.get("component") or {}).get("category") == "side"

    source = admin_client.get(f"/api/builder/components/{component_id}", headers=admin_headers)
    assert source.status_code == 200
    assert ((source.get_json() or {}).get("component") or {}).get("component_name") == "Shared component soup"


def test_single_component_read_endpoint_returns_requested_component() -> None:
    fd, db_path = tempfile.mkstemp(prefix="builder_api_single_component_", suffix=".db")
    os.close(fd)
    app = _app_with_builder_db(db_path)
    admin_client = app.test_client()
    cook_client = app.test_client()
    admin_headers = _headers(role="admin", tenant_id=1, user_id=11)
    cook_headers = _headers(role="cook", tenant_id=1, user_id=42)
    _seed_session(admin_client, role="admin", tenant_id=1, user_id=11)
    _seed_session(cook_client, role="cook", tenant_id=1, user_id=42)

    created = admin_client.post(
        "/api/builder/components",
        json={"component_name": "Scoped soup"},
        headers=admin_headers,
    )
    component_id = str(((created.get_json() or {}).get("component") or {}).get("component_id") or "")
    assert component_id

    readable = cook_client.get(f"/api/builder/components/{component_id}", headers=cook_headers)
    assert readable.status_code == 200
    body = readable.get_json() or {}
    assert body.get("ok") is True
    component = body.get("component") or {}
    assert component.get("component_id") == component_id
    assert component.get("component_name") == "Scoped soup"


def test_single_composition_read_endpoint_returns_requested_composition() -> None:
    fd, db_path = tempfile.mkstemp(prefix="builder_api_single_composition_", suffix=".db")
    os.close(fd)
    app = _app_with_builder_db(db_path)
    admin_client = app.test_client()
    cook_client = app.test_client()
    admin_headers = _headers(role="admin", tenant_id=1, user_id=11)
    cook_headers = _headers(role="cook", tenant_id=1, user_id=42)
    _seed_session(admin_client, role="admin", tenant_id=1, user_id=11)
    _seed_session(cook_client, role="cook", tenant_id=1, user_id=42)

    created = admin_client.post(
        "/api/builder/compositions",
        json={"composition_id": "scoped_plate", "composition_name": "Scoped plate"},
        headers=admin_headers,
    )
    composition_id = str(((created.get_json() or {}).get("composition") or {}).get("composition_id") or "")
    assert composition_id == "scoped_plate"

    readable = cook_client.get(f"/api/builder/compositions/{composition_id}", headers=cook_headers)
    assert readable.status_code == 200
    body = readable.get_json() or {}
    assert body.get("ok") is True
    composition = body.get("composition") or {}
    assert composition.get("composition_id") == composition_id
    assert composition.get("composition_name") == "Scoped plate"
    assert composition.get("use_custom_menu_name") is False
    assert composition.get("menu_name") is None
    assert composition.get("effective_menu_name") == "Scoped plate"


def test_single_component_read_endpoint_respects_private_scope_boundary() -> None:
    fd, db_path = tempfile.mkstemp(prefix="builder_api_single_component_scope_", suffix=".db")
    os.close(fd)
    app = _app_with_builder_db(db_path)
    admin_client = app.test_client()
    cook_a_client = app.test_client()
    cook_b_client = app.test_client()
    admin_headers = _headers(role="admin", tenant_id=1, user_id=11)
    cook_a_headers = _headers(role="cook", tenant_id=1, user_id=42)
    cook_b_headers = _headers(role="cook", tenant_id=1, user_id=43)
    _seed_session(admin_client, role="admin", tenant_id=1, user_id=11)
    _seed_session(cook_a_client, role="cook", tenant_id=1, user_id=42)
    _seed_session(cook_b_client, role="cook", tenant_id=1, user_id=43)

    created = admin_client.post(
        "/api/builder/components",
        json={"component_name": "Boundary soup"},
        headers=admin_headers,
    )
    component_id = str(((created.get_json() or {}).get("component") or {}).get("component_id") or "")
    assert component_id

    forked = cook_a_client.post(f"/api/builder/components/{component_id}/fork", headers=cook_a_headers)
    forked_component_id = str(((forked.get_json() or {}).get("component") or {}).get("component_id") or "")
    assert forked_component_id

    owner_read = cook_a_client.get(f"/api/builder/components/{forked_component_id}", headers=cook_a_headers)
    other_read = cook_b_client.get(f"/api/builder/components/{forked_component_id}", headers=cook_b_headers)

    assert owner_read.status_code == 200
    assert other_read.status_code == 400
    assert (other_read.get_json() or {}).get("error") == "bad_request"


def test_free_create_composition_seeds_persisted_component_links_and_reuses_existing() -> None:
    client = _client()
    existing = client.post(
        "/api/builder/components",
        json={"component_name": "Pannbiff"},
        headers=HEADERS,
    )
    existing_id = ((existing.get_json() or {}).get("component") or {}).get("component_id")

    created = client.post(
        "/api/builder/compositions",
        json={"composition_name": "Pannbiff med potatis"},
        headers=HEADERS,
    )

    assert created.status_code == 201
    body = created.get_json() or {}
    composition = body.get("composition") or {}
    links = composition.get("components") or []
    assert len(links) == 2
    assert links[0].get("component_id") == existing_id

    library = client.get("/api/builder/library", headers=HEADERS)
    library_components = (library.get_json() or {}).get("components") or []
    assert len([item for item in library_components if item.get("component_id") == existing_id]) == 1


def test_create_standalone_component_endpoint() -> None:
    client = _client()

    rv = client.post(
        "/api/builder/components",
        json={"component_name": "Mashed Potatoes"},
        headers=HEADERS,
    )

    assert rv.status_code == 201
    body = rv.get_json() or {}
    assert body.get("ok") is True
    component = body.get("component") or {}
    assert component.get("component_id") == "mashed_potatoes"
    assert component.get("component_name") == "Mashed Potatoes"
    assert component.get("category") == "ovrigt"

    compositions = client.get("/api/builder/compositions", headers=HEADERS).get_json() or {}
    assert compositions.get("count") == 0


def test_create_component_duplicate_normalized_name_reuses_existing_component_with_friendly_message() -> None:
    client = _client()

    first = client.post(
        "/api/builder/components",
        json={"component_name": "Potatismos"},
        headers=HEADERS,
    )
    assert first.status_code == 201
    first_component = (first.get_json() or {}).get("component") or {}
    first_id = str(first_component.get("component_id") or "")
    assert first_id

    duplicate = client.post(
        "/api/builder/components",
        json={"component_name": " .. potatismos !! "},
        headers=HEADERS,
    )
    assert duplicate.status_code == 200
    body = duplicate.get_json() or {}
    assert body.get("ok") is True
    assert body.get("duplicate") is True
    assert body.get("message") == "Komponenten finns redan: Potatismos"
    component = body.get("component") or {}
    assert str(component.get("component_id") or "") == first_id

    listed = client.get("/api/builder/components", headers=HEADERS)
    listed_components = (listed.get_json() or {}).get("components") or []
    names = [str(item.get("component_name") or "") for item in listed_components]
    assert names.count("Potatismos") == 1


def test_component_category_can_be_set_and_cleared() -> None:
    client = _client()

    created = client.post(
        "/api/builder/components",
        json={"component_name": "Tomatsoppa", "category": "main"},
        headers=HEADERS,
    )
    assert created.status_code == 201
    created_body = created.get_json() or {}
    component = created_body.get("component") or {}
    component_id = str(component.get("component_id") or "")
    assert component.get("category") == "main"

    patched = client.patch(
        f"/api/builder/components/{component_id}",
        json={"category": "sauce"},
        headers=HEADERS,
    )
    assert patched.status_code == 200
    patched_body = patched.get_json() or {}
    assert ((patched_body.get("component") or {}).get("category")) == "sauce"

    cleared = client.patch(
        f"/api/builder/components/{component_id}",
        json={"category": None},
        headers=HEADERS,
    )
    assert cleared.status_code == 200
    clear_body = cleared.get_json() or {}
    assert ((clear_body.get("component") or {}).get("category")) == "ovrigt"


def test_component_category_normalization_report_lists_raw_and_suggested_values() -> None:
    client = _client()
    client.post(
        "/api/builder/components",
        json={"component_name": "Tomatsoppa", "category": "main"},
        headers=HEADERS,
    )
    client.post(
        "/api/builder/components",
        json={"component_name": "Okänd kategori komponent"},
        headers=HEADERS,
    )

    rv = client.get("/api/builder/components/category-normalization-report", headers=HEADERS)

    assert rv.status_code == 200
    body = rv.get_json() or {}
    assert body.get("ok") is True
    entries = body.get("entries") or []
    assert any(
        (entry.get("current_category") == "main" and entry.get("suggested_category") == "main")
        for entry in entries
    )
    assert any(
        (entry.get("current_category") is None and entry.get("suggested_category") == "ovrigt")
        for entry in entries
    )


def test_component_name_can_be_renamed_via_patch_name_field() -> None:
    client = _client()

    created = client.post(
        "/api/builder/components",
        json={"component_name": "Köttfärslimpa serveras"},
        headers=HEADERS,
    )
    assert created.status_code == 201
    component_id = str(((created.get_json() or {}).get("component") or {}).get("component_id") or "")
    assert component_id

    patched = client.patch(
        f"/api/builder/components/{component_id}",
        json={"name": "Köttfärslimpa"},
        headers=HEADERS,
    )
    assert patched.status_code == 200
    patched_body = patched.get_json() or {}
    assert ((patched_body.get("component") or {}).get("component_name")) == "Köttfärslimpa"

    listed = client.get("/api/builder/components", headers=HEADERS)
    assert listed.status_code == 200
    listed_components = (listed.get_json() or {}).get("components") or []
    match = next((item for item in listed_components if str(item.get("component_id") or "") == component_id), None)
    assert match is not None
    assert match.get("component_name") == "Köttfärslimpa"


def test_create_standalone_component_endpoint_rejects_empty_name() -> None:
    client = _client()

    rv = client.post(
        "/api/builder/components",
        json={"component_name": "   "},
        headers=HEADERS,
    )

    assert rv.status_code == 400
    body = rv.get_json() or {}
    assert body.get("error") == "bad_request"


def test_delete_component_endpoint_removes_unreferenced_component() -> None:
    client = _client()
    created = client.post(
        "/api/builder/components",
        json={"component_name": "Temp component"},
        headers=HEADERS,
    )
    component_id = ((created.get_json() or {}).get("component") or {}).get("component_id")

    rv = client.delete(f"/api/builder/components/{component_id}", headers=HEADERS)

    assert rv.status_code == 200
    body = rv.get_json() or {}
    assert body.get("ok") is True
    listed = client.get("/api/builder/components", headers=HEADERS).get_json() or {}
    ids = [item.get("component_id") for item in (listed.get("components") or [])]
    assert component_id not in ids


def test_delete_component_endpoint_blocks_when_component_is_used_by_composition() -> None:
    client = _client()
    created = client.post(
        "/api/builder/components",
        json={"component_name": "Shared component"},
        headers=HEADERS,
    )
    component_id = ((created.get_json() or {}).get("component") or {}).get("component_id")
    client.post(
        "/api/builder/compositions",
        json={"composition_id": "plate_used", "composition_name": "Used dish"},
        headers=HEADERS,
    )
    client.post(
        "/api/builder/compositions/plate_used/components/attach",
        json={"component_id": component_id},
        headers=HEADERS,
    )

    rv = client.delete(f"/api/builder/components/{component_id}", headers=HEADERS)

    assert rv.status_code == 409
    body = rv.get_json() or {}
    assert body.get("error") == "conflict"
    refs = body.get("references") or {}
    assert "plate_used" in (refs.get("composition_ids") or [])
    assert "Used dish" in (refs.get("composition_names") or [])
    assert int(refs.get("composition_count") or 0) == len(set(refs.get("composition_names") or []))


def test_component_phrase_fragment_report_lists_links_without_mutating_data() -> None:
    client = _client()
    created = client.post(
        "/api/builder/components",
        json={"component_name": "Med smak av dragon"},
        headers=HEADERS,
    )
    component_id = str(((created.get_json() or {}).get("component") or {}).get("component_id") or "")
    assert component_id

    client.post(
        "/api/builder/compositions",
        json={"composition_id": "dish_a", "composition_name": "Kycklinggryta"},
        headers=HEADERS,
    )
    client.post(
        "/api/builder/compositions/dish_a/components/attach",
        json={"component_id": component_id},
        headers=HEADERS,
    )

    rv = client.get("/api/builder/components/phrase-fragment-report", headers=HEADERS)

    assert rv.status_code == 200
    body = rv.get_json() or {}
    assert body.get("ok") is True
    entries = body.get("entries") or []
    target = next((item for item in entries if str(item.get("component_id") or "") == component_id), None)
    assert target is not None
    assert int(target.get("linked_composition_count") or 0) == 1
    assert "Kycklinggryta" in (target.get("sample_linked_composition_names") or [])
    paths = target.get("link_source_paths") or []
    assert "POST /api/builder/import/publish-drafts" in paths


def test_delete_composition_endpoint_removes_unreferenced_dish() -> None:
    client = _client()
    created = client.post(
        "/api/builder/compositions",
        json={"composition_id": "plate_delete", "composition_name": "Delete me"},
        headers=HEADERS,
    )
    composition_id = ((created.get_json() or {}).get("composition") or {}).get("composition_id")

    rv = client.delete(f"/api/builder/compositions/{composition_id}", headers=HEADERS)

    assert rv.status_code == 200
    body = rv.get_json() or {}
    assert body.get("ok") is True
    listed = client.get("/api/builder/compositions", headers=HEADERS).get_json() or {}
    ids = [item.get("composition_id") for item in (listed.get("compositions") or [])]
    assert composition_id not in ids


def test_delete_composition_endpoint_blocks_when_dish_is_used_by_menu() -> None:
    client = _client()
    client.post(
        "/api/builder/compositions",
        json={"composition_id": "dish_in_menu", "composition_name": "Dish in menu"},
        headers=HEADERS,
    )
    client.post(
        "/api/builder/menus",
        json={"menu_id": "menu_1", "site_id": "site_1", "week_key": "2026-W16"},
        headers=HEADERS,
    )
    client.post(
        "/api/builder/menus/menu_1/import",
        json={"rows": [{"day": "monday", "meal_slot": "lunch", "raw_text": "Dish in menu"}]},
        headers=HEADERS,
    )

    rv = client.delete("/api/builder/compositions/dish_in_menu", headers=HEADERS)

    assert rv.status_code == 409
    body = rv.get_json() or {}
    assert body.get("error") == "conflict"
    refs = body.get("references") or {}
    assert "menu_1" in (refs.get("menu_ids") or [])


def test_list_reusable_components_endpoint_supports_listing_and_search() -> None:
    client = _client()
    client.post("/api/builder/components", json={"component_name": "Mashed Potatoes"}, headers=HEADERS)
    client.post("/api/builder/components", json={"component_name": "Fish Sauce"}, headers=HEADERS)

    list_rv = client.get("/api/builder/components", headers=HEADERS)
    search_rv = client.get("/api/builder/components?q=fish", headers=HEADERS)

    assert list_rv.status_code == 200
    assert search_rv.status_code == 200
    list_body = list_rv.get_json() or {}
    search_body = search_rv.get_json() or {}
    assert list_body.get("ok") is True
    assert len(list_body.get("components") or []) == 2
    first = (list_body.get("components") or [])[0]
    assert "detail_summary" in first
    assert isinstance(first.get("tags"), list)
    summary = first.get("detail_summary") or {}
    assert set(summary.keys()) == {"has_method_data", "has_calculation_data", "has_allergen_data"}
    assert [item.get("component_name") for item in (search_body.get("components") or [])] == ["Fish Sauce"]


def test_component_list_includes_detail_summary_flags_from_component_details() -> None:
    client = _sqlite_client()
    created = client.post(
        "/api/builder/components",
        json={"component_name": "Laxsoppa"},
        headers=HEADERS,
    )
    component_id = str(((created.get_json() or {}).get("component") or {}).get("component_id") or "")
    assert component_id

    saved = client.patch(
        f"/api/builder/components/{component_id}/details",
        json={
            "method_text": "Sjud 10 min.",
            "calculation_notes": "Batch x2",
            "allergens": ["fish"],
        },
        headers=HEADERS,
    )
    assert saved.status_code == 200

    listed = client.get("/api/builder/components", headers=HEADERS)
    assert listed.status_code == 200
    body = listed.get_json() or {}
    components = body.get("components") or []
    target = next((item for item in components if str(item.get("component_id") or "") == component_id), None)
    assert target is not None
    summary = target.get("detail_summary") or {}
    assert summary.get("has_method_data") is True
    assert summary.get("has_calculation_data") is True
    assert summary.get("has_allergen_data") is True


def test_component_list_marks_calculation_flag_when_only_rows_are_present() -> None:
    client = _sqlite_client()
    created = client.post(
        "/api/builder/components",
        json={"component_name": "Kall sas"},
        headers=HEADERS,
    )
    component_id = str(((created.get_json() or {}).get("component") or {}).get("component_id") or "")
    assert component_id

    saved = client.patch(
        f"/api/builder/components/{component_id}/details",
        json={
            "calculation_rows": [
                {
                    "ingredient_name": "Majonnäs",
                    "amount_value": "0.25",
                    "amount_unit": "kg",
                    "price_value": "80",
                    "price_unit": "kr/kg",
                    "calculated_cost": "20",
                }
            ]
        },
        headers=HEADERS,
    )
    assert saved.status_code == 200

    listed = client.get("/api/builder/components", headers=HEADERS)
    assert listed.status_code == 200
    body = listed.get_json() or {}
    components = body.get("components") or []
    target = next((item for item in components if str(item.get("component_id") or "") == component_id), None)
    assert target is not None
    summary = target.get("detail_summary") or {}
    assert summary.get("has_method_data") is False
    assert summary.get("has_calculation_data") is True
    assert summary.get("has_allergen_data") is False


def test_component_alias_endpoints_create_and_list_aliases() -> None:
    client = _client()
    created = client.post(
        "/api/builder/components",
        json={"component_name": "Kokt potatis"},
        headers=HEADERS,
    )
    component_id = ((created.get_json() or {}).get("component") or {}).get("component_id")

    alias_rv = client.post(
        f"/api/builder/components/{component_id}/aliases",
        json={"alias_text": "potatis kokt", "source": "import"},
        headers=HEADERS,
    )
    listed = client.get(
        f"/api/builder/components/{component_id}/aliases",
        headers=HEADERS,
    )

    assert alias_rv.status_code == 201
    alias_body = alias_rv.get_json() or {}
    assert alias_body.get("ok") is True
    assert ((alias_body.get("alias") or {}).get("alias_norm")) == "potatis kokt"
    assert listed.status_code == 200
    listed_body = listed.get_json() or {}
    assert listed_body.get("count") == 1
    assert (listed_body.get("aliases") or [])[0].get("alias_text") == "potatis kokt"


def test_component_details_get_patch_persists_in_backend() -> None:
    client = _sqlite_client()

    created = client.post(
        "/api/builder/components",
        json={"component_name": "Ugnsbakad lax"},
        headers=HEADERS,
    )
    assert created.status_code == 201
    component_id = str(((created.get_json() or {}).get("component") or {}).get("component_id") or "")
    assert component_id

    initial = client.get(f"/api/builder/components/{component_id}/details", headers=HEADERS)
    assert initial.status_code == 200
    initial_body = initial.get_json() or {}
    assert initial_body.get("ok") is True
    initial_details = initial_body.get("details") or {}
    assert initial_details.get("recipe_ingredient_rows") == []
    assert initial_details.get("recipe_ingredients_text") == ""
    assert initial_details.get("method_text") == ""
    assert initial_details.get("calculation_yield") == ""
    assert initial_details.get("calculation_rows") == []
    assert initial_details.get("allergens") == []
    assert initial_details.get("tags") == []
    assert initial_details.get("long_description") == ""

    patch_payload = {
        "recipe_ingredient_rows": [
            {"ingredient_name": "Laxfile", "amount_value": "80", "amount_unit": "g"},
            {"ingredient_name": "Mjolk", "amount_value": "60", "amount_unit": "g"},
            {"ingredient_name": "Smor", "amount_value": "30", "amount_unit": "g"},
        ],
        "recipe_ingredients_text": "Laxfile\nCitron\nDill",
        "method_text": "Baka i 180C i 12 minuter.",
        "method_notes": "Anvand bleck med bakplattspapper.",
        "calculation_yield": "12",
        "calculation_cost": "320.50 SEK",
        "calculation_notes": "Pris baserat pa vecka 22.",
        "calculation_rows": [
            {
                "ingredient_name": "Lax",
                "amount_value": "1.5",
                "amount_unit": "kg",
                "price_value": "149",
                "price_unit": "kr/kg",
                "calculated_cost": "223.50",
            }
        ],
        "allergens": ["fish", "milk"],
        "allergen_notes": "Kan innehalla spar av gluten.",
        "tags": ["fish", "soup"],
        "long_description": "Serveras varm med dill.",
    }
    saved = client.patch(
        f"/api/builder/components/{component_id}/details",
        json=patch_payload,
        headers=HEADERS,
    )
    assert saved.status_code == 200
    saved_body = saved.get_json() or {}
    assert saved_body.get("ok") is True
    saved_details = saved_body.get("details") or {}
    assert saved_details.get("recipe_ingredient_rows") == patch_payload["recipe_ingredient_rows"]
    assert saved_details.get("recipe_ingredients_text") == patch_payload["recipe_ingredients_text"]
    assert saved_details.get("method_text") == patch_payload["method_text"]
    assert saved_details.get("calculation_yield") == patch_payload["calculation_yield"]
    assert saved_details.get("calculation_cost") == patch_payload["calculation_cost"]
    assert saved_details.get("calculation_notes") == patch_payload["calculation_notes"]
    assert saved_details.get("calculation_rows") == patch_payload["calculation_rows"]
    assert saved_details.get("allergens") == ["fish", "milk"]
    assert saved_details.get("allergen_notes") == patch_payload["allergen_notes"]
    assert saved_details.get("tags") == patch_payload["tags"]
    assert saved_details.get("long_description") == patch_payload["long_description"]

    refreshed = client.get(f"/api/builder/components/{component_id}/details", headers=HEADERS)
    assert refreshed.status_code == 200
    refreshed_body = refreshed.get_json() or {}
    assert refreshed_body.get("ok") is True
    refreshed_details = refreshed_body.get("details") or {}
    assert refreshed_details.get("recipe_ingredient_rows") == patch_payload["recipe_ingredient_rows"]
    assert refreshed_details.get("recipe_ingredients_text") == patch_payload["recipe_ingredients_text"]
    assert refreshed_details.get("method_text") == patch_payload["method_text"]
    assert refreshed_details.get("method_notes") == patch_payload["method_notes"]
    assert refreshed_details.get("calculation_yield") == patch_payload["calculation_yield"]
    assert refreshed_details.get("calculation_cost") == patch_payload["calculation_cost"]
    assert refreshed_details.get("calculation_notes") == patch_payload["calculation_notes"]
    assert refreshed_details.get("calculation_rows") == patch_payload["calculation_rows"]
    assert refreshed_details.get("allergens") == ["fish", "milk"]
    assert refreshed_details.get("allergen_notes") == patch_payload["allergen_notes"]
    assert refreshed_details.get("tags") == patch_payload["tags"]
    assert refreshed_details.get("long_description") == patch_payload["long_description"]

    listed = client.get("/api/builder/components", headers=HEADERS)
    assert listed.status_code == 200
    components = (listed.get_json() or {}).get("components") or []
    target = next((item for item in components if str(item.get("component_id") or "") == component_id), None)
    assert target is not None
    assert target.get("tags") == patch_payload["tags"]


def test_component_list_marks_method_flag_when_only_recipe_rows_are_present() -> None:
    client = _sqlite_client()
    created = client.post(
        "/api/builder/components",
        json={"component_name": "Potatismos"},
        headers=HEADERS,
    )
    component_id = str(((created.get_json() or {}).get("component") or {}).get("component_id") or "")
    assert component_id

    saved = client.patch(
        f"/api/builder/components/{component_id}/details",
        json={
            "recipe_ingredient_rows": [
                {
                    "ingredient_name": "Potatis",
                    "amount_value": "80",
                    "amount_unit": "g",
                }
            ]
        },
        headers=HEADERS,
    )
    assert saved.status_code == 200

    listed = client.get("/api/builder/components", headers=HEADERS)
    assert listed.status_code == 200
    body = listed.get_json() or {}
    components = body.get("components") or []
    target = next((item for item in components if str(item.get("component_id") or "") == component_id), None)
    assert target is not None
    summary = target.get("detail_summary") or {}
    assert summary.get("has_method_data") is True
    assert summary.get("has_calculation_data") is False


def test_component_list_does_not_mark_method_flag_when_method_fields_are_empty() -> None:
    client = _sqlite_client()
    created = client.post(
        "/api/builder/components",
        json={"component_name": "Tomat"},
        headers=HEADERS,
    )
    component_id = str(((created.get_json() or {}).get("component") or {}).get("component_id") or "")
    assert component_id

    saved = client.patch(
        f"/api/builder/components/{component_id}/details",
        json={
            "method_text": "   ",
            "method_notes": "\n\t",
            "recipe_ingredient_rows": [],
        },
        headers=HEADERS,
    )
    assert saved.status_code == 200

    listed = client.get("/api/builder/components", headers=HEADERS)
    assert listed.status_code == 200
    components = (listed.get_json() or {}).get("components") or []
    target = next((item for item in components if str(item.get("component_id") or "") == component_id), None)
    assert target is not None
    summary = target.get("detail_summary") or {}
    assert summary.get("has_method_data") is False


def test_component_details_normalizes_blank_recipe_rows_on_save() -> None:
    client = _sqlite_client()
    created = client.post(
        "/api/builder/components",
        json={"component_name": "Morot"},
        headers=HEADERS,
    )
    component_id = str(((created.get_json() or {}).get("component") or {}).get("component_id") or "")
    assert component_id

    saved = client.patch(
        f"/api/builder/components/{component_id}/details",
        json={
            "recipe_ingredient_rows": [
                {"ingredient_name": "", "amount_value": "10", "amount_unit": "g"},
                {"ingredient_name": "   ", "amount_value": "", "amount_unit": ""},
                {"ingredient_name": "  Potatis  ", "amount_value": " 80 ", "amount_unit": " g "},
            ]
        },
        headers=HEADERS,
    )
    assert saved.status_code == 200
    details = (saved.get_json() or {}).get("details") or {}
    assert details.get("recipe_ingredient_rows") == [
        {"ingredient_name": "Potatis", "amount_value": "80", "amount_unit": "g"}
    ]

    listed = client.get("/api/builder/components", headers=HEADERS)
    assert listed.status_code == 200
    components = (listed.get_json() or {}).get("components") or []
    target = next((item for item in components if str(item.get("component_id") or "") == component_id), None)
    assert target is not None
    summary = target.get("detail_summary") or {}
    assert summary.get("has_method_data") is True


def test_component_list_does_not_mark_method_flag_for_rows_with_empty_ingredient_name() -> None:
    client = _sqlite_client()
    created = client.post(
        "/api/builder/components",
        json={"component_name": "Paprika"},
        headers=HEADERS,
    )
    component_id = str(((created.get_json() or {}).get("component") or {}).get("component_id") or "")
    assert component_id

    saved = client.patch(
        f"/api/builder/components/{component_id}/details",
        json={
            "recipe_ingredient_rows": [
                {"ingredient_name": "", "amount_value": "1", "amount_unit": "kg"}
            ]
        },
        headers=HEADERS,
    )
    assert saved.status_code == 200

    details = (saved.get_json() or {}).get("details") or {}
    assert details.get("recipe_ingredient_rows") == []

    listed = client.get("/api/builder/components", headers=HEADERS)
    assert listed.status_code == 200
    components = (listed.get_json() or {}).get("components") or []
    target = next((item for item in components if str(item.get("component_id") or "") == component_id), None)
    assert target is not None
    summary = target.get("detail_summary") or {}
    assert summary.get("has_method_data") is False


def test_component_list_marks_method_flag_when_method_text_exists() -> None:
    client = _sqlite_client()
    created = client.post(
        "/api/builder/components",
        json={"component_name": "Broccoli"},
        headers=HEADERS,
    )
    component_id = str(((created.get_json() or {}).get("component") or {}).get("component_id") or "")
    assert component_id

    saved = client.patch(
        f"/api/builder/components/{component_id}/details",
        json={"method_text": "Koka i 5 minuter."},
        headers=HEADERS,
    )
    assert saved.status_code == 200

    listed = client.get("/api/builder/components", headers=HEADERS)
    assert listed.status_code == 200
    components = (listed.get_json() or {}).get("components") or []
    target = next((item for item in components if str(item.get("component_id") or "") == component_id), None)
    assert target is not None
    summary = target.get("detail_summary") or {}
    assert summary.get("has_method_data") is True


def test_imported_component_without_recipe_data_does_not_get_method_flag() -> None:
    client = _sqlite_client()

    published = client.post(
        "/api/builder/import/publish-drafts",
        json={
            "items": [
                {
                    "selected": True,
                    "item_type": "component",
                    "raw_text": "Kokt potatis",
                    "name": "Kokt potatis",
                    "components": [],
                }
            ]
        },
        headers=HEADERS,
    )
    assert published.status_code == 200

    listed = client.get("/api/builder/components", headers=HEADERS)
    assert listed.status_code == 200
    components = (listed.get_json() or {}).get("components") or []
    target = next(
        (item for item in components if str(item.get("component_name") or "").strip().lower() == "kokt potatis"),
        None,
    )
    assert target is not None
    summary = target.get("detail_summary") or {}
    assert summary.get("has_method_data") is False


def test_component_method_summary_report_lists_field_source_for_method_flag() -> None:
    client = _sqlite_client()
    created = client.post(
        "/api/builder/components",
        json={"component_name": "Ris"},
        headers=HEADERS,
    )
    component_id = str(((created.get_json() or {}).get("component") or {}).get("component_id") or "")
    assert component_id

    saved = client.patch(
        f"/api/builder/components/{component_id}/details",
        json={
            "method_text": "Sjuda i vatten",
            "method_notes": "Ror om",
            "recipe_ingredient_rows": [{"ingredient_name": "Ris", "amount_value": "100", "amount_unit": "g"}],
            "recipe_ingredients_text": "Ris | 100 | g",
        },
        headers=HEADERS,
    )
    assert saved.status_code == 200

    report = client.get("/api/builder/components/method-summary-report", headers=HEADERS)
    assert report.status_code == 200
    body = report.get_json() or {}
    entries = body.get("entries") or []
    target = next((item for item in entries if str(item.get("component_id") or "") == component_id), None)
    assert target is not None
    assert target.get("has_method_data") is True
    sources = target.get("method_sources") or []
    assert "method_text" in sources
    assert "method_notes" in sources
    assert "recipe_ingredient_rows" in sources
    assert "recipe_ingredients_text legacy" not in sources


def test_component_method_summary_report_marks_legacy_recipe_text_when_intentionally_saved() -> None:
    client = _sqlite_client()
    created = client.post(
        "/api/builder/components",
        json={"component_name": "Legacy test"},
        headers=HEADERS,
    )
    component_id = str(((created.get_json() or {}).get("component") or {}).get("component_id") or "")
    assert component_id

    saved = client.patch(
        f"/api/builder/components/{component_id}/details",
        json={
            "recipe_ingredient_rows": [],
            "recipe_ingredients_text": "Morot | 50 | g",
            "method_text": "",
            "method_notes": "",
        },
        headers=HEADERS,
    )
    assert saved.status_code == 200

    listed = client.get("/api/builder/components", headers=HEADERS)
    assert listed.status_code == 200
    components = (listed.get_json() or {}).get("components") or []
    target = next((item for item in components if str(item.get("component_id") or "") == component_id), None)
    assert target is not None
    summary = target.get("detail_summary") or {}
    assert summary.get("has_method_data") is True

    report = client.get("/api/builder/components/method-summary-report", headers=HEADERS)
    assert report.status_code == 200
    entries = (report.get_json() or {}).get("entries") or []
    target_report = next((item for item in entries if str(item.get("component_id") or "") == component_id), None)
    assert target_report is not None
    assert "recipe_ingredients_text legacy" in (target_report.get("method_sources") or [])


def test_library_endpoint_returns_separate_sorted_components_and_compositions() -> None:
    client = _client()
    client.post("/api/builder/components", json={"component_name": "zeta"}, headers=HEADERS)
    client.post("/api/builder/components", json={"component_name": "Alpha"}, headers=HEADERS)
    client.post("/api/builder/compositions", json={"composition_name": "Zulu dish"}, headers=HEADERS)
    client.post("/api/builder/compositions", json={"composition_name": "alpha dish"}, headers=HEADERS)

    rv = client.get("/api/builder/library", headers=HEADERS)

    assert rv.status_code == 200
    body = rv.get_json() or {}
    assert body.get("ok") is True
    components = body.get("components") or []
    compositions = body.get("compositions") or []
    component_names = [item.get("component_name") for item in components]
    assert "Alpha" in component_names
    assert "zeta" in component_names
    assert [item.get("composition_name") for item in compositions] == ["alpha dish", "Zulu dish"]


def test_library_endpoint_no_menu_linkage_required_and_stable_composition_id_reused() -> None:
    client = _client()
    created = client.post(
        "/api/builder/compositions",
        json={"composition_name": "Fish Soup"},
        headers=HEADERS,
    )
    created_id = ((created.get_json() or {}).get("composition") or {}).get("composition_id")
    client.post(
        "/api/builder/components",
        json={"component_name": "Mashed Potatoes"},
        headers=HEADERS,
    )

    rv = client.get("/api/builder/library", headers=HEADERS)

    assert rv.status_code == 200
    body = rv.get_json() or {}
    components = body.get("components") or []
    compositions = body.get("compositions") or []
    assert any(item.get("component_name") == "Mashed Potatoes" for item in components)
    assert any(item.get("composition_id") == created_id for item in compositions)


def test_library_reads_do_not_create_new_compositions() -> None:
    client = _client()
    created = client.post(
        "/api/builder/compositions",
        json={"composition_name": "Open me"},
        headers=HEADERS,
    )
    created_id = ((created.get_json() or {}).get("composition") or {}).get("composition_id")

    before = client.get("/api/builder/compositions", headers=HEADERS)
    before_ids = {
        item.get("composition_id")
        for item in ((before.get_json() or {}).get("compositions") or [])
    }

    first_library = client.get("/api/builder/library", headers=HEADERS)
    second_library = client.get("/api/builder/library", headers=HEADERS)

    assert first_library.status_code == 200
    assert second_library.status_code == 200
    first_ids = {
        item.get("composition_id")
        for item in ((first_library.get_json() or {}).get("compositions") or [])
    }
    second_ids = {
        item.get("composition_id")
        for item in ((second_library.get_json() or {}).get("compositions") or [])
    }
    assert created_id in first_ids
    assert first_ids == second_ids == before_ids


def test_builder_library_import_accepts_lines_without_day_or_meal() -> None:
    client = _client()

    rv = client.post(
        "/api/builder/import",
        json={"lines": ["Kottbullar med potatismos", "Fiskgratang"]},
        headers=HEADERS,
    )

    assert rv.status_code == 200
    body = rv.get_json() or {}
    summary = body.get("summary") or {}
    assert body.get("ok") is True
    assert summary.get("imported_count") == 2
    assert summary.get("created_count") == 2
    assert summary.get("reused_count") == 0


def test_builder_library_import_accepts_multiline_text_and_creates_components() -> None:
    client = _client()

    rv = client.post(
        "/api/builder/import",
        json={"text": "Kottbullar med graddsas och rodbetor"},
        headers=HEADERS,
    )

    assert rv.status_code == 200
    body = rv.get_json() or {}
    summary = body.get("summary") or {}
    row = (summary.get("row_results") or [{}])[0]
    created_id = row.get("composition_id")
    compositions = client.get("/api/builder/compositions", headers=HEADERS).get_json() or {}
    created = next(
        (
            item
            for item in (compositions.get("compositions") or [])
            if item.get("composition_id") == created_id
        ),
        None,
    )

    assert summary.get("created_count") == 1
    assert created is not None
    component_names = [item.get("component_name") for item in (created.get("components") or [])]
    assert component_names == ["Kottbullar", "Graddsas", "Rodbetor"]


def test_builder_library_import_persists_and_reuses_library_components() -> None:
    client = _client()

    first = client.post(
        "/api/builder/import",
        json={"lines": ["Kottbullar med potatismos"]},
        headers=HEADERS,
    )
    assert first.status_code == 200

    second = client.post(
        "/api/builder/import",
        json={"lines": ["Kottbullar med graddsas"]},
        headers=HEADERS,
    )
    assert second.status_code == 200

    library = client.get("/api/builder/library", headers=HEADERS)
    assert library.status_code == 200
    body = library.get_json() or {}
    components = body.get("components") or []
    kottbullar = [item for item in components if (item.get("component_name") or "").lower() == "kottbullar"]
    assert len(kottbullar) == 1


def test_builder_library_import_reuses_alias_without_creating_new_composition() -> None:
    client = _client()
    first = client.post(
        "/api/builder/import",
        json={"lines": ["No Match"]},
        headers=HEADERS,
    )
    first_id = (((first.get_json() or {}).get("summary") or {}).get("row_results") or [{}])[0].get(
        "composition_id"
    )

    second = client.post(
        "/api/builder/import",
        json={"lines": ["No Match"]},
        headers=HEADERS,
    )

    assert second.status_code == 200
    body = second.get_json() or {}
    summary = body.get("summary") or {}
    row = (summary.get("row_results") or [{}])[0]
    assert summary.get("created_count") == 0
    assert summary.get("reused_count") == 1
    assert row.get("composition_id") == first_id


def test_builder_library_import_reports_possible_component_matches_without_blocking() -> None:
    client = _client()
    client.post(
        "/api/builder/components",
        json={"component_name": "Kokt potatis"},
        headers=HEADERS,
    )

    rv = client.post(
        "/api/builder/import",
        json={"lines": ["Kokt potatisar"]},
        headers=HEADERS,
    )

    assert rv.status_code == 200
    body = rv.get_json() or {}
    summary = body.get("summary") or {}
    assert summary.get("created_count") == 1
    assert summary.get("imported_count") == 1
    review_items = summary.get("component_review_items") or []
    assert len(review_items) == 1
    review = review_items[0]
    assert review.get("status") == "possible_match"
    assert review.get("suggested_component_name") == "Kokt potatisar"
    possible = review.get("possible_matches") or []
    assert possible[0].get("component_name") == "Kokt potatis"
    assert isinstance(possible[0].get("score"), float)


def test_builder_library_import_rejects_empty_payload_lines() -> None:
    client = _client()

    rv = client.post(
        "/api/builder/import",
        json={"lines": ["   ", ""]},
        headers=HEADERS,
    )

    assert rv.status_code == 400
    body = rv.get_json() or {}
    assert body.get("error") == "bad_request"


def test_builder_file_import_preview_txt_endpoint() -> None:
    client = _client()

    rv = client.post(
        "/api/builder/import/file/preview",
        data={"file": (io.BytesIO(b"Kottbullar med potatismos\n\nFiskgratang\n"), "library.txt")},
        content_type="multipart/form-data",
        headers=HEADERS,
    )

    assert rv.status_code == 200
    body = rv.get_json() or {}
    preview = body.get("preview") or {}
    assert body.get("ok") is True
    assert preview.get("file_type") == "txt"
    assert preview.get("line_count") == 2
    assert preview.get("lines") == ["Kottbullar med potatismos", "Fiskgratang"]
    assert preview.get("preview_contract_version") == 2
    assert preview.get("importable_items") == [
        {"preview_index": 0, "line": "Kottbullar med potatismos"},
        {"preview_index": 1, "line": "Fiskgratang"},
    ]
    counts = preview.get("counts") or {}
    assert counts.get("importable") == 2
    assert counts.get("ignored") == 1


def test_builder_file_import_preview_classifies_noise_vs_importable() -> None:
    client = _client()

    rv = client.post(
        "/api/builder/import/file/preview",
        data={"file": (io.BytesIO(b"Alt 1\nWeek 12\nFiskgratang\n"), "library.txt")},
        content_type="multipart/form-data",
        headers=HEADERS,
    )

    assert rv.status_code == 200
    body = rv.get_json() or {}
    preview = body.get("preview") or {}
    assert preview.get("importable_lines") == ["Fiskgratang"]
    ignored = preview.get("ignored_lines") or []
    ignored_texts = {item.get("normalized_text") for item in ignored}
    assert "Alt 1" in ignored_texts
    assert "Week 12" in ignored_texts
    counts = preview.get("counts") or {}
    assert counts.get("total_classified") == 3
    assert counts.get("importable") == 1
    assert counts.get("ignored") == 2


def test_builder_file_import_preview_large_payload_keeps_contract_shape() -> None:
    client = _client()
    payload_lines = [f"Dish {index}" for index in range(1, 121)]
    payload = ("\n".join(payload_lines) + "\n").encode("utf-8")

    rv = client.post(
        "/api/builder/import/file/preview",
        data={"file": (io.BytesIO(payload), "library.txt")},
        content_type="multipart/form-data",
        headers=HEADERS,
    )

    assert rv.status_code == 200
    body = rv.get_json() or {}
    preview = body.get("preview") or {}
    assert preview.get("line_count") == 120
    importable_items = preview.get("importable_items") or []
    assert len(importable_items) == 120
    assert importable_items[0] == {"preview_index": 0, "line": "Dish 1"}
    assert importable_items[-1] == {"preview_index": 119, "line": "Dish 120"}
    counts = preview.get("counts") or {}
    assert counts.get("importable") == 120
    assert counts.get("ignored") == 0


def test_builder_file_import_preview_csv_endpoint_detects_text_column() -> None:
    client = _client()
    payload = b"dish_name,category\nKottbullar med potatismos,main\nFiskgratang,main\n"

    rv = client.post(
        "/api/builder/import/file/preview",
        data={"file": (io.BytesIO(payload), "library.csv")},
        content_type="multipart/form-data",
        headers=HEADERS,
    )

    assert rv.status_code == 200
    body = rv.get_json() or {}
    preview = body.get("preview") or {}
    assert preview.get("file_type") == "csv"
    assert preview.get("lines") == ["Kottbullar med potatismos", "Fiskgratang"]
    assert preview.get("csv_column") == "dish_name"
    assert preview.get("csv_column_index") == 0


def test_builder_file_import_preview_xlsx_endpoint_detects_text_column() -> None:
    client = _client()
    payload = _xlsx_bytes(
        [
            ["dish_name", "category"],
            ["Kottbullar med potatismos", "main"],
            ["Fiskgratang", "main"],
        ]
    )

    rv = client.post(
        "/api/builder/import/file/preview",
        data={"file": (io.BytesIO(payload), "library.xlsx")},
        content_type="multipart/form-data",
        headers=HEADERS,
    )

    assert rv.status_code == 200
    body = rv.get_json() or {}
    preview = body.get("preview") or {}
    assert preview.get("file_type") == "xlsx"
    assert preview.get("lines") == ["Kottbullar med potatismos", "Fiskgratang"]
    assert preview.get("csv_column") == "dish_name"
    assert preview.get("csv_column_index") == 0


def test_builder_file_import_preview_docx_endpoint_reads_paragraphs_and_tables() -> None:
    client = _client()
    payload = _docx_bytes(
        paragraphs=["Week 12", "Alt 1", "Fiskgratang"],
        table_rows=[
            ["Dish"],
            ["Kottbullar med potatismos"],
        ],
    )

    rv = client.post(
        "/api/builder/import/file/preview",
        data={"file": (io.BytesIO(payload), "library.docx")},
        content_type="multipart/form-data",
        headers=HEADERS,
    )

    assert rv.status_code == 200
    preview = (rv.get_json() or {}).get("preview") or {}
    assert preview.get("file_type") == "docx"
    assert "Fiskgratang" in (preview.get("importable_lines") or [])
    assert "Kottbullar med potatismos" in (preview.get("importable_lines") or [])


def test_builder_reset_endpoint_clears_all_builder_data() -> None:
    client = _client()
    client.post("/api/builder/components", json={"component_name": "Temp component"}, headers=HEADERS)
    client.post("/api/builder/compositions", json={"composition_name": "Temp dish"}, headers=HEADERS)
    client.post(
        "/api/builder/menus",
        json={"menu_id": "menu_reset", "site_id": "site_1", "week_key": "2026-W16", "title": "Reset menu"},
        headers=HEADERS,
    )
    client.post(
        "/api/builder/menus/menu_reset/import",
        json={"rows": [{"day": "monday", "meal_slot": "lunch", "raw_text": "Temp dish"}]},
        headers=HEADERS,
    )

    rv = client.post("/api/builder/reset", headers=HEADERS)

    assert rv.status_code == 200
    body = rv.get_json() or {}
    assert body.get("ok") is True
    library = client.get("/api/builder/library", headers=HEADERS).get_json() or {}
    menus = client.get("/api/builder/menus", headers=HEADERS).get_json() or {}
    assert (library.get("components") or []) == []
    assert (library.get("compositions") or []) == []
    assert (menus.get("menus") or []) == []


def test_import_review_preview_exposes_cleaned_draft_items() -> None:
    client = _client()

    rv = client.post(
        "/api/builder/import/preview-lines",
        json={"lines": ["Menyval1:köttbullar", "Lördag", "Köttbullar med potatis"]},
        headers=HEADERS,
    )

    assert rv.status_code == 200
    preview = (rv.get_json() or {}).get("preview") or {}
    drafts = preview.get("draft_items") or []
    assert len(drafts) == 3
    first = drafts[0]
    assert first.get("name") == "köttbullar"
    second = drafts[1]
    assert second.get("item_type") == "ignore"
    dish = drafts[2]
    assert dish.get("item_type") == "dish"
    components = dish.get("components") or []
    component_names = [item.get("name") for item in components]
    assert "Köttbullar" in component_names
    assert "potatis" in [str(name).lower() for name in component_names]


def test_import_review_preview_keeps_descriptor_line_as_dish_and_filters_phrase_fragment_component() -> None:
    client = _client()

    rv = client.post(
        "/api/builder/import/preview-lines",
        json={"lines": ["Kycklinggryta med smak av dragon"]},
        headers=HEADERS,
    )

    assert rv.status_code == 200
    preview = (rv.get_json() or {}).get("preview") or {}
    drafts = preview.get("draft_items") or []
    assert len(drafts) == 1
    draft = drafts[0] or {}
    assert draft.get("item_type") == "dish"
    components = draft.get("components") or []
    names = [str(item.get("name") or "") for item in components]
    assert "Kycklinggryta" in names
    assert all("smak av" not in name.lower() for name in names)


def test_import_sessions_create_list_and_open_detail() -> None:
    client = _client()

    created = client.post(
        "/api/builder/import/sessions",
        json={
            "source_name": "Inbox test",
            "import_type": "dish_list",
            "items": [
                {
                    "selected": True,
                    "item_type": "component",
                    "raw_text": "Potatis",
                    "name": "Potatis",
                    "components": [],
                },
                {
                    "selected": False,
                    "item_type": "ignore",
                    "raw_text": "Week 12",
                    "name": "Week 12",
                    "components": [],
                },
            ],
        },
        headers=HEADERS,
    )

    assert created.status_code == 201
    created_body = created.get_json() or {}
    session = created_body.get("session") or {}
    session_id = str(session.get("session_id") or "")
    assert session_id
    assert session.get("pending_review_count") == 1

    listed = client.get("/api/builder/import/sessions", headers=HEADERS)
    assert listed.status_code == 200
    listed_body = listed.get_json() or {}
    assert listed_body.get("count", 0) >= 1
    assert listed_body.get("pending_count", 0) >= 1

    detail = client.get(f"/api/builder/import/sessions/{session_id}", headers=HEADERS)
    assert detail.status_code == 200
    detail_body = detail.get_json() or {}
    grouped = detail_body.get("grouped") or {}
    assert len(grouped.get("components") or []) == 1


def test_import_sessions_item_update_and_publish_selected() -> None:
    client = _client()

    created = client.post(
        "/api/builder/import/sessions",
        json={
            "source_name": "Publish session",
            "import_type": "component_list",
            "items": [
                {
                    "selected": True,
                    "item_type": "component",
                    "raw_text": "Potatis",
                    "name": "Potatis",
                    "components": [],
                }
            ],
        },
        headers=HEADERS,
    )
    assert created.status_code == 201
    created_body = created.get_json() or {}
    session = created_body.get("session") or {}
    items = created_body.get("items") or []
    session_id = str(session.get("session_id") or "")
    item_id = str((items[0] or {}).get("item_id") or "")
    assert session_id and item_id

    patched = client.patch(
        f"/api/builder/import/sessions/{session_id}/items/{item_id}",
        json={"name": "Potatis mos", "selected": True, "components": ["Potatis", "Smor"]},
        headers=HEADERS,
    )
    assert patched.status_code == 200

    detail_after_patch = client.get(f"/api/builder/import/sessions/{session_id}", headers=HEADERS)
    assert detail_after_patch.status_code == 200
    detail_after_patch_items = (detail_after_patch.get_json() or {}).get("items") or []
    patched_item = next(item for item in detail_after_patch_items if str(item.get("item_id") or "") == item_id)
    assert patched_item.get("components") == ["Potatis", "Smor"]

    published = client.post(
        f"/api/builder/import/sessions/{session_id}/publish-selected",
        json={},
        headers=HEADERS,
    )
    assert published.status_code == 200
    summary = (published.get_json() or {}).get("summary") or {}
    assert summary.get("imported_count") == 1

    detail = client.get(f"/api/builder/import/sessions/{session_id}", headers=HEADERS)
    detail_items = (detail.get_json() or {}).get("items") or []
    assert any(str(item.get("item_status") or "") == "published" for item in detail_items)

    library = client.get("/api/builder/library", headers=HEADERS).get_json() or {}
    names = [item.get("component_name") for item in (library.get("components") or [])]
    assert "Potatis mos" in names


def test_builder_reset_clears_import_sessions() -> None:
    client = _client()

    created = client.post(
        "/api/builder/import/sessions",
        json={
            "source_name": "Reset me",
            "import_type": "dish_list",
            "items": [
                {
                    "selected": True,
                    "item_type": "component",
                    "raw_text": "Ris",
                    "name": "Ris",
                    "components": [],
                }
            ],
        },
        headers=HEADERS,
    )
    assert created.status_code == 201

    reset = client.post("/api/builder/reset", headers=HEADERS)
    assert reset.status_code == 200
    reset_body = reset.get_json() or {}
    counts = reset_body.get("cleared_counts") or {}
    assert "builder_import_sessions" in counts

    listed = client.get("/api/builder/import/sessions", headers=HEADERS)
    listed_body = listed.get_json() or {}
    assert listed_body.get("count") == 0
    assert listed_body.get("pending_count") == 0


def test_import_sessions_duplicate_payload_twice_succeeds() -> None:
    client = _sqlite_client()

    payload = {
        "import_type": "dish_list",
        "items": [
            {
                "item_id": "tmp_0",
                "selected": True,
                "item_type": "component",
                "raw_text": "Potatis",
                "name": "Potatis",
                "components": [],
            }
        ],
    }

    first = client.post("/api/builder/import/sessions", json=payload, headers=HEADERS)
    second = client.post("/api/builder/import/sessions", json=payload, headers=HEADERS)

    assert first.status_code == 201
    assert second.status_code == 201

    first_session = (first.get_json() or {}).get("session") or {}
    second_session = (second.get_json() or {}).get("session") or {}
    assert first_session.get("session_id")
    assert second_session.get("session_id")
    assert first_session.get("session_id") != second_session.get("session_id")


def test_import_sessions_noisy_and_malformed_rows_do_not_crash() -> None:
    client = _sqlite_client()

    rv = client.post(
        "/api/builder/import/sessions",
        json={
            "import_type": "dish_list",
            "items": [
                {"selected": True, "item_type": "ignore", "raw_text": "Week 14", "name": "Week 14", "components": []},
                {"selected": True, "item_type": "ignore", "raw_text": "Monday", "name": "Monday", "components": []},
                {"selected": True, "item_type": "dish", "raw_text": "Alt 1", "name": "Alt 1", "components": []},
                {"selected": True, "item_type": "dish", "raw_text": "Fiskgratang", "name": "Fiskgratang", "components": [{"name": "Fisk"}, {"name": "Potatis"}]},
                {"selected": True, "item_type": "dish", "raw_text": "Fiskgratang", "name": "Fiskgratang", "components": [{"name": "Fisk"}, {"name": "Potatis"}]},
                {"selected": True, "item_type": "component", "raw_text": "Potatis", "name": "Potatis", "components": []},
                {"selected": True, "item_type": "component", "raw_text": "Potatis", "name": "Potatis", "components": []},
                {},
                {"selected": "badbool", "item_type": None, "raw_text": "", "name": None, "components": "bad"},
            ],
        },
        headers=HEADERS,
    )

    assert rv.status_code == 201
    body = rv.get_json() or {}
    assert body.get("ok") is True
    session = body.get("session") or {}
    items = body.get("items") or []
    assert session.get("session_id")
    assert isinstance(items, list)


def test_import_sessions_invalid_items_payload_returns_structured_400() -> None:
    client = _sqlite_client()

    rv = client.post(
        "/api/builder/import/sessions",
        json={"items": "not-a-list"},
        headers=HEADERS,
    )

    assert rv.status_code == 400
    body = rv.get_json() or {}
    assert body.get("ok") is False
    assert body.get("error") == "bad_request"
    details = body.get("details") or {}
    assert details.get("field") == "items"


def test_import_review_publish_uses_inline_edited_component_name() -> None:
    client = _client()

    rv = client.post(
        "/api/builder/import/publish-drafts",
        json={
            "items": [
                {
                    "selected": True,
                    "item_type": "component",
                    "raw_text": "Menyval1:köttbullar",
                    "name": "Köttbullar",
                    "components": [],
                }
            ]
        },
        headers=HEADERS,
    )

    assert rv.status_code == 200
    library = client.get("/api/builder/library", headers=HEADERS).get_json() or {}
    names = [item.get("component_name") for item in (library.get("components") or [])]
    assert "Köttbullar" in names


def test_import_review_publish_skips_ignored_items() -> None:
    client = _client()

    rv = client.post(
        "/api/builder/import/publish-drafts",
        json={
            "items": [
                {
                    "selected": True,
                    "item_type": "ignore",
                    "raw_text": "Lördag",
                    "name": "Lördag",
                    "components": [],
                }
            ]
        },
        headers=HEADERS,
    )

    assert rv.status_code == 200
    summary = (rv.get_json() or {}).get("summary") or {}
    assert summary.get("imported_count") == 0
    assert summary.get("ignored_count") == 1
    library = client.get("/api/builder/library", headers=HEADERS).get_json() or {}
    assert (library.get("components") or []) == []
    assert (library.get("compositions") or []) == []


def test_import_review_publish_only_selected_items() -> None:
    client = _client()

    rv = client.post(
        "/api/builder/import/publish-drafts",
        json={
            "items": [
                {
                    "selected": True,
                    "item_type": "component",
                    "raw_text": "Potatis",
                    "name": "Potatis",
                    "components": [],
                },
                {
                    "selected": False,
                    "item_type": "component",
                    "raw_text": "Ris",
                    "name": "Ris",
                    "components": [],
                },
            ]
        },
        headers=HEADERS,
    )

    assert rv.status_code == 200
    summary = (rv.get_json() or {}).get("summary") or {}
    assert summary.get("imported_count") == 1
    library = client.get("/api/builder/library", headers=HEADERS).get_json() or {}
    names = [item.get("component_name") for item in (library.get("components") or [])]
    assert "Potatis" in names
    assert "Ris" not in names


def test_import_review_publish_dish_decomposes_to_components_not_single_component_row() -> None:
    client = _client()

    rv = client.post(
        "/api/builder/import/publish-drafts",
        json={
            "items": [
                {
                    "selected": True,
                    "item_type": "dish",
                    "raw_text": "Köttbullar med potatis",
                    "name": "Köttbullar med potatis",
                    "components": [{"name": "Köttbullar"}, {"name": "Potatis"}],
                }
            ]
        },
        headers=HEADERS,
    )

    assert rv.status_code == 200
    compositions = client.get("/api/builder/compositions", headers=HEADERS).get_json() or {}
    items = compositions.get("compositions") or []
    assert len(items) == 1
    components = (items[0].get("components") or [])
    component_names = [item.get("component_name") for item in components]
    assert "Köttbullar" in component_names
    assert "Potatis" in component_names
    assert "Köttbullar med potatis" not in component_names


def test_import_review_publish_reuses_existing_component_and_avoids_duplicates() -> None:
    client = _client()

    created = client.post(
        "/api/builder/components",
        json={"component_name": "Potatismos"},
        headers=HEADERS,
    )
    assert created.status_code == 201

    rv = client.post(
        "/api/builder/import/publish-drafts",
        json={
            "items": [
                {
                    "selected": True,
                    "item_type": "dish",
                    "raw_text": "Köttfärslimpa med potatismos",
                    "name": "Köttfärslimpa med potatismos",
                    "components": [{"name": "Köttfärslimpa"}, {"name": "Potatismos"}, {"name": " Potatismos "}],
                }
            ]
        },
        headers=HEADERS,
    )

    assert rv.status_code == 200
    library = client.get("/api/builder/library", headers=HEADERS).get_json() or {}
    components = library.get("components") or []
    potatismos = [item for item in components if str(item.get("component_name") or "").strip().lower() == "potatismos"]
    assert len(potatismos) == 1

    compositions = client.get("/api/builder/compositions", headers=HEADERS).get_json() or {}
    items = compositions.get("compositions") or []
    assert len(items) == 1
    linked = items[0].get("components") or []
    linked_ids = [str(item.get("component_id") or "") for item in linked if str(item.get("component_name") or "").strip().lower() == "potatismos"]
    assert len(linked_ids) == 1


def test_import_review_publish_skips_phrase_fragment_component_from_payload() -> None:
    client = _client()

    rv = client.post(
        "/api/builder/import/publish-drafts",
        json={
            "items": [
                {
                    "selected": True,
                    "item_type": "dish",
                    "raw_text": "Kycklinggryta med smak av dragon",
                    "name": "Kycklinggryta med smak av dragon",
                    "components": [{"name": "Kycklinggryta"}, {"name": "Med smak av dragon"}],
                }
            ]
        },
        headers=HEADERS,
    )

    assert rv.status_code == 200
    compositions = client.get("/api/builder/compositions", headers=HEADERS).get_json() or {}
    items = compositions.get("compositions") or []
    assert len(items) == 1
    components = (items[0].get("components") or [])
    component_names = [str(item.get("component_name") or "") for item in components]
    assert "Kycklinggryta" in component_names
    assert all("smak av" not in name.lower() for name in component_names)


def test_import_review_publish_does_not_auto_link_existing_phrase_fragment_component() -> None:
    client = _client()

    existing = client.post(
        "/api/builder/components",
        json={"component_name": "Med smak av dragon"},
        headers=HEADERS,
    )
    assert existing.status_code == 201

    rv = client.post(
        "/api/builder/import/publish-drafts",
        json={
            "items": [
                {
                    "selected": True,
                    "item_type": "dish",
                    "raw_text": "Kycklinggryta med smak av dragon",
                    "name": "Kycklinggryta med smak av dragon",
                    "components": [],
                }
            ]
        },
        headers=HEADERS,
    )

    assert rv.status_code == 200
    compositions = client.get("/api/builder/compositions", headers=HEADERS).get_json() or {}
    items = compositions.get("compositions") or []
    assert len(items) == 1
    components = (items[0].get("components") or [])
    names = [str(item.get("component_name") or "") for item in components]
    assert "Kycklinggryta" in names
    assert all("med smak av" not in value.lower() for value in names)


def test_builder_file_import_preview_xlsx_supports_explicit_column_name() -> None:
    client = _client()
    payload = _xlsx_bytes(
        [
            ["id", "text", "tag"],
            ["1", "Kottbullar med potatismos", "A"],
            ["2", "Fiskgratang", "B"],
        ]
    )

    rv = client.post(
        "/api/builder/import/file/preview",
        data={
            "file": (io.BytesIO(payload), "library.xlsx"),
            "csv_column": "text",
        },
        content_type="multipart/form-data",
        headers=HEADERS,
    )

    assert rv.status_code == 200
    preview = (rv.get_json() or {}).get("preview") or {}
    assert preview.get("file_type") == "xlsx"
    assert preview.get("importable_lines") == ["Kottbullar med potatismos", "Fiskgratang"]
    assert preview.get("csv_column") == "text"
    assert preview.get("csv_column_index") == 1


def test_builder_file_import_preview_xlsx_classifies_noise_vs_importable() -> None:
    client = _client()
    payload = _xlsx_bytes(
        [
            ["text"],
            ["Week 12"],
            ["Alt 1"],
            ["Fiskgratang"],
        ]
    )

    rv = client.post(
        "/api/builder/import/file/preview",
        data={"file": (io.BytesIO(payload), "library.xlsx")},
        content_type="multipart/form-data",
        headers=HEADERS,
    )

    assert rv.status_code == 200
    preview = (rv.get_json() or {}).get("preview") or {}
    assert preview.get("importable_lines") == ["Fiskgratang"]
    ignored = preview.get("ignored_lines") or []
    ignored_texts = {item.get("normalized_text") for item in ignored}
    assert "Week 12" in ignored_texts
    assert "Alt 1" in ignored_texts


def test_builder_file_import_confirm_reuses_hardened_pipeline() -> None:
    client = _client()
    preview_payload = b"text\nKottbullar med potatismos\n"

    preview = client.post(
        "/api/builder/import/file/preview",
        data={"file": (io.BytesIO(preview_payload), "library.csv")},
        content_type="multipart/form-data",
        headers=HEADERS,
    )
    lines = ((preview.get_json() or {}).get("preview") or {}).get("lines") or []

    first = client.post(
        "/api/builder/import/file/confirm",
        json={"lines": lines},
        headers=HEADERS,
    )
    second = client.post(
        "/api/builder/import/file/confirm",
        json={"lines": lines},
        headers=HEADERS,
    )

    assert first.status_code == 200
    assert second.status_code == 200

    first_summary = ((first.get_json() or {}).get("summary") or {})
    second_summary = ((second.get_json() or {}).get("summary") or {})
    assert first_summary.get("created_count") == 1
    assert second_summary.get("created_count") == 0
    assert second_summary.get("reused_count") == 1
    assert first_summary.get("created_composition_count") == 1
    assert second_summary.get("reused_composition_count") == 1
    assert first_summary.get("ignored_noise_count") == 0

    row = (first_summary.get("row_results") or [{}])[0]
    assert row.get("kind") == "composition"
    assert row.get("composition_id")


def test_builder_file_import_confirm_imports_only_preview_importable_lines() -> None:
    client = _client()
    preview_payload = b"Alt 1\nFiskgratang\nAlt 2\n"

    preview = client.post(
        "/api/builder/import/file/preview",
        data={"file": (io.BytesIO(preview_payload), "library.txt")},
        content_type="multipart/form-data",
        headers=HEADERS,
    )
    preview_body = preview.get_json() or {}
    importable_lines = ((preview_body.get("preview") or {}).get("importable_lines") or [])
    selected_lines = importable_lines[:1]

    confirmed = client.post(
        "/api/builder/import/file/confirm",
        json={"lines": selected_lines, "ignored_noise_count": 2},
        headers=HEADERS,
    )

    assert confirmed.status_code == 200
    summary = ((confirmed.get_json() or {}).get("summary") or {})
    assert summary.get("imported_count") == 1
    assert summary.get("ignored_noise_count") == 2
    rows = summary.get("row_results") or []
    assert len(rows) == 1
    assert rows[0].get("raw_text") == "Fiskgratang"


def test_builder_file_import_confirm_reuses_pipeline_for_xlsx_preview_lines() -> None:
    client = _client()
    preview_payload = _xlsx_bytes(
        [
            ["text"],
            ["Kottbullar med potatismos"],
        ]
    )

    preview = client.post(
        "/api/builder/import/file/preview",
        data={"file": (io.BytesIO(preview_payload), "library.xlsx")},
        content_type="multipart/form-data",
        headers=HEADERS,
    )
    lines = ((preview.get_json() or {}).get("preview") or {}).get("lines") or []

    first = client.post(
        "/api/builder/import/file/confirm",
        json={"lines": lines},
        headers=HEADERS,
    )
    second = client.post(
        "/api/builder/import/file/confirm",
        json={"lines": lines},
        headers=HEADERS,
    )

    assert first.status_code == 200
    assert second.status_code == 200

    first_summary = ((first.get_json() or {}).get("summary") or {})
    second_summary = ((second.get_json() or {}).get("summary") or {})
    assert first_summary.get("created_count") == 1
    assert second_summary.get("created_count") == 0
    assert second_summary.get("reused_count") == 1

    row = (first_summary.get("row_results") or [{}])[0]
    assert row.get("kind") == "composition"
    assert row.get("composition_id")
    assert "day" not in row
    assert "meal_slot" not in row


def test_builder_file_import_confirm_summary_reports_component_creation_and_reuse() -> None:
    client = _client()

    rv = client.post(
        "/api/builder/import/file/confirm",
        json={
            "lines": [
                "Kottbullar med potatismos",
                "Kottbullar med graddsas",
            ],
            "ignored_noise_count": 3,
        },
        headers=HEADERS,
    )

    assert rv.status_code == 200
    summary = ((rv.get_json() or {}).get("summary") or {})
    assert summary.get("imported_count") == 2
    assert summary.get("created_composition_count") == 2
    assert summary.get("reused_composition_count") == 0
    assert summary.get("created_component_count") == 3
    assert summary.get("reused_component_count") == 1
    assert summary.get("ignored_noise_count") == 3


def test_builder_file_import_confirm_response_remains_library_only() -> None:
    client = _client()

    rv = client.post(
        "/api/builder/import/file/confirm",
        json={"lines": ["Fiskgratang"]},
        headers=HEADERS,
    )

    assert rv.status_code == 200
    summary = ((rv.get_json() or {}).get("summary") or {})
    assert "day" not in summary
    assert "meal_slot" not in summary
    assert "menu_detail_id" not in summary
    rows = summary.get("row_results") or []
    assert len(rows) == 1
    row = rows[0]
    assert "day" not in row
    assert "meal_slot" not in row
    assert "menu_detail_id" not in row


def test_list_compositions_endpoint() -> None:
    client = _client()
    client.post(
        "/api/builder/compositions",
        json={"composition_id": "plate_1", "composition_name": "Fish Plate"},
        headers=HEADERS,
    )

    rv = client.get("/api/builder/compositions", headers=HEADERS)

    assert rv.status_code == 200
    body = rv.get_json() or {}
    assert body.get("ok") is True
    assert body.get("count") == 1
    compositions = body.get("compositions") or []
    assert compositions[0]["composition_id"] == "plate_1"
    assert compositions[0]["use_custom_menu_name"] is False
    assert compositions[0]["menu_name"] is None
    assert compositions[0]["effective_menu_name"] == "Fish Plate"


def test_add_component_to_composition_endpoint() -> None:
    client = _client()
    client.post(
        "/api/builder/compositions",
        json={"composition_id": "plate_1", "composition_name": "Fish Plate"},
        headers=HEADERS,
    )

    rv = client.post(
        "/api/builder/compositions/plate_1/components",
        json={"component_name": "Fisk", "role": "component"},
        headers=HEADERS,
    )

    assert rv.status_code == 200
    body = rv.get_json() or {}
    assert body.get("ok") is True
    components = body.get("composition", {}).get("components") or []
    assert len(components) == 1
    assert components[0]["component_name"] == "Fisk"
    assert components[0]["component_id"] == "fisk"
    assert components[0]["role"] == "component"


def test_add_component_to_composition_endpoint_supports_connector_role() -> None:
    client = _client()
    client.post(
        "/api/builder/compositions",
        json={"composition_id": "plate_2", "composition_name": "Fish Plate 2"},
        headers=HEADERS,
    )

    rv = client.post(
        "/api/builder/compositions/plate_2/components",
        json={"component_name": "med", "role": "connector"},
        headers=HEADERS,
    )

    assert rv.status_code == 200
    body = rv.get_json() or {}
    assert body.get("ok") is True
    components = body.get("composition", {}).get("components") or []
    assert len(components) == 1
    assert components[0]["component_id"] == "med"
    assert components[0]["role"] == "connector"


def test_add_component_to_composition_endpoint_allows_empty_role() -> None:
    client = _client()
    client.post(
        "/api/builder/compositions",
        json={"composition_id": "plate_no_role", "composition_name": "Role Free Plate"},
        headers=HEADERS,
    )

    rv = client.post(
        "/api/builder/compositions/plate_no_role/components",
        json={"component_name": "Fisk"},
        headers=HEADERS,
    )

    assert rv.status_code == 200
    components = ((rv.get_json() or {}).get("composition") or {}).get("components") or []
    assert len(components) == 1
    assert components[0].get("role") is None


def test_attach_existing_component_endpoint_reuses_component_id_and_does_not_create_duplicate_entity() -> None:
    client = _client()
    created_component = client.post(
        "/api/builder/components",
        json={"component_name": "Mashed Potatoes"},
        headers=HEADERS,
    )
    component_id = ((created_component.get_json() or {}).get("component") or {}).get("component_id")
    client.post(
        "/api/builder/compositions",
        json={"composition_id": "plate_attach", "composition_name": "Fish Plate"},
        headers=HEADERS,
    )

    rv = client.post(
        "/api/builder/compositions/plate_attach/components/attach",
        json={"component_id": component_id, "role": "component"},
        headers=HEADERS,
    )

    assert rv.status_code == 200
    body = rv.get_json() or {}
    components = body.get("composition", {}).get("components") or []
    assert len(components) == 1
    assert components[0].get("component_id") == component_id

    listed = client.get("/api/builder/components", headers=HEADERS).get_json() or {}
    assert len([item for item in (listed.get("components") or []) if item.get("component_id") == component_id]) == 1


def test_attach_existing_component_endpoint_rejects_empty_or_invalid_id() -> None:
    client = _client()
    client.post(
        "/api/builder/compositions",
        json={"composition_id": "plate_invalid", "composition_name": "Fish Plate"},
        headers=HEADERS,
    )

    empty_rv = client.post(
        "/api/builder/compositions/plate_invalid/components/attach",
        json={"component_id": "   "},
        headers=HEADERS,
    )
    invalid_rv = client.post(
        "/api/builder/compositions/plate_invalid/components/attach",
        json={"component_id": "unknown"},
        headers=HEADERS,
    )

    assert empty_rv.status_code == 400
    assert invalid_rv.status_code == 400


def test_attach_existing_component_endpoint_no_menu_linkage_required() -> None:
    client = _client()
    component = client.post(
        "/api/builder/components",
        json={"component_name": "Gravy"},
        headers=HEADERS,
    )
    component_id = ((component.get_json() or {}).get("component") or {}).get("component_id")
    composition = client.post(
        "/api/builder/compositions",
        json={"composition_name": "Fish Soup"},
        headers=HEADERS,
    )
    composition_id = ((composition.get_json() or {}).get("composition") or {}).get("composition_id")

    rv = client.post(
        "/api/builder/compositions/" + str(composition_id) + "/components/attach",
        json={"component_id": component_id, "role": "component"},
        headers=HEADERS,
    )

    assert rv.status_code == 200
    body = rv.get_json() or {}
    assert body.get("composition", {}).get("composition_id") == composition_id


def test_attach_existing_component_endpoint_prevents_duplicate_attach() -> None:
    client = _client()
    component = client.post(
        "/api/builder/components",
        json={"component_name": "Rice"},
        headers=HEADERS,
    )
    component_id = ((component.get_json() or {}).get("component") or {}).get("component_id")
    client.post(
        "/api/builder/compositions",
        json={"composition_id": "plate_dupe", "composition_name": "Rice Plate"},
        headers=HEADERS,
    )
    first = client.post(
        "/api/builder/compositions/plate_dupe/components/attach",
        json={"component_id": component_id, "role": "component"},
        headers=HEADERS,
    )
    second = client.post(
        "/api/builder/compositions/plate_dupe/components/attach",
        json={"component_id": component_id, "role": "component"},
        headers=HEADERS,
    )

    assert first.status_code == 200
    assert second.status_code == 400


def test_reorder_components_endpoint_persists_order_and_sort_order() -> None:
    client = _client()
    client.post(
        "/api/builder/compositions",
        json={"composition_id": "plate_reorder", "composition_name": "Reorder Plate"},
        headers=HEADERS,
    )
    a = client.post(
        "/api/builder/compositions/plate_reorder/components",
        json={"component_name": "Potato", "role": "side"},
        headers=HEADERS,
    )
    client.post(
        "/api/builder/compositions/plate_reorder/components",
        json={"component_name": "Fish", "role": "main"},
        headers=HEADERS,
    )

    components = ((a.get_json() or {}).get("composition") or {}).get("components") or []
    assert len(components) == 1
    listed_before = client.get("/api/builder/compositions", headers=HEADERS).get_json() or {}
    plate_before = next(
        item
        for item in (listed_before.get("compositions") or [])
        if item.get("composition_id") == "plate_reorder"
    )
    entries = plate_before.get("components") or []

    rv = client.patch(
        "/api/builder/compositions/plate_reorder/components/reorder",
        json={
            "ordered_entries": [
                {"component_id": entries[1].get("component_id"), "sort_order": entries[1].get("sort_order")},
                {"component_id": entries[0].get("component_id"), "sort_order": entries[0].get("sort_order")},
            ]
        },
        headers=HEADERS,
    )

    assert rv.status_code == 200
    body = rv.get_json() or {}
    after = ((body.get("composition") or {}).get("components") or [])
    assert [item.get("component_name") for item in after] == ["Fish", "Potato"]
    assert [item.get("sort_order") for item in after] == [10, 20]
    assert "primary_recipe_id" not in (after[0] if after else {})

    listed_after = client.get("/api/builder/compositions", headers=HEADERS).get_json() or {}
    plate_after = next(
        item
        for item in (listed_after.get("compositions") or [])
        if item.get("composition_id") == "plate_reorder"
    )
    reloaded = plate_after.get("components") or []
    assert [item.get("component_name") for item in reloaded] == ["Fish", "Potato"]
    assert [item.get("sort_order") for item in reloaded] == [10, 20]


def test_render_composition_text_endpoint_uses_persisted_order() -> None:
    client = _client()
    client.post(
        "/api/builder/compositions",
        json={"composition_id": "plate_text", "composition_name": "Text Plate"},
        headers=HEADERS,
    )
    client.post(
        "/api/builder/compositions/plate_text/components",
        json={"component_name": "Potato", "role": "side"},
        headers=HEADERS,
    )
    client.post(
        "/api/builder/compositions/plate_text/components",
        json={"component_name": "Fish", "role": "main"},
        headers=HEADERS,
    )

    listed = client.get("/api/builder/compositions", headers=HEADERS).get_json() or {}
    plate = next(
        item
        for item in (listed.get("compositions") or [])
        if item.get("composition_id") == "plate_text"
    )
    entries = plate.get("components") or []
    reorder = client.patch(
        "/api/builder/compositions/plate_text/components/reorder",
        json={
            "ordered_entries": [
                {"component_id": entries[1].get("component_id"), "sort_order": entries[1].get("sort_order")},
                {"component_id": entries[0].get("component_id"), "sort_order": entries[0].get("sort_order")},
            ]
        },
        headers=HEADERS,
    )
    assert reorder.status_code == 200

    rv = client.get("/api/builder/compositions/plate_text/render/text", headers=HEADERS)

    assert rv.status_code == 200
    body = rv.get_json() or {}
    rendered = body.get("rendered") or {}
    assert body.get("ok") is True
    assert rendered.get("text") == "Text Plate: Fish (main), Potato (side)"
    names = [item.get("component_name") for item in (rendered.get("components") or [])]
    assert names == ["Fish", "Potato"]
    assert [item.get("sort_order") for item in (rendered.get("components") or [])] == [10, 20]


def test_render_composition_text_endpoint_is_deterministic_and_isolated() -> None:
    client = _client()
    client.post(
        "/api/builder/compositions",
        json={"composition_id": "plate_a", "composition_name": "Plate A"},
        headers=HEADERS,
    )
    client.post(
        "/api/builder/compositions",
        json={"composition_id": "plate_b", "composition_name": "Plate B"},
        headers=HEADERS,
    )
    client.post(
        "/api/builder/compositions/plate_a/components",
        json={"component_name": "Fish"},
        headers=HEADERS,
    )
    client.post(
        "/api/builder/compositions/plate_b/components",
        json={"component_name": "Soup"},
        headers=HEADERS,
    )

    first = client.get("/api/builder/compositions/plate_a/render/text", headers=HEADERS)
    second = client.get("/api/builder/compositions/plate_a/render/text", headers=HEADERS)

    assert first.status_code == 200
    assert second.status_code == 200
    first_rendered = (first.get_json() or {}).get("rendered") or {}
    second_rendered = (second.get_json() or {}).get("rendered") or {}
    assert first_rendered.get("text") == "Plate A: Fish"
    assert second_rendered.get("text") == "Plate A: Fish"
    assert first_rendered.get("text") == second_rendered.get("text")
    assert [item.get("component_name") for item in (first_rendered.get("components") or [])] == ["Fish"]


def test_remove_component_from_composition_endpoint() -> None:
    client = _client()
    client.post(
        "/api/builder/compositions",
        json={"composition_id": "plate_3", "composition_name": "Fish Plate 3"},
        headers=HEADERS,
    )
    client.post(
        "/api/builder/compositions/plate_3/components",
        json={"component_name": "Fisk", "role": "component"},
        headers=HEADERS,
    )
    client.post(
        "/api/builder/compositions/plate_3/components",
        json={"component_name": "Potatis", "role": "component"},
        headers=HEADERS,
    )

    rv = client.delete(
        "/api/builder/compositions/plate_3/components/fisk",
        headers=HEADERS,
    )

    assert rv.status_code == 200
    body = rv.get_json() or {}
    assert body.get("ok") is True
    components = body.get("composition", {}).get("components") or []
    assert len(components) == 1
    assert components[0]["component_id"] == "potatis"


def test_rename_component_in_composition_endpoint() -> None:
    client = _client()
    client.post(
        "/api/builder/compositions",
        json={"composition_id": "plate_4", "composition_name": "Fish Plate 4"},
        headers=HEADERS,
    )
    client.post(
        "/api/builder/compositions/plate_4/components",
        json={"component_name": "Fisk", "role": "connector"},
        headers=HEADERS,
    )
    client.post(
        "/api/builder/compositions/plate_4/components",
        json={"component_name": "Potatis", "role": "component"},
        headers=HEADERS,
    )

    rv = client.patch(
        "/api/builder/compositions/plate_4/components/fisk",
        json={"component_name": "Lax"},
        headers=HEADERS,
    )

    assert rv.status_code == 200
    body = rv.get_json() or {}
    assert body.get("ok") is True
    components = body.get("composition", {}).get("components") or []
    assert [item["component_id"] for item in components] == ["fisk", "potatis"]
    assert [item["component_name"] for item in components] == ["Lax", "Potatis"]
    assert components[0]["role"] == "connector"

    composition_rv = client.get("/api/builder/compositions", headers=HEADERS)
    assert composition_rv.status_code == 200
    compositions = (composition_rv.get_json() or {}).get("compositions") or []
    target = next((item for item in compositions if item.get("composition_id") == "plate_4"), None)
    assert target is not None
    composition_components = target.get("components") or []
    assert [item["component_id"] for item in composition_components] == ["fisk", "potatis"]
    assert [item["component_name"] for item in composition_components] == ["Lax", "Potatis"]


def test_rename_component_in_composition_endpoint_preserves_swedish_component_name() -> None:
    client = _client()
    client.post(
        "/api/builder/compositions",
        json={"composition_id": "plate_6", "composition_name": "Fish Plate 6"},
        headers=HEADERS,
    )
    client.post(
        "/api/builder/compositions/plate_6/components",
        json={"component_name": "Fisk", "role": "component"},
        headers=HEADERS,
    )

    rv = client.patch(
        "/api/builder/compositions/plate_6/components/fisk",
        json={"component_name": "Köttbullar"},
        headers=HEADERS,
    )

    assert rv.status_code == 200
    body = rv.get_json() or {}
    components = body.get("composition", {}).get("components") or []
    assert len(components) == 1
    assert components[0]["component_name"] == "Köttbullar"
    assert components[0]["component_id"] == "fisk"


def test_rename_then_add_component_persists_in_list_reload_for_coarse_component_flow() -> None:
    client = _client()
    client.post(
        "/api/builder/compositions",
        json={"composition_id": "plate_coarse", "composition_name": "Plate coarse"},
        headers=HEADERS,
    )
    client.post(
        "/api/builder/compositions/plate_coarse/components",
        json={"component_name": "stekt lök,potatis", "role": "side"},
        headers=HEADERS,
    )

    renamed = client.patch(
        "/api/builder/compositions/plate_coarse/components/stekt_lok_potatis",
        json={"component_name": "stekt lök"},
        headers=HEADERS,
    )
    assert renamed.status_code == 200

    added = client.post(
        "/api/builder/compositions/plate_coarse/components",
        json={"component_name": "kokt potatis", "role": "side"},
        headers=HEADERS,
    )
    assert added.status_code == 200

    listed = client.get("/api/builder/compositions", headers=HEADERS)
    assert listed.status_code == 200

    compositions = (listed.get_json() or {}).get("compositions") or []
    target = next((item for item in compositions if item.get("composition_id") == "plate_coarse"), None)
    assert target is not None
    names = [str(item.get("component_name") or "") for item in (target.get("components") or [])]
    assert "stekt lök,potatis" not in names
    assert "stekt lök" in names
    assert "kokt potatis" in names


def test_rename_component_in_composition_endpoint_rejects_empty_name() -> None:
    client = _client()
    client.post(
        "/api/builder/compositions",
        json={"composition_id": "plate_5", "composition_name": "Fish Plate 5"},
        headers=HEADERS,
    )
    client.post(
        "/api/builder/compositions/plate_5/components",
        json={"component_name": "Fisk", "role": "component"},
        headers=HEADERS,
    )

    rv = client.patch(
        "/api/builder/compositions/plate_5/components/fisk",
        json={"component_name": "   "},
        headers=HEADERS,
    )

    assert rv.status_code == 400
    body = rv.get_json() or {}
    assert body.get("error") == "bad_request"


def test_update_component_role_in_composition_endpoint_role_only() -> None:
    client = _client()
    client.post(
        "/api/builder/compositions",
        json={"composition_id": "plate_role_patch", "composition_name": "Fish Plate"},
        headers=HEADERS,
    )
    client.post(
        "/api/builder/compositions/plate_role_patch/components",
        json={"component_name": "Fisk", "role": "component"},
        headers=HEADERS,
    )

    rv = client.patch(
        "/api/builder/compositions/plate_role_patch/components/fisk",
        json={"role": "main"},
        headers=HEADERS,
    )

    assert rv.status_code == 200
    components = ((rv.get_json() or {}).get("composition") or {}).get("components") or []
    assert len(components) == 1
    assert components[0].get("component_id") == "fisk"
    assert components[0].get("role") == "main"


def test_update_component_role_in_composition_endpoint_allows_clearing_role() -> None:
    client = _client()
    client.post(
        "/api/builder/compositions",
        json={"composition_id": "plate_role_clear", "composition_name": "Fish Plate"},
        headers=HEADERS,
    )
    client.post(
        "/api/builder/compositions/plate_role_clear/components",
        json={"component_name": "Fisk", "role": "main"},
        headers=HEADERS,
    )

    rv = client.patch(
        "/api/builder/compositions/plate_role_clear/components/fisk",
        json={"role": "   "},
        headers=HEADERS,
    )

    assert rv.status_code == 200
    components = ((rv.get_json() or {}).get("composition") or {}).get("components") or []
    assert len(components) == 1
    assert components[0].get("role") is None


def test_rename_component_in_composition_endpoint_can_update_name_and_role_together() -> None:
    client = _client()
    client.post(
        "/api/builder/compositions",
        json={"composition_id": "plate_name_role_patch", "composition_name": "Fish Plate"},
        headers=HEADERS,
    )
    client.post(
        "/api/builder/compositions/plate_name_role_patch/components",
        json={"component_name": "Fisk", "role": "component"},
        headers=HEADERS,
    )

    rv = client.patch(
        "/api/builder/compositions/plate_name_role_patch/components/fisk",
        json={"component_name": "Lax", "role": "main"},
        headers=HEADERS,
    )

    assert rv.status_code == 200
    components = ((rv.get_json() or {}).get("composition") or {}).get("components") or []
    assert len(components) == 1
    assert components[0].get("component_id") == "fisk"
    assert components[0].get("component_name") == "Lax"
    assert components[0].get("role") == "main"


def test_component_rename_keeps_shared_dishes_on_same_component_id() -> None:
    client = _client()
    client.post(
        "/api/builder/components",
        json={"component_name": "Hoisinsås og ris"},
        headers=HEADERS,
    )

    client.post(
        "/api/builder/compositions",
        json={"composition_id": "plate_shared_1", "composition_name": "Dish 1"},
        headers=HEADERS,
    )
    client.post(
        "/api/builder/compositions",
        json={"composition_id": "plate_shared_2", "composition_name": "Dish 2"},
        headers=HEADERS,
    )
    client.post(
        "/api/builder/compositions/plate_shared_1/components",
        json={"component_name": "Hoisinsås og ris", "role": "component"},
        headers=HEADERS,
    )
    client.post(
        "/api/builder/compositions/plate_shared_2/components/attach",
        json={"component_id": "hoisinsas_og_ris", "role": "component"},
        headers=HEADERS,
    )

    renamed = client.patch(
        "/api/builder/components/hoisinsas_og_ris",
        json={"name": "Hoisinsås"},
        headers=HEADERS,
    )
    assert renamed.status_code == 200

    listed_compositions = client.get("/api/builder/compositions", headers=HEADERS)
    assert listed_compositions.status_code == 200
    compositions = (listed_compositions.get_json() or {}).get("compositions") or []
    for composition_id in ("plate_shared_1", "plate_shared_2"):
        target = next((item for item in compositions if item.get("composition_id") == composition_id), None)
        assert target is not None
        components = target.get("components") or []
        assert [item.get("component_id") for item in components] == ["hoisinsas_og_ris"]
        assert [item.get("component_name") for item in components] == ["Hoisinsås"]

    listed = client.get("/api/builder/components", headers=HEADERS)
    assert listed.status_code == 200
    ids = [item.get("component_id") for item in ((listed.get_json() or {}).get("components") or [])]
    assert ids.count("hoisinsas_og_ris") == 1


def test_update_component_endpoint_requires_component_name_or_role() -> None:
    client = _client()
    client.post(
        "/api/builder/compositions",
        json={"composition_id": "plate_patch_required", "composition_name": "Fish Plate"},
        headers=HEADERS,
    )
    client.post(
        "/api/builder/compositions/plate_patch_required/components",
        json={"component_name": "Fisk", "role": "component"},
        headers=HEADERS,
    )

    rv = client.patch(
        "/api/builder/compositions/plate_patch_required/components/fisk",
        json={},
        headers=HEADERS,
    )

    assert rv.status_code == 400
    body = rv.get_json() or {}
    assert body.get("error") == "bad_request"


def test_invalid_payload_handling_returns_400() -> None:
    client = _client()

    rv1 = client.post(
        "/api/builder/compositions",
        json={"composition_id": "plate_1"},
        headers=HEADERS,
    )
    assert rv1.status_code == 400
    body1 = rv1.get_json() or {}
    assert body1.get("error") == "bad_request"


def test_create_component_recipe_endpoint_requires_yield_portions_and_structured_lines() -> None:
    client = _client()
    component_rv = client.post(
        "/api/builder/components",
        json={"component_name": "Meatballs"},
        headers=HEADERS,
    )
    component_id = ((component_rv.get_json() or {}).get("component") or {}).get("component_id")

    rv = client.post(
        f"/api/builder/components/{component_id}/recipes",
        json={
            "recipe_name": "Base",
            "visibility": "site",
            "yield_portions": 24,
            "is_primary": True,
            "ingredient_lines": [
                {
                    "ingredient_name": "Potato",
                    "amount_value": 900,
                    "amount_unit": "g",
                    "note": "peeled",
                    "sort_order": 10,
                }
            ],
        },
        headers=HEADERS,
    )

    assert rv.status_code == 201
    body = rv.get_json() or {}
    recipe = body.get("recipe") or {}
    assert recipe.get("yield_portions") == 24
    lines = body.get("ingredient_lines") or []
    assert len(lines) == 1
    assert lines[0].get("ingredient_name") == "Potato"
    assert lines[0].get("amount_value") == 900.0
    assert lines[0].get("amount_unit") == "g"
    assert lines[0].get("note") == "peeled"

    list_components = client.get("/api/builder/components", headers=HEADERS)
    components = (list_components.get_json() or {}).get("components") or []
    linked = next(item for item in components if item.get("component_id") == component_id)
    assert linked.get("primary_recipe_id") == recipe.get("recipe_id")


def test_create_component_recipe_endpoint_rejects_missing_yield_portions() -> None:
    client = _client()
    component_rv = client.post(
        "/api/builder/components",
        json={"component_name": "Fish"},
        headers=HEADERS,
    )
    component_id = ((component_rv.get_json() or {}).get("component") or {}).get("component_id")

    rv = client.post(
        f"/api/builder/components/{component_id}/recipes",
        json={"recipe_name": "Base", "visibility": "private"},
        headers=HEADERS,
    )

    assert rv.status_code == 400
    body = rv.get_json() or {}
    assert body.get("error") == "bad_request"


def test_add_recipe_ingredient_endpoint_accepts_structured_amount_and_reads_back() -> None:
    client = _client()
    component_rv = client.post(
        "/api/builder/components",
        json={"component_name": "Sauce"},
        headers=HEADERS,
    )
    component_id = ((component_rv.get_json() or {}).get("component") or {}).get("component_id")

    recipe_rv = client.post(
        f"/api/builder/components/{component_id}/recipes",
        json={"recipe_name": "Sauce Base", "yield_portions": 10, "visibility": "private"},
        headers=HEADERS,
    )
    recipe_id = ((recipe_rv.get_json() or {}).get("recipe") or {}).get("recipe_id")

    add_rv = client.post(
        f"/api/builder/components/{component_id}/recipes/{recipe_id}/ingredients",
        json={
            "ingredient_name": "Cream",
            "amount_value": 2.5,
            "amount_unit": "dl",
            "note": "warm",
        },
        headers=HEADERS,
    )
    get_rv = client.get(
        f"/api/builder/components/{component_id}/recipes/{recipe_id}",
        headers=HEADERS,
    )

    assert add_rv.status_code == 201
    assert get_rv.status_code == 200
    lines = (get_rv.get_json() or {}).get("ingredient_lines") or []
    assert len(lines) == 1
    assert lines[0].get("ingredient_name") == "Cream"
    assert lines[0].get("amount_value") == 2.5
    assert lines[0].get("amount_unit") == "dl"
    assert lines[0].get("note") == "warm"


def test_set_component_primary_recipe_endpoint_rejects_recipe_from_other_component() -> None:
    client = _client()
    c1 = client.post("/api/builder/components", json={"component_name": "Fish"}, headers=HEADERS)
    c2 = client.post("/api/builder/components", json={"component_name": "Potato"}, headers=HEADERS)
    c1_id = ((c1.get_json() or {}).get("component") or {}).get("component_id")
    c2_id = ((c2.get_json() or {}).get("component") or {}).get("component_id")

    recipe_rv = client.post(
        f"/api/builder/components/{c1_id}/recipes",
        json={"recipe_name": "Fish Base", "yield_portions": 10, "visibility": "private"},
        headers=HEADERS,
    )
    recipe_id = ((recipe_rv.get_json() or {}).get("recipe") or {}).get("recipe_id")

    rv = client.patch(
        f"/api/builder/components/{c2_id}/recipes/primary",
        json={"recipe_id": recipe_id},
        headers=HEADERS,
    )

    assert rv.status_code == 400
    body = rv.get_json() or {}
    assert body.get("error") == "bad_request"


def test_list_component_recipes_endpoint_returns_component_scoped_deterministic_list() -> None:
    client = _client()
    c1 = client.post("/api/builder/components", json={"component_name": "Fish"}, headers=HEADERS)
    c2 = client.post("/api/builder/components", json={"component_name": "Potato"}, headers=HEADERS)
    c1_id = ((c1.get_json() or {}).get("component") or {}).get("component_id")
    c2_id = ((c2.get_json() or {}).get("component") or {}).get("component_id")

    r_b = client.post(
        f"/api/builder/components/{c1_id}/recipes",
        json={"recipe_name": "B Recipe", "yield_portions": 10, "visibility": "private"},
        headers=HEADERS,
    )
    r_a = client.post(
        f"/api/builder/components/{c1_id}/recipes",
        json={"recipe_name": "A Recipe", "yield_portions": 12, "visibility": "private"},
        headers=HEADERS,
    )
    client.post(
        f"/api/builder/components/{c2_id}/recipes",
        json={"recipe_name": "Other Component Recipe", "yield_portions": 8, "visibility": "private"},
        headers=HEADERS,
    )

    primary_id = ((r_b.get_json() or {}).get("recipe") or {}).get("recipe_id")
    client.patch(
        f"/api/builder/components/{c1_id}/recipes/primary",
        json={"recipe_id": primary_id},
        headers=HEADERS,
    )

    rv = client.get(f"/api/builder/components/{c1_id}/recipes", headers=HEADERS)

    assert rv.status_code == 200
    body = rv.get_json() or {}
    assert body.get("ok") is True
    recipes = body.get("recipes") or []
    assert len(recipes) == 2
    assert recipes[0].get("recipe_id") == primary_id
    assert recipes[0].get("is_primary") is True
    assert recipes[1].get("recipe_name") == "A Recipe"
    assert body.get("component", {}).get("component_id") == c1_id

    all_ids = {item.get("recipe_id") for item in recipes}
    assert ((r_a.get_json() or {}).get("recipe") or {}).get("recipe_id") in all_ids
    assert not any(item.get("recipe_name") == "Other Component Recipe" for item in recipes)


def test_update_recipe_metadata_endpoint() -> None:
    client = _client()
    c = client.post("/api/builder/components", json={"component_name": "Fish"}, headers=HEADERS)
    component_id = ((c.get_json() or {}).get("component") or {}).get("component_id")
    created = client.post(
        f"/api/builder/components/{component_id}/recipes",
        json={"recipe_name": "Old", "yield_portions": 10, "visibility": "private", "notes": "old"},
        headers=HEADERS,
    )
    recipe_id = ((created.get_json() or {}).get("recipe") or {}).get("recipe_id")

    rv = client.patch(
        f"/api/builder/components/{component_id}/recipes/{recipe_id}",
        json={"recipe_name": "New", "yield_portions": 24, "visibility": "site", "notes": "new"},
        headers=HEADERS,
    )

    assert rv.status_code == 200
    recipe = (rv.get_json() or {}).get("recipe") or {}
    assert recipe.get("recipe_name") == "New"
    assert recipe.get("yield_portions") == 24
    assert recipe.get("visibility") == "site"
    assert recipe.get("notes") == "new"


def test_update_and_delete_recipe_ingredient_endpoint() -> None:
    client = _client()
    c = client.post("/api/builder/components", json={"component_name": "Soup"}, headers=HEADERS)
    component_id = ((c.get_json() or {}).get("component") or {}).get("component_id")
    created = client.post(
        f"/api/builder/components/{component_id}/recipes",
        json={"recipe_name": "Base", "yield_portions": 10, "visibility": "private"},
        headers=HEADERS,
    )
    recipe_id = ((created.get_json() or {}).get("recipe") or {}).get("recipe_id")
    added = client.post(
        f"/api/builder/components/{component_id}/recipes/{recipe_id}/ingredients",
        json={"ingredient_name": "Salt", "amount_value": 10, "amount_unit": "g", "note": "initial"},
        headers=HEADERS,
    )
    line_id = ((added.get_json() or {}).get("ingredient_line") or {}).get("recipe_ingredient_line_id")

    update_rv = client.patch(
        f"/api/builder/components/{component_id}/recipes/{recipe_id}/ingredients/{line_id}",
        json={
            "ingredient_name": "Sea salt",
            "amount_value": 12,
            "amount_unit": "g",
            "note": "updated",
            "sort_order": 20,
        },
        headers=HEADERS,
    )
    delete_rv = client.delete(
        f"/api/builder/components/{component_id}/recipes/{recipe_id}/ingredients/{line_id}",
        headers=HEADERS,
    )
    detail_rv = client.get(
        f"/api/builder/components/{component_id}/recipes/{recipe_id}",
        headers=HEADERS,
    )

    assert update_rv.status_code == 200
    line = (update_rv.get_json() or {}).get("ingredient_line") or {}
    assert line.get("ingredient_name") == "Sea salt"
    assert line.get("amount_value") == 12.0
    assert line.get("sort_order") == 20
    assert delete_rv.status_code == 200
    lines = (detail_rv.get_json() or {}).get("ingredient_lines") or []
    assert lines == []


def test_delete_recipe_endpoint_guard_for_primary_recipe() -> None:
    client = _client()
    c = client.post("/api/builder/components", json={"component_name": "Stew"}, headers=HEADERS)
    component_id = ((c.get_json() or {}).get("component") or {}).get("component_id")
    created = client.post(
        f"/api/builder/components/{component_id}/recipes",
        json={"recipe_name": "Base", "yield_portions": 8, "visibility": "private", "is_primary": True},
        headers=HEADERS,
    )
    recipe_id = ((created.get_json() or {}).get("recipe") or {}).get("recipe_id")

    blocked = client.delete(
        f"/api/builder/components/{component_id}/recipes/{recipe_id}",
        headers=HEADERS,
    )
    clear_primary = client.patch(
        f"/api/builder/components/{component_id}/recipes/primary",
        json={"recipe_id": ""},
        headers=HEADERS,
    )
    deleted = client.delete(
        f"/api/builder/components/{component_id}/recipes/{recipe_id}",
        headers=HEADERS,
    )

    assert blocked.status_code == 400
    assert (blocked.get_json() or {}).get("error") == "bad_request"
    assert clear_primary.status_code == 200
    assert deleted.status_code == 200


def test_update_recipe_ingredient_enforces_component_ownership() -> None:
    client = _client()
    c1 = client.post("/api/builder/components", json={"component_name": "Fish"}, headers=HEADERS)
    c2 = client.post("/api/builder/components", json={"component_name": "Potato"}, headers=HEADERS)
    c1_id = ((c1.get_json() or {}).get("component") or {}).get("component_id")
    c2_id = ((c2.get_json() or {}).get("component") or {}).get("component_id")

    created = client.post(
        f"/api/builder/components/{c1_id}/recipes",
        json={"recipe_name": "Fish Base", "yield_portions": 10, "visibility": "private"},
        headers=HEADERS,
    )
    recipe_id = ((created.get_json() or {}).get("recipe") or {}).get("recipe_id")
    added = client.post(
        f"/api/builder/components/{c1_id}/recipes/{recipe_id}/ingredients",
        json={"ingredient_name": "Salt", "amount_value": 10, "amount_unit": "g"},
        headers=HEADERS,
    )
    line_id = ((added.get_json() or {}).get("ingredient_line") or {}).get("recipe_ingredient_line_id")

    rv = client.patch(
        f"/api/builder/components/{c2_id}/recipes/{recipe_id}/ingredients/{line_id}",
        json={"ingredient_name": "Salt", "amount_value": 11, "amount_unit": "g"},
        headers=HEADERS,
    )

    assert rv.status_code == 400
    assert (rv.get_json() or {}).get("error") == "bad_request"


def test_recipe_scaling_preview_endpoint_returns_scaled_rows_and_metadata() -> None:
    client = _client()
    c = client.post("/api/builder/components", json={"component_name": "Soup"}, headers=HEADERS)
    component_id = ((c.get_json() or {}).get("component") or {}).get("component_id")

    created = client.post(
        f"/api/builder/components/{component_id}/recipes",
        json={"recipe_name": "Base", "yield_portions": 4, "visibility": "private", "notes": "v1"},
        headers=HEADERS,
    )
    recipe_id = ((created.get_json() or {}).get("recipe") or {}).get("recipe_id")

    client.post(
        f"/api/builder/components/{component_id}/recipes/{recipe_id}/ingredients",
        json={
            "ingredient_name": "Water",
            "amount_value": 2,
            "amount_unit": "l",
            "note": "cold",
            "sort_order": 10,
        },
        headers=HEADERS,
    )
    client.post(
        f"/api/builder/components/{component_id}/recipes/{recipe_id}/ingredients",
        json={"ingredient_name": "Salt", "amount_value": 8, "amount_unit": "g", "sort_order": 20},
        headers=HEADERS,
    )

    rv = client.get(
        f"/api/builder/components/{component_id}/recipes/{recipe_id}/scaling-preview?target_portions=10",
        headers=HEADERS,
    )

    assert rv.status_code == 200
    body = rv.get_json() or {}
    preview = body.get("preview") or {}
    assert body.get("ok") is True
    recipe = preview.get("recipe") or {}
    assert recipe.get("recipe_id") == recipe_id
    assert recipe.get("component_id") == component_id
    assert recipe.get("recipe_name") == "Base"
    assert recipe.get("notes") == "v1"
    assert preview.get("source_yield_portions") == 4
    assert preview.get("target_portions") == 10
    assert preview.get("scaling_factor") == "2.5"
    lines = preview.get("ingredient_lines") or []
    assert [item.get("ingredient_name") for item in lines] == ["Water", "Salt"]
    assert [item.get("amount_unit") for item in lines] == ["l", "g"]
    assert [item.get("original_amount_value") for item in lines] == ["2", "8"]
    assert [item.get("scaled_amount_value") for item in lines] == ["5.0", "20.0"]
    assert [item.get("note") for item in lines] == ["cold", None]


def test_recipe_scaling_preview_endpoint_rejects_invalid_target_portions() -> None:
    client = _client()
    c = client.post("/api/builder/components", json={"component_name": "Soup"}, headers=HEADERS)
    component_id = ((c.get_json() or {}).get("component") or {}).get("component_id")
    created = client.post(
        f"/api/builder/components/{component_id}/recipes",
        json={"recipe_name": "Base", "yield_portions": 4, "visibility": "private"},
        headers=HEADERS,
    )
    recipe_id = ((created.get_json() or {}).get("recipe") or {}).get("recipe_id")

    missing = client.get(
        f"/api/builder/components/{component_id}/recipes/{recipe_id}/scaling-preview",
        headers=HEADERS,
    )
    invalid = client.get(
        f"/api/builder/components/{component_id}/recipes/{recipe_id}/scaling-preview?target_portions=0",
        headers=HEADERS,
    )

    assert missing.status_code == 400
    assert (missing.get_json() or {}).get("error") == "bad_request"
    assert invalid.status_code == 400
    assert (invalid.get_json() or {}).get("error") == "bad_request"


def test_recipe_scaling_preview_endpoint_enforces_component_ownership() -> None:
    client = _client()
    c1 = client.post("/api/builder/components", json={"component_name": "Fish"}, headers=HEADERS)
    c2 = client.post("/api/builder/components", json={"component_name": "Potato"}, headers=HEADERS)
    c1_id = ((c1.get_json() or {}).get("component") or {}).get("component_id")
    c2_id = ((c2.get_json() or {}).get("component") or {}).get("component_id")
    created = client.post(
        f"/api/builder/components/{c1_id}/recipes",
        json={"recipe_name": "Fish Base", "yield_portions": 6, "visibility": "private"},
        headers=HEADERS,
    )
    recipe_id = ((created.get_json() or {}).get("recipe") or {}).get("recipe_id")

    rv = client.get(
        f"/api/builder/components/{c2_id}/recipes/{recipe_id}/scaling-preview?target_portions=12",
        headers=HEADERS,
    )

    assert rv.status_code == 400
    assert (rv.get_json() or {}).get("error") == "bad_request"


def test_recipe_ingredient_trait_signals_readback_and_preview_endpoint() -> None:
    client = _client()
    c = client.post("/api/builder/components", json={"component_name": "Soup"}, headers=HEADERS)
    component_id = ((c.get_json() or {}).get("component") or {}).get("component_id")
    created = client.post(
        f"/api/builder/components/{component_id}/recipes",
        json={"recipe_name": "Base", "yield_portions": 6, "visibility": "private"},
        headers=HEADERS,
    )
    recipe_id = ((created.get_json() or {}).get("recipe") or {}).get("recipe_id")

    add_a = client.post(
        f"/api/builder/components/{component_id}/recipes/{recipe_id}/ingredients",
        json={
            "ingredient_name": "Cream",
            "amount_value": 2,
            "amount_unit": "dl",
            "trait_signals": ["lactose", "lactose"],
            "sort_order": 10,
        },
        headers=HEADERS,
    )
    add_b = client.post(
        f"/api/builder/components/{component_id}/recipes/{recipe_id}/ingredients",
        json={
            "ingredient_name": "Cod",
            "amount_value": 500,
            "amount_unit": "g",
            "trait_signals": ["fish"],
            "sort_order": 20,
        },
        headers=HEADERS,
    )
    detail = client.get(
        f"/api/builder/components/{component_id}/recipes/{recipe_id}",
        headers=HEADERS,
    )
    preview = client.get(
        f"/api/builder/components/{component_id}/recipes/{recipe_id}/trait-signals",
        headers=HEADERS,
    )

    assert add_a.status_code == 201
    assert add_b.status_code == 201
    assert detail.status_code == 200
    lines = (detail.get_json() or {}).get("ingredient_lines") or []
    assert [item.get("trait_signals") for item in lines] == [["lactose"], ["fish"]]

    assert preview.status_code == 200
    body = preview.get_json() or {}
    data = body.get("preview") or {}
    assert body.get("ok") is True
    assert data.get("trait_signals_present") == ["fish", "lactose"]
    line_signals = [item.get("trait_signals") for item in (data.get("ingredient_lines") or [])]
    assert line_signals == [["lactose"], ["fish"]]


def test_recipe_trait_signals_preview_enforces_component_ownership() -> None:
    client = _client()
    c1 = client.post("/api/builder/components", json={"component_name": "Fish"}, headers=HEADERS)
    c2 = client.post("/api/builder/components", json={"component_name": "Potato"}, headers=HEADERS)
    c1_id = ((c1.get_json() or {}).get("component") or {}).get("component_id")
    c2_id = ((c2.get_json() or {}).get("component") or {}).get("component_id")

    created = client.post(
        f"/api/builder/components/{c1_id}/recipes",
        json={"recipe_name": "Fish Base", "yield_portions": 6, "visibility": "private"},
        headers=HEADERS,
    )
    recipe_id = ((created.get_json() or {}).get("recipe") or {}).get("recipe_id")

    rv = client.get(
        f"/api/builder/components/{c2_id}/recipes/{recipe_id}/trait-signals",
        headers=HEADERS,
    )

    assert rv.status_code == 400
    assert (rv.get_json() or {}).get("error") == "bad_request"


def test_component_declaration_readiness_endpoint_returns_trait_sources() -> None:
    client = _client()
    created_component = client.post(
        "/api/builder/components",
        json={"component_name": "Fish Sauce"},
        headers=HEADERS,
    )
    component_id = ((created_component.get_json() or {}).get("component") or {}).get("component_id")
    created_recipe = client.post(
        f"/api/builder/components/{component_id}/recipes",
        json={"recipe_name": "Base", "yield_portions": 8, "visibility": "private"},
        headers=HEADERS,
    )
    recipe_id = ((created_recipe.get_json() or {}).get("recipe") or {}).get("recipe_id")
    client.post(
        f"/api/builder/components/{component_id}/recipes/{recipe_id}/ingredients",
        json={
            "ingredient_name": "Cod",
            "amount_value": 500,
            "amount_unit": "g",
            "trait_signals": ["fish"],
            "sort_order": 10,
        },
        headers=HEADERS,
    )

    rv = client.get(
        f"/api/builder/components/{component_id}/declaration-readiness?include_declaration=1",
        headers=HEADERS,
    )

    assert rv.status_code == 200
    body = rv.get_json() or {}
    readiness = body.get("readiness") or {}
    assert body.get("ok") is True
    assert body.get("declaration_enabled") is True
    assert readiness.get("component_id") == component_id
    assert readiness.get("trait_signals_present") == ["fish"]
    assert (readiness.get("conflict_preview") or {}).get("conflicts_present") == ["fish_relevant"]
    conflict_sources = (readiness.get("conflict_preview") or {}).get("conflict_sources") or []
    assert len(conflict_sources) == 1
    assert conflict_sources[0].get("conflict_key") == "fish_relevant"
    assert conflict_sources[0].get("triggering_trait_signals") == ["fish"]
    sources = readiness.get("ingredient_sources") or []
    assert len(sources) == 1
    assert sources[0].get("ingredient_name") == "Cod"
    assert sources[0].get("trait_signals") == ["fish"]


def test_composition_declaration_readiness_endpoint_aggregates_component_signals() -> None:
    client = _client()
    created = client.post(
        "/api/builder/compositions",
        json={"composition_id": "plate_1", "composition_name": "Fish Plate"},
        headers=HEADERS,
    )
    assert created.status_code == 201

    added = client.post(
        "/api/builder/compositions/plate_1/components",
        json={"component_name": "Fish", "role": "main", "sort_order": 10},
        headers=HEADERS,
    )
    component_id = (((added.get_json() or {}).get("composition") or {}).get("components") or [{}])[0].get(
        "component_id"
    )
    recipe = client.post(
        f"/api/builder/components/{component_id}/recipes",
        json={"recipe_name": "Fish Base", "yield_portions": 8, "visibility": "private"},
        headers=HEADERS,
    )
    recipe_id = ((recipe.get_json() or {}).get("recipe") or {}).get("recipe_id")
    client.post(
        f"/api/builder/components/{component_id}/recipes/{recipe_id}/ingredients",
        json={"ingredient_name": "Cod", "amount_value": 500, "amount_unit": "g", "trait_signals": ["fish"]},
        headers=HEADERS,
    )

    rv = client.get(
        "/api/builder/compositions/plate_1/declaration-readiness?include_declaration=true",
        headers=HEADERS,
    )

    assert rv.status_code == 200
    body = rv.get_json() or {}
    readiness = body.get("readiness") or {}
    assert body.get("declaration_enabled") is True
    assert readiness.get("composition_id") == "plate_1"
    assert readiness.get("trait_signals_present") == ["fish"]
    assert (readiness.get("conflict_preview") or {}).get("conflicts_present") == ["fish_relevant"]
    components = readiness.get("components") or []
    assert len(components) == 1
    assert components[0].get("trait_signals_present") == ["fish"]
    assert (components[0].get("conflict_preview") or {}).get("conflicts_present") == ["fish_relevant"]


def test_component_declaration_readiness_endpoint_can_be_disabled_by_toggle() -> None:
    client = _client()
    created = client.post(
        "/api/builder/components",
        json={"component_name": "Soup"},
        headers=HEADERS,
    )
    component_id = ((created.get_json() or {}).get("component") or {}).get("component_id")

    rv = client.get(
        f"/api/builder/components/{component_id}/declaration-readiness?include_declaration=0",
        headers=HEADERS,
    )

    assert rv.status_code == 200
    body = rv.get_json() or {}
    assert body.get("ok") is True
    assert body.get("declaration_enabled") is False
    assert body.get("readiness") is None


def test_cook_can_read_org_recipe_and_edit_own_private_fork_via_api() -> None:
    fd, db_path = tempfile.mkstemp(prefix="builder_api_cook_", suffix=".db")
    os.close(fd)
    app = _app_with_builder_db(db_path)
    admin_client = app.test_client()
    cook_client = app.test_client()
    admin_headers = _headers(role="admin", tenant_id=1, user_id=11)
    cook_headers = _headers(role="cook", tenant_id=1, user_id=42)
    _seed_session(admin_client, role="admin", tenant_id=1, user_id=11)
    _seed_session(cook_client, role="cook", tenant_id=1, user_id=42)

    component_rv = admin_client.post(
        "/api/builder/components",
        json={"component_name": "Cook API Soup"},
        headers=admin_headers,
    )
    component_id = ((component_rv.get_json() or {}).get("component") or {}).get("component_id")
    recipe_rv = admin_client.post(
        f"/api/builder/components/{component_id}/recipes",
        json={"recipe_name": "Base", "yield_portions": 4, "visibility": "private"},
        headers=admin_headers,
    )
    recipe_id = ((recipe_rv.get_json() or {}).get("recipe") or {}).get("recipe_id")
    line_rv = admin_client.post(
        f"/api/builder/components/{component_id}/recipes/{recipe_id}/ingredients",
        json={"ingredient_name": "Salt", "amount_value": 10, "amount_unit": "g"},
        headers=admin_headers,
    )
    line_id = ((line_rv.get_json() or {}).get("ingredient_line") or {}).get("recipe_ingredient_line_id")

    listed = cook_client.get(f"/api/builder/components/{component_id}/recipes", headers=cook_headers)
    detail = cook_client.get(f"/api/builder/components/{component_id}/recipes/{recipe_id}", headers=cook_headers)
    scaling = cook_client.get(
        f"/api/builder/components/{component_id}/recipes/{recipe_id}/scaling-preview?target_portions=8",
        headers=cook_headers,
    )
    traits = cook_client.get(
        f"/api/builder/components/{component_id}/recipes/{recipe_id}/trait-signals",
        headers=cook_headers,
    )

    fork_component = cook_client.post(f"/api/builder/components/{component_id}/fork", headers=cook_headers)
    fork_component_id = ((fork_component.get_json() or {}).get("component") or {}).get("component_id")
    fork_recipe = cook_client.post(
        f"/api/builder/components/{fork_component_id}/recipes",
        json={"recipe_name": "Fork Base", "yield_portions": 6, "visibility": "private"},
        headers=cook_headers,
    )
    fork_recipe_id = ((fork_recipe.get_json() or {}).get("recipe") or {}).get("recipe_id")
    fork_line = cook_client.post(
        f"/api/builder/components/{fork_component_id}/recipes/{fork_recipe_id}/ingredients",
        json={"ingredient_name": "Pepper", "amount_value": 1, "amount_unit": "g"},
        headers=cook_headers,
    )
    fork_line_id = ((fork_line.get_json() or {}).get("ingredient_line") or {}).get("recipe_ingredient_line_id")
    update_recipe = cook_client.patch(
        f"/api/builder/components/{fork_component_id}/recipes/{fork_recipe_id}",
        json={"recipe_name": "Fork Updated", "yield_portions": 12, "visibility": "private"},
        headers=cook_headers,
    )
    update_line = cook_client.patch(
        f"/api/builder/components/{fork_component_id}/recipes/{fork_recipe_id}/ingredients/{fork_line_id}",
        json={"ingredient_name": "Black pepper", "amount_value": 2, "amount_unit": "g"},
        headers=cook_headers,
    )
    delete_line = cook_client.delete(
        f"/api/builder/components/{fork_component_id}/recipes/{fork_recipe_id}/ingredients/{fork_line_id}",
        headers=cook_headers,
    )
    delete_recipe = cook_client.delete(
        f"/api/builder/components/{fork_component_id}/recipes/{fork_recipe_id}",
        headers=cook_headers,
    )

    composition_rv = admin_client.post(
        "/api/builder/compositions",
        json={"composition_id": "cook_api_plate", "composition_name": "Cook API Plate"},
        headers=admin_headers,
    )
    composition_id = ((composition_rv.get_json() or {}).get("composition") or {}).get("composition_id")
    composition_fork = cook_client.post(
        f"/api/builder/compositions/{composition_id}/fork",
        headers=cook_headers,
    )
    fork_composition_id = ((composition_fork.get_json() or {}).get("composition") or {}).get("composition_id")
    add_component = cook_client.post(
        f"/api/builder/compositions/{fork_composition_id}/components",
        json={"component_name": "Side Dish", "role": "side"},
        headers=cook_headers,
    )

    assert listed.status_code == 200
    assert detail.status_code == 200
    assert scaling.status_code == 200
    assert traits.status_code == 200
    assert (listed.get_json() or {}).get("count") == 1
    assert ((detail.get_json() or {}).get("recipe") or {}).get("recipe_id") == recipe_id
    assert fork_component.status_code == 201
    assert fork_recipe.status_code == 201
    assert fork_line.status_code == 201
    assert update_recipe.status_code == 200
    assert update_line.status_code == 200
    assert delete_line.status_code == 200
    assert delete_recipe.status_code == 200
    assert composition_fork.status_code == 201
    assert add_component.status_code == 200
    assert len(((add_component.get_json() or {}).get("composition") or {}).get("components") or []) >= 1


def test_cook_cannot_edit_org_original_or_other_cook_private_fork_via_api() -> None:
    fd, db_path = tempfile.mkstemp(prefix="builder_api_cook_blocked_", suffix=".db")
    os.close(fd)
    app = _app_with_builder_db(db_path)
    admin_client = app.test_client()
    cook_a_client = app.test_client()
    cook_b_client = app.test_client()
    admin_headers = _headers(role="admin", tenant_id=1, user_id=11)
    cook_a_headers = _headers(role="cook", tenant_id=1, user_id=42)
    cook_b_headers = _headers(role="cook", tenant_id=1, user_id=43)
    _seed_session(admin_client, role="admin", tenant_id=1, user_id=11)
    _seed_session(cook_a_client, role="cook", tenant_id=1, user_id=42)
    _seed_session(cook_b_client, role="cook", tenant_id=1, user_id=43)

    component_rv = admin_client.post(
        "/api/builder/components",
        json={"component_name": "Blocked API Soup"},
        headers=admin_headers,
    )
    component_id = ((component_rv.get_json() or {}).get("component") or {}).get("component_id")
    recipe_rv = admin_client.post(
        f"/api/builder/components/{component_id}/recipes",
        json={"recipe_name": "Base", "yield_portions": 4, "visibility": "private"},
        headers=admin_headers,
    )
    recipe_id = ((recipe_rv.get_json() or {}).get("recipe") or {}).get("recipe_id")

    composition_rv = admin_client.post(
        "/api/builder/compositions",
        json={"composition_id": "blocked_api_plate", "composition_name": "Blocked API Plate"},
        headers=admin_headers,
    )
    composition_id = ((composition_rv.get_json() or {}).get("composition") or {}).get("composition_id")

    fork_component = cook_a_client.post(f"/api/builder/components/{component_id}/fork", headers=cook_a_headers)
    fork_component_id = ((fork_component.get_json() or {}).get("component") or {}).get("component_id")
    fork_composition = cook_a_client.post(
        f"/api/builder/compositions/{composition_id}/fork",
        headers=cook_a_headers,
    )
    fork_composition_id = ((fork_composition.get_json() or {}).get("composition") or {}).get("composition_id")

    org_recipe_attempt = cook_a_client.post(
        f"/api/builder/components/{component_id}/recipes",
        json={"recipe_name": "Cook Blocked", "yield_portions": 2, "visibility": "private"},
        headers=cook_a_headers,
    )
    org_composition_attempt = cook_a_client.post(
        f"/api/builder/compositions/{composition_id}/components",
        json={"component_name": "Blocked Side"},
        headers=cook_a_headers,
    )
    own_recipe_attempt = cook_a_client.post(
        f"/api/builder/components/{fork_component_id}/recipes",
        json={"recipe_name": "Cook Allowed", "yield_portions": 2, "visibility": "private"},
        headers=cook_a_headers,
    )
    cross_recipe_attempt = cook_b_client.post(
        f"/api/builder/components/{fork_component_id}/recipes",
        json={"recipe_name": "Cook B Blocked", "yield_portions": 2, "visibility": "private"},
        headers=cook_b_headers,
    )
    own_composition_attempt = cook_a_client.post(
        f"/api/builder/compositions/{fork_composition_id}/components",
        json={"component_name": "Allowed Side"},
        headers=cook_a_headers,
    )
    cross_composition_attempt = cook_b_client.post(
        f"/api/builder/compositions/{fork_composition_id}/components",
        json={"component_name": "Blocked Side"},
        headers=cook_b_headers,
    )

    assert org_recipe_attempt.status_code == 400
    assert org_composition_attempt.status_code == 400
    assert own_recipe_attempt.status_code == 201
    assert cross_recipe_attempt.status_code == 400
    assert own_composition_attempt.status_code == 200
    assert cross_composition_attempt.status_code == 400
    assert recipe_id is not None
