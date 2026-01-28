from __future__ import annotations

from io import BytesIO

import docx  # type: ignore
from flask.testing import FlaskClient

ADMIN_HEADERS = {"X-User-Role": "admin", "X-Tenant-Id": "1"}


def _build_docx() -> bytes:
    d = docx.Document()
    p = d.add_paragraph()
    p.add_run("v. 8")
    d.add_paragraph("Söndag :")
    d.add_paragraph("Lunch: Biff Lindström med sås & potatis")
    d.add_paragraph("Dessert: Smördegsbakade äpplen med vaniljsås")
    d.add_paragraph("Kväll: Mannagrynspudding med sylt & grädde")
    d.add_paragraph("Med reservation för ändringar. Ni når oss på telefon: 0701486879")
    d.add_paragraph("Allt med röd text tillhandahåller ni.")
    bio = BytesIO()
    d.save(bio)
    return bio.getvalue()


def test_ui_menu_import_upload_uses_new_parser_and_ignores_footer(client_admin: FlaskClient):
    # Upload DOCX via real UI endpoint
    data = {
        "menu_file": (BytesIO(_build_docx()), "meny-vecka8.docx"),
    }
    r = client_admin.post(
        "/ui/admin/menu-import/upload",
        data=data,
        headers=ADMIN_HEADERS,
        content_type="multipart/form-data",
        follow_redirects=True,
    )
    assert r.status_code == 200
    # Verify week 8 Sunday lunch is correct and footer not present via API
    r_week = client_admin.get(
        "/menu/week",
        query_string={"week": 8, "year": 2026},
        headers=ADMIN_HEADERS,
    )
    assert r_week.status_code == 200
    payload = r_week.get_json()
    assert payload["ok"] is True
    days = payload["menu"]["days"]
    assert "Sun" in days, f"days keys: {list(days.keys())}"
    assert "Lunch" in days["Sun"], f"Sun keys: {list(days['Sun'].keys())}"
    assert "alt1" in days["Sun"]["Lunch"], f"Lunch variants: {list(days['Sun']['Lunch'].keys())}"
    dish = (days["Sun"]["Lunch"]["alt1"] or {}).get("dish_name")
    assert dish and "Biff Lindström" in dish
    assert "Med reservation" not in dish
    assert "Allt med röd text" not in dish
