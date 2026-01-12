import io


def build_complex_docx():
    try:
        from docx import Document  # type: ignore
    except Exception:
        return None
    doc = Document()
    doc.add_paragraph("Matsedel v. 8-9")
    # Single paragraph with multiple sections and adjacent day header after Kväll
    doc.add_paragraph(
        "Måndag: Lunch: Alt 1: Pasta Alt 2: Sallad Dessert: Kaka Kväll: Alt 1: Soppa Tisdag: Lunch: Alt 1: Fisk Kväll: Alt 1: Gryta"
    )
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def test_docx_parser_handles_combined_paragraphs_and_week_range():
    buf = build_complex_docx()
    if buf is None:
        import pytest
        pytest.skip("python-docx not installed")
    from core.menu_docx_parser import parse_menu_docx
    out = parse_menu_docx(buf)
    assert sorted(out["weeks"].keys()) == [8, 9]
    d1 = out["weeks"][8]["days"][1]
    assert d1["lunch"]["alt1_text"] == "Pasta"
    assert d1["lunch"]["alt2_text"] == "Sallad"
    assert d1["lunch"]["dessert"] == "Kaka"
    assert d1["dinner"]["alt1_text"] == "Soppa"
    d2 = out["weeks"][9]["days"][2]
    assert d2["lunch"]["alt1_text"] == "Fisk"
    assert d2["dinner"]["alt1_text"] == "Gryta"
