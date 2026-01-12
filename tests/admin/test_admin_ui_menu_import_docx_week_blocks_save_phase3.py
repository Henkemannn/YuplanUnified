import io
import json


def build_week_blocks_docx():
    try:
        from docx import Document  # type: ignore
    except Exception:
        return None
    doc = Document()
    doc.add_paragraph("Matsedel")
    # Explicit headings on their own lines
    doc.add_paragraph("v. 8")
    doc.add_paragraph("Måndag: Lunch: Alt 1: Pasta Dessert: Kaka Kväll: Alt 1: Soppa")
    doc.add_paragraph("Tisdag: Lunch: Alt 1: Kyckling Kväll: Alt 1: Gryta")
    doc.add_paragraph("v. 9")
    doc.add_paragraph("Måndag: Lunch: Alt 1: Fisk Kväll: Alt 1: Lasagne")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def test_docx_week_blocks_parsed_and_persisted_per_week(client_admin):
    buf = build_week_blocks_docx()
    if buf is None:
        import pytest
        pytest.skip("python-docx not installed")
    # Parse to weeks_json
    from core.menu_docx_parser import parse_menu_docx
    parsed = parse_menu_docx(buf)
    weeks = parsed.get("weeks") or {}
    assert 8 in weeks and 9 in weeks
    # Persist via save endpoint
    with client_admin.session_transaction() as sess:
        sess["CSRF_TOKEN"] = "tok"
        sess["site_id"] = "site-x"
    year = __import__("datetime").date.today().isocalendar()[0]
    resp = client_admin.post(
        "/ui/admin/menu-import/preview/save",
        data={
            "year": str(year),
            "weeks_json": json.dumps(weeks),
            "csrf_token": "tok",
        },
        headers={"X-User-Role": "admin", "X-Tenant-Id": "1", "X-Site-Id": "site-x", "X-CSRF-Token": "tok"},
    )
    assert resp.status_code in (200, 302)
    # Verify week 8 Monday
    r8 = client_admin.get(
        f"/api/menu/day?year={year}&week=8&day=1",
        headers={"X-User-Role": "admin", "X-Tenant-Id": "1", "X-Site-Id": "site-x"},
    )
    assert r8.status_code == 200
    d8 = r8.get_json()
    assert d8["lunch"]["alt1_text"] == "Pasta"
    assert d8["dinner"]["alt1_text"] == "Soppa"
    # Verify week 9 Monday differs
    r9 = client_admin.get(
        f"/api/menu/day?year={year}&week=9&day=1",
        headers={"X-User-Role": "admin", "X-Tenant-Id": "1", "X-Site-Id": "site-x"},
    )
    assert r9.status_code == 200
    d9 = r9.get_json()
    assert d9["lunch"]["alt1_text"] == "Fisk"
    assert d9["dinner"]["alt1_text"] == "Lasagne"
