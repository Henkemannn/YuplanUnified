import pytest
from flask.testing import FlaskClient
from io import BytesIO

try:
    import docx  # type: ignore
except Exception:  # pragma: no cover
    docx = None  # type: ignore

ADMIN_HEADERS = {"X-User-Role": "admin", "X-Tenant-Id": "1"}


def _build_minimal_docx() -> bytes:
    if docx is None:  # pragma: no cover
        pytest.skip("python-docx not installed")
    d = docx.Document()  # type: ignore[attr-defined]
    d.add_paragraph("Vecka 15")
    d.add_paragraph("Måndag")
    d.add_paragraph("Alt 1: Köttbullar")
    d.add_paragraph("Alt 2: Fiskgratäng")
    d.add_paragraph("Dessert: Glass")
    buf = BytesIO()
    d.save(buf)  # type: ignore[attr-defined]
    return buf.getvalue()


@pytest.mark.skipif(docx is None, reason="python-docx not installed")
def test_docx_upload_returns_success(app_session) -> None:
    client: FlaskClient = app_session.test_client()
    data = {"menu_file": (BytesIO(_build_minimal_docx()), "test_menu.docx")}
    resp = client.post(
        "/ui/admin/menu-import/upload",
        data=data,
        content_type="multipart/form-data",
        headers=ADMIN_HEADERS,
        follow_redirects=True,
    )
    assert resp.status_code == 200
    # Verify persistence by querying menus for week 15
    from sqlalchemy import text
    from core.db import get_session
    db = get_session()
    try:
        row = db.execute(text("SELECT COUNT(1) FROM menus WHERE tenant_id=1 AND week=15"), {}).fetchone()
        assert row and int(row[0]) >= 1
    finally:
        db.close()
