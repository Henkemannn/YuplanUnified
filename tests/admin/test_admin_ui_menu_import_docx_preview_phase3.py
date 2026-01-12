import io


def build_simple_docx():
    try:
        from docx import Document  # type: ignore
    except Exception:
        return None
    doc = Document()
    doc.add_paragraph("Matsedel v. 8-9")
    doc.add_paragraph("Måndag")
    doc.add_paragraph("Lunch")
    doc.add_paragraph("Alt 1: Pasta")
    doc.add_paragraph("Alt 2: Sallad")
    doc.add_paragraph("Dessert: Kaka")
    doc.add_paragraph("Kväll")
    doc.add_paragraph("Alt 1: Soppa")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def test_docx_upload_shows_preview(client_admin):
    buf = build_simple_docx()
    if buf is None:
        import pytest
        pytest.skip("python-docx not installed")
    with client_admin.session_transaction() as sess:
        sess["CSRF_TOKEN"] = "tok-docx"
    data = {"menu_file": (buf, "menu.docx"), "csrf_token": "tok-docx"}
    resp = client_admin.post(
        "/ui/admin/menu-import/upload",
        data=data,
        content_type="multipart/form-data",
        headers={"X-User-Role": "admin", "X-Tenant-Id": "1", "X-CSRF-Token": "tok-docx"},
    )
    assert resp.status_code == 200
    html = resp.get_data(as_text=True)
    assert "Förhandsgranskning" in html
    assert "v. 8–9" in html or "v. 8-9" in html
    assert "Pasta" in html
