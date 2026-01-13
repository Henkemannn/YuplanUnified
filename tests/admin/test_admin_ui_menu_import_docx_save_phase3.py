import io
import json


def build_simple_docx():
    try:
        from docx import Document  # type: ignore
    except Exception:
        return None
    doc = Document()
    doc.add_paragraph("Matsedel v. 8-9")
    doc.add_paragraph("Måndag")
    doc.add_paragraph("Lunch")
    doc.add_paragraph("Alt 1: Pasta")
    doc.add_paragraph("Alt 2: Sallad")
    doc.add_paragraph("Dessert: Kaka")
    doc.add_paragraph("Kväll")
    doc.add_paragraph("Alt 1: Soppa")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def test_docx_preview_and_save_persists_menu_and_api(client_admin):
    buf = build_simple_docx()
    if buf is None:
        import pytest
        pytest.skip("python-docx not installed")
    # Preview upload
    with client_admin.session_transaction() as sess:
        sess["CSRF_TOKEN"] = "tok-docx"
        sess["site_id"] = "site-x"
    resp = client_admin.post(
        "/ui/admin/menu-import/upload",
        data={"menu_file": (buf, "menu.docx"), "csrf_token": "tok-docx"},
        content_type="multipart/form-data",
        headers={"X-User-Role": "admin", "X-Tenant-Id": "1", "X-Site-Id": "site-x", "X-CSRF-Token": "tok-docx"},
    )
    assert resp.status_code == 200
    html = resp.get_data(as_text=True)
    assert "Förhandsgranskning" in html
    # Prepare weeks json matching the DOCX we built (distinct weeks)
    weeks = {
        8: {
            "days": {
                1: {"lunch": {"alt1_text": "Pasta", "alt2_text": "Sallad", "dessert": "Kaka"},
                    "dinner": {"alt1_text": "Soppa", "alt2_text": "", "dessert": ""}}
            }
        },
        9: {
            "days": {
                1: {"lunch": {"alt1_text": "Fisk", "alt2_text": "", "dessert": ""},
                    "dinner": {"alt1_text": "Gryta", "alt2_text": "", "dessert": ""}}
            }
        }
    }
    year = __import__("datetime").date.today().isocalendar()[0]
    # Save preview to DB
    resp2 = client_admin.post(
        "/ui/admin/menu-import/preview/save",
        data={
            "year": str(year),
            "weeks_json": json.dumps(weeks),
            "csrf_token": "tok-docx",
        },
        headers={"X-User-Role": "admin", "X-Tenant-Id": "1", "X-Site-Id": "site-x", "X-CSRF-Token": "tok-docx"},
    )
    assert resp2.status_code in (302, 200)
    # Verify DB rows exist via API (week 8)
    resp3 = client_admin.get(
        f"/api/menu/day?year={year}&week=8&day=1",
        headers={"X-User-Role": "admin", "X-Tenant-Id": "1", "X-Site-Id": "site-x"},
    )
    assert resp3.status_code == 200
    data = resp3.get_json()
    assert data["lunch"]["alt1_text"] == "Pasta"
    assert data["lunch"]["alt2_text"] == "Sallad"
    assert data["lunch"]["dessert"] == "Kaka"
    assert data["dinner"]["alt1_text"] == "Soppa"
    # Verify week 9 persists distinct values
    resp4 = client_admin.get(
        f"/api/menu/day?year={year}&week=9&day=1",
        headers={"X-User-Role": "admin", "X-Tenant-Id": "1", "X-Site-Id": "site-x"},
    )
    assert resp4.status_code == 200
    data9 = resp4.get_json()
    assert data9["lunch"]["alt1_text"] == "Fisk"
