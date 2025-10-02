# Copyright (c) 2025 Henrik Jonsson, Yuplan. All rights reserved.
# Proprietary and confidential. Unauthorized copying, distribution or use is strictly prohibited.
# Menyimport-funktionalitet för Yuplan
import re

from docx import Document


class MenyImporter:
    """Hanterar import av menyer från Word-dokument"""
    
    def __init__(self):
        pass
        
    def parse_word_document(self, file_path: str) -> dict:
        """Läser Word-dokument och extraherar menyinformation för flera veckor"""
        try:
            doc = Document(file_path)
            content = []
            # Extrahera all text från dokumentet, splitta alltid på \n
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    for line in paragraph.text.strip().split("\n"):
                        if line.strip():
                            content.append(line.strip())
            # Extrahera text från tabeller också
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            for line in cell.text.strip().split("\n"):
                                if line.strip():
                                    content.append(line.strip())
            print("DEBUG: Inläst content från docx:")
            for row in content:
                print(repr(row))
            return self.analyze_multiweek_menu_content(content)
        except Exception as e:
            return {"error": f"Kunde inte läsa Word-dokument: {str(e)}"}
    def split_content_by_weeks(self, content: list[str]) -> list[tuple[int, list[str]]]:
        """Splitta content till [(week_number, [lines...]), ...]"""
        week_patterns = [
            r"vecka\s*[:\.]?\s*(\d+)",
            r"v[\.:\s]*(\d+)",
            r"v\.\s*(\d+)",
            r"week\s+(\d+)",
        ]
        week_sections = []
        current_week = None
        current_lines = []
        for line in content:
            found = False
            for pattern in week_patterns:
                m = re.search(pattern, line.lower())
                if m:
                    week_num = int(m.group(1))
                    if 1 <= week_num <= 52:
                        if current_week is not None and current_lines:
                            week_sections.append((current_week, current_lines))
                        current_week = week_num
                        current_lines = []
                        found = True
                        break
            if not found:
                if current_week is not None:
                    current_lines.append(line)
        if current_week is not None and current_lines:
            week_sections.append((current_week, current_lines))
        return week_sections

    def analyze_multiweek_menu_content(self, content: list[str]) -> dict:
        """Analysera menyinnehåll för flera veckor"""
        week_sections = self.split_content_by_weeks(content)
        if not week_sections:
            return {
                "success": False,
                "error": "Kunde inte hitta något veckonummer i dokumentet. Ange vecka manuellt."
            }
        all_weeks = {}
        for week_num, week_lines in week_sections:
            daily_menus = self.extract_daily_menus(week_lines)
            all_weeks[week_num] = daily_menus
        return {
            "weeks": all_weeks,
            "success": True,
            "message": f"Framgångsrikt importerade menyer för veckor: {', '.join(str(w) for w in all_weeks)}"
        }
    
    def detect_week_number(self, content: list[str]) -> int | None:
        """Identifierar veckonummer från menytext, t.ex. 'V.1', 'v1', 'vecka 1'. Returnerar None om inget hittas."""
        week_patterns = [
            r"vecka\s*[:\.]?\s*(\d+)",
            r"v[\.:\s]*(\d+)",
            r"v\.\s*(\d+)",
            r"week\s+(\d+)",
        ]
        for line in content:
            line_lower = line.lower()
            for pattern in week_patterns:
                match = re.search(pattern, line_lower)
                if match:
                    week_num = int(match.group(1))
                    if 1 <= week_num <= 52:
                        return week_num
        return None
    
    def extract_daily_menus(self, content: list[str]) -> dict[str, dict[str, str]]:
        days = ["måndag", "tisdag", "onsdag", "torsdag", "fredag", "lördag", "söndag"]
        day_abbrev = ["mån", "tis", "ons", "tor", "fre", "lör", "sön"]
        menus = {day: {} for day in days}
        current_day = None
        current_menu = {}
        meny_types = [
            ("Alt1", re.compile(r"^(Alt\s*1:|Alternativ\s*1:|Lunch:)", re.IGNORECASE)),
            ("Alt2", re.compile(r"^(Alt\s*2:|Alternativ\s*2:)", re.IGNORECASE)),
            ("Dessert", re.compile(r"^(Dessert:)", re.IGNORECASE)),
            ("Kväll", re.compile(r"^(Kväll:)", re.IGNORECASE)),
        ]
        def normalize_day_line(s):
            s = s.replace("\u00a0", " ")
            s = s.replace("\u200b", "")
            s = s.replace("\uff1a", ":")
            s = s.replace("\u0589", ":")
            s = s.replace("\u2236", ":")
            s = s.replace("\u02d0", ":")
            s = s.replace("\u05c3", ":")
            s = s.replace("\ufe55", ":")
            s = s.replace("\u003a", ":")
            s = re.sub(r"\s+", " ", s)
            return s.strip().lower()
        norm_days = [normalize_day_line(day) for day in days]
        norm_abbrs = [normalize_day_line(abbr) for abbr in day_abbrev]
        for idx, line in enumerate(content):
            print(f"PARSER DEBUG: Rad {idx}: {repr(line)}")
            for subline in line.split("\n"):
                line_stripped = subline.rstrip("\r\n")
                line_clean = line_stripped.strip()
                norm_line = normalize_day_line(line_clean)
                print(f"  Subrad: {repr(line_clean)} | Normaliserad: {repr(norm_line)}")
                day_found = None
                day_match = None
                for i, norm_day in enumerate(norm_days):
                    m = re.match(rf"^{re.escape(norm_day)}\s*[:.\-–—]?\s*(.*)$", norm_line)
                    if not m:
                        m = re.match(rf"^{re.escape(norm_abbrs[i])}\s*[:.\-–—]?\s*(.*)$", norm_line)
                    if m:
                        day_found = days[i]
                        day_match = m
                        print(f"    MATCH DAG: {days[i]} på subrad {repr(line_clean)}")
                        break
                if day_found:
                    if current_day is not None:
                        print(f"    SPARAR DAG: {current_day} -> {current_menu}")
                        menus[current_day] = current_menu if current_menu else {}
                    current_day = day_found
                    current_menu = {}
                    rest = day_match.group(1).strip() if day_match else ""
                    if rest:
                        line_clean2 = rest
                        found_type = None
                        for menytyp, regex in meny_types:
                            m2 = regex.match(line_clean2)
                            if m2:
                                found_type = menytyp
                                value = line_clean2[m2.end():].strip()
                                if value:
                                    current_menu[menytyp] = value
                                break
                        if not found_type:
                            if "Alt1" not in current_menu:
                                current_menu["Alt1"] = line_clean2
                            elif "Kväll" not in current_menu:
                                current_menu["Kväll"] = line_clean2
                            elif "Dessert" not in current_menu:
                                current_menu["Dessert"] = line_clean2
                    continue
                if current_day is not None and line_clean:
                    found_type = None
                    for menytyp, regex in meny_types:
                        m = regex.match(line_clean)
                        if m:
                            found_type = menytyp
                            value = line_clean[m.end():].strip()
                            if value:
                                current_menu[menytyp] = value
                            break
                    if not found_type:
                        if "Alt1" not in current_menu:
                            current_menu["Alt1"] = line_clean
                        elif "Kväll" not in current_menu:
                            current_menu["Kväll"] = line_clean
                        elif "Dessert" not in current_menu:
                            current_menu["Dessert"] = line_clean
        if current_day is not None:
            menus[current_day] = current_menu if current_menu else {}
        return {day: menus.get(day, {}) for day in days}
    
    def analyze_menu_content(self, content: list[str]) -> dict:
        """Huvudfunktion som analyserar menyinnehåll"""
        week_number = self.detect_week_number(content)
        if week_number is None:
            return {
                "success": False,
                "error": "Kunde inte hitta något veckonummer i dokumentet. Ange vecka manuellt."
            }
        daily_menus = self.extract_daily_menus(content)
        return {
            "week": week_number,
            "menus": daily_menus,
            "success": True,
            "message": f"Framgångsrikt importerade menyer för vecka {week_number}"
        }

# Hjälpfunktioner för Flask-integration
def save_imported_menus(conn, week: int, menus: dict[str, dict[str, str]]):
    """Sparar importerade menyer till databasen"""
    try:
        # Ta bort befintliga menyer för veckan
        conn.execute("DELETE FROM veckomeny WHERE vecka = ?", (week,))

        # Map full day names to abbreviations
        day_map = {
            "måndag": "Mån",
            "tisdag": "Tis",
            "onsdag": "Ons",
            "torsdag": "Tors",
            "fredag": "Fre",
            "lördag": "Lör",
            "söndag": "Sön",
        }

        # Lägg till nya menyer
        # Tillåt även att dagarna i menus är redan förkortade (Mån, Tis, ...)
        reverse_day_map = {abbr: abbr for abbr in day_map.values()}
        for dag, menu_data in menus.items():
            dag_lc = dag.lower()
            # Om dag redan är förkortad, använd den, annars mappa
            dag_abbr = day_map.get(dag_lc, reverse_day_map.get(dag, dag))
            if menu_data.get("Alt1"):
                conn.execute(
                    "INSERT INTO veckomeny (vecka, dag, alt_typ, menytext) VALUES (?, ?, ?, ?)",
                    (week, dag_abbr, "Alt1", menu_data["Alt1"].strip())
                )
            if menu_data.get("Alt2"):
                conn.execute(
                    "INSERT INTO veckomeny (vecka, dag, alt_typ, menytext) VALUES (?, ?, ?, ?)",
                    (week, dag_abbr, "Alt2", menu_data["Alt2"].strip())
                )
            if menu_data.get("Dessert"):
                conn.execute(
                    "INSERT INTO veckomeny (vecka, dag, alt_typ, menytext) VALUES (?, ?, ?, ?)",
                    (week, dag_abbr, "Dessert", menu_data["Dessert"].strip())
                )
            if menu_data.get("Kväll"):
                conn.execute(
                    "INSERT INTO veckomeny (vecka, dag, alt_typ, menytext) VALUES (?, ?, ?, ?)",
                    (week, dag_abbr, "Kväll", menu_data["Kväll"].strip())
                )

        conn.commit()
        return True
    except Exception as e:
        conn.rollback()
        print(f"Fel vid sparande av menyer: {e}")
        return False

def process_uploaded_menu(file_path: str, db_connection):
    """Processerar en uppladdad menyfil"""
    importer = MenyImporter()
    result = importer.parse_word_document(file_path)
    
    if result.get("success"):
        success = save_imported_menus(
            db_connection, 
            result["week"], 
            result["menus"]
        )
        if success:
            result["saved"] = True
        else:
            result["error"] = "Kunde inte spara till databas"
    
    return result
