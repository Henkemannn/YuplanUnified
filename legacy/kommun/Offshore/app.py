"""


Clean Flask app exposing turnus endpoints and admin/superuser panels.
Implements a simple 2-week base schedule builder and apply-to-6-cooks flow.
"""
from __future__ import annotations

import csv
import io
import json
import os
import secrets
import sqlite3
import time
from datetime import datetime, timedelta
from pathlib import Path
from tempfile import NamedTemporaryFile

import rotation
import rotation_simple as rs
from flask import (
    Flask,
    flash,
    jsonify,
    redirect,
    render_template,
    request,
    send_file,
    session,
    url_for,
)
from werkzeug.utils import secure_filename

# Lazy import blueprint after app init later (avoid circular at top) - we'll register after app object creation.
try:
    from docx import Document  # python-docx for DOCX generation
    try:
        from docx.shared import Cm, Inches, Pt
    except Exception:
        Inches = Pt = Cm = None
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        OxmlElement = None
        qn = None
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception:
        WD_ALIGN_PARAGRAPH = None
except Exception:
    Document = None

app = Flask(__name__)
# Force template auto-reload even when debug=False so UI edits (dashboard/planning hub) show immediately after restart
app.config.setdefault("TEMPLATES_AUTO_RELOAD", True)
try:
    app.jinja_env.auto_reload = True
except Exception:
    pass
app.secret_key = os.environ.get("FLASK_SECRET_KEY", "dev-secret-key")
app.config["DEMO_MODE"] = os.environ.get("DEMO_MODE", "0") == "1"

def is_demo():
    return bool(app.config.get("DEMO_MODE"))
""" NOTE:
Previously the waste/service blueprint was imported & registered here before get_db was defined.
Because waste.py did `from app import get_db` at module import time, this caused an AttributeError
while app.py was still executing (get_db not yet defined). The broad try/except swallowed it and
the blueprint never registered, leading to missing /service/* endpoints in tests.
We now defer blueprint import/registration until AFTER get_db is defined (see below) so routes load.
"""

# --- DB helper ---
def get_db():
    db_path = None
    try:
        dbp = getattr(rotation, "DB_PATH", None)
        db_path = dbp.as_posix() if dbp is not None else None
    except Exception:
        db_path = None
    if not db_path:
        db_path = os.path.join(os.getcwd(), "app.db")
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    try:
        conn.execute("PRAGMA foreign_keys = ON;")
    except Exception:
        pass
    return conn

# --- Demo synthetic seed ---
def _demo_seed_week(db: sqlite3.Connection):
    if not is_demo():
        return
    try:
        # Check minimal presence
        cur = db.execute("SELECT COUNT(*) FROM dish_catalog")
        if cur.fetchone()[0] > 0:
            return  # already seeded
    except Exception:
        return
    try:
        today = datetime.utcnow().date()
        rig_id = 1
        # Ensure rig and demo user exist (light)
        try:
            db.execute("CREATE TABLE IF NOT EXISTS rigs(id INTEGER PRIMARY KEY AUTOINCREMENT, name TEXT)")
            db.execute("CREATE TABLE IF NOT EXISTS users(id INTEGER PRIMARY KEY AUTOINCREMENT, name TEXT, rig_id INTEGER, role TEXT, email TEXT)")
            if db.execute("SELECT COUNT(*) FROM rigs").fetchone()[0] == 0:
                db.execute("INSERT INTO rigs(name) VALUES('Demo Rig')")
            if db.execute("SELECT COUNT(*) FROM users").fetchone()[0] == 0:
                db.execute("INSERT INTO users(name, rig_id, role, email) VALUES('Demo Bruker',1,'user','demo@example.com')")
        except Exception:
            pass
        dishes = [
            ("Fisk Grateng","fisk"),
            ("Laks Ovnsbakt","fisk"),
            ("Biff Stroganoff","kott"),
            ("Kylling Gryte","kott"),
            ("Grønnsakssuppe","soppa"),
            ("Tomatsuppe","soppa"),
            ("Quinoa Salat","extra"),
            ("Ovnsbakt Blomkål","extra"),
        ]
        for name, cat in dishes:
            slug = _slugify_dish(name)
            try:
                db.execute("INSERT OR IGNORE INTO dish_catalog(rig_id, slug, name, first_seen_date) VALUES(?,?,?,?)", (rig_id, slug, name, today.isoformat()))
            except Exception:
                pass
        # Map a simple 5-day rotation (menu_index=1)
        try:
            for weekday in range(5):  # Mon-Fri
                for meal in ("lunsj","middag"):
                    # Pick deterministic dish per category
                    for cat in ("soppa","fisk","kott","extra"):
                        row = db.execute("SELECT id FROM dish_catalog WHERE rig_id=? AND LOWER(name) LIKE ? LIMIT 1", (rig_id, f"%{cat[0:3]}%"))
                        drow = row.fetchone()
                        if not drow:
                            drow = db.execute("SELECT id FROM dish_catalog WHERE rig_id=? LIMIT 1", (rig_id,)).fetchone()
                        if drow:
                            try:
                                db.execute("INSERT OR IGNORE INTO menu_dish_map(rig_id, menu_index, weekday, meal, category, dish_id) VALUES(?,?,?,?,?,?)", (rig_id, 1, weekday, meal, cat, drow[0]))
                            except Exception:
                                pass
        except Exception:
            pass
        # Seed service_metrics over last 7 days for two meals
        try:
            db.execute("CREATE TABLE IF NOT EXISTS service_metrics (id INTEGER PRIMARY KEY AUTOINCREMENT, rig_id INTEGER NOT NULL, date TEXT NOT NULL, meal TEXT NOT NULL, dish_id INTEGER, category TEXT, guest_count INTEGER NOT NULL, produced_qty_kg REAL, served_qty_kg REAL, leftover_qty_kg REAL, served_g_per_guest REAL, notes TEXT, created_at TEXT NOT NULL DEFAULT (datetime('now')))")
        except Exception:
            pass
        for day_offset in range(7):
            d = today - timedelta(days=day_offset)
            guest_base = 60 - day_offset*2
            for meal in ("lunsj","middag"):
                for cat, base_portion in (("fisk",0.145),("kott",0.165),("soppa",0.20),("extra",0.11)):
                    produced = round(base_portion * guest_base * 1.05, 2)
                    served = round(base_portion * guest_base, 2)
                    leftover = max(0.0, round(produced - served, 2))
                    served_g_per_guest = (served * 1000.0) / guest_base
                    try:
                        db.execute("INSERT INTO service_metrics(rig_id,date,meal,category,guest_count,produced_qty_kg,served_qty_kg,leftover_qty_kg,served_g_per_guest) VALUES(?,?,?,?,?,?,?,?,?)", (rig_id, d.isoformat(), meal, cat, guest_base, produced, served, leftover, served_g_per_guest))
                    except Exception:
                        pass
        db.commit()
    except Exception:
        pass

@app.before_request
def _demo_guard():
    if not is_demo():
        return
    # Seed once per process early (cheap check)
    try:
        db = get_db()
        _demo_seed_week(db)
    except Exception:
        pass
    # Anonymize session derived display name
    if session.get("user_name") and session["user_name"] != "Demo Bruker":
        session["user_name"] = "Demo Bruker"
    # Block destructive endpoints (placeholder names, adjust as real endpoints exist)
    destructive_paths = ("/reset", "/admin/delete", "/superuser", "/run-migration")
    if request.path.startswith(destructive_paths):
        return jsonify({"ok": False, "error": "Demo-läge: blokkert", "demo": True}), 403
    # Fake write responses for unsafe modifications if not explicitly allowed
    if request.method in ("POST","DELETE","PUT","PATCH"):
        # Allow login attempts to proceed; block others except service/log which we fake-store
        if request.path.endswith("/service/log"):
            return jsonify({"ok": True, "demo": True, "note": "Ingen reell lagring"}), 200
        # Allow /login /register
        if "/login" in request.path or "/register" in request.path:
            return None
        return jsonify({"ok": False, "error": "Demo-läge: skriveblokk", "demo": True}), 403

@app.context_processor
def inject_demo_flag():
    return {"DEMO_MODE": is_demo()}

"""Marketing landing + health.

We previously added a defensive fallback /landing and an early root_router.
Those created duplicate endpoints (/landing -> landing_page_fallback & landing_page,
and '/' -> root_router & root). This confused Flask's rule map ordering and
made debugging 404 reports harder. We collapse to a single /landing and a
single '/' route (the canonical one at bottom handling auth redirects).
"""

@app.route("/ping")
def ping():
    return jsonify({"ok": True, "msg": "alive"})

# --- Deferred blueprint registration to avoid circular import timing issues ---
try:
    from waste import waste_bp  # waste.py can now safely import get_db if needed
    if not any(bp.url_prefix == "/service" for bp in app.blueprints.values()):  # minimal idempotency guard
        app.register_blueprint(waste_bp)
except Exception as _bp_err:
    # Surface a hint during testing/debug instead of silent pass
    try:
        print(f"[init] Failed registering waste blueprint: {_bp_err}")
    except Exception:
        pass

@app.route("/landing")
def landing_page():
    if session.get("user_id"):
        try:
            return redirect(url_for("dashboard"))
        except Exception:
            pass
    return render_template("landing.html")

@app.route("/coming-soon")
def coming_soon():
    """Public marketing placeholder for combined Offshore/Kommun offering."""
    return render_template("coming_soon.html")

@app.route("/_routes")
def _list_routes():
    out = []
    for r in app.url_map.iter_rules():
        r_methods = getattr(r, "methods", None)
        if r_methods:
            try:
                methods_list = sorted([m for m in r_methods if isinstance(m, str)])
            except Exception:
                methods_list = []
        else:
            methods_list = []
        out.append({"rule": str(r), "endpoint": r.endpoint, "methods": methods_list})
    return jsonify({"ok": True, "count": len(out), "routes": out})

@app.route("/clean")
def clean_redirect():
    # Helper for screenshot / presentation mode (adds ?clean=1)
    if session.get("user_id"):
        return redirect(url_for("dashboard", clean=1))
    return redirect(url_for("landing_page", clean=1))


# --- Lightweight table ensure helpers (incremental, avoids large refactor now) ---
def _ensure_prep_tables(db: sqlite3.Connection):
    """Ensure private prep task table exists.

    We deliberately keep this small & idempotent until a broader init consolidation step.
    """
    try:
        db.execute(
            """
            CREATE TABLE IF NOT EXISTS prep_tasks_private (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                rig_id INTEGER NOT NULL,
                date TEXT NOT NULL,            -- YYYY-MM-DD
                meal TEXT NOT NULL,            -- 'lunsj' | 'middag'
                category TEXT NOT NULL,        -- 'soppa' | 'fisk' | 'kott' | 'extra'
                text TEXT NOT NULL,
                done INTEGER DEFAULT 0,
                dish_id INTEGER,               -- kobling til dish_catalog (kan være NULL inntil backfill)
                created_at TEXT DEFAULT CURRENT_TIMESTAMP
            );
            """
        )
        # Helpful indices for lookups
        db.execute("CREATE INDEX IF NOT EXISTS idx_prep_tasks_private_lookup ON prep_tasks_private(user_id, rig_id, date, meal, category);")
        # Frysplock table (separate to allow qty/unit semantics)
        db.execute(
            """
            CREATE TABLE IF NOT EXISTS frys_items_private (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                rig_id INTEGER NOT NULL,
                date TEXT NOT NULL,
                meal TEXT NOT NULL,
                category TEXT NOT NULL,
                item_name TEXT NOT NULL,
                qty REAL,
                unit TEXT,
                done INTEGER DEFAULT 0,
                dish_id INTEGER,               -- kobling til dish_catalog
                created_at TEXT DEFAULT CURRENT_TIMESTAMP
            );
            """
        )
        db.execute("CREATE INDEX IF NOT EXISTS idx_frys_items_private_lookup ON frys_items_private(user_id, rig_id, date, meal, category);")
        # Try add columns if table existed tidligere uten dish_id
        try:
            db.execute("ALTER TABLE prep_tasks_private ADD COLUMN dish_id INTEGER")
        except Exception:
            pass
        try:
            db.execute("ALTER TABLE frys_items_private ADD COLUMN dish_id INTEGER")
        except Exception:
            pass
    except Exception:
        pass


# --- Canonical dish mapping helpers ---
def _slugify_dish(name: str) -> str:
    if not name:
        return ""
    import re as _re
    import unicodedata
    n = unicodedata.normalize("NFKD", name).encode("ascii", "ignore").decode("ascii")
    n = n.lower()
    n = _re.sub(r"[^a-z0-9]+", "-", n).strip("-")
    return n[:120]


def _ensure_dish_catalog(db: sqlite3.Connection):
    try:
        db.execute(
            """
            CREATE TABLE IF NOT EXISTS dish_catalog (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              rig_id INTEGER NOT NULL,
              slug TEXT NOT NULL,
              name TEXT NOT NULL,
              first_seen_date TEXT,
              active INTEGER DEFAULT 1,
                            recipe_id INTEGER, -- optional link to recipes.id
              UNIQUE(rig_id, slug)
            )
            """
        )
        db.execute(
            """
            CREATE TABLE IF NOT EXISTS menu_dish_map (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              rig_id INTEGER NOT NULL,
              menu_index INTEGER NOT NULL,
              weekday INTEGER NOT NULL,     -- 0=Mon
              meal TEXT NOT NULL,           -- lunsj|middag
              category TEXT NOT NULL,       -- soppa|fisk|kott|extra
              dish_id INTEGER NOT NULL,
              UNIQUE(rig_id, menu_index, weekday, meal, category)
            )
            """
        )
        # Helpful index for lookups by date context
        db.execute("CREATE INDEX IF NOT EXISTS idx_menu_dish_lookup ON menu_dish_map(rig_id, menu_index, weekday, meal, category)")
        # Backfill recipe_id column if newly added (SQLite can't add column constraints easily; ignore errors)
        try:
            db.execute("ALTER TABLE dish_catalog ADD COLUMN recipe_id INTEGER")
        except Exception:
            pass
    except Exception:
        pass


def _ensure_recipes(db: sqlite3.Connection):
    """Ensure core recipe tables exist (simple MVP)."""
    try:
        db.execute(
            """
            CREATE TABLE IF NOT EXISTS recipes (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              rig_id INTEGER NOT NULL,
              title TEXT NOT NULL,
              source_url TEXT,
              raw_text TEXT,             -- imported freeform blob
              notes TEXT,
              base_portions INTEGER,
              categories TEXT,           -- comma separated
              method_type TEXT,          -- generic, base, garnish, technique, component
              created_at TEXT DEFAULT CURRENT_TIMESTAMP,
              updated_at TEXT DEFAULT CURRENT_TIMESTAMP
            )
            """
        )
        db.execute("CREATE INDEX IF NOT EXISTS idx_recipes_rig ON recipes(rig_id, title)")
        # Backfill newly added columns if table existed earlier
        try:
            db.execute("ALTER TABLE recipes ADD COLUMN categories TEXT")
        except Exception:
            pass
        try:
            db.execute("ALTER TABLE recipes ADD COLUMN method_type TEXT")
        except Exception:
            pass
    except Exception:
        pass


def _normalize_categories(raw: str) -> str:
    if not raw:
        return ""
    parts = [p.strip().lower() for p in raw.replace(";", ",").split(",") if p.strip()]
    # De-duplicate while preserving order
    seen = set()
    ordered = []
    for p in parts:
        if p not in seen:
            seen.add(p)
            ordered.append(p)
    return ",".join(ordered[:24])  # safety cap


@app.get("/recipes")
def recipe_list():
    if not session.get("user_id"):
        return redirect(url_for("login"))
    db = get_db(); _ensure_recipes(db)
    user = db.execute("SELECT id, rig_id FROM users WHERE id=?", (session["user_id"],)).fetchone()
    if not user or not user["rig_id"]:
        flash("Ingen rigg tilknyttet.", "warning"); return redirect(url_for("dashboard"))
    rig_id = user["rig_id"]
    q = (request.args.get("q") or "").strip()
    filter_missing = request.args.get("filter") == "missing"
    base_sql = "SELECT r.*, (SELECT COUNT(*) FROM dish_catalog d WHERE d.rig_id=r.rig_id AND d.recipe_id=r.id) AS link_count FROM recipes r WHERE r.rig_id=?"
    params = [rig_id]
    if q:
        base_sql += " AND r.title LIKE ?"
        params.append(f"%{q}%")
    base_sql += " ORDER BY r.updated_at DESC LIMIT 400"
    rows = db.execute(base_sql, params).fetchall()
    recs = [dict(id=r["id"], title=r["title"], link_count=r["link_count"], updated_at=r["updated_at"], categories=(r["categories"] or "")) for r in rows]
    # If filter=missing, remove those with links
    if filter_missing:
        recs = [r for r in recs if not r["link_count"]]
    return render_template("recipe_list.html", recipes=recs, q=q, filter_missing=filter_missing)


@app.get("/recipes/new")
def recipe_new_form():
    if not session.get("user_id"):
        return redirect(url_for("login"))
    return render_template("recipe_form.html", recipe=None)


@app.post("/recipes/new")
def recipe_new_submit():
    if not session.get("user_id"):
        return redirect(url_for("login"))
    db = get_db(); _ensure_recipes(db)
    user = db.execute("SELECT id, rig_id FROM users WHERE id=?", (session["user_id"],)).fetchone()
    if not user or not user["rig_id"]:
        flash("Ingen rigg.", "warning"); return redirect(url_for("dashboard"))
    rig_id = user["rig_id"]
    title = (request.form.get("title") or "").strip()
    if not title:
        flash("Tittel kreves.", "danger"); return redirect(url_for("recipe_new_form"))
    raw_text = (request.form.get("raw_text") or "").strip()
    source_url = (request.form.get("source_url") or "").strip() or None
    notes = (request.form.get("notes") or "").strip() or None
    base_portions = request.form.get("base_portions") or None
    try:
        base_portions_val = int(base_portions) if base_portions else None
    except Exception:
        base_portions_val = None
    raw_categories = request.form.get("categories")
    categories = _normalize_categories(raw_categories) if raw_categories else ""
    method_type = (request.form.get("method_type") or "base").strip().lower()
    db.execute("INSERT INTO recipes(rig_id,title,source_url,raw_text,notes,base_portions,categories,method_type) VALUES(?,?,?,?,?,?,?,?)",
               (rig_id, title, source_url, raw_text, notes, base_portions_val, categories, method_type))
    rid = db.execute("SELECT last_insert_rowid()").fetchone()[0]
    db.commit()
    flash("Recept opprettet.", "success")
    # Optional immediate dish link
    link_dish_id = request.form.get("link_dish_id")
    if link_dish_id:
        try:
            db.execute("UPDATE dish_catalog SET recipe_id=? WHERE id=? AND rig_id=?", (rid, int(link_dish_id), rig_id))
            db.commit()
        except Exception:
            pass
    return redirect(url_for("recipe_detail", recipe_id=rid))


@app.get("/recipes/<int:recipe_id>")
def recipe_detail(recipe_id: int):
    if not session.get("user_id"):
        return redirect(url_for("login"))
    db = get_db(); _ensure_recipes(db)
    user = db.execute("SELECT id, rig_id FROM users WHERE id=?", (session["user_id"],)).fetchone()
    if not user or not user["rig_id"]:
        flash("Ingen rigg.", "warning"); return redirect(url_for("dashboard"))
    rig_id = user["rig_id"]
    row = db.execute("SELECT *, (SELECT COUNT(*) FROM dish_catalog d WHERE d.rig_id=? AND d.recipe_id=recipes.id) AS link_count FROM recipes WHERE id=? AND rig_id=?", (rig_id, recipe_id, rig_id)).fetchone()
    if not row:
        flash("Recept ikke funnet.", "warning"); return redirect(url_for("recipe_list"))
    linked_dishes = db.execute("SELECT id, name FROM dish_catalog WHERE rig_id=? AND recipe_id=? ORDER BY name COLLATE NOCASE", (rig_id, recipe_id)).fetchall()
    rec = dict(id=row["id"], title=row["title"], source_url=row["source_url"], raw_text=row["raw_text"], notes=row["notes"], base_portions=row["base_portions"], categories=row["categories"] or "", method_type=row["method_type"] or "base", link_count=row["link_count"], created_at=row["created_at"], updated_at=row["updated_at"], linked_dishes=[dict(id=d["id"], name=d["name"]) for d in linked_dishes])
    return render_template("recipe_detail.html", recipe=rec)


@app.get("/recipes/<int:recipe_id>/edit")
def recipe_edit_form(recipe_id: int):
    if not session.get("user_id"):
        return redirect(url_for("login"))
    db = get_db(); _ensure_recipes(db)
    user = db.execute("SELECT id, rig_id FROM users WHERE id=?", (session["user_id"],)).fetchone()
    if not user or not user["rig_id"]:
        flash("Ingen rigg.", "warning"); return redirect(url_for("dashboard"))
    row = db.execute("SELECT * FROM recipes WHERE id=? AND rig_id=?", (recipe_id, user["rig_id"])).fetchone()
    if not row:
        flash("Recept ikke funnet.", "warning"); return redirect(url_for("recipe_list"))
    rec = dict(id=row["id"], title=row["title"], source_url=row["source_url"], raw_text=row["raw_text"], notes=row["notes"], base_portions=row["base_portions"], categories=row["categories"] or "", method_type=row["method_type"] or "base")
    return render_template("recipe_form.html", recipe=rec)


@app.post("/recipes/<int:recipe_id>/edit")
def recipe_edit_submit(recipe_id: int):
    if not session.get("user_id"):
        return redirect(url_for("login"))
    db = get_db(); _ensure_recipes(db)
    user = db.execute("SELECT id, rig_id FROM users WHERE id=?", (session["user_id"],)).fetchone()
    if not user or not user["rig_id"]:
        flash("Ingen rigg.", "warning"); return redirect(url_for("dashboard"))
    rig_id = user["rig_id"]
    row = db.execute("SELECT id FROM recipes WHERE id=? AND rig_id=?", (recipe_id, rig_id)).fetchone()
    if not row:
        flash("Recept ikke funnet.", "warning"); return redirect(url_for("recipe_list"))
    title = (request.form.get("title") or "").strip()
    if not title:
        flash("Tittel kreves.", "danger"); return redirect(url_for("recipe_edit_form", recipe_id=recipe_id))
    raw_text = (request.form.get("raw_text") or "").strip()
    source_url = (request.form.get("source_url") or "").strip() or None
    notes = (request.form.get("notes") or "").strip() or None
    base_portions = request.form.get("base_portions") or None
    try:
        base_portions_val = int(base_portions) if base_portions else None
    except Exception:
        base_portions_val = None
    raw_categories = request.form.get("categories")
    categories = _normalize_categories(raw_categories) if raw_categories else ""
    method_type = (request.form.get("method_type") or "base").strip().lower()
    db.execute("UPDATE recipes SET title=?, source_url=?, raw_text=?, notes=?, base_portions=?, categories=?, method_type=?, updated_at=CURRENT_TIMESTAMP WHERE id=? AND rig_id=?",
               (title, source_url, raw_text, notes, base_portions_val, categories, method_type, recipe_id, rig_id))
    db.commit()
    flash("Recept oppdatert.", "success")
    return redirect(url_for("recipe_detail", recipe_id=recipe_id))


def _load_menu_settings(db: sqlite3.Connection, rig_id: int):
    row = db.execute("SELECT start_week, start_index, menu_json FROM menu_settings WHERE rig_id=?", (rig_id,)).fetchone()
    if not row:
        return None, None, []
    start_week = row[0] or 1
    start_index = row[1] or 1
    try:
        menu_sheets = json.loads(row[2] or "[]")
    except Exception:
        menu_sheets = []
    return start_week, start_index, menu_sheets


def _norm_category(raw: str):
    if not raw:
        return None
    r = str(raw).strip().lower()
    repl = r.replace("ø", "o").replace("å", "a").replace("ä", "a").replace("ö", "o").replace("æ", "a")
    if ("sopp" in repl) or ("suppe" in repl) or ("soup" in repl) or ("sopa" in repl):
        return "soppa"
    if ("fisk" in repl) or ("fish" in repl):
        return "fisk"
    if ("kjott" in repl) or ("kjot" in repl) or ("kott" in repl) or ("meat" in repl):
        return "kott"
    return "extra"  # treat unknown as extra for now


_DAY_NAME_MAP = {
    "mandag":0,"monday":0,"maandag":0,"må":0,"måndag":0,
    "tirsdag":1,"tuesday":1,"ti":1,"tisdag":1,
    "onsdag":2,"wednesday":2,"ons":2,"onsdag":2,
    "torsdag":3,"thursday":3,"tor":3,"torsdag":3,"thurs":3,
    "fredag":4,"friday":4,"fre":4,"fredag":4,
    "lørdag":5,"lordag":5,"lördag":5,"lørdag":5,"saturday":5,"sat":5,
    "søndag":6,"sondag":6,"söndag":6,"sunday":6,"sun":6
}


def _weekday_from_label(label: str):
    if not label:
        return None
    k = label.strip().lower()
    return _DAY_NAME_MAP.get(k)


def _rebuild_menu_dish_map(db: sqlite3.Connection, rig_id: int):
    """Parse menu_settings.menu_json and build dish_catalog + menu_dish_map (idempotent)."""
    _ensure_dish_catalog(db)
    start_week, start_index, menu_sheets = _load_menu_settings(db, rig_id)
    if not menu_sheets:
        return
    sheets_by_uke = {str(sh.get("uke")): sh for sh in menu_sheets}
    # Insert dishes
    for uke, sheet in sheets_by_uke.items():
        for meal_key in ("lunsj","middag"):
            for it in (sheet.get(meal_key) or []):
                dag = it.get("dag")
                wd = _weekday_from_label(dag)
                if wd is None:
                    continue
                cat = _norm_category(it.get("kategori")) or "extra"
                dish_name = (it.get("rett") or "").strip()
                if not dish_name:
                    continue
                slug = _slugify_dish(dish_name)
                if not slug:
                    continue
                db.execute(
                    "INSERT OR IGNORE INTO dish_catalog(rig_id, slug, name, first_seen_date) VALUES(?,?,?,date('now'))",
                    (rig_id, slug, dish_name)
                )
                # resolve id
                drow = db.execute("SELECT id FROM dish_catalog WHERE rig_id=? AND slug=?", (rig_id, slug)).fetchone()
                if not drow:
                    continue
                dish_id = drow[0]
                # uke value may be string like '1' or 'Uke 1' -> extract first int
                try:
                    import re as _re
                    m = _re.search(r"(\d+)", str(uke))
                    menu_index = int(m.group(1)) if m else int(uke)
                except Exception:
                    continue
                try:
                    db.execute(
                        "INSERT OR IGNORE INTO menu_dish_map(rig_id, menu_index, weekday, meal, category, dish_id) VALUES(?,?,?,?,?,?)",
                        (rig_id, menu_index, wd, meal_key, cat, dish_id)
                    )
                except Exception:
                    pass
    db.commit()


def _menu_index_for_date(date_obj, start_week: int, start_index: int):
    iso_year, iso_week, _ = date_obj.isocalendar()
    return ((iso_week - int(start_week) + (int(start_index) - 1)) % 4) + 1


def _backfill_task_dish_ids(db: sqlite3.Connection, rig_id: int, user_id: int | None = None):
    """Populate dish_id in prep_tasks_private / frys_items_private based on date+meal+category.
    If user_id given, limit to that user's rows to keep it fast for dashboard loads.
    """
    _ensure_prep_tables(db)
    _ensure_dish_catalog(db)
    _ensure_recipes(db)
    start_week, start_index, _menu_sheets = _load_menu_settings(db, rig_id)
    if start_week is None:
        return
    try:
        _sw = int(start_week or 1)
    except Exception:
        _sw = 1
    try:
        _si = int(start_index or 1)
    except Exception:
        _si = 1
    # Query rows needing dish_id
    base_where = "rig_id=? AND (dish_id IS NULL OR dish_id=0)"
    params = [rig_id]
    if user_id:
        base_where += " AND user_id=?"
        params.append(user_id)
    rows = db.execute(f"SELECT id, date, meal, category FROM prep_tasks_private WHERE {base_where} LIMIT 500", tuple(params)).fetchall()
    frows = db.execute(f"SELECT id, date, meal, category FROM frys_items_private WHERE {base_where} LIMIT 500", tuple(params)).fetchall()
    if not rows and not frows:
        return
    # Build quick lookup for menu_dish_map
    # We rely on menu_dish_map being present; rebuild if empty
    cnt = db.execute("SELECT COUNT(*) FROM menu_dish_map WHERE rig_id=?", (rig_id,)).fetchone()[0]
    if cnt == 0:
        _rebuild_menu_dish_map(db, rig_id)
    for r in rows:
        try:
            from datetime import datetime as _dt
            dt = _dt.strptime(r["date"], "%Y-%m-%d").date()
            m_idx = _menu_index_for_date(dt, _sw, _si)
            wd = dt.weekday()
            maprow = db.execute(
                "SELECT dish_id FROM menu_dish_map WHERE rig_id=? AND menu_index=? AND weekday=? AND meal=? AND category=?",
                (rig_id, m_idx, wd, r["meal"], r["category"])
            ).fetchone()
            if maprow:
                db.execute("UPDATE prep_tasks_private SET dish_id=? WHERE id=?", (maprow[0], r["id"]))
        except Exception:
            continue
    for r in frows:
        try:
            from datetime import datetime as _dt
            dt = _dt.strptime(r["date"], "%Y-%m-%d").date()
            m_idx = _menu_index_for_date(dt, _sw, _si)
            wd = dt.weekday()
            maprow = db.execute(
                "SELECT dish_id FROM menu_dish_map WHERE rig_id=? AND menu_index=? AND weekday=? AND meal=? AND category=?",
                (rig_id, m_idx, wd, r["meal"], r["category"])
            ).fetchone()
            if maprow:
                db.execute("UPDATE frys_items_private SET dish_id=? WHERE id=?", (maprow[0], r["id"]))
        except Exception:
            continue
    db.commit()


def _compute_recipe_coverage(db: sqlite3.Connection, rig_id: int):
    """Return dict with recipe coverage stats for dashboard.

    Coverage definition MVP:
      numerator: COUNT(dishes with non-null recipe_id AND recipe exists)
      denominator: COUNT(distinct dishes that have any prep or frys activity OR appear in menu_dish_map)
    Falls back gracefully on missing tables.
    """
    try:
        # ensure tables exist (idempotent)
        _ensure_dish_catalog(db)
        _ensure_recipes(db)
        # base: dishes that appear in menu mapping
        menu_dishes = db.execute("SELECT DISTINCT dish_id FROM menu_dish_map WHERE rig_id=?", (rig_id,)).fetchall()
        dish_ids = {r[0] for r in menu_dishes if r[0] is not None}
        # include dishes referenced by prep/frys tasks (backfilled or manual)
        try:
            prep_ids = db.execute("SELECT DISTINCT dish_id FROM prep_tasks_private WHERE rig_id=? AND dish_id IS NOT NULL", (rig_id,)).fetchall()
            for r in prep_ids: dish_ids.add(r[0])
        except Exception:
            pass
        try:
            frys_ids = db.execute("SELECT DISTINCT dish_id FROM frys_items_private WHERE rig_id=? AND dish_id IS NOT NULL", (rig_id,)).fetchall()
            for r in frys_ids: dish_ids.add(r[0])
        except Exception:
            pass
        if not dish_ids:
            return {"numerator": 0, "denominator": 0, "percent": 0}
        placeholders = ",".join(["?"] * len(dish_ids))
        rows = db.execute(f"SELECT COUNT(*) FROM dish_catalog WHERE rig_id=? AND id IN ({placeholders}) AND recipe_id IS NOT NULL", (rig_id, *dish_ids)).fetchone()
        numerator = rows[0] if rows else 0
        denominator = len(dish_ids)
        percent = int((numerator / denominator) * 100) if denominator else 0
        return {"numerator": numerator, "denominator": denominator, "percent": percent}
    except Exception:
        return {"numerator": 0, "denominator": 0, "percent": 0}


@app.get("/planning/overview")
def planning_overview():
    if not session.get("user_id"):
        flash("Logg inn kreves.", "warning")
        return redirect(url_for("login"))
    db = get_db()
    user_id = session["user_id"]
    user = db.execute("SELECT id, rig_id FROM users WHERE id=?", (user_id,)).fetchone()
    if not user or not user["rig_id"]:
        flash("Ingen rigg tilknyttet.", "warning")
        return redirect(url_for("dashboard"))
    rig_id = user["rig_id"]
    _ensure_prep_tables(db)
    _ensure_dish_catalog(db)
    _rebuild_menu_dish_map(db, rig_id)
    _backfill_task_dish_ids(db, rig_id, user_id=user_id)
    rows = db.execute(
        """
        SELECT d.id AS dish_id, d.name,
               SUM(CASE WHEN pt.done=0 THEN 1 ELSE 0 END) AS open_prep,
               COUNT(pt.id) AS total_prep,
               SUM(CASE WHEN fi.done=0 THEN 1 ELSE 0 END) AS open_frys,
               COUNT(fi.id) AS total_frys
        FROM dish_catalog d
        LEFT JOIN prep_tasks_private pt ON pt.dish_id = d.id AND pt.user_id=? AND pt.rig_id=?
        LEFT JOIN frys_items_private fi ON fi.dish_id = d.id AND fi.user_id=? AND fi.rig_id=?
        WHERE d.rig_id=?
        GROUP BY d.id, d.name
        HAVING total_prep > 0 OR total_frys > 0
        ORDER BY d.name COLLATE NOCASE
        """, (user_id, rig_id, user_id, rig_id, rig_id)
    ).fetchall()
    dishes = [dict(dish_id=r["dish_id"], name=r["name"], open_prep=r["open_prep"], total_prep=r["total_prep"], open_frys=r["open_frys"], total_frys=r["total_frys"]) for r in rows]
    return render_template("planning_overview.html", dishes=dishes)


@app.get("/planning/dish/<int:dish_id>")
def planning_dish_detail(dish_id: int):
    """Detaljvy for en rett: viser alle prepp- og frysplock-linjer knyttet til dish_id for innlogget bruker.
    Viser melding hvis ingen finnes ennå."""
    if not session.get("user_id"):
        flash("Logg inn kreves.", "warning")
        return redirect(url_for("login"))
    db = get_db()
    user_id = session["user_id"]
    user = db.execute("SELECT id, rig_id FROM users WHERE id=?", (user_id,)).fetchone()
    if not user or not user["rig_id"]:
        flash("Ingen rigg tilknyttet.", "warning")
        return redirect(url_for("dashboard"))
    rig_id = user["rig_id"]
    # Sikre katalog og backfill slik at dish_id blir populert for nyeste data
    _ensure_prep_tables(db)
    _ensure_dish_catalog(db)
    dish_row = db.execute("SELECT id, name, rig_id, recipe_id FROM dish_catalog WHERE id=?", (dish_id,)).fetchone()
    if not dish_row or dish_row["rig_id"] != rig_id:
        flash("Rett ikke funnet.", "warning")
        return redirect(url_for("planning_overview"))
    _backfill_task_dish_ids(db, rig_id, user_id=user_id)
    # Hent prepp og frys
    prep_rows = db.execute(
        "SELECT id,date,meal,category,text,done FROM prep_tasks_private WHERE user_id=? AND rig_id=? AND dish_id=? ORDER BY date, meal, category, id",
        (user_id, rig_id, dish_id)
    ).fetchall()
    frys_rows = db.execute(
        "SELECT id,date,meal,category,item_name,qty,unit,done FROM frys_items_private WHERE user_id=? AND rig_id=? AND dish_id=? ORDER BY date, meal, category, id",
        (user_id, rig_id, dish_id)
    ).fetchall()
    # Gruppér
    from collections import OrderedDict
    groups: OrderedDict[str, dict] = OrderedDict()
    def key(date, meal, category):
        return f"{date}|{meal}|{category}"
    for r in prep_rows:
        k = key(r["date"], r["meal"], r["category"])
        if k not in groups:
            groups[k] = {
                "date": r["date"], "meal": r["meal"], "category": r["category"],
                "prep": [], "frys": []
            }
        groups[k]["prep"].append(dict(id=r["id"], text=r["text"], done=bool(r["done"])))
    for r in frys_rows:
        k = key(r["date"], r["meal"], r["category"])
        if k not in groups:
            groups[k] = {
                "date": r["date"], "meal": r["meal"], "category": r["category"],
                "prep": [], "frys": []
            }
        groups[k]["frys"].append(dict(id=r["id"], item_name=r["item_name"], qty=r["qty"], unit=r["unit"], done=bool(r["done"])))
    grouped = list(groups.values())
    # Receptinfo
    _ensure_recipes(db)
    recipe = None
    if dish_row.get("recipe_id"):
        r = db.execute("SELECT id, title FROM recipes WHERE id=?", (dish_row["recipe_id"],)).fetchone()
        if r:
            recipe = dict(id=r["id"], title=r["title"])
    # Kandidater (senaste 8 uppdaterade)
    candidates = db.execute("SELECT id, title FROM recipes WHERE rig_id=? ORDER BY updated_at DESC LIMIT 8", (rig_id,)).fetchall()
    cand_list = [dict(id=c["id"], title=c["title"]) for c in candidates]
    return render_template("planning_dish_detail.html", dish=dish_row, groups=grouped, has_any=(len(prep_rows)+len(frys_rows))>0, recipe=recipe, candidates=cand_list)


@app.post("/dish/<int:dish_id>/link_recipe")
def link_recipe_to_dish(dish_id: int):
    if not session.get("user_id"):
        return redirect(url_for("login"))
    db = get_db(); _ensure_recipes(db); _ensure_dish_catalog(db)
    user = db.execute("SELECT id, rig_id FROM users WHERE id=?", (session["user_id"],)).fetchone()
    if not user or not user["rig_id"]:
        flash("Ingen rigg.", "warning"); return redirect(url_for("dashboard"))
    rig_id = user["rig_id"]
    dish = db.execute("SELECT id FROM dish_catalog WHERE id=? AND rig_id=?", (dish_id, rig_id)).fetchone()
    if not dish:
        flash("Rett ikke funnet.", "warning"); return redirect(url_for("planning_overview"))
    rid = request.form.get("recipe_id") or ""
    if not rid.isdigit():
        flash("Ugyldig recipe id.", "warning"); return redirect(url_for("planning_dish_detail", dish_id=dish_id))
    db.execute("UPDATE dish_catalog SET recipe_id=? WHERE id=? AND rig_id=?", (int(rid), dish_id, rig_id))
    db.commit()
    flash("Recept kopplat.", "success")
    return redirect(url_for("planning_dish_detail", dish_id=dish_id))



# --- DOCX helpers ---
def _build_docx_modern(rig_name: str, dt, lunch_picked, lunch_all, lunch_glu, lunch_nuts, dinner_picked, dinner_all, dinner_glu, dinner_nuts):
    """Build a clean daily menu DOCX without relying on the template, focusing on readable layout.

    Returns a python-docx Document instance.
    """
    if Document is None:
        raise RuntimeError("DOCX library missing")
    doc = Document()
    # Page margins (safe defaults)
    try:
        if Cm:
            for s in doc.sections:
                s.top_margin = Cm(1.5)
                s.bottom_margin = Cm(1.5)
                s.left_margin = Cm(1.5)
                s.right_margin = Cm(1.5)
    except Exception:
        pass

    # Header table: Place of business / Date
    hdr = doc.add_table(rows=2, cols=2)
    hdr.autofit = True
    p = hdr.cell(0, 0).paragraphs[0]
    p.add_run("Place of business: ").bold = True
    p.add_run(rig_name or "")
    p2 = hdr.cell(0, 1).paragraphs[0]
    p2.add_run("Date: ").bold = True
    p2.add_run(dt.strftime("%Y-%m-%d"))
    # empty row for spacing
    hdr.cell(1, 0).merge(hdr.cell(1, 1))
    hdr.cell(1, 0).paragraphs[0].add_run("")

    # Helper to suffix allergens
    def suf(cat: str, alls: dict, glu: dict, nuts: dict) -> str:
        alg = (alls.get(cat) or "").strip().replace(" ", "")
        g = (glu.get(cat) or "").strip()
        n = (nuts.get(cat) or "").strip()
        if alg and "1" in [x for x in alg.split(",") if x]:
            if g:
                alg = alg + f" [gluten: {g}]"
        if alg and "8" in [x for x in alg.split(",") if x]:
            if n:
                alg = alg + f" [nuts: {n}]"
        return f" ({alg})" if alg else ""

    # Two-column main table
    t = doc.add_table(rows=2, cols=2)
    t.autofit = True
    # Lunch cell (0,0)
    lc = t.cell(0, 0)
    for label, cat in (("soup of the day:", "soppa"), ("today’s fish:", "fisk"), ("today’s meat:", "kott")):
        para = lc.add_paragraph()
        r = para.add_run(label + " ")
        r.bold = True
    para.add_run((lunch_picked.get(cat) or "").strip() + suf(cat, lunch_all, lunch_glu, lunch_nuts))
    # extra
    para = lc.add_paragraph()
    r = para.add_run("3. dish/green dish: ")
    r.bold = True
    para.add_run((lunch_picked.get("extra") or "").strip() + suf("extra", lunch_all, lunch_glu, lunch_nuts))

    # Dinner cell (0,1)
    rc = t.cell(0, 1)
    for label, cat in (("soup of the day:", "soppa"), ("today’s fish:", "fisk"), ("today’s meat:", "kott")):
        para = rc.add_paragraph()
        r = para.add_run(label + " ")
        r.bold = True
    para.add_run((dinner_picked.get(cat) or "").strip() + suf(cat, dinner_all, dinner_glu, dinner_nuts))
    # extra
    para = rc.add_paragraph()
    r = para.add_run("3. dish/green dish: ")
    r.bold = True
    para.add_run((dinner_picked.get("extra") or "").strip() + suf("extra", dinner_all, dinner_glu, dinner_nuts))

    # Add a compact allergen legend
    try:
        icons = {
            "1": "🥖",  # Gluten
            "2": "🦐",  # Crustaceans
            "3": "🥚",  # Eggs
            "4": "🐟",  # Fish
            "5": "🥜",  # Peanuts
            "6": "🌱",  # Soybeans
            "7": "🥛",  # Milk
            "8": "🌰",  # Nuts
            "9": "🥬",  # Celery
            "10": "🌶️", # Mustard
            "11": "⚪",  # Sesame
            "12": "SO₂", # Sulphites
            "13": "🌼",  # Lupin
            "14": "🦑",  # Molluscs
        }
        gluten_sub = ["1A Wheat", "1B Rye", "1C Barley", "1D Oats", "1E Spelt", "1F Kamut", "1G Hybrid strains"]
        nuts_sub = [
            "8A Almonds", "8B Hazelnuts", "8C Walnuts", "8D Cashews",
            "8E Pecans", "8F Brazil nuts", "8G Pistachios", "8H Macadamia"
        ]
        doc.add_paragraph("")
        title = doc.add_paragraph("Allergen guide")
        try:
            title.runs[0].bold = True
        except Exception:
            pass
        legend_lines = [
            f"{icons['1']} 1 Gluten — " + ", ".join(gluten_sub),
            f"{icons['2']} 2 Crustaceans",
            f"{icons['3']} 3 Eggs",
            f"{icons['4']} 4 Fish",
            f"{icons['5']} 5 Peanuts",
            f"{icons['6']} 6 Soybeans",
            f"{icons['7']} 7 Milk",
            f"{icons['8']} 8 Nuts — " + ", ".join(nuts_sub),
            f"{icons['9']} 9 Celery",
            f"{icons['10']} 10 Mustard",
            f"{icons['11']} 11 Sesame",
            f"{icons['12']} 12 Sulphur dioxide & sulphites",
            f"{icons['13']} 13 Lupin",
            f"{icons['14']} 14 Molluscs",
        ]
        for line in legend_lines:
            p = doc.add_paragraph(line)
            try:
                if Pt:
                    for r in p.runs:
                        if r.font.size is None or (r.font.size and r.font.size.pt > 10):
                            r.font.size = Pt(9)
            except Exception:
                pass
    except Exception:
        pass

    return doc


## removed legacy template-like docx builder and logo finder


def _parse_ts(ts: str) -> datetime:
    return datetime.strptime(ts, "%Y-%m-%dT%H:%M")


# --- Auth helpers ---
def admin_required(fn):
    from functools import wraps
    @wraps(fn)
    def wrapper(*args, **kwargs):
        if session.get("role") != "admin":
            flash("Admin kreves.", "warning")
            return redirect(url_for("admin_login"))
        return fn(*args, **kwargs)
    return wrapper


def superuser_required(fn):
    from functools import wraps
    @wraps(fn)
    def wrapper(*args, **kwargs):
        if not session.get("superuser"):
            flash("Superuser kreves.", "warning")
            return redirect(url_for("superuser_login"))
        return fn(*args, **kwargs)
    return wrapper


def _current_rig(db):
    """Return (rig_id, rig_name) for current session user if any."""
    rig_id = None
    rig_name = None
    if session.get("user_id"):
        row = db.execute("SELECT rig_id FROM users WHERE id=?", (session["user_id"],)).fetchone()
        if row and row["rig_id"]:
            rig_id = row["rig_id"]
            r = db.execute("SELECT name FROM rigs WHERE id=?", (rig_id,)).fetchone()
            rig_name = r["name"] if r else None
    return rig_id, rig_name


def _build_interval_from_params():
    start = (request.args.get("start") or request.form.get("start") or "").strip()
    end = (request.args.get("end") or request.form.get("end") or "").strip()
    if not start or not end:
        raise ValueError("start og end krävs")
    # Accept YYYY-MM-DD or full TS
    def norm_start(x: str) -> str:
        # Date-only -> start of day
        return x if "T" in x else f"{x}T00:00"
    def norm_end(x: str) -> str:
        # Date-only -> end of day (inclusive)
        return x if "T" in x else f"{x}T23:59"
    return norm_start(start), norm_end(end)


# --- User dashboard (auto period) ---
@app.route("/dashboard", methods=["GET"])  # simplified: no manual date picking
def dashboard():
    db = get_db()
    user = None
    rig_id = None
    rig_name = None
    if session.get("user_id"):
        user = db.execute("SELECT id, name, email, rig_id FROM users WHERE id=?", (session["user_id"],)).fetchone()
        if user and user["rig_id"]:
            rig_id = user["rig_id"]
            r = db.execute("SELECT name FROM rigs WHERE id=?", (rig_id,)).fetchone()
            rig_name = r["name"] if r else None

    # Settings for menu rotation from DB (per rig), fallback to session for admins preview
    s = None
    menu_sheets_db = []
    if rig_id:
        try:
            db.execute(
                """
                CREATE TABLE IF NOT EXISTS menu_settings (
                  rig_id INTEGER PRIMARY KEY,
                  start_week INTEGER,
                  start_index INTEGER,
                  menu_json TEXT,
                  updated_at TEXT DEFAULT (datetime('now'))
                )
                """
            )
            row = db.execute("SELECT start_week, start_index, menu_json FROM menu_settings WHERE rig_id=?", (rig_id,)).fetchone()
            if row:
                s = {"start_week": row[0] or 1, "start_index": row[1] or 1}
                try:
                    menu_sheets_db = json.loads(row[2] or "[]")
                except Exception:
                    menu_sheets_db = []
        except Exception:
            pass
    if s is None:
        s = session.get("menu_start_info")

    arbeidsperiod_start = None
    week_start = None
    week_end = None
    slots_by_day = None
    menus_by_day = None
    menu_days_period = None

    quick_menu_idx = None
    schedule_preview = []
    try:
        if rig_id and user:
            # Fetch all published, bound slots for this user on this rig from today forward
            today_ts = datetime.now().strftime("%Y-%m-%d") + "T00:00"
            rows = db.execute(
                """
                SELECT ts.id, ts.role, ts.start_ts, ts.end_ts, ts.status, ts.notes,
                       u.name AS user_name
                FROM turnus_slots ts
                INNER JOIN turnus_account_binding tb ON tb.slot_id = ts.id
                INNER JOIN users u ON u.id = tb.user_id
                WHERE ts.rig_id = ? AND ts.status='published' AND tb.user_id = ? AND ts.end_ts >= ?
                ORDER BY ts.start_ts ASC
                """,
                (rig_id, user["id"], today_ts),
            ).fetchall()

            # If none in the future, look at the most recent past period containing today
            if not rows:
                rows = db.execute(
                    """
                    SELECT ts.id, ts.role, ts.start_ts, ts.end_ts, ts.status, ts.notes,
                           u.name AS user_name
                    FROM turnus_slots ts
                    INNER JOIN turnus_account_binding tb ON tb.slot_id = ts.id
                    INNER JOIN users u ON u.id = tb.user_id
                    WHERE ts.rig_id = ? AND ts.status='published' AND tb.user_id = ?
                    ORDER BY ts.start_ts ASC
                    """,
                    (rig_id, user["id"]),
                ).fetchall()

            # Group consecutive rows into ~14-day periods and pick the first period whose end >= today
            from collections import defaultdict
            periods = []
            cur = None
            for r in rows:
                st = _parse_ts(r["start_ts"])  # datetime
                et = _parse_ts(r["end_ts"])    # datetime
                if cur is None:
                    cur = {"start": st, "end": et, "rows": [r]}
                    continue
                anchor = cur["start"]
                if st <= anchor + timedelta(days=14, hours=1):
                    cur["rows"].append(r)
                    if et > cur["end"]:
                        cur["end"] = et
                else:
                    periods.append(cur)
                    cur = {"start": st, "end": et, "rows": [r]}
            if cur is not None:
                periods.append(cur)

            # choose period: one that includes today or the next upcoming one
            now = datetime.now()
            chosen = None
            for p in periods:
                if p["start"] <= now <= p["end"]:
                    chosen = p
                    break
                if now < p["start"]:
                    chosen = p
                    break
            if chosen:
                d0 = chosen["start"].date()
                d1 = (d0 + timedelta(days=13))
                arbeidsperiod_start = d0.strftime("%Y-%m-%d")
                week_start = d0.strftime("%Y-%m-%d")
                week_end = d1.strftime("%Y-%m-%d")
                # schedule preview: next 3 upcoming slots (from now)
                schedule_preview = [
                    {
                        "role": r["role"],
                        "start_ts": r["start_ts"],
                        "end_ts": r["end_ts"],
                    }
                    for r in rows
                    if _parse_ts(r["end_ts"]) >= datetime.now()
                ][:3]

                # Build day list for the 14-day period
                day_list = []
                curd = d0
                while curd <= d1:
                    day_list.append(curd)
                    curd += timedelta(days=1)

                # Group user's slots by day
                slots_by_day = defaultdict(list)
                def overlaps_day(r, day_date):
                    st = _parse_ts(r["start_ts"])  # datetime
                    en = _parse_ts(r["end_ts"])    # datetime
                    ds = datetime.combine(day_date, datetime.min.time())
                    de = datetime.combine(day_date, datetime.max.time().replace(hour=23, minute=59))
                    return (st <= de) and (en >= ds)
                for r in chosen["rows"]:
                    for day_date in day_list:
                        if overlaps_day(r, day_date):
                            day_key = day_date.strftime("%Y-%m-%d")
                            slots_by_day[day_key].append({
                                "id": r["id"],
                                "role": r["role"],
                                "start_ts": r["start_ts"],
                                "end_ts": r["end_ts"],
                                "status": r["status"],
                                "notes": r["notes"],
                                "user_name": r["user_name"],
                            })

                # Build menus_by_day using 4-week rotation settings and parsed menu_sheets
                menu_sheets = menu_sheets_db or (session.get("menu_sheets") or [])
                start_index = (s.get("start_index") if isinstance(s, dict) else None) or 1
                start_week = (s.get("start_week") if isinstance(s, dict) else None) or 1
                weekday_no_to_nor = {0: "Mandag", 1: "Tirsdag", 2: "Onsdag", 3: "Torsdag", 4: "Fredag", 5: "Lørdag", 6: "Søndag"}

                def menu_index_for_date(dt_date):
                    iso_year, iso_week, _ = dt_date.isocalendar()
                    idx = ((iso_week - int(start_week) + (int(start_index) - 1)) % 4) + 1
                    return idx

                sheets_by_uke = {str(sh.get("uke")): sh for sh in menu_sheets}
                def _norm_cat(raw):
                    if not raw:
                        return None
                    r = str(raw).strip().lower()
                    repl = r.replace("ø", "o").replace("å", "a").replace("ä", "a").replace("ö", "o").replace("æ", "a")
                    # soup variants: soppa (sv), suppe (no), soup (en), sopa (es)
                    if ("sopp" in repl) or ("suppe" in repl) or ("soup" in repl) or ("sopa" in repl):
                        return "soppa"
                    # fish variants
                    if ("fisk" in repl) or ("fish" in repl):
                        return "fisk"
                    # meat variants
                    if ("kjott" in repl) or ("kjot" in repl) or ("kott" in repl) or ("meat" in repl):
                        return "kott"
                    return None

                # Ensure override tables exist
                try:
                    db.execute(
                        """
                        CREATE TABLE IF NOT EXISTS daily_menu_overrides (
                          id INTEGER PRIMARY KEY AUTOINCREMENT,
                          rig_id INTEGER NOT NULL,
                          date TEXT NOT NULL,
                          meal TEXT NOT NULL,
                          category TEXT NOT NULL,
                          dish TEXT NOT NULL,
                          UNIQUE(rig_id, date, meal, category)
                        )
                        """
                    )
                    db.execute(
                        """
                                                CREATE TABLE IF NOT EXISTS daily_menu_allergens (
                          id INTEGER PRIMARY KEY AUTOINCREMENT,
                          rig_id INTEGER NOT NULL,
                          date TEXT NOT NULL,
                          meal TEXT NOT NULL,
                          category TEXT NOT NULL,
                          allergens TEXT,
                                                    gluten_types TEXT,
                                                    nuts_types TEXT,
                          UNIQUE(rig_id, date, meal, category)
                        )
                        """
                    )
                except Exception:
                    pass

                def _base_picks_for_date(meal_key: str, day_date):
                    idx = menu_index_for_date(day_date)
                    sh = sheets_by_uke.get(str(idx))
                    dn = weekday_no_to_nor[day_date.weekday()]
                    cats = {"soppa": [], "fisk": [], "kott": [], "extra": []}
                    if sh:
                        items = (sh.get(meal_key) or [])
                        for it in items:
                            if it.get("dag") != dn:
                                continue
                            cat = _norm_cat(it.get("kategori"))
                            if cat in ("soppa", "fisk", "kott"):
                                cats[cat].append(it.get("rett") or "")
                            else:
                                cats["extra"].append(it.get("rett") or "")
                    # pick first of each category for display
                    return {k: (v[0] if v else "") for k, v in cats.items()}

                def _apply_overrides_for_date(date_s: str, meal_key: str, picks: dict) -> dict:
                    """Returnerer private overrides for innlogget bruker; viser ikke andre kockers endringer."""
                    out = {
                        "soppa": picks.get("soppa") or "",
                        "fisk": picks.get("fisk") or "",
                        "kott": picks.get("kott") or "",
                        "extra": "",
                    }
                    uid = session.get("user_id")
                    if not uid:
                        return out
                    try:
                        db.execute(
                            """
                            CREATE TABLE IF NOT EXISTS daily_menu_overrides_private (
                              id INTEGER PRIMARY KEY AUTOINCREMENT,
                              rig_id INTEGER NOT NULL,
                              user_id INTEGER NOT NULL,
                              date TEXT NOT NULL,
                              meal TEXT NOT NULL,
                              category TEXT NOT NULL,
                              dish TEXT NOT NULL,
                              UNIQUE(rig_id,user_id,date,meal,category)
                            )
                            """
                        )
                        rows_priv = db.execute(
                            "SELECT category, dish FROM daily_menu_overrides_private WHERE rig_id=? AND user_id=? AND date=? AND meal=?",
                            (rig_id, uid, date_s, meal_key),
                        ).fetchall()
                        for r in rows_priv:
                            cat, dish = r[0], r[1]
                            if cat in out and dish:
                                out[cat] = dish
                    except Exception:
                        pass
                    return out
                menus_by_day = {}
                menu_days_period = []
                for day_date in day_list:
                    day_key = day_date.strftime("%Y-%m-%d")
                    idx = menu_index_for_date(day_date)
                    # compute base picks per category and apply overrides; then flatten to list for UI
                    base_l = _base_picks_for_date("lunsj", day_date)
                    base_d = _base_picks_for_date("middag", day_date)
                    picks_l = _apply_overrides_for_date(day_key, "lunsj", base_l)
                    picks_d = _apply_overrides_for_date(day_key, "middag", base_d)
                    # Finn hvilke kategorier som er endret av brukeren (privat forskjell)
                    lunsj_overridden = [c for c in ["soppa","fisk","kott","extra"] if picks_l.get(c) and picks_l.get(c) != (base_l.get(c) or "")]
                    middag_overridden = [c for c in ["soppa","fisk","kott","extra"] if picks_d.get(c) and picks_d.get(c) != (base_d.get(c) or "")]
                    lunsj_list = [picks_l.get("soppa") or "", picks_l.get("fisk") or "", picks_l.get("kott") or "", picks_l.get("extra") or ""]
                    lunsj_list = [x for x in lunsj_list if x]
                    middag_list = [picks_d.get("soppa") or "", picks_d.get("fisk") or "", picks_d.get("kott") or "", picks_d.get("extra") or ""]
                    middag_list = [x for x in middag_list if x]
                    menus_by_day[day_key] = {
                        "index": idx,
                        "lunsj": lunsj_list,
                        "middag": middag_list,
                        "lunsj_cats": {
                            "soppa": picks_l.get("soppa") or "",
                            "fisk": picks_l.get("fisk") or "",
                            "kott": picks_l.get("kott") or "",
                            "extra": picks_l.get("extra") or "",
                        },
                        "middag_cats": {
                            "soppa": picks_d.get("soppa") or "",
                            "fisk": picks_d.get("fisk") or "",
                            "kott": picks_d.get("kott") or "",
                            "extra": picks_d.get("extra") or "",
                        },
                        "lunsj_overridden": lunsj_overridden,
                        "middag_overridden": middag_overridden,
                    }
                    menu_days_period.append({
                        "date": day_key,
                        "lunsj": lunsj_list,
                        "middag": middag_list,
                        "lunsj_cats": menus_by_day[day_key]["lunsj_cats"],
                        "middag_cats": menus_by_day[day_key]["middag_cats"],
                        "lunsj_overridden": lunsj_overridden,
                        "middag_overridden": middag_overridden,
                    })

                # Build a clean "menykort for dagvakt":
                # - Friday (d0+7) dinner only (snu_till_dag)
                # - Sat..Thu (d0+8..d0+13) lunch + dinner
                # - Next Friday (d0+14) lunch only
                menykort_dagvakt = []
                def day_key_of(dd):
                    return dd.strftime("%Y-%m-%d")
                # Friday dinner
                d8 = d0 + timedelta(days=7)
                k8 = day_key_of(d8)
                if k8 in menus_by_day:
                    menykort_dagvakt.append({
                        "date": k8,
                        "weekday": d8.strftime("%A"),
                        "lunsj": [],
                        "middag": menus_by_day[k8]["middag"] or [],
                        "lunsj_cats": {"soppa": "", "fisk": "", "kott": "", "extra": ""},
                        "middag_cats": menus_by_day[k8].get("middag_cats") or {"soppa": "", "fisk": "", "kott": "", "extra": ""},
                        "label": "Middag",
                    })
                # Sat..Thu (full days)
                for off in range(8, 14):
                    dd = d0 + timedelta(days=off)
                    kk = day_key_of(dd)
                    v = menus_by_day.get(kk) or {"lunsj": [], "middag": []}
                    menykort_dagvakt.append({
                        "date": kk,
                        "weekday": dd.strftime("%A"),
                        "lunsj": v.get("lunsj") or [],
                        "middag": v.get("middag") or [],
                        "lunsj_cats": (v.get("lunsj_cats") or {"soppa": "", "fisk": "", "kott": "", "extra": ""}),
                        "middag_cats": (v.get("middag_cats") or {"soppa": "", "fisk": "", "kott": "", "extra": ""}),
                        "label": "Lunsj + Middag",
                    })
                # Departure Friday lunch (d0+14) – compute menu for that date too
                d14 = d0 + timedelta(days=14)
                idx14 = menu_index_for_date(d14)
                sh14 = sheets_by_uke.get(str(idx14))
                lunsj14 = []
                if sh14:
                    # pick per-category with overrides for that day
                    base14 = _base_picks_for_date("lunsj", d14)
                    picks14 = _apply_overrides_for_date(d14.strftime("%Y-%m-%d"), "lunsj", base14)
                    lunsj14 = [picks14.get("soppa") or "", picks14.get("fisk") or "", picks14.get("kott") or "", picks14.get("extra") or ""]
                    lunsj14 = [x for x in lunsj14 if x]
                menykort_dagvakt.append({
                    "date": d14.strftime("%Y-%m-%d"),
                    "weekday": d14.strftime("%A"),
                    "lunsj": lunsj14,
                    "middag": [],
                    "lunsj_cats": (picks14 if sh14 else {"soppa": "", "fisk": "", "kott": "", "extra": ""}),
                    "middag_cats": {"soppa": "", "fisk": "", "kott": "", "extra": ""},
                    "label": "Lunsj",
                })
                # expose in template
                menu_days_period = menykort_dagvakt

                # Personalisert menyuke:
                now = datetime.now()
                current_shift_date = None
                next_shift_date = None
                for r in rows:
                    st = _parse_ts(r["start_ts"])
                    en = _parse_ts(r["end_ts"])
                    if st <= now <= en:
                        current_shift_date = st.date()
                        break
                    if st >= now and next_shift_date is None:
                        next_shift_date = st.date()
                anchor_user_date = current_shift_date or next_shift_date or now.date()
                try:
                    user_menu_idx = menu_index_for_date(anchor_user_date)
                except Exception:
                    user_menu_idx = None
                from datetime import date as _date
                try:
                    quick_menu_idx_global = menu_index_for_date(_date.today())
                except Exception:
                    quick_menu_idx_global = user_menu_idx
                if current_shift_date:
                    next_shift_menu_idx = user_menu_idx
                else:
                    try:
                        next_shift_menu_idx = menu_index_for_date(next_shift_date) if next_shift_date else user_menu_idx
                    except Exception:
                        next_shift_menu_idx = None
                quick_menu_idx = user_menu_idx  # primær visning: brukerens menyuke
    except Exception as e:
        flash(f"Kunne ikke bygge visning: {e}", "danger")

    # Fallback: hvis ingen periode ble valgt (ingen slots bundet) – lag en enkel 14-dagers liste fra i dag
    if not menu_days_period:
        try:
            day_list = [datetime.now().date() + timedelta(days=i) for i in range(14)]
            # Rebruk enkel rotasjonslogikk hvis vi har menyark
            fallback_days = []
            # Hent meny-innstillinger (start_week/start_index) fra s hvis finnes
            start_week_fb = (s or {}).get("start_week", 1) if isinstance(s, dict) else 1
            start_index_fb = (s or {}).get("start_index", 1) if isinstance(s, dict) else 1
            menu_sheets = menu_sheets_db or (session.get("menu_sheets") or [])
            sheets_by_uke = {str(sh.get("uke")): sh for sh in menu_sheets}
            weekday_no_to_nor = {0: "Mandag", 1: "Tirsdag", 2: "Onsdag", 3: "Torsdag", 4: "Fredag", 5: "Lørdag", 6: "Søndag"}
            def menu_index_for_date_fb(dt_date):
                iso_year, iso_week, _ = dt_date.isocalendar()
                return ((iso_week - int(start_week_fb) + (int(start_index_fb) - 1)) % 4) + 1
            def _norm_cat_fb(raw):
                if not raw:
                    return None
                r = str(raw).strip().lower()
                repl = r.replace("ø","o").replace("å","a").replace("ä","a").replace("ö","o").replace("æ","a")
                if ("sopp" in repl) or ("suppe" in repl) or ("soup" in repl) or ("sopa" in repl):
                    return "soppa"
                if ("fisk" in repl) or ("fish" in repl):
                    return "fisk"
                if ("kjott" in repl) or ("kjot" in repl) or ("kott" in repl) or ("meat" in repl):
                    return "kott"
                return None
            for day_date in day_list:
                idx = menu_index_for_date_fb(day_date)
                sh = sheets_by_uke.get(str(idx))
                dn = weekday_no_to_nor[day_date.weekday()]
                lunsj_cats = {"soppa":"","fisk":"","kott":"","extra":""}
                middag_cats = {"soppa":"","fisk":"","kott":"","extra":""}
                if sh:
                    for meal_key, target in (("lunsj", lunsj_cats),("middag", middag_cats)):
                        items = (sh.get(meal_key) or [])
                        for it in items:
                            if it.get("dag") != dn: continue
                            cat = _norm_cat_fb(it.get("kategori"))
                            if cat in target and not target[cat]:
                                target[cat] = it.get("rett") or ""
                lunsj = [v for v in [lunsj_cats.get("soppa"), lunsj_cats.get("fisk"), lunsj_cats.get("kott"), lunsj_cats.get("extra")] if v]
                middag = [v for v in [middag_cats.get("soppa"), middag_cats.get("fisk"), middag_cats.get("kott"), middag_cats.get("extra")] if v]
                fallback_days.append({
                    "date": day_date.strftime("%Y-%m-%d"),
                    "weekday": day_date.strftime("%A"),
                    "lunsj": lunsj,
                    "middag": middag,
                    "lunsj_cats": lunsj_cats,
                    "middag_cats": middag_cats,
                    "label": "Lunsj + Middag" if lunsj or middag else ""
                })
            menu_days_period = fallback_days
        except Exception as _e:
            pass


    # Hard fallback: hvis quick_menu_idx fortsatt None men vi har start parametre -> beregn direkte
    if quick_menu_idx is None and isinstance(s, dict):
        try:
            from datetime import date as _date
            iso_year, iso_week, _ = _date.today().isocalendar()
            sw = int(s.get("start_week") or 1)
            si = int(s.get("start_index") or 1)
            quick_menu_idx = ((iso_week - sw + (si - 1)) % 4) + 1
        except Exception:
            quick_menu_idx = None

    # --- Period progress + today's task summary ---
    period_progress = None  # dict: day_index (0-based), total_days, percent
    today_task_summary = None  # dict with counts
    try:
        next_period_start = None
        if arbeidsperiod_start:
            d0 = datetime.strptime(arbeidsperiod_start, "%Y-%m-%d").date()
            today_d = datetime.now().date()
            end_d = d0 + timedelta(days=13)
            if today_d < d0:
                # Not started yet
                next_period_start = d0.strftime("%Y-%m-%d")
            else:
                # In progress or finished -> clamp idx
                if today_d <= end_d:
                    idx = (today_d - d0).days
                else:
                    idx = 13
                percent = int(((idx+1)/14)*100)
                if percent < 35:
                    color_class = "bg-danger"
                elif percent < 70:
                    color_class = "bg-warning text-dark"
                else:
                    color_class = "bg-success"
                period_progress = {"day_index": idx, "total_days": 14, "percent": percent, "color_class": color_class}
        # Aggregate today's prep/frys tasks for quick glance
        if session.get("user_id") and user and user.get("id") and rig_id:
            td_key = datetime.now().strftime("%Y-%m-%d")
            dbh = db  # same connection
            prep_row = dbh.execute("SELECT COUNT(*) FROM prep_tasks_private WHERE user_id=? AND rig_id=? AND date=? AND done=0", (user["id"], rig_id, td_key)).fetchone()
            prep_done_row = dbh.execute("SELECT COUNT(*) FROM prep_tasks_private WHERE user_id=? AND rig_id=? AND date=? AND done=1", (user["id"], rig_id, td_key)).fetchone()
            frys_row = dbh.execute("SELECT COUNT(*) FROM frys_items_private WHERE user_id=? AND rig_id=? AND date=? AND done=0", (user["id"], rig_id, td_key)).fetchone()
            frys_done_row = dbh.execute("SELECT COUNT(*) FROM frys_items_private WHERE user_id=? AND rig_id=? AND date=? AND done=1", (user["id"], rig_id, td_key)).fetchone()
            today_task_summary = {
                "prep_open": prep_row[0] if prep_row else 0,
                "prep_done": prep_done_row[0] if prep_done_row else 0,
                "frys_open": frys_row[0] if frys_row else 0,
                "frys_done": frys_done_row[0] if frys_done_row else 0,
                "date": td_key
            }
        # Add explicit current (today) rotation index for simpler UI label independent of shift logic
        try:
            # Reuse menu_index_for_date if defined in local scope
            from datetime import date as _date
            if "menu_index_for_date" in locals():
                current_menu_week = menu_index_for_date(_date.today())
            else:
                current_menu_week = None
        except Exception:
            current_menu_week = None
        # Cache banner data in session for global top bar reuse
        session["current_menu_week"] = current_menu_week
        session["period_progress"] = period_progress
        session["next_period_start"] = next_period_start
        session["today_task_summary"] = today_task_summary
        session["iso_week"] = datetime.now().isocalendar()[1]
    except Exception:
        # Silently ignore any banner/session population errors to not break dashboard
        pass

    return render_template(
        "dashboard.html",
        user=user or {},
        rig_name=rig_name,
        s=s,
        arbeidsperiod_start=arbeidsperiod_start,
        menu_days_period=menu_days_period,
        slots_by_day=slots_by_day,
        week_start=week_start,
        week_end=week_end,
        menus_by_day=menus_by_day,
    quick_menu_idx=quick_menu_idx,  # personlig (nåværende eller neste skift)
    quick_menu_idx_global=locals().get("quick_menu_idx_global"),
    next_shift_menu_idx=locals().get("next_shift_menu_idx"),
    user_menu_anchor_date=locals().get("anchor_user_date"),
        schedule_preview=schedule_preview,
        today_str=datetime.now().strftime("%Y-%m-%d"),
        period_progress=period_progress,
        today_task_summary=today_task_summary,
        iso_week=datetime.now().isocalendar()[1],  # ekte ISO uke (1-53)
        current_menu_week=locals().get("current_menu_week"),
        recipe_coverage=_compute_recipe_coverage(db, rig_id) if rig_id else None,
    )

# --- Global context processor for banner variables ---
@app.context_processor
def inject_global_banner():
    try:
        return {
            "today_str": datetime.now().strftime("%Y-%m-%d"),
            "iso_week": session.get("iso_week") or datetime.now().isocalendar()[1],
            "current_menu_week": session.get("current_menu_week"),
            "period_progress": session.get("period_progress"),
            "today_task_summary": session.get("today_task_summary"),
            "next_period_start": session.get("next_period_start")
        }
    except Exception:
        return {}

    # (Unreachable) NOTE: If we need runtime debug, move logging above return.


# --- User schedule: next six shifts ---
@app.get("/me/schedule")
def me_schedule():
    db = get_db()
    if not session.get("user_id"):
        flash("Logg inn kreves.", "warning")
        return redirect(url_for("login"))
    user_id = session["user_id"]
    rig_id, rig_name = _current_rig(db)
    today_ts = datetime.now().strftime("%Y-%m-%dT%H:%M")
    rows = db.execute(
        """
        SELECT ts.id, ts.role, ts.start_ts, ts.end_ts
        FROM turnus_slots ts
        INNER JOIN turnus_account_binding tb ON tb.slot_id = ts.id
        WHERE ts.rig_id = ? AND ts.status='published' AND tb.user_id = ? AND ts.end_ts >= ?
        ORDER BY ts.start_ts ASC
        LIMIT 6
        """,
        (rig_id, user_id, today_ts),
    ).fetchall()
    return render_template("user_schedule.html", rows=rows, rig_name=rig_name)


# --- Upcoming planning (prep/frysplock/notes) ---
@app.get("/prep/upcoming")
def prep_upcoming():
    db = get_db()
    if not session.get("user_id"):
        flash("Logg inn kreves.", "warning")
        return redirect(url_for("login"))
    user = db.execute("SELECT id, rig_id FROM users WHERE id=?", (session["user_id"],)).fetchone()
    rig_id = user["rig_id"] if user else None
    if not rig_id:
        flash("Mangler rigg.", "warning")
        return redirect(url_for("dashboard"))

    # Collect periods similar to dashboard, then pick next up to 3
    today_ts = datetime.now().strftime("%Y-%m-%d") + "T00:00"
    rows = db.execute(
        """
        SELECT ts.id, ts.role, ts.start_ts, ts.end_ts
        FROM turnus_slots ts
        INNER JOIN turnus_account_binding tb ON tb.slot_id = ts.id
        WHERE ts.rig_id = ? AND ts.status='published' AND tb.user_id = ? AND ts.end_ts >= ?
        ORDER BY ts.start_ts ASC
        """,
        (rig_id, user["id"], today_ts),
    ).fetchall()
    periods = []
    cur = None
    for r in rows:
        st = _parse_ts(r["start_ts"]).date()
        et = _parse_ts(r["end_ts"]).date()
        if cur is None:
            cur = {"start": st, "end": et, "rows": [r]}
            continue
        anchor = cur["start"]
        if st <= anchor + timedelta(days=14):
            cur["rows"].append(r)
            if et > cur["end"]:
                cur["end"] = et
        else:
            periods.append(cur)
            cur = {"start": st, "end": et, "rows": [r]}
    if cur is not None:
        periods.append(cur)

    # Menu settings
    s = None
    menu_sheets_db = []
    row = db.execute("SELECT start_week, start_index, menu_json FROM menu_settings WHERE rig_id=?", (rig_id,)).fetchone()
    if row:
        s = {"start_week": row[0] or 1, "start_index": row[1] or 1}
        try:
            menu_sheets_db = json.loads(row[2] or "[]")
        except Exception:
            menu_sheets_db = []
    menu_sheets = menu_sheets_db or (session.get("menu_sheets") or [])
    start_index = (s.get("start_index") if isinstance(s, dict) else None) or 1
    start_week = (s.get("start_week") if isinstance(s, dict) else None) or 1
    weekday_no_to_nor = {0: "Mandag", 1: "Tirsdag", 2: "Onsdag", 3: "Torsdag", 4: "Fredag", 5: "Lørdag", 6: "Søndag"}
    def menu_index_for_date(dt_date):
        iso_year, iso_week, _ = dt_date.isocalendar()
        idx = ((iso_week - int(start_week) + (int(start_index) - 1)) % 4) + 1
        return idx
    sheets_by_uke = {str(sh.get("uke")): sh for sh in menu_sheets}

    def menykort_for_period(p):
        d0 = p["start"]
        d1 = d0 + timedelta(days=13)
        menus_by_day = {}
        day_list = [d0 + timedelta(days=i) for i in range(14)]
        for day_date in day_list:
            day_key = day_date.strftime("%Y-%m-%d")
            idx = menu_index_for_date(day_date)
            sh = sheets_by_uke.get(str(idx))
            lunsj_list = []
            middag_list = []
            if sh:
                dn = weekday_no_to_nor[day_date.weekday()]
                lunsj_list = [it["rett"] for it in (sh.get("lunsj") or []) if it.get("dag") == dn]
                middag_list = [it["rett"] for it in (sh.get("middag") or []) if it.get("dag") == dn]
            menus_by_day[day_key] = {"index": idx, "lunsj": lunsj_list, "middag": middag_list}
        # dagvakt menykort
        out = []
        d8 = d0 + timedelta(days=7)
        k8 = d8.strftime("%Y-%m-%d")
        if k8 in menus_by_day:
            out.append({"date": k8, "weekday": d8.strftime("%A"), "lunsj": [], "middag": menus_by_day[k8]["middag"] or [], "label": "Middag"})
        for off in range(8, 14):
            dd = d0 + timedelta(days=off)
            kk = dd.strftime("%Y-%m-%d")
            v = menus_by_day.get(kk) or {"lunsj": [], "middag": []}
            out.append({"date": kk, "weekday": dd.strftime("%A"), "lunsj": v.get("lunsj") or [], "middag": v.get("middag") or [], "label": "Lunsj + Middag"})
        d14 = d0 + timedelta(days=14)
        idx14 = menu_index_for_date(d14)
        sh14 = sheets_by_uke.get(str(idx14))
        lunsj14 = []
        if sh14:
            dn14 = weekday_no_to_nor[d14.weekday()]
            lunsj14 = [it["rett"] for it in (sh14.get("lunsj") or []) if it.get("dag") == dn14]
        out.append({"date": d14.strftime("%Y-%m-%d"), "weekday": d14.strftime("%A"), "lunsj": lunsj14, "middag": [], "label": "Lunsj"})
        return out

    # Pagination: allow viewing further periods via ?offset= (0-based index into periods list)
    try:
        offset = int(request.args.get("offset", "0") or 0)
        if offset < 0:
            offset = 0
    except ValueError:
        offset = 0
    page_size = 6  # show 6 periods (brukarens önskan: se kommande 6 turer)
    total_periods = len(periods)
    slice_start = offset
    slice_end = offset + page_size
    periods_slice = periods[slice_start:slice_end]
    upcoming = []
    for p in periods_slice:
        upcoming.append({
            "start": p["start"].strftime("%Y-%m-%d"),
            "end": (p["start"] + timedelta(days=13)).strftime("%Y-%m-%d"),
            "menykort": menykort_for_period(p)
        })
    has_prev = slice_start > 0
    has_next = slice_end < total_periods
    prev_offset = max(0, slice_start - page_size)
    next_offset = slice_end if has_next else slice_start

    return render_template(
        "prep_upcoming.html",
        upcoming=upcoming,
        offset=offset,
        page_size=page_size,
        total_periods=total_periods,
        has_prev=has_prev,
        has_next=has_next,
        prev_offset=prev_offset,
        next_offset=next_offset,
    )


# --- Prep tasks (private per bruker) ---
@app.get("/prep/tasks")
def prep_tasks_list():
    if not session.get("user_id"):
        return jsonify({"ok": False, "error": "Auth"}), 401
    user_id = session["user_id"]
    db = get_db()
    _ensure_prep_tables(db)
    user = db.execute("SELECT id, rig_id FROM users WHERE id=?", (user_id,)).fetchone()
    if not user or not user["rig_id"]:
        return jsonify({"ok": False, "error": "Ingen rigg"}), 400
    date = request.args.get("date", "").strip()
    meal = request.args.get("meal", "").strip()
    category = request.args.get("category", "").strip()
    if not (date and meal and category):
        return jsonify({"ok": False, "error": "Parametre mangler"}), 400
    rows = db.execute(
        """SELECT id, text, done FROM prep_tasks_private
            WHERE user_id=? AND rig_id=? AND date=? AND meal=? AND category=?
            ORDER BY done ASC, id ASC""",
        (user_id, user["rig_id"], date, meal, category),
    ).fetchall()
    tasks = [dict(id=r["id"], text=r["text"], done=bool(r["done"])) for r in rows]
    return jsonify({"ok": True, "tasks": tasks})


@app.post("/prep/tasks/add")
def prep_tasks_add():
    if not session.get("user_id"):
        return jsonify({"ok": False, "error": "Auth"}), 401
    user_id = session["user_id"]
    db = get_db()
    _ensure_prep_tables(db)
    user = db.execute("SELECT id, rig_id FROM users WHERE id=?", (user_id,)).fetchone()
    if not user or not user["rig_id"]:
        return jsonify({"ok": False, "error": "Ingen rigg"}), 400
    try:
        payload = request.get_json(force=True, silent=True) or {}
    except Exception:
        payload = {}
    date = (payload.get("date") or "").strip()
    meal = (payload.get("meal") or "").strip()
    category = (payload.get("category") or "").strip()
    text = (payload.get("text") or "").strip()
    if not text:
        return jsonify({"ok": False, "error": "Tom tekst"}), 400
    if not (date and meal and category):
        return jsonify({"ok": False, "error": "Parametre mangler"}), 400
    if meal not in ("lunsj", "middag"):
        return jsonify({"ok": False, "error": "Ugyldig måltid"}), 400
    if category not in ("soppa", "fisk", "kott", "extra"):
        return jsonify({"ok": False, "error": "Ugyldig kategori"}), 400
    db.execute(
        """INSERT INTO prep_tasks_private(user_id, rig_id, date, meal, category, text)
               VALUES(?,?,?,?,?,?)""",
        (user_id, user["rig_id"], date, meal, category, text),
    )
    db.commit()
    return jsonify({"ok": True})


@app.post("/prep/tasks/toggle")
def prep_tasks_toggle():
    if not session.get("user_id"):
        return jsonify({"ok": False, "error": "Auth"}), 401
    user_id = session["user_id"]
    db = get_db()
    _ensure_prep_tables(db)
    user = db.execute("SELECT id, rig_id FROM users WHERE id=?", (user_id,)).fetchone()
    if not user or not user["rig_id"]:
        return jsonify({"ok": False, "error": "Ingen rigg"}), 400
    try:
        payload = request.get_json(force=True, silent=True) or {}
    except Exception:
        payload = {}
    task_id = payload.get("id")
    if not task_id:
        return jsonify({"ok": False, "error": "Mangler id"}), 400
    row = db.execute(
        "SELECT id, done FROM prep_tasks_private WHERE id=? AND user_id=? AND rig_id=?",
        (task_id, user_id, user["rig_id"]),
    ).fetchone()
    if not row:
        return jsonify({"ok": False, "error": "Fant ikke"}), 404
    new_done = 0 if row["done"] else 1
    db.execute("UPDATE prep_tasks_private SET done=? WHERE id=?", (new_done, task_id))
    db.commit()
    return jsonify({"ok": True, "done": bool(new_done)})


# --- Frysplock endpoints (private) ---
@app.get("/frys/items")
def frys_items_list():
    if not session.get("user_id"):
        return jsonify({"ok": False, "error": "Auth"}), 401
    user_id = session["user_id"]
    db = get_db(); _ensure_prep_tables(db)
    user = db.execute("SELECT id, rig_id FROM users WHERE id=?", (user_id,)).fetchone()
    if not user or not user["rig_id"]:
        return jsonify({"ok": False, "error": "Ingen rigg"}), 400
    date = request.args.get("date", "").strip()
    meal = request.args.get("meal", "").strip()
    category = request.args.get("category", "").strip()
    if not (date and meal and category):
        return jsonify({"ok": False, "error": "Parametre mangler"}), 400
    rows = db.execute(
        """SELECT id, item_name, qty, unit, done FROM frys_items_private
            WHERE user_id=? AND rig_id=? AND date=? AND meal=? AND category=?
            ORDER BY done ASC, id ASC""",
        (user_id, user["rig_id"], date, meal, category),
    ).fetchall()
    items = [dict(id=r["id"], item_name=r["item_name"], qty=r["qty"], unit=r["unit"], done=bool(r["done"])) for r in rows]
    return jsonify({"ok": True, "items": items})

@app.post("/frys/items/add")
def frys_items_add():
    if not session.get("user_id"):
        return jsonify({"ok": False, "error": "Auth"}), 401
    user_id = session["user_id"]
    db = get_db(); _ensure_prep_tables(db)
    user = db.execute("SELECT id, rig_id FROM users WHERE id=?", (user_id,)).fetchone()
    if not user or not user["rig_id"]:
        return jsonify({"ok": False, "error": "Ingen rigg"}), 400
    payload = request.get_json(silent=True) or {}
    date = (payload.get("date") or "").strip()
    meal = (payload.get("meal") or "").strip()
    category = (payload.get("category") or "").strip()
    item_name = (payload.get("item_name") or "").strip()
    qty = payload.get("qty")
    unit = (payload.get("unit") or "kg").strip() or "kg"
    if not item_name:
        return jsonify({"ok": False, "error": "Tomt namn"}), 400
    if not (date and meal and category):
        return jsonify({"ok": False, "error": "Parametre mangler"}), 400
    try:
        if qty is not None:
            qty = float(qty)
    except Exception:
        return jsonify({"ok": False, "error": "Ugyldig mengde"}), 400
    db.execute(
        """INSERT INTO frys_items_private(user_id, rig_id, date, meal, category, item_name, qty, unit)
               VALUES(?,?,?,?,?,?,?,?)""",
        (user_id, user["rig_id"], date, meal, category, item_name, qty, unit),
    )
    db.commit(); return jsonify({"ok": True})

@app.post("/frys/items/toggle")
def frys_items_toggle():
    if not session.get("user_id"):
        return jsonify({"ok": False, "error": "Auth"}), 401
    user_id = session["user_id"]
    db = get_db(); _ensure_prep_tables(db)
    user = db.execute("SELECT id, rig_id FROM users WHERE id=?", (user_id,)).fetchone()
    if not user or not user["rig_id"]:
        return jsonify({"ok": False, "error": "Ingen rigg"}), 400
    payload = request.get_json(silent=True) or {}
    item_id = payload.get("id")
    if not item_id:
        return jsonify({"ok": False, "error": "Mangler id"}), 400
    row = db.execute(
        "SELECT id, done FROM frys_items_private WHERE id=? AND user_id=? AND rig_id=?",
        (item_id, user_id, user["rig_id"]),
    ).fetchone()
    if not row:
        return jsonify({"ok": False, "error": "Fant ikke"}), 404
    new_done = 0 if row["done"] else 1
    db.execute("UPDATE frys_items_private SET done=? WHERE id=?", (new_done, item_id))
    db.commit(); return jsonify({"ok": True, "done": bool(new_done)})


# Batch status for indicators
@app.post("/prep/status/batch")
def prep_status_batch():
    if not session.get("user_id"):
        return jsonify({"ok": False, "error": "Auth"}), 401
    user_id = session["user_id"]
    db = get_db(); _ensure_prep_tables(db)
    user = db.execute("SELECT id, rig_id FROM users WHERE id=?", (user_id,)).fetchone()
    if not user or not user["rig_id"]:
        return jsonify({"ok": False, "error": "Ingen rigg"}), 400
    payload = request.get_json(silent=True) or {}
    items = payload.get("items") or []
    rig_id = user["rig_id"]
    out = {}
    for it in items:
        try:
            date = it.get("date"); meal = it.get("meal"); category = it.get("category")
        except Exception:
            continue
        if not (date and meal and category):
            continue
        key = f"{date}|{meal}|{category}"
        # counts
        prep_row = db.execute(
            "SELECT COUNT(*) AS total, SUM(CASE WHEN done=0 THEN 1 ELSE 0 END) AS open FROM prep_tasks_private WHERE user_id=? AND rig_id=? AND date=? AND meal=? AND category=?",
            (user_id, rig_id, date, meal, category),
        ).fetchone()
        frys_row = db.execute(
            "SELECT COUNT(*) AS total, SUM(CASE WHEN done=0 THEN 1 ELSE 0 END) AS open FROM frys_items_private WHERE user_id=? AND rig_id=? AND date=? AND meal=? AND category=?",
            (user_id, rig_id, date, meal, category),
        ).fetchone()
        out[key] = {
            "prep_total": prep_row["total"],
            "prep_open": prep_row["open"] or 0,
            "frys_total": frys_row["total"],
            "frys_open": frys_row["open"] or 0,
        }
    return jsonify({"ok": True, "statuses": out})


# --- Year calendar (stub) ---
@app.get("/me/calendar")
def me_calendar():
    db = get_db()
    if not session.get("user_id"):
        flash("Logg inn kreves.", "warning")
        return redirect(url_for("login"))
    rig_id, rig_name = _current_rig(db)
    user_id = session["user_id"]
    year = datetime.now().year
    start = datetime(year, 1, 1)
    end = datetime(year, 12, 31, 23, 59)
    rows = db.execute(
        """
        SELECT ts.start_ts, ts.end_ts
        FROM turnus_slots ts
        INNER JOIN turnus_account_binding tb ON tb.slot_id = ts.id
        WHERE ts.rig_id = ? AND tb.user_id = ? AND ts.end_ts >= ? AND ts.start_ts <= ?
        """,
        (rig_id, user_id, start.strftime("%Y-%m-%dT%H:%M"), end.strftime("%Y-%m-%dT%H:%M")),
    ).fetchall()
    # Collect days with any coverage
    covered = set()
    for r in rows:
        st = _parse_ts(r["start_ts"]).date()
        en = _parse_ts(r["end_ts"]).date()
        cur = st
        while cur <= en:
            covered.add(cur.isoformat())
            cur += timedelta(days=1)
    return render_template("calendar.html", year=year, covered=covered, rig_name=rig_name)


# --- Single day drilldown: menu + prepp/frys for dato ---
@app.get("/day/<date_s>")
def day_detail(date_s: str):
    """Viser meny (med private overrides) samt prepp- og frysplocklinjer for gitt dato for innlogget bruker."""
    if not session.get("user_id"):
        flash("Logg inn kreves.", "warning")
        return redirect(url_for("login"))
    # Valider datoformat
    try:
        dt = datetime.strptime(date_s, "%Y-%m-%d").date()
    except Exception:
        flash("Ugyldig dato.", "warning")
        return redirect(url_for("me_calendar"))
    db = get_db(); _ensure_prep_tables(db)
    user = db.execute("SELECT id, rig_id FROM users WHERE id=?", (session["user_id"],)).fetchone()
    if not user or not user["rig_id"]:
        flash("Ingen rigg tilknyttet.", "warning")
        return redirect(url_for("dashboard"))
    rig_id = user["rig_id"]
    # Hent menyinnstillinger + sheets
    row = db.execute("SELECT start_week, start_index, menu_json FROM menu_settings WHERE rig_id=?", (rig_id,)).fetchone()
    if row:
        start_week = int(row[0] or 1); start_index = int(row[1] or 1)
        try:
            menu_sheets = json.loads(row[2] or "[]")
        except Exception:
            menu_sheets = []
    else:
        start_week = 1; start_index = 1; menu_sheets = []
    sheets_by_uke = {str(sh.get("uke")): sh for sh in menu_sheets}
    weekday_no_to_nor = {0: "Mandag",1:"Tirsdag",2:"Onsdag",3:"Torsdag",4:"Fredag",5:"Lørdag",6:"Søndag"}
    def menu_index_for_date(d):
        iso_year, iso_week, _ = d.isocalendar()
        return ((iso_week - start_week + (start_index - 1)) % 4) + 1
    def _norm_cat(raw: str):
        if not raw: return None
        r = raw.strip().lower()
        if "sopp" in r or "soup" in r: return "soppa"
        if "fisk" in r or "fish" in r: return "fisk"
        if "kjott" in r or "kjot" in r or "kott" in r or "meat" in r: return "kott"
        return None
    def base_picks(meal_key: str):
        idx = menu_index_for_date(dt)
        sh = sheets_by_uke.get(str(idx))
        dn = weekday_no_to_nor[dt.weekday()]
        cats = {"soppa": [], "fisk": [], "kott": [], "extra": []}
        if sh:
            items = (sh.get(meal_key) or [])
            for it in items:
                if it.get("dag") != dn: continue
                cat = _norm_cat(it.get("kategori"))
                if cat in ("soppa","fisk","kott"):
                    cats[cat].append(it.get("rett") or "")
                else:
                    cats["extra"].append(it.get("rett") or "")
        return {k: (v[0] if v else "") for k,v in cats.items()}
    def apply_private_overrides(meal_key: str, picks: dict):
        out = {k: picks.get(k) or "" for k in ["soppa","fisk","kott"]}; out["extra"] = ""
        try:
            rows_priv = db.execute("SELECT category,dish FROM daily_menu_overrides_private WHERE rig_id=? AND user_id=? AND date=? AND meal=?", (rig_id, user["id"], date_s, meal_key)).fetchall()
            for r in rows_priv:
                if r[0] in out and r[1]: out[r[0]] = r[1]
        except Exception:
            pass
        return out
    base_l = base_picks("lunsj"); base_d = base_picks("middag")
    picks_l = apply_private_overrides("lunsj", base_l)
    picks_d = apply_private_overrides("middag", base_d)
    # Prep/frys per kategori
    def fetch_tasks(meal, cat):
        rows = db.execute("SELECT id,text,done FROM prep_tasks_private WHERE user_id=? AND rig_id=? AND date=? AND meal=? AND category=? ORDER BY done ASC, id ASC", (user["id"], rig_id, date_s, meal, cat)).fetchall()
        return [dict(id=r["id"], text=r["text"], done=bool(r["done"])) for r in rows]
    def fetch_frys(meal, cat):
        rows = db.execute("SELECT id,item_name,qty,unit,done FROM frys_items_private WHERE user_id=? AND rig_id=? AND date=? AND meal=? AND category=? ORDER BY done ASC, id ASC", (user["id"], rig_id, date_s, meal, cat)).fetchall()
        return [dict(id=r["id"], item_name=r["item_name"], qty=r["qty"], unit=r["unit"], done=bool(r["done"])) for r in rows]
    cat_order = ["soppa","fisk","kott","extra"]
    meals = []
    for meal_key, picks, base in [("lunsj", picks_l, base_l), ("middag", picks_d, base_d)]:
        cats = []
        for c in cat_order:
            dish = picks.get(c) or ""
            if not dish and c != "extra":
                # still include empty categories for consistency
                pass
            tasks = fetch_tasks(meal_key, c)
            frys_items = fetch_frys(meal_key, c)
            cats.append({
                "category": c,
                "dish": dish,
                "overridden": (dish and dish != (base.get(c) or "")),
                "base": base.get(c) or "",
                "tasks": tasks,
                "frys": frys_items,
            })
        meals.append({"meal": meal_key, "categories": cats})
    return render_template("day_detail.html", date=date_s, weekday=weekday_no_to_nor[dt.weekday()], meals=meals)


# --- Day notes (private) ---
def _ensure_day_notes(db):
    try:
        db.execute(
            """CREATE TABLE IF NOT EXISTS daily_notes_private (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                rig_id INTEGER NOT NULL,
                user_id INTEGER NOT NULL,
                date TEXT NOT NULL,
                note TEXT,
                updated_at TEXT,
                UNIQUE(rig_id,user_id,date)
            )"""
        )
    except Exception:
        pass

@app.get("/day/note")
def day_note_get():
    if not session.get("user_id"):
        return jsonify({"ok": False, "error": "Auth"}), 401
    date_s = request.args.get("date","").strip()
    if not date_s:
        return jsonify({"ok": False, "error": "date required"}), 400
    db = get_db(); _ensure_day_notes(db)
    user = db.execute("SELECT id, rig_id FROM users WHERE id=?", (session["user_id"],)).fetchone()
    if not user or not user["rig_id"]:
        return jsonify({"ok": False, "error": "Ingen rigg"}), 400
    row = db.execute("SELECT note FROM daily_notes_private WHERE rig_id=? AND user_id=? AND date=?", (user["rig_id"], user["id"], date_s)).fetchone()
    return jsonify({"ok": True, "note": (row["note"] if row else "")})

@app.post("/day/note")
def day_note_set():
    if not session.get("user_id"):
        return jsonify({"ok": False, "error": "Auth"}), 401
    try:
        payload = request.get_json(force=True, silent=True) or {}
    except Exception:
        payload = {}
    date_s = (payload.get("date") or "").strip()
    note = (payload.get("note") or "").strip()
    if not date_s:
        return jsonify({"ok": False, "error": "date required"}), 400
    db = get_db(); _ensure_day_notes(db)
    user = db.execute("SELECT id, rig_id FROM users WHERE id=?", (session["user_id"],)).fetchone()
    if not user or not user["rig_id"]:
        return jsonify({"ok": False, "error": "Ingen rigg"}), 400
    from datetime import datetime as _dt
    ts = _dt.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")
    db.execute(
        "INSERT INTO daily_notes_private(rig_id,user_id,date,note,updated_at) VALUES(?,?,?,?,?) ON CONFLICT(rig_id,user_id,date) DO UPDATE SET note=excluded.note, updated_at=excluded.updated_at",
        (user["rig_id"], user["id"], date_s, note, ts)
    )
    db.commit()
    return jsonify({"ok": True, "saved_at": ts})


# --- Week prep aggregation (private) ---
@app.get("/week/<date_s>")
def week_prep_view(date_s: str):
    """Aggregert veckovy: viser meny (med private overrides) + prepp og frysplock for en uke (mandag-søndag)
    som inneholder gitt dato. Toggle skjer via eksisterende endpoints (prep_tasks_toggle / frys_items_toggle)."""
    if not session.get("user_id"):
        flash("Logg inn kreves.", "warning")
        return redirect(url_for("login"))
    try:
        base_date = datetime.strptime(date_s, "%Y-%m-%d").date()
    except Exception:
        flash("Ugyldig dato.", "warning")
        return redirect(url_for("me_calendar"))
    db = get_db(); _ensure_prep_tables(db)
    user = db.execute("SELECT id, rig_id FROM users WHERE id=?", (session["user_id"],)).fetchone()
    if not user or not user["rig_id"]:
        flash("Ingen rigg tilknyttet.", "warning")
        return redirect(url_for("dashboard"))
    rig_id = user["rig_id"]
    # Menu settings
    row = db.execute("SELECT start_week, start_index, menu_json FROM menu_settings WHERE rig_id=?", (rig_id,)).fetchone()
    if row:
        start_week = int(row[0] or 1); start_index = int(row[1] or 1)
        try:
            menu_sheets = json.loads(row[2] or "[]")
        except Exception:
            menu_sheets = []
    else:
        start_week = 1; start_index = 1; menu_sheets = []
    sheets_by_uke = {str(sh.get("uke")): sh for sh in menu_sheets}
    weekday_no_to_nor = {0: "Mandag",1:"Tirsdag",2:"Onsdag",3:"Torsdag",4:"Fredag",5:"Lørdag",6:"Søndag"}
    def menu_index_for_date(d):
        iso_year, iso_week, _ = d.isocalendar(); return ((iso_week - start_week + (start_index - 1)) % 4) + 1
    def _norm_cat(raw: str):
        if not raw: return None
        r = raw.strip().lower()
        if "sopp" in r or "soup" in r: return "soppa"
        if "fisk" in r or "fish" in r: return "fisk"
        if "kjott" in r or "kjot" in r or "kott" in r or "meat" in r: return "kott"
        return None
    def base_picks(day_date, meal_key: str):
        idx = menu_index_for_date(day_date)
        sh = sheets_by_uke.get(str(idx))
        dn = weekday_no_to_nor[day_date.weekday()]
        cats = {"soppa": [], "fisk": [], "kott": [], "extra": []}
        if sh:
            items = (sh.get(meal_key) or [])
            for it in items:
                if it.get("dag") != dn: continue
                cat = _norm_cat(it.get("kategori"))
                if cat in ("soppa","fisk","kott"):
                    cats[cat].append(it.get("rett") or "")
                else:
                    cats["extra"].append(it.get("rett") or "")
        return {k: (v[0] if v else "") for k,v in cats.items()}
    def apply_private_overrides(day_key: str, meal_key: str, picks: dict):
        out = {k: picks.get(k) or "" for k in ["soppa","fisk","kott"]}; out["extra"] = ""
        try:
            rows_priv = db.execute("SELECT category,dish FROM daily_menu_overrides_private WHERE rig_id=? AND user_id=? AND date=? AND meal=?", (rig_id, user["id"], day_key, meal_key)).fetchall()
            for r in rows_priv:
                if r[0] in out and r[1]: out[r[0]] = r[1]
        except Exception:
            pass
        return out
    def fetch_tasks(day_key, meal, cat):
        rows = db.execute("SELECT id,text,done FROM prep_tasks_private WHERE user_id=? AND rig_id=? AND date=? AND meal=? AND category=? ORDER BY done ASC, id ASC", (user["id"], rig_id, day_key, meal, cat)).fetchall()
        return [dict(id=r["id"], text=r["text"], done=bool(r["done"])) for r in rows]
    def fetch_frys(day_key, meal, cat):
        rows = db.execute("SELECT id,item_name,qty,unit,done FROM frys_items_private WHERE user_id=? AND rig_id=? AND date=? AND meal=? AND category=? ORDER BY done ASC, id ASC", (user["id"], rig_id, day_key, meal, cat)).fetchall()
        return [dict(id=r["id"], item_name=r["item_name"], qty=r["qty"], unit=r["unit"], done=bool(r["done"])) for r in rows]
    # Compute monday start
    monday = base_date - timedelta(days=base_date.weekday())
    days = []
    for i in range(7):
        dd = monday + timedelta(days=i)
        day_key = dd.strftime("%Y-%m-%d")
        base_l = base_picks(dd, "lunsj"); base_d = base_picks(dd, "middag")
        picks_l = apply_private_overrides(day_key, "lunsj", base_l)
        picks_d = apply_private_overrides(day_key, "middag", base_d)
        cat_order = ["soppa","fisk","kott","extra"]
        meals = []
        for meal_key, picks, base in [("lunsj", picks_l, base_l), ("middag", picks_d, base_d)]:
            cats = []
            for c in cat_order:
                dish = picks.get(c) or ""
                tasks = fetch_tasks(day_key, meal_key, c)
                frys_items = fetch_frys(day_key, meal_key, c)
                cats.append({
                    "category": c,
                    "dish": dish,
                    "overridden": (dish and dish != (base.get(c) or "")),
                    "base": base.get(c) or "",
                    "tasks": tasks,
                    "frys": frys_items,
                    "day_key": day_key,
                    "meal": meal_key
                })
            meals.append({"meal": meal_key, "categories": cats})
        days.append({"date": day_key, "weekday": weekday_no_to_nor[dd.weekday()], "meals": meals})
    prev_week = (monday - timedelta(days=7)).strftime("%Y-%m-%d")
    next_week = (monday + timedelta(days=7)).strftime("%Y-%m-%d")
    return render_template("week_prep.html", week_start=monday.strftime("%Y-%m-%d"), days=days, prev_week=prev_week, next_week=next_week, anchor_date=date_s)


@app.get("/period/<date_s>/aggregate")
def period_aggregate(date_s: str):
    """Flat aggregat for hele 14-dagers arbeidsperioden som inneholder datoen.
    Henter samme periodelogikk som dashboard: grupper slots til perioder og velger den som inneholder eller kommer etter dato.
    Returnerer to lister (prep_rows, frys_rows) i template med kontekst (dato, ukedag, meal, kategori, dish, tekst)."""
    if not session.get("user_id"):
        flash("Logg inn kreves.", "warning"); return redirect(url_for("login"))
    try:
        anchor_date = datetime.strptime(date_s, "%Y-%m-%d").date()
    except Exception:
        flash("Ugyldig dato.", "warning"); return redirect(url_for("dashboard"))
    db = get_db(); _ensure_prep_tables(db)
    user = db.execute("SELECT id, rig_id FROM users WHERE id=?", (session["user_id"],)).fetchone()
    if not user or not user["rig_id"]:
        flash("Ingen rigg tilknyttet.", "warning"); return redirect(url_for("dashboard"))
    rig_id = user["rig_id"]
    # Hent slots fra i dag minus litt buffer til fremover for å bygge perioder
    all_rows = db.execute(
        """
        SELECT ts.id, ts.start_ts, ts.end_ts
        FROM turnus_slots ts
        INNER JOIN turnus_account_binding tb ON tb.slot_id = ts.id
        WHERE ts.rig_id=? AND tb.user_id=?
        ORDER BY ts.start_ts ASC
        """, (rig_id, user["id"])
    ).fetchall()
    def parse_ts(s):
        try: return datetime.strptime(s, "%Y-%m-%dT%H:%M")
        except Exception: return None
    periods = []
    cur = None
    for r in all_rows:
        st = parse_ts(r["start_ts"]); et = parse_ts(r["end_ts"])
        if not (st and et):
            continue
        if cur is None:
            cur = {"start": st, "end": et, "rows": [r]}
            continue
        anchor = cur["start"]
        if st <= anchor + timedelta(days=14, hours=1):
            cur["rows"].append(r)
            if et > cur["end"]: cur["end"] = et
        else:
            periods.append(cur); cur = {"start": st, "end": et, "rows": [r]}
    if cur is not None:
        periods.append(cur)
    chosen = None
    for p in periods:
        if p["start"].date() <= anchor_date <= p["end"].date():
            chosen = p; break
        if anchor_date < p["start"].date() and chosen is None:
            chosen = p; break
    if not chosen:
        flash("Fant ingen periodedata.", "info"); return render_template("period_aggregate.html", prep_rows=[], frys_rows=[], period_start=None, period_end=None, anchor_date=date_s)
    d0 = chosen["start"].date(); d1 = d0 + timedelta(days=13)
    # Menyinnstillinger for dish-navn/overrides
    row = db.execute("SELECT start_week, start_index, menu_json FROM menu_settings WHERE rig_id=?", (rig_id,)).fetchone()
    if row:
        start_week = int(row[0] or 1); start_index = int(row[1] or 1)
        try: menu_sheets = json.loads(row[2] or "[]")
        except Exception: menu_sheets = []
    else:
        start_week = 1; start_index = 1; menu_sheets = []
    sheets_by_uke = {str(sh.get("uke")): sh for sh in menu_sheets}
    weekday_no_to_short = {0:"Ma",1:"Ti",2:"On",3:"To",4:"Fr",5:"Lø",6:"Sø"}
    weekday_no_to_full = {0:"Mandag",1:"Tirsdag",2:"Onsdag",3:"Torsdag",4:"Fredag",5:"Lørdag",6:"Søndag"}
    def menu_index_for_date(d):
        iso_year, iso_week, _ = d.isocalendar(); return ((iso_week - start_week + (start_index - 1)) % 4) + 1
    def _norm_cat(raw: str):
        if not raw: return None
        r = raw.strip().lower()
        if "sopp" in r or "soup" in r: return "soppa"
        if "fisk" in r or "fish" in r: return "fisk"
        if "kjott" in r or "kjot" in r or "kott" in r or "meat" in r: return "kott"
        return None
    def base_picks(day_date, meal_key: str):
        idx = menu_index_for_date(day_date)
        sh = sheets_by_uke.get(str(idx))
        dn = weekday_no_to_full[day_date.weekday()]
        cats = {"soppa": [], "fisk": [], "kott": [], "extra": []}
        if sh:
            for it in (sh.get(meal_key) or []):
                if it.get("dag") != dn: continue
                cat = _norm_cat(it.get("kategori"))
                if cat in ("soppa","fisk","kott"): cats[cat].append(it.get("rett") or "")
                else: cats["extra"].append(it.get("rett") or "")
        return {k:(v[0] if v else "") for k,v in cats.items()}
    def apply_private_overrides(day_key, meal_key, picks):
        out = {k: picks.get(k) or "" for k in ["soppa","fisk","kott"]}; out["extra"]=""
        try:
            rows_priv = db.execute("SELECT category,dish FROM daily_menu_overrides_private WHERE rig_id=? AND user_id=? AND date=? AND meal=?", (rig_id, user["id"], day_key, meal_key)).fetchall()
            for r in rows_priv:
                if r[0] in out and r[1]: out[r[0]] = r[1]
        except Exception: pass
        return out
    prep_rows = []
    frys_rows = []
    cat_order = ["soppa","fisk","kott","extra"]
    curd = d0
    while curd <= d1:
        day_key = curd.strftime("%Y-%m-%d")
        base_l = base_picks(curd, "lunsj"); base_d = base_picks(curd, "middag")
        picks_l = apply_private_overrides(day_key, "lunsj", base_l)
        picks_d = apply_private_overrides(day_key, "middag", base_d)
        # Fetch tasks / frys for alle kategorier begge måltider
        for meal_key, picks, base in [("lunsj", picks_l, base_l), ("middag", picks_d, base_d)]:
            for c in cat_order:
                dish = picks.get(c) or ""
                # Prep
                for r in db.execute("SELECT id,text,done FROM prep_tasks_private WHERE user_id=? AND rig_id=? AND date=? AND meal=? AND category=? ORDER BY id", (user["id"], rig_id, day_key, meal_key, c)).fetchall():
                    prep_rows.append({
                        "id": r["id"], "date": day_key, "weekday": weekday_no_to_short[curd.weekday()], "meal": meal_key,
                        "category": c, "dish": dish or (base.get(c) or ""), "text": r["text"], "done": bool(r["done"])
                    })
                # Frys
                for r in db.execute("SELECT id,item_name,qty,unit,done FROM frys_items_private WHERE user_id=? AND rig_id=? AND date=? AND meal=? AND category=? ORDER BY id", (user["id"], rig_id, day_key, meal_key, c)).fetchall():
                    frys_rows.append({
                        "id": r["id"], "date": day_key, "weekday": weekday_no_to_short[curd.weekday()], "meal": meal_key,
                        "category": c, "dish": dish or (base.get(c) or ""), "item_name": r["item_name"], "qty": r["qty"], "unit": r["unit"], "done": bool(r["done"])
                    })
        curd += timedelta(days=1)
    # Sort
    def sort_key_p(x): return (x["date"], x["meal"], {"soppa":0,"fisk":1,"kott":2,"extra":3}.get(x["category"],9), x["id"])
    prep_rows.sort(key=sort_key_p)
    frys_rows.sort(key=sort_key_p)
    return render_template("period_aggregate.html", prep_rows=prep_rows, frys_rows=frys_rows, period_start=d0.strftime("%Y-%m-%d"), period_end=d1.strftime("%Y-%m-%d"), anchor_date=date_s)


# --- Unified planning hub ---
@app.get("/planning")
def planning_hub():
    if not session.get("user_id"):
        flash("Logg inn kreves.", "warning"); return redirect(url_for("login"))
    db = get_db(); user = db.execute("SELECT id, rig_id FROM users WHERE id=?", (session["user_id"],)).fetchone()
    rig_id = user["rig_id"] if user else None
    # Reuse dashboard period detection (lightweight): find arbeidsperiod_start via existing helper logic by calling dashboard core calculation indirectly would be heavy; replicate minimal part.
    arbeidsperiod_start = None
    anchor_user_date = None
    menu_days_period = []
    today_str = datetime.now().strftime("%Y-%m-%d")
    if rig_id:
        try:
            today_ts = today_str + "T00:00"
            rows = db.execute("""SELECT ts.start_ts, ts.end_ts FROM turnus_slots ts INNER JOIN turnus_account_binding tb ON tb.slot_id=ts.id WHERE ts.rig_id=? AND ts.status='published' AND tb.user_id=? ORDER BY ts.start_ts ASC""", (rig_id, user["id"])).fetchall()
            # reuse grouping logic
            periods=[]; cur=None
            for r in rows:
                st=_parse_ts(r["start_ts"]); et=_parse_ts(r["end_ts"])
                if cur is None:
                    cur={"start":st,"end":et,"rows":[r]}; continue
                if st <= cur["start"] + timedelta(days=14, hours=1):
                    cur["rows"].append(r); cur["end"]=max(cur["end"], et)
                else:
                    periods.append(cur); cur={"start":st,"end":et,"rows":[r]}
            if cur: periods.append(cur)
            now=datetime.now(); chosen=None
            for p in periods:
                if p["start"] <= now <= p["end"] or now < p["start"]:
                    chosen=p; break
            if chosen:
                d0=chosen["start"].date(); arbeidsperiod_start=d0.strftime("%Y-%m-%d")
                # minimal menu_days_period stub (first day only) just for anchor fallback
                menu_days_period=[{"date": arbeidsperiod_start}]
        except Exception:
            pass
    anchor = arbeidsperiod_start or anchor_user_date or (menu_days_period[0]["date"] if menu_days_period else today_str)
    return render_template("planning_hub.html", anchor=anchor, arbeidsperiod_start=arbeidsperiod_start, today_str=today_str)


# --- Recipes (stub) ---
@app.get("/recipes")
def recipes():
    return render_template("recipes.html")


# Legacy Turnus page now redirects
@app.get("/turnus/admin")
@admin_required
def turnus_admin_home():
    return redirect(url_for("turnus_simple"))


# --- Simple turnus generator ---
@app.get("/turnus/simple")
@admin_required
def turnus_simple():
    db = get_db()
    rig_id, rig_name = _current_rig(db)
    if not rig_id:
        flash("Din konto mangler tilknyttet rigg.", "warning")
        return redirect(url_for("admin_dashboard_protected"))
    return render_template("turnus_simple.html", rig_id=rig_id, rig_name=rig_name, base_slots=None, preview_slots=None, start_friday=None, preview_weeks=12)


@app.post("/turnus/simple/build_base")
@admin_required
def turnus_simple_build_base():
    db = get_db()
    rig_id, rig_name = _current_rig(db)
    if not rig_id:
        flash("Din konto mangler tilknyttet rigg.", "warning")
        return redirect(url_for("admin_dashboard_protected"))
    start_friday = (request.form.get("start_friday") or "").strip()
    preview_weeks = request.form.get("preview_weeks", type=int) or 12
    if not start_friday:
        flash("Mangler start-fredag.", "warning")
        return redirect(url_for("turnus_simple"))
    try:
        base_slots = rs.generate_two_week_block_for_cook(start_friday, "BASE")
        preview_slots = rs.generate_baseline_schedule_6_cooks(start_friday, weeks=preview_weeks)
        return render_template(
            "turnus_simple.html",
            rig_id=rig_id,
            rig_name=rig_name,
            base_slots=base_slots,
            preview_slots=preview_slots,
            start_friday=start_friday,
            preview_weeks=preview_weeks,
        )
    except Exception as e:
        flash(f"Feil ved bygging: {e}", "danger")
        return redirect(url_for("turnus_simple"))


@app.post("/turnus/simple/apply")
@admin_required
def turnus_simple_apply():
    db = get_db()
    rig_id, rig_name = _current_rig(db)
    if not rig_id:
        flash("Din konto mangler tilknyttet rigg.", "warning")
        return redirect(url_for("admin_dashboard_protected"))
    start_friday = (request.form.get("start_friday") or "").strip()
    weeks = request.form.get("weeks", type=int) or 12
    if not start_friday:
        flash("Mangler start-fredag.", "warning")
        return redirect(url_for("turnus_simple"))
    try:
        slots = rs.generate_baseline_schedule_6_cooks(start_friday, weeks=weeks)
        # persist
        # Resolve DB path via rotation.DB_PATH if available, otherwise app.db in CWD
        db_path = None
        try:
            dbp = getattr(rotation, "DB_PATH", None)
            db_path = dbp.as_posix() if dbp is not None else None
        except Exception:
            db_path = None
        if not db_path:
            db_path = os.path.join(os.getcwd(), "app.db")
        created = rs.persist_slots_sqlite(db_path=db_path, rig_id=rig_id, slots=slots, status="published")
        # Ensure binding table exists and auto-apply mapping if configured
        try:
            dbc = sqlite3.connect(db_path)
            dbc.execute(
                """
                CREATE TABLE IF NOT EXISTS turnus_account_binding (
                  id INTEGER PRIMARY KEY AUTOINCREMENT,
                  slot_id INTEGER NOT NULL,
                  user_id INTEGER NOT NULL,
                  notes TEXT,
                  bound_at TEXT DEFAULT (datetime('now'))
                )
                """
            )
            dbc.commit()
        finally:
            try:
                dbc.close()
            except Exception:
                pass
        try:
            bound = rotation.apply_virtual_mapping(rig_id)
            if bound:
                flash(f"Auto-koblet {bound} pass til brukere (virt.kokk → bruker).", "info")
        except Exception:
            # Mapping might not be configured yet; it's okay
            pass
        flash(f"Opprettet {created} slots for 6 virtuelle kokker (uker generert: {weeks}).", "success")
    except Exception as e:
        flash(f"Klarte ikke å skrive slots: {e}", "danger")
    return redirect(url_for("turnus_simple"))


# --- Admin Dashboard + Users ---
@app.get("/admin")
def admin_dashboard():
    return redirect(url_for("admin_dashboard_protected"))


@app.get("/admin/dashboard")
@admin_required
def admin_dashboard_protected():
    db = get_db()
    rigs = db.execute("SELECT id, name, description FROM rigs ORDER BY id").fetchall()
    rig_name = None
    if session.get("user_id"):
        row = db.execute("SELECT rig_id FROM users WHERE id=?", (session["user_id"],)).fetchone()
        if row and row["rig_id"]:
            rrow = db.execute("SELECT name FROM rigs WHERE id=?", (row["rig_id"],)).fetchone()
            rig_name = rrow["name"] if rrow else None
    users = db.execute("SELECT id, name, email, role, rig_id FROM users ORDER BY id").fetchall()
    return render_template("admin_dashboard.html", users=users, rigs=rigs, rig_name=rig_name)


def _parse_menu_csv(file_storage) -> list[dict]:
    # Returns a list of sheets: [{ 'uke': 'Uke 1', 'lunsj': [...], 'middag': [...] }]
    raw = file_storage.read()
    # Try common encodings
    text = None
    for enc in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            text = raw.decode(enc)
            break
        except Exception:
            continue
    if text is None:
        text = raw.decode("utf-8", errors="replace")
    # Detect delimiter (; or ,)
    sample = text.splitlines()[:5]
    has_semicolon = sum(1 for ln in sample if ";" in ln) > sum(1 for ln in sample if "," in ln)
    delimiter = ";" if has_semicolon else ","
    f = io.StringIO(text)
    reader = csv.DictReader(f, delimiter=delimiter)
    # Normalize headers
    def norm(k: str) -> str:
        return (k or "").strip().strip('"').lower()
    rows = []
    for r in reader:
        rr = {norm(k): (v.strip().strip('"') if isinstance(v, str) else v) for k, v in r.items()}
        uke = rr.get("uke") or rr.get("u") or rr.get("week")
        dag = rr.get("dag") or rr.get("day")
        maltid = rr.get("måltid") or rr.get("maltid") or rr.get("meal")
        kategori = rr.get("kategori") or rr.get("category")
        rett = rr.get("rett") or rr.get("dish")
        if not (uke and dag and maltid and kategori and rett):
            continue
        rows.append({
            "uke": str(uke),
            "dag": dag,
            "maltid": maltid,
            "kategori": kategori,
            "rett": rett,
        })
    # Group by uke and meal
    from collections import defaultdict
    by_uke = defaultdict(list)
    for r in rows:
        by_uke[r["uke"]].append(r)
    sheets: list[dict] = []
    # Keep uke order by sorted numeric if possible
    def uke_key(u):
        try:
            return int(u)
        except Exception:
            return 999999
    for uke in sorted(by_uke.keys(), key=uke_key):
        items = by_uke[uke]
        lunsj = [
            {"dag": it["dag"], "kategori": it["kategori"], "rett": it["rett"]}
            for it in items if (it["maltid"].lower().startswith("luns")) or (it["maltid"].lower().startswith("lun"))
        ]
        middag = [
            {"dag": it["dag"], "kategori": it["kategori"], "rett": it["rett"]}
            for it in items if it["maltid"].lower().startswith("mid")
        ]
        sheets.append({"uke": uke, "lunsj": lunsj, "middag": middag})
    return sheets


@app.route("/admin/menu", methods=["GET", "POST"])
@admin_required
def admin_menu():
    # Ensure settings table and columns exist
    db = get_db()
    rig_id, _ = _current_rig(db)
    try:
        db.execute(
            """
            CREATE TABLE IF NOT EXISTS menu_settings (
              rig_id INTEGER PRIMARY KEY,
              start_week INTEGER,
              start_index INTEGER,
              menu_json TEXT,
              updated_at TEXT DEFAULT (datetime('now')),
              start_date TEXT,
              last_upload_at TEXT,
              last_upload_name TEXT
            )
            """
        )
        # Backfill columns for existing DBs
        for col in ("start_date", "last_upload_at", "last_upload_name"):
            try:
                db.execute(f"ALTER TABLE menu_settings ADD COLUMN {col} TEXT")
            except Exception:
                pass
        db.commit()
    except Exception:
        pass

    # Read current settings from DB (source of truth across sessions)
    row = None
    if rig_id:
        try:
            row = db.execute(
                "SELECT start_week, start_index, menu_json, updated_at, start_date, last_upload_at, last_upload_name FROM menu_settings WHERE rig_id=?",
                (rig_id,),
            ).fetchone()
        except Exception:
            row = None

    # Build start info from DB, fallback to session
    startvecka_info = session.get("menu_start_info")
    if (not startvecka_info) and row:
        startvecka_info = {
            "start_week": row[0] or 1,
            "start_index": row[1] or 1,
            "start_date": (row[4] or ""),
        }

    # Compute current active week index (aktuell_uke)
    aktuell_uke = None
    try:
        from datetime import date
        iso_year, iso_week, _ = date.today().isocalendar()
        sw = int((startvecka_info or {}).get("start_week") or (row[0] if row else 1) or 1)
        si = int((startvecka_info or {}).get("start_index") or (row[1] if row else 1) or 1)
        aktuell_uke = ((iso_week - sw + (si - 1)) % 4) + 1
    except Exception:
        aktuell_uke = None

    # Load menus: prefer session (freshly uploaded this session), else DB
    menu_sheets = session.get("menu_sheets")
    if (not menu_sheets) and row:
        try:
            menu_sheets = json.loads(row[2] or "[]")
        except Exception:
            menu_sheets = []
    if request.method == "POST":
        # Decide which form
        if request.files.get("menuFile"):
            f = request.files["menuFile"]
            if not f or not getattr(f, "filename", None):
                flash("Ingen fil valgt.", "warning")
                return redirect(url_for("admin_menu"))
            try:
                # persist file
                os.makedirs("uploads", exist_ok=True)
                ts = datetime.now().strftime("%Y%m%d-%H%M%S")
                safe_input_name = f.filename or "menu.csv"
                filename = secure_filename(safe_input_name)
                save_name = f"menu_{ts}_{filename}"
                save_path = os.path.join("uploads", save_name)
                f.stream.seek(0)
                f.save(save_path)
                # parse
                with open(save_path, "rb") as fh:
                    sheets = _parse_menu_csv(fh)
                session["menu_sheets"] = sheets
                menu_sheets = sheets
                # persist to DB for rig
                db = get_db()
                rig_id, _ = _current_rig(db)
                if rig_id:
                    try:
                        db.execute(
                            """
                            CREATE TABLE IF NOT EXISTS menu_settings (
                              rig_id INTEGER PRIMARY KEY,
                              start_week INTEGER,
                              start_index INTEGER,
                              menu_json TEXT,
                              updated_at TEXT DEFAULT (datetime('now'))
                            )
                            """
                        )
                        db.execute(
                            "INSERT INTO menu_settings(rig_id,start_week,start_index,menu_json,updated_at) VALUES(?,?,?,?,datetime('now'))\n                             ON CONFLICT(rig_id) DO UPDATE SET menu_json=excluded.menu_json, updated_at=datetime('now')",
                            (rig_id, (startvecka_info or {}).get("start_week") or 1, (startvecka_info or {}).get("start_index") or 1, json.dumps(sheets)),
                        )
                        db.commit()
                    except Exception:
                        pass
                flash(f"Meny importert: {save_name}", "success")
            except Exception as e:
                flash(f"Kunne ikke importere meny: {e}", "danger")
        elif all(k in request.form for k in ("startMenu", "startWeek")):
            try:
                start_menu_s = request.form.get("startMenu")
                start_week_s = request.form.get("startWeek")
                if not start_menu_s or not start_week_s:
                    raise ValueError("Mangler startmeny eller startuke")
                start_index = int(start_menu_s)
                start_week = int(start_week_s)
                info = {
                    "start_index": start_index,
                    "start_week": start_week,
                    "start_date": datetime.now().strftime("%Y-%m-%d"),
                }
                session["menu_start_info"] = info
                startvecka_info = info
                flash("Startmeny og uke lagret.", "success")
                # recompute aktuell_uke using both start_week and start_index
                from datetime import date
                iso_year, iso_week, _ = date.today().isocalendar()
                aktuell_uke = ((iso_week - start_week + (start_index - 1)) % 4) + 1
                # persist to DB for rig
                if rig_id:
                    try:
                        db.execute(
                            """
                            CREATE TABLE IF NOT EXISTS menu_settings (
                              rig_id INTEGER PRIMARY KEY,
                              start_week INTEGER,
                              start_index INTEGER,
                              menu_json TEXT,
                              updated_at TEXT DEFAULT (datetime('now')),
                              start_date TEXT,
                              last_upload_at TEXT,
                              last_upload_name TEXT
                            )
                            """
                        )
                        for col in ("start_date", "last_upload_at", "last_upload_name"):
                            try:
                                db.execute(f"ALTER TABLE menu_settings ADD COLUMN {col} TEXT")
                            except Exception:
                                pass
                        db.execute(
                            "INSERT INTO menu_settings(rig_id,start_week,start_index,menu_json,start_date,updated_at) VALUES(?,?,?,?,?,datetime('now'))\n                             ON CONFLICT(rig_id) DO UPDATE SET start_week=excluded.start_week, start_index=excluded.start_index, start_date=excluded.start_date, updated_at=datetime('now')",
                            (rig_id, start_week, start_index, json.dumps(session.get("menu_sheets") or []), info.get("start_date")),
                        )
                        db.commit()
                    except Exception:
                        pass
            except Exception as e:
                flash(f"Ugyldige verdier: {e}", "warning")

    return render_template(
        "admin_menu.html",
        startvecka_info=startvecka_info,
        aktuell_uke=aktuell_uke,
        menu_sheets=menu_sheets,
        last_upload_at=(row[5] if row else None),
        last_upload_name=(row[6] if row else None),
    )


@app.post("/admin/menu/upload")
@admin_required
def admin_menu_upload():
    # Upload and parse menu CSV/Excel
    if not request.files.get("menuFile"):
        flash("Ingen fil valgt.", "warning")
        return redirect(url_for("admin_menu"))
    f = request.files["menuFile"]
    try:
        os.makedirs("uploads", exist_ok=True)
        ts = datetime.now().strftime("%Y%m%d-%H%M%S")
        safe_input_name = f.filename or "menu.csv"
        filename = secure_filename(safe_input_name)
        save_name = f"menu_{ts}_{filename}"
        save_path = os.path.join("uploads", save_name)
        f.stream.seek(0)
        f.save(save_path)
        with open(save_path, "rb") as fh:
            sheets = _parse_menu_csv(fh)
        session["menu_sheets"] = sheets
        flash(f"Meny importert: {save_name}", "success")
        # Persist to DB for rig as the source of truth and capture upload metadata
        db = get_db()
        rig_id, _ = _current_rig(db)
        if rig_id:
            try:
                db.execute(
                    """
                    CREATE TABLE IF NOT EXISTS menu_settings (
                      rig_id INTEGER PRIMARY KEY,
                      start_week INTEGER,
                      start_index INTEGER,
                      menu_json TEXT,
                      updated_at TEXT DEFAULT (datetime('now')),
                      start_date TEXT,
                      last_upload_at TEXT,
                      last_upload_name TEXT
                    )
                    """
                )
                for col in ("start_date", "last_upload_at", "last_upload_name"):
                    try:
                        db.execute(f"ALTER TABLE menu_settings ADD COLUMN {col} TEXT")
                    except Exception:
                        pass
                db.execute(
                    "INSERT INTO menu_settings(rig_id,start_week,start_index,menu_json,last_upload_at,last_upload_name,updated_at) VALUES(?,?,?,?,?,?,datetime('now'))\n                     ON CONFLICT(rig_id) DO UPDATE SET menu_json=excluded.menu_json, last_upload_at=excluded.last_upload_at, last_upload_name=excluded.last_upload_name, updated_at=datetime('now')",
                    (rig_id, 1, 1, json.dumps(sheets), datetime.now().strftime("%Y-%m-%d %H:%M"), save_name),
                )
                db.commit()
            except Exception:
                pass
    except Exception as e:
        flash(f"Kunne ikke importere meny: {e}", "danger")
    return redirect(url_for("admin_menu"))


@app.post("/admin/menu/update")
@admin_required
def admin_menu_update():
    uke = request.form.get("uke")
    dag = request.form.get("dag")
    maltid = request.form.get("maltid")
    kategori = request.form.get("kategori")
    rett = request.form.get("rett") or ""
    if not all([uke, dag, maltid, kategori]):
        return jsonify({"ok": False, "error": "Mangler felt"}), 400
    sheets = session.get("menu_sheets") or []
    updated = False
    m = (maltid or "").lower()
    for sh in sheets:
        if str(sh.get("uke")) == str(uke):
            arr = sh.get("lunsj") if m.startswith("lun") else sh.get("middag")
            if not arr:
                continue
            for it in arr:
                if (it.get("dag") == dag) and (it.get("kategori") == kategori):
                    it["rett"] = rett
                    updated = True
                    break
    session["menu_sheets"] = sheets
    return jsonify({"ok": True, "updated": updated})


@app.route("/admin/users", methods=["GET", "POST"])
@admin_required
def admin_users():
    db = get_db()
    error = None
    new_user_info = None
    rig_id = None
    rig_name = None
    if session.get("user_id"):
        row = db.execute("SELECT rig_id FROM users WHERE id=?", (session["user_id"],)).fetchone()
        if row:
            rig_id = row["rig_id"]
    if not rig_id:
        flash("Din konto mangler tilknyttet rigg.", "warning")
        return redirect(url_for("admin_dashboard_protected"))
    rrow = db.execute("SELECT name FROM rigs WHERE id=?", (rig_id,)).fetchone()
    rig_name = rrow["name"] if rrow else None

    if request.method == "POST":
        if request.form.get("delete_user_id"):
            uid = request.form.get("delete_user_id", type=int)
            try:
                db.execute("DELETE FROM users WHERE id=? AND rig_id=?", (uid, rig_id))
                db.commit()
                flash("Bruker slettet.", "success")
                return redirect(url_for("admin_users"))
            except Exception as e:
                error = f"Kunne ikke slette bruker: {e}"
        elif request.form.get("reset_pw_user_id"):
            uid = request.form.get("reset_pw_user_id", type=int)
            try:
                u = db.execute("SELECT email FROM users WHERE id=? AND rig_id=?", (uid, rig_id)).fetchone()
                if not u:
                    error = "Bruker ikke funnet."
                else:
                    from werkzeug.security import generate_password_hash
                    temp_pw = secrets.token_urlsafe(6)
                    db.execute("UPDATE users SET password_hash=? WHERE id=?", (generate_password_hash(temp_pw), uid))
                    db.commit()
                    flash(f"Engangspassord for {u['email']}: {temp_pw}", "info")
                    return redirect(url_for("admin_users"))
            except Exception as e:
                error = f"Kunne ikke resette passord: {e}"
        elif request.form.get("update_user_id"):
            uid = request.form.get("update_user_id", type=int)
            name = (request.form.get("edit_name") or "").strip()
            email = (request.form.get("edit_email") or "").strip().lower()
            role_in = (request.form.get("edit_role") or "").strip().lower()
            role = "admin" if role_in == "admin" else "user"
            if not (uid and name and email):
                error = "Navn og e-post kreves."
            else:
                try:
                    db.execute(
                        "UPDATE users SET name=?, email=?, role=? WHERE id=? AND rig_id=?",
                        (name, email, role, uid, rig_id),
                    )
                    db.commit()
                    flash("Bruker oppdatert.", "success")
                    return redirect(url_for("admin_users"))
                except Exception as e:
                    error = f"Kunne ikke oppdatere bruker: {e}"
        elif request.form.get("map_user_id") and request.form.get("virtual_role"):
            # Map selected user to chosen virtual Kokk role (rig-level mapping)
            uid = request.form.get("map_user_id", type=int)
            role_label = (request.form.get("virtual_role") or "").strip()
            try:
                existing = {}
                try:
                    existing = rotation.get_virtual_mapping(rig_id)
                except Exception:
                    existing = {}
                # Update mapping for that role only
                if role_label and uid is not None:
                    existing[role_label] = uid
                    rotation.set_virtual_mapping(rig_id, existing)
                    # Optionally apply to unbound slots now
                    if request.form.get("apply_now"):
                        rotation.apply_virtual_mapping(rig_id)
                    flash(f"Koblet {role_label} til bruker id {uid}.", "success")
                else:
                    flash("Ugyldig valg av virtuell kokk.", "warning")
                return redirect(url_for("admin_users"))
            except Exception as e:
                error = f"Kunne ikke lagre kobling: {e}"
        elif all(k in request.form for k in ("name", "email", "role")):
            name = (request.form.get("name") or "").strip()
            email = (request.form.get("email") or "").strip().lower()
            role_in = (request.form.get("role") or "").strip().lower()
            role = "admin" if role_in == "admin" else "user"
            if not (name and email and role):
                error = "Alle felt må fylles ut."
            else:
                from werkzeug.security import generate_password_hash
                temp_pw = secrets.token_urlsafe(6)
                try:
                    db.execute(
                        "INSERT INTO users (email, name, password_hash, role, rig_id) VALUES (?,?,?,?,?)",
                        (email, name, generate_password_hash(temp_pw), role, rig_id),
                    )
                    db.commit()
                    new_user_info = type("Obj", (), {"email": email, "role": role, "temp_password": temp_pw})
                    flash("Ny bruker opprettet.", "success")
                except Exception as e:
                    error = f"Kunne ikke opprette bruker: {e}"

    rigs = db.execute("SELECT id, name FROM rigs ORDER BY id").fetchall()
    users = db.execute(
        """
        SELECT id, name, email, role, rig_id,
               NULL AS phone,
               NULL AS emergency_contact_name,
               NULL AS emergency_contact_phone,
               NULL AS emergency_contact_relation,
               NULL AS notes
        FROM users
        WHERE rig_id = ?
        ORDER BY id
        """,
        (rig_id,),
    ).fetchall()
    # Fetch current virtual mapping to show selected role per user
    try:
        virtual_mapping = rotation.get_virtual_mapping(rig_id)
    except Exception:
        virtual_mapping = {}
    # Compute reverse mapping: user_id -> label for quick lookup in template
    reverse_map = {uid: label for label, uid in virtual_mapping.items()}
    roles = [f"Kokk {i}" for i in range(1, 7)]
    edit_user_id = request.args.get("edit", type=int)
    return render_template(
        "admin_users.html",
        users=users,
        rigs=rigs,
        edit_user_id=edit_user_id,
        new_user_info=new_user_info,
        error=error,
        rig_name=rig_name,
        roles=roles,
        virtual_mapping=virtual_mapping,
        reverse_map=reverse_map,
    )


@app.get("/turnus/virtual")
@admin_required
def turnus_virtual_view():
    """Show upcoming periods for a specific virtual role (e.g., Kokk 3)."""
    db = get_db()
    rig_id, rig_name = _current_rig(db)
    if not rig_id:
        flash("Din konto mangler tilknyttet rigg.", "warning")
        return redirect(url_for("admin_dashboard_protected"))
    role = request.args.get("role", default="").strip() or None
    if not role:
        flash("Mangler rolle.", "warning")
        return redirect(url_for("admin_users"))
    today_s = datetime.now().strftime("%Y-%m-%d") + "T00:00"
    limit = request.args.get("limit", type=int) or 6
    rows = db.execute(
        """
        SELECT start_ts, end_ts, status, notes
        FROM turnus_slots
        WHERE rig_id = ? AND role = ? AND end_ts >= ?
        ORDER BY start_ts ASC
        """,
        (rig_id, role, today_s),
    ).fetchall()

    # Group into periods the same way as overview
    periods = []
    cur = None
    for r in rows:
        st = _parse_ts(r[0])
        et = _parse_ts(r[1])
        note = (r[3] or "").lower()
        if cur is None:
            cur = {"start": st, "end": et, "count": 1, "has_published": (r[2] == "published")}
        else:
            anchor = cur["start"]
            if st <= anchor + timedelta(days=14, hours=1):
                if et > cur["end"]:
                    cur["end"] = et
                cur["count"] += 1
                if r[2] == "published":
                    cur["has_published"] = True
            else:
                periods.append(cur)
                cur = {"start": st, "end": et, "count": 1, "has_published": (r[2] == "published")}
    if cur is not None:
        periods.append(cur)

    # Only full periods (>=15 slots) and limit
    full_periods = [p for p in periods if p.get("count", 0) >= 15][:limit]
    items = [{
        "start_date": p["start"].strftime("%Y-%m-%d"),
        "end_date": p["end"].strftime("%Y-%m-%d"),
        "status": "published" if p["has_published"] else "planned",
    } for p in full_periods]
    return render_template("turnus_virtual.html", rig_name=rig_name, role=role, items=items, limit=limit)


# --- Menus (stubs) ---
@app.post("/menus/import_menu_file")
@admin_required
def import_menu_file():
    # Save the uploaded menu file (CSV/PDF/DOCX/XLSX) into uploads/ for later processing
    f = request.files.get("menuFile") or request.files.get("file")
    if not f or not getattr(f, "filename", None):
        flash("Ingen fil valgt.", "warning")
        return redirect(url_for("admin_menu"))
    try:
        os.makedirs("uploads", exist_ok=True)
        ts = datetime.now().strftime("%Y%m%d-%H%M%S")
        safe_input_name = f.filename or "uploaded_file"
        filename = secure_filename(safe_input_name)
        save_name = f"menu_{ts}_{filename}"
        save_path = os.path.join("uploads", save_name)
        f.save(save_path)
        flash(f"Fil lastet opp: {save_name}", "success")
    except Exception as e:
        flash(f"Kunne ikke laste opp fil: {e}", "danger")
    return redirect(url_for("admin_menu"))


@app.post("/menus/import_rotation")
@admin_required
def import_rotation():
    # Save the uploaded rotation Excel file into uploads/ for later processing
    f = request.files.get("file")
    year = request.form.get("year")
    if not f or not getattr(f, "filename", None):
        flash("Ingen fil valgt.", "warning")
        return redirect(url_for("admin_menu"))
    try:
        os.makedirs("uploads", exist_ok=True)
        ts = datetime.now().strftime("%Y%m%d-%H%M%S")
        safe_input_name = f.filename or "rotation.xlsx"
        filename = secure_filename(safe_input_name)
        prefix = f"rotation_{year}_" if year else "rotation_"
        save_name = f"{prefix}{ts}_{filename}"
        save_path = os.path.join("uploads", save_name)
        f.save(save_path)
        flash(f"Rotasjonsfil lagret: {save_name}", "success")
    except Exception as e:
        flash(f"Kunne ikke lagre rotasjonsfil: {e}", "danger")
    return redirect(url_for("admin_menu"))


## Removed legacy admin-only /menus routes to avoid collision with public /menus overview


@app.get("/export/week/<int:work_cycle_id>")
@admin_required
def export_week(work_cycle_id: int):
    flash("Export uke er ikke implementert ennå.", "info")
    return redirect(url_for("admin_menu"))


@app.get("/export/shopping/<int:work_cycle_id>")
@admin_required
def export_shopping(work_cycle_id: int):
    flash("Handleliste-eksport er ikke implementert ennå.", "info")
    return redirect(url_for("admin_menu"))


@app.post("/note/<int:instance_id>")
@admin_required
def update_note(instance_id: int):
    flash("Oppdatere notat er ikke implementert ennå.", "info")
    return redirect(url_for("admin_menu"))


## duplicate dashboard implementation removed; using the GET-only auto period version above


# --- Turnus overview (admin) ---
@app.get("/turnus/overview")
@admin_required
def turnus_overview():
    db = get_db()
    rig_id, rig_name = _current_rig(db)
    if not rig_id:
        flash("Din konto mangler tilknyttet rigg.", "warning")
        return redirect(url_for("admin_dashboard_protected"))

    # Horizon: show future periods from today
    today_s = datetime.now().strftime("%Y-%m-%d") + "T00:00"
    rows = db.execute(
        """
        SELECT role, start_ts, end_ts, status, notes
        FROM turnus_slots
        WHERE rig_id = ? AND end_ts >= ? AND role LIKE 'Kokk %'
        ORDER BY role ASC, start_ts ASC
        """,
        (rig_id, today_s),
    ).fetchall()

    # Group into ~14-day periods per role
    from collections import defaultdict
    grouped = defaultdict(list)  # role -> list of periods
    cur = None
    last_role = None
    for r in rows:
        role = r["role"] or ""
        st = _parse_ts(r["start_ts"])  # datetime
        et = _parse_ts(r["end_ts"])    # datetime
        note = (r["notes"] or "").lower()
        if role != last_role:
            # finalize previous
            if cur is not None:
                grouped[last_role].append(cur)
            # start new group for this role
            cur = {
                "start": st,
                "end": et,
                "count": 1,
                "breakdown": {"dag": 0, "natt": 0, "snu": 0},
                "has_published": (r["status"] == "published"),
            }
            if note.startswith("snu"):
                cur["breakdown"]["snu"] += 1
            elif note == "natt":
                cur["breakdown"]["natt"] += 1
            elif note == "dag":
                cur["breakdown"]["dag"] += 1
            else:
                # treat unknown as a day shift length-wise? keep out of breakdown
                pass
            last_role = role
            continue
        # same role, decide if within 14 days from group start anchor
        if cur is None:
            cur = {
                "start": st,
                "end": et,
                "count": 1,
                "breakdown": {"dag": 0, "natt": 0, "snu": 0},
                "has_published": (r["status"] == "published"),
            }
            if note.startswith("snu"):
                cur["breakdown"]["snu"] += 1
            elif note == "natt":
                cur["breakdown"]["natt"] += 1
            elif note == "dag":
                cur["breakdown"]["dag"] += 1
            last_role = role
            continue
        anchor = cur["start"]
        if st <= anchor + timedelta(days=14, hours=1):  # allow a bit of slack
            # extend current period
            if et > cur["end"]:
                cur["end"] = et
            cur["count"] += 1
            if note.startswith("snu"):
                cur["breakdown"]["snu"] += 1
            elif note == "natt":
                cur["breakdown"]["natt"] += 1
            elif note == "dag":
                cur["breakdown"]["dag"] += 1
            if r["status"] == "published":
                cur["has_published"] = True
        else:
            # close previous, start new
            grouped[last_role].append(cur)
            cur = {
                "start": st,
                "end": et,
                "count": 1,
                "breakdown": {"dag": 0, "natt": 0, "snu": 0},
                "has_published": (r["status"] == "published"),
            }
            if note.startswith("snu"):
                cur["breakdown"]["snu"] += 1
            elif note == "natt":
                cur["breakdown"]["natt"] += 1
            elif note == "dag":
                cur["breakdown"]["dag"] += 1
    if cur is not None and last_role is not None:
        grouped[last_role].append(cur)

    # Limit to next few periods per role (default 2) and only show full periods
    limit = request.args.get("limit", type=int) or 2
    roles = [f"Kokk {i}" for i in range(1, 7)]
    overview = []
    for role in roles:
        # keep only full periods (generator creates 15 pass per 14-dagars period)
        full_periods = [p for p in grouped.get(role, []) if p.get("count", 0) >= 15]
        periods = full_periods[:limit]
        nice = []
        for p in periods:
            nice.append({
                "start_date": p["start"].strftime("%Y-%m-%d"),
                "end_date": p["end"].strftime("%Y-%m-%d"),
                "days": (p["end"] - p["start"]).days + 1,
                "slots": p["count"],
                "status": "published" if p["has_published"] else "planned",
                "breakdown": p.get("breakdown", {}),
            })
        overview.append({"role": role, "periods": nice})

    return render_template("turnus_overview.html", rig_name=rig_name, overview=overview)


@app.post("/turnus/cleanup")
@admin_required
def turnus_cleanup():
    """Utility endpoint to deduplicate slots and/or clear future slots for this rig."""
    db = get_db()
    rig_id, _ = _current_rig(db)
    if not rig_id:
        flash("Din konto mangler tilknyttet rigg.", "warning")
        return redirect(url_for("admin_dashboard_protected"))

    action = (request.form.get("action") or "").lower()
    cleared = 0
    deduped_groups = 0

    try:
        if action == "clear_future":
            from_date = request.form.get("from_date")
            if not from_date:
                from_date = datetime.now().strftime("%Y-%m-%d")
            start_ts = f"{from_date}T00:00"
            # Collect slot ids to clear
            rows = db.execute(
                "SELECT id FROM turnus_slots WHERE rig_id = ? AND start_ts >= ?",
                (rig_id, start_ts),
            ).fetchall()
            slot_ids = [r[0] for r in rows]
            if slot_ids:
                placeholders = ",".join(["?"] * len(slot_ids))
                # Remove bindings first if table exists
                try:
                    db.execute(f"DELETE FROM turnus_account_binding WHERE slot_id IN ({placeholders})", slot_ids)
                except Exception:
                    pass
                db.execute(f"DELETE FROM turnus_slots WHERE id IN ({placeholders})", slot_ids)
                db.commit()
                cleared = len(slot_ids)

        elif action == "dedupe":
            # Find duplicate (role, start_ts, end_ts)
            rows = db.execute(
                """
                SELECT role, start_ts, end_ts, COUNT(*) AS c, GROUP_CONCAT(id) AS ids
                FROM turnus_slots
                WHERE rig_id = ?
                GROUP BY role, start_ts, end_ts
                HAVING c > 1
                """,
                (rig_id,),
            ).fetchall()
            for r in rows:
                ids = [int(x) for x in (r["ids"] or "").split(",") if x]
                if len(ids) <= 1:
                    continue
                keep = min(ids)
                to_del = [i for i in ids if i != keep]
                placeholders = ",".join(["?"] * len(to_del))
                try:
                    db.execute(f"DELETE FROM turnus_account_binding WHERE slot_id IN ({placeholders})", to_del)
                except Exception:
                    pass
                db.execute(f"DELETE FROM turnus_slots WHERE id IN ({placeholders})", to_del)
                deduped_groups += 1
            if deduped_groups:
                db.commit()

        else:
            flash("Ukjent handling.", "warning")
            return redirect(url_for("turnus_overview"))

    except Exception as e:
        flash(f"Rydde-feil: {e}", "danger")
        return redirect(url_for("turnus_overview"))

    if action == "clear_future":
        flash(f"Slettet {cleared} framtidige slots fra og med valgt dato.", "success")
    elif action == "dedupe":
        flash(f"Fjernet duplikater i {deduped_groups} grupper.", "success")
    return redirect(url_for("turnus_overview"))


@app.get("/turnus/mapping")
@admin_required
def turnus_mapping():
    """UI to map virtual cooks (Kokk 1–6) to real users on this rig."""
    db = get_db()
    rig_id, rig_name = _current_rig(db)
    if not rig_id:
        flash("Din konto mangler tilknyttet rigg.", "warning")
        return redirect(url_for("admin_dashboard_protected"))
    users = db.execute(
        "SELECT id, name, email FROM users WHERE rig_id = ? ORDER BY name, email",
        (rig_id,),
    ).fetchall()
    try:
        existing = rotation.get_virtual_mapping(rig_id)
    except Exception:
        existing = {}
    roles = [f"Kokk {i}" for i in range(1, 7)]
    return render_template("turnus_mapping.html", rig_name=rig_name, users=users, mapping=existing, roles=roles)


@app.post("/turnus/mapping")
@admin_required
def turnus_mapping_save():
    db = get_db()
    rig_id, _ = _current_rig(db)
    if not rig_id:
        flash("Din konto mangler tilknyttet rigg.", "warning")
        return redirect(url_for("admin_dashboard_protected"))

    # Build mapping from form fields kokk_1..kokk_6
    mapping = {}
    for i in range(1, 7):
        val = request.form.get(f"kokk_{i}")
        try:
            uid = int(val) if val else None
        except Exception:
            uid = None
        if uid:
            mapping[f"Kokk {i}"] = uid

    try:
        rotation.set_virtual_mapping(rig_id, mapping)
        # Optional: rebind all (delete existing bindings for Kokk % roles then re-apply)
        applied = 0
        if request.form.get("apply_now") or request.form.get("rebind_all"):
            if request.form.get("rebind_all"):
                try:
                    db.execute(
                        """
                        DELETE FROM turnus_account_binding
                        WHERE slot_id IN (
                            SELECT id FROM turnus_slots WHERE rig_id = ? AND role LIKE 'Kokk %'
                        )
                        """,
                        (rig_id,),
                    )
                    db.commit()
                except Exception as e:
                    flash(f"Kunne ikke fjerne eksisterende bindinger: {e}", "warning")
            applied = rotation.apply_virtual_mapping(rig_id)
        flash("Kobling lagret. " + (f"Opprettet {applied} bindinger." if applied else ""), "success")
    except Exception as e:
        flash(f"Kunne ikke lagre kobling: {e}", "danger")

    return redirect(url_for("turnus_mapping"))


@app.get("/turnus/rebind")
@admin_required
def turnus_rebind():
    """Form to bind a specific user to all slots of a virtual Kokk label within a date range (vikar handling)."""
    db = get_db()
    rig_id, rig_name = _current_rig(db)
    if not rig_id:
        flash("Din konto mangler tilknyttet rigg.", "warning")
        return redirect(url_for("admin_dashboard_protected"))
    users = db.execute(
        "SELECT id, name, email FROM users WHERE rig_id = ? ORDER BY name, email",
        (rig_id,),
    ).fetchall()
    roles = [f"Kokk {i}" for i in range(1, 7)]
    return render_template("turnus_rebind.html", rig_name=rig_name, users=users, roles=roles)


@app.post("/turnus/rebind")
@admin_required
def turnus_rebind_apply():
    db = get_db()
    rig_id, _ = _current_rig(db)
    if not rig_id:
        flash("Din konto mangler tilknyttet rigg.", "warning")
        return redirect(url_for("admin_dashboard_protected"))

    role = (request.form.get("role") or "").strip()
    user_id = request.form.get("user_id", type=int)
    start_s = (request.form.get("start") or "").strip()
    end_s = (request.form.get("end") or "").strip()
    if not (role and user_id and start_s and end_s):
        flash("Mangler felt.", "warning")
        return redirect(url_for("turnus_rebind"))
    try:
        d0 = datetime.strptime(start_s, "%Y-%m-%d").date()
        d1 = datetime.strptime(end_s, "%Y-%m-%d").date()
    except Exception:
        flash("Ugyldig datoformat.", "danger")
        return redirect(url_for("turnus_rebind"))
    if d1 < d0:
        flash("Sluttdato kan ikke være før startdato.", "warning")
        return redirect(url_for("turnus_rebind"))
    start_ts = f"{d0.strftime('%Y-%m-%d')}T00:00"
    end_ts = f"{d1.strftime('%Y-%m-%d')}T23:59"

    # Find all slot ids for role overlapping the window
    rows = db.execute(
        """
        SELECT id FROM turnus_slots
        WHERE rig_id = ? AND role = ? AND start_ts < ? AND end_ts > ?
        ORDER BY start_ts ASC
        """,
        (rig_id, role, end_ts, start_ts),
    ).fetchall()
    slot_ids = [r[0] for r in rows]
    if not slot_ids:
        flash("Fant ingen slots i perioden for valgt rolle.", "info")
        return redirect(url_for("turnus_rebind"))

    # Delete existing bindings for these slots, then insert new ones
    # Use executemany for inserts
    placeholders = ",".join(["?"] * len(slot_ids))
    try:
        db.execute(f"DELETE FROM turnus_account_binding WHERE slot_id IN ({placeholders})", slot_ids)
        to_insert = [(sid, user_id, f"vikar for {role}") for sid in slot_ids]
        db.executemany(
            "INSERT INTO turnus_account_binding (slot_id, user_id, notes, bound_at) VALUES (?,?,?, datetime('now'))",
            to_insert,
        )
        db.commit()
        flash(f"Bundet {len(slot_ids)} slots til valgt bruker for {role}.", "success")
    except Exception as e:
        flash(f"Feil ved rebind: {e}", "danger")

    return redirect(url_for("turnus_overview"))

@app.route("/turnus/create_real_cooks", methods=["GET", "POST"])
@admin_required
def turnus_create_real_cooks():
    """Create six named user accounts and map Kokk 1-6 to them for the current rig, then apply bindings."""
    db = get_db()
    rig_id, rig_name = _current_rig(db)
    if not rig_id:
        flash("Din konto mangler tilknyttet rigg.", "warning")
        return redirect(url_for("admin_dashboard_protected"))

    # Desired mapping order: 1->Henrik, 2->Alex, 3->Stefan, 4->Micke, 5->Helge, 6->Cedric
    names = ["Henrik", "Alex", "Stefan", "Micke", "Helge", "Cedric"]
    emails = [f"{n.lower()}@yuplan.local" for n in names]
    default_pw = os.environ.get("DEFAULT_TEST_PW", "kokk1234")

    from werkzeug.security import generate_password_hash

    created = 0
    ensured_ids: list[int] = []
    for name, email in zip(names, emails, strict=False):
        row = db.execute("SELECT id, rig_id FROM users WHERE email=?", (email,)).fetchone()
        if row:
            uid = row["id"]
            # Ensure name and rig match
            try:
                db.execute(
                    "UPDATE users SET name=?, role='user', rig_id=? WHERE id=?",
                    (name, rig_id, uid),
                )
                db.commit()
            except Exception:
                pass
            ensured_ids.append(uid)
        else:
            try:
                db.execute(
                    "INSERT INTO users (email, name, password_hash, role, rig_id) VALUES (?,?,?,?,?)",
                    (email, name, generate_password_hash(default_pw), "user", rig_id),
                )
                db.commit()
                uid = db.execute("SELECT id FROM users WHERE email=?", (email,)).fetchone()["id"]
                ensured_ids.append(uid)
                created += 1
            except Exception as e:
                flash(f"Kunne ikke opprette bruker {name}: {e}", "danger")

    # Build mapping dict Kokk 1..6 -> user_id
    mapping = {f"Kokk {i}": ensured_ids[i-1] for i in range(1, 7) if i-1 < len(ensured_ids)}
    try:
        rotation.set_virtual_mapping(rig_id, mapping)
        # Apply to all slots without binding (future and past)
        applied = rotation.apply_virtual_mapping(rig_id)
        flash(
            f"Brukere sikret (nye: {created}). Mapping satt for Kokk 1–6. Opprettet {applied} bindinger.",
            "success",
        )
    except Exception as e:
        flash(f"Kunne ikke sette mapping/binde slots: {e}", "danger")

    # Send admin to overview to inspect result
    return redirect(url_for("turnus_overview"))

# --- Misc helpers ---
@app.get("/turnus/preview")
def api_turnus_preview():
    rig_id = request.args.get("rig_id", type=int)
    if not rig_id:
        return jsonify({"ok": False, "error": "rig_id krävs"}), 400
    try:
        start_ts, end_ts = _build_interval_from_params()
    except ValueError as e:
        return jsonify({"ok": False, "error": str(e)}), 400
    db = get_db()
    rows = db.execute(
        """
        SELECT id, rig_id, start_ts, end_ts, role, status, notes
        FROM turnus_slots
        WHERE rig_id = ? AND start_ts < ? AND end_ts > ?
        ORDER BY start_ts ASC
        """,
        (rig_id, end_ts, start_ts),
    ).fetchall()
    items = [dict(r) for r in rows]
    return jsonify({"ok": True, "count": len(items), "items": items})


@app.get("/turnus/view")
def api_turnus_view():
    rig_id = request.args.get("rig_id", type=int)
    if not rig_id:
        return jsonify({"ok": False, "error": "rig_id krävs"}), 400
    try:
        start_ts, end_ts = _build_interval_from_params()
    except ValueError as e:
        return jsonify({"ok": False, "error": str(e)}), 400
    db = get_db()
    rows = db.execute(
        """
        SELECT id, rig_id, start_ts, end_ts, role, status, notes
        FROM turnus_slots
        WHERE rig_id = ? AND start_ts < ? AND end_ts > ? AND status='published'
        ORDER BY start_ts ASC
        """,
        (rig_id, end_ts, start_ts),
    ).fetchall()
    items = [dict(r) for r in rows]
    return jsonify({"ok": True, "count": len(items), "items": items})


# --- Menus overview (all 4 uka) ---
@app.get("/public/menus")
def menus_overview():
    db = get_db()
    rig_id, rig_name = _current_rig(db)
    if not rig_id:
        # allow anonymous to see nothing
        flash("Ingen rigg knyttet til kontoen.", "warning")
        return redirect(url_for("dashboard"))
    # Load menu settings
    start_week = 1
    start_index = 1
    menu_sheets = []
    try:
        row = db.execute("SELECT start_week, start_index, menu_json FROM menu_settings WHERE rig_id=?", (rig_id,)).fetchone()
        if row:
            start_week = row[0] or 1
            start_index = row[1] or 1
            try:
                menu_sheets = json.loads(row[2] or "[]")
            except Exception:
                menu_sheets = []
    except Exception:
        pass
    # Sort by uke 1..4 if present
    try:
        menu_sheets = sorted(menu_sheets, key=lambda sh: int(sh.get("uke") or 999))
    except Exception:
        pass
    # Weekday order
    order = ["Mandag","Tirsdag","Onsdag","Torsdag","Fredag","Lørdag","Søndag"]
    # Normalize to dict uke -> { lunsj_by_day, middag_by_day }
    menus = []
    def _norm_cat(raw):
        if not raw:
            return None
        r = str(raw).strip().lower()
        repl = r.replace("ø", "o").replace("å", "a").replace("ä", "a").replace("ö", "o").replace("æ", "a")
        if ("sopp" in repl) or ("suppe" in repl) or ("soup" in repl) or ("sopa" in repl):
            return "soppa"
        if ("fisk" in repl) or ("fish" in repl):
            return "fisk"
        if ("kjott" in repl) or ("kjot" in repl) or ("kott" in repl) or ("meat" in repl):
            return "kott"
        return None

    for sh in menu_sheets:
        uke = str(sh.get("uke"))
        lunsj = sh.get("lunsj") or []
        middag = sh.get("middag") or []
        lb = {d: [] for d in order}
        mb = {d: [] for d in order}
        for it in lunsj:
            d = it.get("dag")
            if d in lb:
                lb[d].append({
                    "rett": it.get("rett"),
                    "kategori": _norm_cat(it.get("kategori"))
                })
        for it in middag:
            d = it.get("dag")
            if d in mb:
                mb[d].append({
                    "rett": it.get("rett"),
                    "kategori": _norm_cat(it.get("kategori"))
                })
        menus.append({"uke": uke, "lunsj": lb, "middag": mb})
    return render_template("menus_overview.html", rig_name=rig_name, start_week=start_week, start_index=start_index, menus=menus)


# --- Daily printable menu (DOCX) ---

 
@app.get("/menus/daily_docx")
def menus_daily_docx():
    """Generate a printable DOCX for today's (or selected) menu using a clean built-in layout (no external template).

    Columns: left=lunch, right=dinner. Language: English labels.
    """
    if Document is None:
        flash("DOCX library not available. Please install python-docx.", "danger")
        return redirect(url_for("dashboard"))

    db = get_db()
    rig_id, rig_name = _current_rig(db)
    if not rig_id:
        flash("Ingen rigg tilknyttet.", "warning")
        return redirect(url_for("dashboard"))

    # Resolve date param (YYYY-MM-DD)
    date_s = (request.args.get("date") or datetime.now().strftime("%Y-%m-%d")).strip()
    try:
        dt = datetime.strptime(date_s, "%Y-%m-%d").date()
    except Exception:
        flash("Ugyldig datoformat (forventet YYYY-MM-DD).", "danger")
        return redirect(url_for("dashboard"))

    # Load menu settings for the rig
    start_week = 1
    start_index = 1
    menu_sheets = []
    try:
        row = db.execute("SELECT start_week, start_index, menu_json FROM menu_settings WHERE rig_id=?", (rig_id,)).fetchone()
        if row:
            start_week = row[0] or 1
            start_index = row[1] or 1
            try:
                menu_sheets = json.loads(row[2] or "[]")
            except Exception:
                menu_sheets = []
    except Exception:
        pass

    # Helper: pick dishes per category for the given date
    weekday_no_to_nor = {0: "Mandag", 1: "Tirsdag", 2: "Onsdag", 3: "Torsdag", 4: "Fredag", 5: "Lørdag", 6: "Søndag"}
    def menu_index_for_date(dt_date):
        iso_year, iso_week, _ = dt_date.isocalendar()
        idx = ((iso_week - int(start_week) + (int(start_index) - 1)) % 4) + 1
        return str(idx)

    def _norm_cat(raw: str | None):
        if not raw:
            return None
        r = str(raw).strip().lower()
        repl = r.replace("ø", "o").replace("å", "a").replace("ä", "a").replace("ö", "o").replace("æ", "a")
        if ("sopp" in repl) or ("suppe" in repl) or ("soup" in repl) or ("sopa" in repl):
            return "soppa"
        if ("fisk" in repl) or ("fish" in repl):
            return "fisk"
        if ("kjott" in repl) or ("kjot" in repl) or ("kott" in repl) or ("meat" in repl):
            return "kott"
        return None

    # Build index for quick lookup
    sheets_by_uke = {str(sh.get("uke")): sh for sh in (menu_sheets or [])}
    uke_key = menu_index_for_date(dt)
    sh = sheets_by_uke.get(uke_key)
    if not sh:
        flash("Ingen meny funnet for valgt dato.", "warning")
        return redirect(url_for("dashboard"))
    dn = weekday_no_to_nor[dt.weekday()]

    # Collect lunch/dinner per normalized category
    def collect_for(meal_key: str):
        items = sh.get(meal_key) or []
        cats: dict[str, list[str]] = {"soppa": [], "fisk": [], "kott": [], "extra": []}
        for it in items:
            if it.get("dag") != dn:
                continue
            cat = _norm_cat(it.get("kategori"))
            if cat in ("soppa", "fisk", "kott"):
                cats[cat].append(it.get("rett") or "")
            else:
                cats["extra"].append(it.get("rett") or "")
        return cats

    lunch = collect_for("lunsj")
    dinner = collect_for("middag")

    # Apply per-day overrides and allergens from DB (if any)
    try:
        db.execute(
            """
            CREATE TABLE IF NOT EXISTS daily_menu_overrides (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              rig_id INTEGER NOT NULL,
              date TEXT NOT NULL,
              meal TEXT NOT NULL,           -- 'lunsj' | 'middag'
              category TEXT NOT NULL,       -- 'soppa' | 'fisk' | 'kott' | 'extra'
              dish TEXT NOT NULL,
              UNIQUE(rig_id, date, meal, category)
            )
            """
        )
        db.execute(
            """
            CREATE TABLE IF NOT EXISTS daily_menu_allergens (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              rig_id INTEGER NOT NULL,
              date TEXT NOT NULL,
              meal TEXT NOT NULL,           -- 'lunsj' | 'middag'
              category TEXT NOT NULL,       -- 'soppa' | 'fisk' | 'kott' | 'extra'
              allergens TEXT,               -- comma-separated numbers, e.g. '1,3,7'
              UNIQUE(rig_id, date, meal, category)
            )
            """
        )
        db.commit()
    except Exception:
        pass

    def apply_overrides(meal_key: str, cats: dict[str, list[str]]):
        # Per category pick override if present, else first item
        out: dict[str, str] = {}
        for cat in ("soppa", "fisk", "kott", "extra"):
            row = db.execute(
                "SELECT dish FROM daily_menu_overrides WHERE rig_id=? AND date=? AND meal=? AND category=?",
                (rig_id, date_s, meal_key, cat),
            ).fetchone()
            if row and row[0]:
                out[cat] = row[0]
            else:
                lst = cats.get(cat) or []
                out[cat] = (lst[0] if lst else "")
        return out

    # Ensure gluten_types and nuts_types columns exist
    try:
        db.execute("ALTER TABLE daily_menu_allergens ADD COLUMN gluten_types TEXT")
        db.commit()
    except Exception:
        pass
    try:
        db.execute("ALTER TABLE daily_menu_allergens ADD COLUMN nuts_types TEXT")
        db.commit()
    except Exception:
        pass

    def read_allergens(meal_key: str):
        m_codes: dict[str, str] = {}
        m_gluten: dict[str, str] = {}
        m_nuts: dict[str, str] = {}
        rows = db.execute(
            "SELECT category, allergens, gluten_types, nuts_types FROM daily_menu_allergens WHERE rig_id=? AND date=? AND meal=?",
            (rig_id, date_s, meal_key),
        ).fetchall()
        for r in rows:
            m_codes[r[0]] = (r[1] or "")
            m_gluten[r[0]] = (r[2] or "")
            m_nuts[r[0]] = (r[3] or "")
        return m_codes, m_gluten, m_nuts

    lunch_picked = apply_overrides("lunsj", lunch)
    dinner_picked = apply_overrides("middag", dinner)
    lunch_all, lunch_glu, lunch_nuts = read_allergens("lunsj")
    dinner_all, dinner_glu, dinner_nuts = read_allergens("middag")

    # Ensure default text for 'extra' when nothing is specified
    DEFAULT_EXTRA = "Varied types of side dishes."
    if not (lunch_picked.get("extra") or "").strip():
        lunch_picked["extra"] = DEFAULT_EXTRA
    if not (dinner_picked.get("extra") or "").strip():
        dinner_picked["extra"] = DEFAULT_EXTRA

    # Always use the modern builder and avoid template/token handling
    doc = _build_docx_modern(str(rig_name or ""), dt, lunch_picked, lunch_all, lunch_glu, lunch_nuts, dinner_picked, dinner_all, dinner_glu, dinner_nuts)

    # No template/token helpers needed for modern builder

    # Compact slightly to fit text on a single page with fixed scaling
    do_compact = True
    scale = 0.92
    if do_compact:
        try:
            # Only scale fonts to preserve layout exactly
            def _scale_run_fonts(p, scale: float, min_pt: float = 9.0):
                if not Pt:
                    return
                for r in getattr(p, "runs", []) or []:
                    try:
                        f = r.font
                        if f is None:
                            continue
                        if f.size is not None:
                            new_pt = max(min_pt, float(f.size.pt) * scale)
                            f.size = Pt(new_pt)
                    except Exception:
                        continue

            def _scale_container_fonts(obj, scale: float):
                try:
                    for p in getattr(obj, "paragraphs", []) or []:
                        _scale_run_fonts(p, scale)
                except Exception:
                    pass
                try:
                    for tbl in getattr(obj, "tables", []) or []:
                        for row in tbl.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    _scale_run_fonts(p, scale)
                except Exception:
                    pass

            # Scale document styles (paragraph/character) so text without explicit run sizes also shrinks
            try:
                if Pt and hasattr(doc, "styles"):
                    for st in list(doc.styles):
                        try:
                            fnt = getattr(st, "font", None)
                            if fnt is not None and fnt.size is not None:
                                fnt.size = Pt(max(8.0, float(fnt.size.pt) * scale))
                        except Exception:
                            continue
            except Exception:
                pass

            # Enforce A4 page size to match expected output; margins unchanged
            try:
                if Cm:
                    A4_W, A4_H = Cm(21.0), Cm(29.7)
                    for s in doc.sections:
                        try:
                            s.page_width = A4_W
                            s.page_height = A4_H
                        except Exception:
                            pass
            except Exception:
                pass

            # Apply font scaling to main doc and headers/footers
            _scale_container_fonts(doc, scale)
            try:
                for sec in doc.sections:
                    for part in [getattr(sec, "header", None), getattr(sec, "first_page_header", None), getattr(sec, "even_page_header", None),
                                 getattr(sec, "footer", None), getattr(sec, "first_page_footer", None), getattr(sec, "even_page_footer", None)]:
                        if part:
                            _scale_container_fonts(part, scale)
            except Exception:
                pass
        except Exception:
            pass

    # Write to temp file and return
    os.makedirs("uploads/daily_menus", exist_ok=True)
    with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(Path(os.getcwd())/"uploads"/"daily_menus")) as tmp:
        out_path = Path(tmp.name)
    doc.save(str(out_path))
    return send_file(str(out_path), as_attachment=True, download_name=f"daily_menu_{date_s}.docx")




# --- Daily HTML print (no Word required) ---
@app.get("/menus/daily_print")
def menus_daily_print():
    """Render a printable HTML page for the daily menu without requiring Word.

    Uses the same live data (overrides + allergens + gluten subtypes) as DOCX.
    """
    db = get_db()
    rig_id, rig_name = _current_rig(db)
    if not rig_id:
        flash("Ingen rigg tilknyttet.", "warning")
        return redirect(url_for("dashboard"))
    date_s = (request.args.get("date") or datetime.now().strftime("%Y-%m-%d")).strip()
    try:
        dt = datetime.strptime(date_s, "%Y-%m-%d").date()
    except Exception:
        flash("Ugyldig datoformat (YYYY-MM-DD).", "danger")
        return redirect(url_for("dashboard"))

    # Load menu settings
    start_week = 1
    start_index = 1
    menu_sheets = []
    try:
        row = db.execute("SELECT start_week, start_index, menu_json FROM menu_settings WHERE rig_id=?", (rig_id,)).fetchone()
        if row:
            start_week = row[0] or 1
            start_index = row[1] or 1
            try:
                menu_sheets = json.loads(row[2] or "[]")
            except Exception:
                menu_sheets = []
    except Exception:
        pass
    weekday_no_to_nor = {0: "Mandag", 1: "Tirsdag", 2: "Onsdag", 3: "Torsdag", 4: "Fredag", 5: "Lørdag", 6: "Søndag"}
    def menu_index_for_date(dt_date):
        iso_year, iso_week, _ = dt_date.isocalendar()
        idx = ((iso_week - int(start_week) + (int(start_index) - 1)) % 4) + 1
        return str(idx)
    def _norm_cat(raw: str | None):
        if not raw:
            return None
        r = str(raw).strip().lower()
        repl = r.replace("ø", "o").replace("å", "a").replace("ä", "a").replace("ö", "o").replace("æ", "a")
        if ("sopp" in repl) or ("suppe" in repl) or ("soup" in repl) or ("sopa" in repl):
            return "soppa"
        if ("fisk" in repl) or ("fish" in repl):
            return "fisk"
        if ("kjott" in repl) or ("kjot" in repl) or ("kott" in repl) or ("meat" in repl):
            return "kott"
        return None
    sheets_by_uke = {str(sh.get("uke")): sh for sh in (menu_sheets or [])}
    sh = sheets_by_uke.get(menu_index_for_date(dt))
    dn = weekday_no_to_nor[dt.weekday()]
    def collect_for(meal_key: str):
        items = (sh.get(meal_key) if sh else []) or []
        cats: dict[str, list[str]] = {"soppa": [], "fisk": [], "kott": [], "extra": []}
        for it in items:
            if it.get("dag") != dn:
                continue
            cat = _norm_cat(it.get("kategori"))
            if cat in ("soppa", "fisk", "kott"):
                cats[cat].append(it.get("rett") or "")
            else:
                cats["extra"].append(it.get("rett") or "")
        return cats
    lunch = collect_for("lunsj")
    dinner = collect_for("middag")

    # Ensure allergen table has gluten_types and nuts_types and read overrides/allergens
    try:
        db.execute("ALTER TABLE daily_menu_allergens ADD COLUMN gluten_types TEXT")
        db.commit()
    except Exception:
        pass
    try:
        db.execute("ALTER TABLE daily_menu_allergens ADD COLUMN nuts_types TEXT")
        db.commit()
    except Exception:
        pass
    def apply_overrides(meal_key: str, cats: dict[str, list[str]]):
        out: dict[str, str] = {}
        for cat in ("soppa", "fisk", "kott", "extra"):
            row = db.execute(
                "SELECT dish FROM daily_menu_overrides WHERE rig_id=? AND date=? AND meal=? AND category=?",
                (rig_id, date_s, meal_key, cat),
            ).fetchone()
            if row and row[0]:
                out[cat] = row[0]
            else:
                lst = cats.get(cat) or []
                out[cat] = (lst[0] if lst else "")
        return out
    def read_allergens(meal_key: str):
        m_codes: dict[str, str] = {}
        m_gluten: dict[str, str] = {}
        m_nuts: dict[str, str] = {}
        rows = db.execute(
            "SELECT category, allergens, gluten_types, nuts_types FROM daily_menu_allergens WHERE rig_id=? AND date=? AND meal=?",
            (rig_id, date_s, meal_key),
        ).fetchall()
        for r in rows:
            m_codes[r[0]] = (r[1] or "")
            m_gluten[r[0]] = (r[2] or "")
            m_nuts[r[0]] = (r[3] or "")
        return m_codes, m_gluten, m_nuts
    lunch_picked = apply_overrides("lunsj", lunch)
    dinner_picked = apply_overrides("middag", dinner)
    lunch_all, lunch_glu, lunch_nuts = read_allergens("lunsj")
    dinner_all, dinner_glu, dinner_nuts = read_allergens("middag")

    def line(label: str, cat: str, picked: dict[str, str], alls: dict[str, str], glu: dict[str, str], nuts: dict[str, str]) -> str:
        dish = (picked.get(cat) or "").strip()
        if cat == "extra" and not dish:
            dish = "Varied types of side dishes."
        alg = (alls.get(cat) or "").strip().replace(" ", "")
        glu_s = (glu.get(cat) or "").strip()
        nuts_s = (nuts.get(cat) or "").strip()
        if alg and "1" in [x for x in alg.split(",") if x]:
            if glu_s:
                alg = alg + f" [gluten: {glu_s}]"
        if alg and "8" in [x for x in alg.split(",") if x]:
            if nuts_s:
                alg = alg + f" [nuts: {nuts_s}]"
        suffix = f" ({alg})" if alg else ""
        return f"{label} {dish}{suffix}" if dish else label

    lunch_lines = {
        "soup": line("soup of the day:", "soppa", lunch_picked, lunch_all, lunch_glu, lunch_nuts),
        "fish": line("today’s fish:", "fisk", lunch_picked, lunch_all, lunch_glu, lunch_nuts),
        "meat": line("today’s meat:", "kott", lunch_picked, lunch_all, lunch_glu, lunch_nuts),
        "extra": line("3. dish/green dish:", "extra", lunch_picked, lunch_all, lunch_glu, lunch_nuts),
    }
    dinner_lines = {
        "soup": line("soup of the day:", "soppa", dinner_picked, dinner_all, dinner_glu, dinner_nuts),
        "fish": line("today’s fish:", "fisk", dinner_picked, dinner_all, dinner_glu, dinner_nuts),
        "meat": line("today’s meat:", "kott", dinner_picked, dinner_all, dinner_glu, dinner_nuts),
        "extra": line("3. dish/green dish:", "extra", dinner_picked, dinner_all, dinner_glu, dinner_nuts),
    }

    return render_template(
        "daily_print.html",
        rig_name=rig_name or "",
        date_str=dt.strftime("%Y-%m-%d"),
        lunch_lines=lunch_lines,
        dinner_lines=dinner_lines,
    )


## removed legacy template generator


@app.post("/menus/daily/update")
def menus_daily_update():
    """Save an override and/or allergens for a specific date+meal+category.

    JSON body: { date: 'YYYY-MM-DD', meal: 'lunsj'|'middag', category: 'soppa'|'fisk'|'kott'|'extra', dish?: str, allergens?: [int]|comma-string }
    """
    db = get_db()
    rig_id, _ = _current_rig(db)
    if not rig_id:
        return jsonify({"ok": False, "error": "Ingen rigg"}), 400
    uid = session.get("user_id")
    if not uid:
        return jsonify({"ok": False, "error": "Ingen bruker"}), 403
    try:
        data = request.get_json(force=True) or {}
    except Exception:
        data = {}
    date_s = (data.get("date") or "").strip()
    meal = (data.get("meal") or "").strip().lower()
    category = (data.get("category") or "").strip().lower()
    dish = (data.get("dish") or "").strip()
    allergens_in = data.get("allergens")
    gluten_types_in = data.get("gluten_types")  # optional list or comma string for gluten subtypes
    nuts_types_in = data.get("nuts_types")      # optional list or comma string for nuts subtypes
    if isinstance(allergens_in, list):
        allergens = ",".join(str(int(x)) for x in allergens_in if str(x).strip())
    else:
        allergens = (str(allergens_in).strip() if allergens_in is not None else None)
    if isinstance(gluten_types_in, list):
        gluten_types = ",".join(str(x).strip() for x in gluten_types_in if str(x).strip())
    else:
        gluten_types = (str(gluten_types_in).strip() if gluten_types_in is not None else None)
    if isinstance(nuts_types_in, list):
        nuts_types = ",".join(str(x).strip() for x in nuts_types_in if str(x).strip())
    else:
        nuts_types = (str(nuts_types_in).strip() if nuts_types_in is not None else None)
    if not (date_s and meal in ("lunsj", "middag") and category in ("soppa", "fisk", "kott", "extra")):
        return jsonify({"ok": False, "error": "Ugyldige felt"}), 400
    try:
        # Private tabeller per bruker
        db.execute(
            """
            CREATE TABLE IF NOT EXISTS daily_menu_overrides_private (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              rig_id INTEGER NOT NULL,
              user_id INTEGER NOT NULL,
              date TEXT NOT NULL,
              meal TEXT NOT NULL,
              category TEXT NOT NULL,
              dish TEXT NOT NULL,
              UNIQUE(rig_id,user_id,date,meal,category)
            )
            """
        )
        db.execute(
            """
            CREATE TABLE IF NOT EXISTS daily_menu_allergens_private (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              rig_id INTEGER NOT NULL,
              user_id INTEGER NOT NULL,
              date TEXT NOT NULL,
              meal TEXT NOT NULL,
              category TEXT NOT NULL,
              allergens TEXT,
              gluten_types TEXT,
              nuts_types TEXT,
              UNIQUE(rig_id,user_id,date,meal,category)
            )
            """
        )
        if dish:
            db.execute(
                "INSERT INTO daily_menu_overrides_private(rig_id,user_id,date,meal,category,dish) VALUES(?,?,?,?,?,?)\n                 ON CONFLICT(rig_id,user_id,date,meal,category) DO UPDATE SET dish=excluded.dish",
                (rig_id, uid, date_s, meal, category, dish),
            )
        if allergens is not None or gluten_types is not None or nuts_types is not None:
            db.execute(
                "INSERT INTO daily_menu_allergens_private(rig_id,user_id,date,meal,category,allergens,gluten_types,nuts_types) VALUES(?,?,?,?,?,?,?,?)\n                 ON CONFLICT(rig_id,user_id,date,meal,category) DO UPDATE SET allergens=COALESCE(excluded.allergens, daily_menu_allergens_private.allergens), gluten_types=COALESCE(excluded.gluten_types, daily_menu_allergens_private.gluten_types), nuts_types=COALESCE(excluded.nuts_types, daily_menu_allergens_private.nuts_types)",
                (rig_id, uid, date_s, meal, category, allergens, gluten_types, nuts_types),
            )
        db.commit()
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.post("/menus/daily/reset")
def menus_daily_reset():
    """Reset (delete) any saved overrides and allergen data for one day (or a date range) for current rig.

    Accepts JSON or form fields:
      - date: YYYY-MM-DD (required) start date
      - date_to / end: optional end date (inclusive). If omitted only one date is cleared.

    Returns JSON: { ok: bool, cleared_overrides: int, cleared_allergens: int }
    """
    db = get_db()
    rig_id, _ = _current_rig(db)
    if not rig_id:
        return jsonify({"ok": False, "error": "Ingen rigg"}), 400
    uid = session.get("user_id")
    if not uid:
        return jsonify({"ok": False, "error": "Ingen bruker"}), 403
    # Parse input (JSON preferred)
    try:
        data = request.get_json(force=True) or {}
    except Exception:
        data = {}
    date_s = (data.get("date") or request.form.get("date") or "").strip()
    date_to_s = (data.get("date_to") or data.get("end") or request.form.get("date_to") or request.form.get("end") or "").strip()
    if not date_s:
        return jsonify({"ok": False, "error": "Mangler dato"}), 400
    # Validate format (basic)
    def _valid(d: str) -> bool:
        try:
            datetime.strptime(d, "%Y-%m-%d")
            return True
        except Exception:
            return False
    if not _valid(date_s):
        return jsonify({"ok": False, "error": "Ugyldig dato"}), 400
    if date_to_s and not _valid(date_to_s):
        return jsonify({"ok": False, "error": "Ugyldig sluttdato"}), 400
    if not date_to_s:
        date_to_s = date_s
    # Ensure correct ordering
    if date_to_s < date_s:
        date_s, date_to_s = date_to_s, date_s
    # Build BETWEEN clause (dates stored as YYYY-MM-DD strings so lexical between works)
    try:
        # Ensure tables exist (safe no-op if already there)
        db.execute(
            """
            CREATE TABLE IF NOT EXISTS daily_menu_overrides (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              rig_id INTEGER NOT NULL,
              date TEXT NOT NULL,
              meal TEXT NOT NULL,
              category TEXT NOT NULL,
              dish TEXT NOT NULL,
              UNIQUE(rig_id, date, meal, category)
            )
            """
        )
        db.execute(
            """
            CREATE TABLE IF NOT EXISTS daily_menu_allergens (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              rig_id INTEGER NOT NULL,
              date TEXT NOT NULL,
              meal TEXT NOT NULL,
              category TEXT NOT NULL,
              allergens TEXT,
              gluten_types TEXT,
              nuts_types TEXT,
              UNIQUE(rig_id, date, meal, category)
            )
            """
        )
    except Exception:
        pass
    try:
        # Private tables (create if missing)
        db.execute(
            """
            CREATE TABLE IF NOT EXISTS daily_menu_overrides_private (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              rig_id INTEGER NOT NULL,
              user_id INTEGER NOT NULL,
              date TEXT NOT NULL,
              meal TEXT NOT NULL,
              category TEXT NOT NULL,
              dish TEXT NOT NULL,
              UNIQUE(rig_id,user_id,date,meal,category)
            )
            """
        )
        db.execute(
            """
            CREATE TABLE IF NOT EXISTS daily_menu_allergens_private (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              rig_id INTEGER NOT NULL,
              user_id INTEGER NOT NULL,
              date TEXT NOT NULL,
              meal TEXT NOT NULL,
              category TEXT NOT NULL,
              allergens TEXT,
              gluten_types TEXT,
              nuts_types TEXT,
              UNIQUE(rig_id,user_id,date,meal,category)
            )
            """
        )
        cur1 = db.execute(
            "DELETE FROM daily_menu_overrides_private WHERE rig_id=? AND user_id=? AND date BETWEEN ? AND ?",
            (rig_id, uid, date_s, date_to_s),
        )
        cleared_overrides = cur1.rowcount if hasattr(cur1, "rowcount") else 0
        cur2 = db.execute(
            "DELETE FROM daily_menu_allergens_private WHERE rig_id=? AND user_id=? AND date BETWEEN ? AND ?",
            (rig_id, uid, date_s, date_to_s),
        )
        cleared_allergens = cur2.rowcount if hasattr(cur2, "rowcount") else 0
        db.commit()
        return jsonify({
            "ok": True,
            "cleared_overrides": cleared_overrides,
            "cleared_allergens": cleared_allergens,
            "date_start": date_s,
            "date_end": date_to_s,
        })
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.get("/messages/admins")
def list_admins_for_rig():
    db = get_db()
    rig_id, _ = _current_rig(db)
    if not rig_id:
        return jsonify({"ok": False, "error": "Ingen rigg"}), 400
    rows = db.execute("SELECT id, name, email FROM users WHERE role='admin' AND rig_id=? ORDER BY name, email", (rig_id,)).fetchall()
    admins = [{"id": r[0], "name": r[1] or r[2], "email": r[2]} for r in rows]
    return jsonify({"ok": True, "admins": admins})


@app.get("/downloads/daily_menu/<path:filename>")
def download_daily_menu(filename: str):
    # Serve files only from uploads/daily_menus
    safe = secure_filename(filename)
    base = Path(os.getcwd()) / "uploads" / "daily_menus"
    file_path = base / safe
    if not file_path.exists():
        return jsonify({"ok": False, "error": "Filen finnes ikke"}), 404
    return send_file(str(file_path), as_attachment=True, download_name=safe)


@app.post("/messages/send_daily_menu")
def send_daily_menu_message():
    """Generate and save daily DOCX, then create a message to a selected admin with a download link."""
    db = get_db()
    if not session.get("user_id"):
        return jsonify({"ok": False, "error": "Innlogging kreves"}), 401
    user_id = session["user_id"]
    rig_id, rig_name = _current_rig(db)
    if not rig_id:
        return jsonify({"ok": False, "error": "Ingen rigg"}), 400
    try:
        payload = request.get_json(force=True) or {}
    except Exception:
        payload = {}
    date_s = (payload.get("date") or datetime.now().strftime("%Y-%m-%d")).strip()
    admin_id = payload.get("admin_id")
    if not admin_id:
        return jsonify({"ok": False, "error": "Mangler admin"}), 400
    # Validate admin belongs to rig
    row = db.execute("SELECT id FROM users WHERE id=? AND role='admin' AND rig_id=?", (admin_id, rig_id)).fetchone()
    if not row:
        return jsonify({"ok": False, "error": "Ugyldig admin"}), 400

    # Generate DOCX (save to uploads/daily_menus with predictable name) using modern builder only
    if Document is None:
        return jsonify({"ok": False, "error": "DOCX-bibliotek mangler"}), 500

    # Reuse logic by internally calling the builder; replicate minimal pieces here
    # We'll call menus_daily_docx's inner steps: collect lunch/dinner with overrides and allergens
    # Load menu settings
    start_week = 1
    start_index = 1
    menu_sheets = []
    try:
        r = db.execute("SELECT start_week, start_index, menu_json FROM menu_settings WHERE rig_id=?", (rig_id,)).fetchone()
        if r:
            start_week = r[0] or 1
            start_index = r[1] or 1
            try:
                menu_sheets = json.loads(r[2] or "[]")
            except Exception:
                menu_sheets = []
    except Exception:
        pass
    weekday_no_to_nor = {0: "Mandag", 1: "Tirsdag", 2: "Onsdag", 3: "Torsdag", 4: "Fredag", 5: "Lørdag", 6: "Søndag"}
    def menu_index_for_date(dt_date):
        iso_year, iso_week, _ = dt_date.isocalendar()
        idx = ((iso_week - int(start_week) + (int(start_index) - 1)) % 4) + 1
        return str(idx)
    def _norm_cat(raw: str | None):
        if not raw:
            return None
        r = str(raw).strip().lower()
        repl = r.replace("ø", "o").replace("å", "a").replace("ä", "a").replace("ö", "o").replace("æ", "a")
        if ("sopp" in repl) or ("suppe" in repl) or ("soup" in repl) or ("sopa" in repl):
            return "soppa"
        if ("fisk" in repl) or ("fish" in repl):
            return "fisk"
        if ("kjott" in repl) or ("kjot" in repl) or ("kott" in repl) or ("meat" in repl):
            return "kott"
        return None
    try:
        dt = datetime.strptime(date_s, "%Y-%m-%d").date()
    except Exception:
        return jsonify({"ok": False, "error": "Ugyldig dato"}), 400
    sheets_by_uke = {str(sh.get("uke")): sh for sh in (menu_sheets or [])}
    sh = sheets_by_uke.get(menu_index_for_date(dt))
    if not sh:
        return jsonify({"ok": False, "error": "Ingen meny for dato"}), 400
    dn = weekday_no_to_nor[dt.weekday()]
    def collect_for(meal_key: str):
        items = sh.get(meal_key) or []
        cats: dict[str, list[str]] = {"soppa": [], "fisk": [], "kott": [], "extra": []}
        for it in items:
            if it.get("dag") != dn:
                continue
            cat = _norm_cat(it.get("kategori"))
            if cat in ("soppa", "fisk", "kott"):
                cats[cat].append(it.get("rett") or "")
            else:
                cats["extra"].append(it.get("rett") or "")
        return cats
    lunch = collect_for("lunsj")
    dinner = collect_for("middag")
    # Ensure allergen table has gluten_types and nuts_types
    try:
        db.execute("ALTER TABLE daily_menu_allergens ADD COLUMN gluten_types TEXT")
        db.commit()
    except Exception:
        pass
    try:
        db.execute("ALTER TABLE daily_menu_allergens ADD COLUMN nuts_types TEXT")
        db.commit()
    except Exception:
        pass
    def apply_overrides(meal_key: str, cats: dict[str, list[str]]):
        out: dict[str, str] = {}
        for cat in ("soppa", "fisk", "kott", "extra"):
            row = db.execute(
                "SELECT dish FROM daily_menu_overrides WHERE rig_id=? AND date=? AND meal=? AND category=?",
                (rig_id, date_s, meal_key, cat),
            ).fetchone()
            if row and row[0]:
                out[cat] = row[0]
            else:
                lst = cats.get(cat) or []
                out[cat] = (lst[0] if lst else "")
        return out
    def read_allergens(meal_key: str):
        m_codes: dict[str, str] = {}
        m_gluten: dict[str, str] = {}
        m_nuts: dict[str, str] = {}
        rows = db.execute(
            "SELECT category, allergens, gluten_types, nuts_types FROM daily_menu_allergens WHERE rig_id=? AND date=? AND meal=?",
            (rig_id, date_s, meal_key),
        ).fetchall()
        for r in rows:
            m_codes[r[0]] = (r[1] or "")
            m_gluten[r[0]] = (r[2] or "")
            m_nuts[r[0]] = (r[3] or "")
        return m_codes, m_gluten, m_nuts
    lunch_picked = apply_overrides("lunsj", lunch)
    dinner_picked = apply_overrides("middag", dinner)
    lunch_all, lunch_glu, lunch_nuts = read_allergens("lunsj")
    dinner_all, dinner_glu, dinner_nuts = read_allergens("middag")

    # Ensure default text for 'extra' when nothing is specified
    DEFAULT_EXTRA = "Varied types of side dishes."
    if not (lunch_picked.get("extra") or "").strip():
        lunch_picked["extra"] = DEFAULT_EXTRA
    if not (dinner_picked.get("extra") or "").strip():
        dinner_picked["extra"] = DEFAULT_EXTRA

    # Build DOCX using the same logic as the /menus/daily_docx endpoint
    doc = _build_docx_modern(str(rig_name or ""), dt, lunch_picked, lunch_all, lunch_glu, lunch_nuts, dinner_picked, dinner_all, dinner_glu, dinner_nuts)

    # Save
    out_dir = Path(os.getcwd()) / "uploads" / "daily_menus"
    out_dir.mkdir(parents=True, exist_ok=True)
    filename = f"daily_menu_{date_s}_{int(time.time())}.docx"
    out_path = out_dir / filename
    doc.save(str(out_path))

    # Messages table
    try:
        db.execute(
            """
            CREATE TABLE IF NOT EXISTS messages (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              rig_id INTEGER NOT NULL,
              from_user_id INTEGER,
              to_user_id INTEGER,
              is_bulletin INTEGER DEFAULT 0,
              subject TEXT,
              body TEXT,
              attachment_path TEXT,
              created_at TEXT DEFAULT (datetime('now')),
              read_at TEXT
            )
            """
        )
        subject = f"Daily menu {date_s}"
        body = f"Daily menu attached. Download: /downloads/daily_menu/{filename}"
        db.execute(
            "INSERT INTO messages(rig_id,from_user_id,to_user_id,is_bulletin,subject,body,attachment_path) VALUES(?,?,?,?,?,?,?)",
            (rig_id, user_id, admin_id, 0, subject, body, f"{filename}"),
        )
        db.commit()
    except Exception as e:
        return jsonify({"ok": False, "error": f"Klarte ikke å lagre melding: {e}"}), 500

    return jsonify({"ok": True, "download_url": f"/downloads/daily_menu/{filename}"})


# --- Messages Center ---
@app.get("/messages")
def messages_center():
    if not session.get("user_id"):
        flash("Logg inn kreves.", "warning")
        return redirect(url_for("login"))
    db = get_db()
    rig_id, rig_name = _current_rig(db)
    user_id = session.get("user_id")
    # Ensure table exists
    try:
        db.execute(
            """
            CREATE TABLE IF NOT EXISTS messages (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              rig_id INTEGER NOT NULL,
              from_user_id INTEGER,
              to_user_id INTEGER,
              is_bulletin INTEGER DEFAULT 0,
              subject TEXT,
              body TEXT,
              attachment_path TEXT,
              created_at TEXT DEFAULT (datetime('now')),
              read_at TEXT
            )
            """
        )
    except Exception:
        pass
    bulletins = db.execute(
        "SELECT m.id, m.subject, m.body, m.created_at, u.name as from_name FROM messages m LEFT JOIN users u ON u.id=m.from_user_id WHERE m.rig_id=? AND m.is_bulletin=1 ORDER BY m.id DESC LIMIT 50",
        (rig_id,),
    ).fetchall()
    inbox = db.execute(
        "SELECT m.id, m.subject, m.body, m.created_at, u.name as from_name, m.attachment_path FROM messages m LEFT JOIN users u ON u.id=m.from_user_id WHERE m.rig_id=? AND m.to_user_id=? ORDER BY m.id DESC LIMIT 50",
        (rig_id, user_id),
    ).fetchall()
    outbox = db.execute(
        "SELECT m.id, m.subject, m.body, m.created_at, u.name as to_name, m.attachment_path FROM messages m LEFT JOIN users u ON u.id=m.to_user_id WHERE m.rig_id=? AND m.from_user_id=? ORDER BY m.id DESC LIMIT 50",
        (rig_id, user_id),
    ).fetchall()
    # Users for direct send
    users = db.execute("SELECT id, name, email FROM users WHERE rig_id=? ORDER BY name, email", (rig_id,)).fetchall()
    return render_template("messages.html", rig_name=rig_name, bulletins=bulletins, inbox=inbox, outbox=outbox, users=users)


@app.post("/messages/post_bulletin")
def post_bulletin():
    if not session.get("user_id"):
        flash("Logg inn kreves.", "warning")
        return redirect(url_for("login"))
    db = get_db()
    rig_id, _ = _current_rig(db)
    user_id = session.get("user_id")
    subject = (request.form.get("subject") or "").strip()
    body = (request.form.get("body") or "").strip()
    if not subject and not body:
        flash("Melding eller emne må fylles ut.", "warning")
        return redirect(url_for("messages_center"))
    try:
        db.execute(
            "INSERT INTO messages(rig_id,from_user_id,is_bulletin,subject,body) VALUES(?,?,?,?,?)",
            (rig_id, user_id, 1, subject, body),
        )
        db.commit()
        flash("Bulletin publisert.", "success")
    except Exception as e:
        flash(f"Kunne ikke publisere: {e}", "danger")
    return redirect(url_for("messages_center"))


@app.post("/messages/send_direct")
def send_direct():
    if not session.get("user_id"):
        flash("Logg inn kreves.", "warning")
        return redirect(url_for("login"))
    db = get_db()
    rig_id, _ = _current_rig(db)
    user_id = session.get("user_id")
    to_user_id = request.form.get("to_user_id", type=int)
    subject = (request.form.get("subject") or "").strip()
    body = (request.form.get("body") or "").strip()
    if not to_user_id:
        flash("Velg mottaker.", "warning")
        return redirect(url_for("messages_center"))
    try:
        # Ensure recipient is in same rig
        row = db.execute("SELECT id FROM users WHERE id=? AND rig_id=?", (to_user_id, rig_id)).fetchone()
        if not row:
            flash("Ugyldig mottaker.", "danger")
            return redirect(url_for("messages_center"))
        db.execute(
            "INSERT INTO messages(rig_id,from_user_id,to_user_id,is_bulletin,subject,body) VALUES(?,?,?,?,?,?)",
            (rig_id, user_id, to_user_id, 0, subject, body),
        )
        db.commit()
        flash("Melding sendt.", "success")
    except Exception as e:
        flash(f"Kunne ikke sende: {e}", "danger")
    return redirect(url_for("messages_center"))


@app.get("/menus/daily/get")
def menus_daily_get():
    """Return base menu and saved overrides/allergens for a given date.

    Query: ?date=YYYY-MM-DD
    Response: { date, lunch: {cat:{dish,allergens,gluten_types}}, dinner: {...}, base:{...} }
    """
    db = get_db()
    rig_id, _ = _current_rig(db)
    if not rig_id:
        return jsonify({"ok": False, "error": "Ingen rigg"}), 400
    # Ensure tables exist for overrides and allergens
    try:
        db.execute(
            """
            CREATE TABLE IF NOT EXISTS daily_menu_overrides (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              rig_id INTEGER NOT NULL,
              date TEXT NOT NULL,
              meal TEXT NOT NULL,
              category TEXT NOT NULL,
              dish TEXT NOT NULL,
              UNIQUE(rig_id, date, meal, category)
            )
            """
        )
        db.execute(
            """
            CREATE TABLE IF NOT EXISTS daily_menu_allergens (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              rig_id INTEGER NOT NULL,
              date TEXT NOT NULL,
              meal TEXT NOT NULL,
              category TEXT NOT NULL,
              allergens TEXT,
              gluten_types TEXT,
              nuts_types TEXT,
              UNIQUE(rig_id, date, meal, category)
            )
            """
        )
        try:
            db.execute("ALTER TABLE daily_menu_allergens ADD COLUMN gluten_types TEXT")
            db.commit()
        except Exception:
            pass
        try:
            db.execute("ALTER TABLE daily_menu_allergens ADD COLUMN nuts_types TEXT")
            db.commit()
        except Exception:
            pass
    except Exception:
        pass
    date_s = (request.args.get("date") or datetime.now().strftime("%Y-%m-%d")).strip()
    try:
        dt = datetime.strptime(date_s, "%Y-%m-%d").date()
    except Exception:
        return jsonify({"ok": False, "error": "Ugyldig dato"}), 400

    # Load settings and base menu
    start_week = 1
    start_index = 1
    menu_sheets = []
    try:
        row = db.execute("SELECT start_week, start_index, menu_json FROM menu_settings WHERE rig_id=?", (rig_id,)).fetchone()
        if row:
            start_week = row[0] or 1
            start_index = row[1] or 1
            try:
                menu_sheets = json.loads(row[2] or "[]")
            except Exception:
                menu_sheets = []
    except Exception:
        pass
    weekday_no_to_nor = {0: "Mandag", 1: "Tirsdag", 2: "Onsdag", 3: "Torsdag", 4: "Fredag", 5: "Lørdag", 6: "Søndag"}
    def menu_index_for_date(dt_date):
        iso_year, iso_week, _ = dt_date.isocalendar()
        idx = ((iso_week - int(start_week) + (int(start_index) - 1)) % 4) + 1
        return str(idx)
    def _norm_cat(raw: str | None):
        if not raw:
            return None
        r = str(raw).strip().lower()
        repl = r.replace("ø", "o").replace("å", "a").replace("ä", "a").replace("ö", "o").replace("æ", "a")
        if ("sopp" in repl) or ("suppe" in repl) or ("soup" in repl) or ("sopa" in repl):
            return "soppa"
        if ("fisk" in repl) or ("fish" in repl):
            return "fisk"
        if ("kjott" in repl) or ("kjot" in repl) or ("kott" in repl) or ("meat" in repl):
            return "kott"
        return None
    sheets_by_uke = {str(sh.get("uke")): sh for sh in (menu_sheets or [])}
    sh = sheets_by_uke.get(menu_index_for_date(dt))
    dn = weekday_no_to_nor[dt.weekday()]
    def collect_for(meal_key: str):
        items = (sh.get(meal_key) if sh else []) or []
        cats: dict[str, list[str]] = {"soppa": [], "fisk": [], "kott": [], "extra": []}
        for it in items:
            if it.get("dag") != dn:
                continue
            cat = _norm_cat(it.get("kategori"))
            if cat in ("soppa", "fisk", "kott"):
                cats[cat].append(it.get("rett") or "")
            else:
                cats["extra"].append(it.get("rett") or "")
        # base pick = first in each
        base = {k: (v[0] if v else "") for k, v in cats.items()}
        return base
    base_l = collect_for("lunsj")
    base_d = collect_for("middag")
    # read overrides
    def read_over(meal_key: str):
        rows = db.execute(
            "SELECT category, dish FROM daily_menu_overrides WHERE rig_id=? AND date=? AND meal=?",
            (rig_id, date_s, meal_key),
        ).fetchall()
        m = {r[0]: r[1] for r in rows}
        return m
    def read_all(meal_key: str):
        try:
            db.execute("ALTER TABLE daily_menu_allergens ADD COLUMN gluten_types TEXT")
            db.commit()
        except Exception:
            pass
        rows = db.execute(
            "SELECT category, allergens, gluten_types, nuts_types FROM daily_menu_allergens WHERE rig_id=? AND date=? AND meal=?",
            (rig_id, date_s, meal_key),
        ).fetchall()
        m = {r[0]: {"allergens": r[1] or "", "gluten_types": r[2] or "", "nuts_types": r[3] or ""} for r in rows}
        return m
    ov_l = read_over("lunsj")
    ov_d = read_over("middag")
    al_l = read_all("lunsj")
    al_d = read_all("middag")
    def merge(base: dict, ov: dict, al: dict):
        out = {}
        for cat in ("soppa", "fisk", "kott", "extra"):
            out[cat] = {
                "dish": ov.get(cat) or base.get(cat) or "",
                "allergens": (al.get(cat, {}).get("allergens") or ""),
                "gluten_types": (al.get(cat, {}).get("gluten_types") or ""),
                "nuts_types": (al.get(cat, {}).get("nuts_types") or ""),
            }
        return out
    return jsonify({
        "ok": True,
        "date": date_s,
        "lunch": merge(base_l, ov_l, al_l),
        "dinner": merge(base_d, ov_d, al_d),
        "base": {"lunch": base_l, "dinner": base_d},
    })


@app.post("/menus/daily/bulk_update")
def menus_daily_bulk_update():
    """Save overrides + allergens for both lunch and dinner in one call.

    JSON body: { date, lunch:{cat:{dish,allergens,gluten_types}}, dinner:{...} }
    """
    db = get_db()
    rig_id, _ = _current_rig(db)
    if not rig_id:
        return jsonify({"ok": False, "error": "Ingen rigg"}), 400
    try:
        data = request.get_json(force=True) or {}
    except Exception:
        return jsonify({"ok": False, "error": "Ingen JSON"}), 400
    date_s = (data.get("date") or "").strip()
    lunch = data.get("lunch") or {}
    dinner = data.get("dinner") or {}
    if not date_s:
        return jsonify({"ok": False, "error": "Mangler dato"}), 400
    try:
        db.execute(
            """
            CREATE TABLE IF NOT EXISTS daily_menu_overrides (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              rig_id INTEGER NOT NULL,
              date TEXT NOT NULL,
              meal TEXT NOT NULL,
              category TEXT NOT NULL,
              dish TEXT NOT NULL,
              UNIQUE(rig_id, date, meal, category)
            )
            """
        )
        db.execute(
            """
            CREATE TABLE IF NOT EXISTS daily_menu_allergens (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              rig_id INTEGER NOT NULL,
              date TEXT NOT NULL,
              meal TEXT NOT NULL,
              category TEXT NOT NULL,
              allergens TEXT,
              gluten_types TEXT,
              nuts_types TEXT,
              UNIQUE(rig_id, date, meal, category)
            )
            """
        )
        try:
            db.execute("ALTER TABLE daily_menu_allergens ADD COLUMN gluten_types TEXT")
        except Exception:
            pass
        try:
            db.execute("ALTER TABLE daily_menu_allergens ADD COLUMN nuts_types TEXT")
        except Exception:
            pass
        # helper upserts
        def upsert(meal_key: str, obj: dict):
            for cat, vals in obj.items():
                if not isinstance(vals, dict):
                    continue
                dish = (vals.get("dish") or "").strip()
                allergens = vals.get("allergens")
                gluten_types = vals.get("gluten_types")
                nuts_types = vals.get("nuts_types")
                if dish:
                    db.execute(
                        "INSERT INTO daily_menu_overrides(rig_id,date,meal,category,dish) VALUES(?,?,?,?,?)\n                         ON CONFLICT(rig_id,date,meal,category) DO UPDATE SET dish=excluded.dish",
                        (rig_id, date_s, meal_key, cat, dish),
                    )
                if allergens is not None or gluten_types is not None or nuts_types is not None:
                    db.execute(
                        "INSERT INTO daily_menu_allergens(rig_id,date,meal,category,allergens,gluten_types,nuts_types) VALUES(?,?,?,?,?,?,?)\n                         ON CONFLICT(rig_id,date,meal,category) DO UPDATE SET allergens=COALESCE(excluded.allergens, daily_menu_allergens.allergens), gluten_types=COALESCE(excluded.gluten_types, daily_menu_allergens.gluten_types), nuts_types=COALESCE(excluded.nuts_types, daily_menu_allergens.nuts_types)",
                        (rig_id, date_s, meal_key, cat, (allergens or ""), (gluten_types or ""), (nuts_types or "")),
                    )
        upsert("lunsj", lunch)
        upsert("middag", dinner)
        db.commit()
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


# --- Auth routes ---
@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        email = (request.form.get("email") or "").strip().lower()
        password = request.form.get("password") or ""
        db = get_db()
        user = db.execute("SELECT * FROM users WHERE email=?", (email,)).fetchone()
        if not user:
            flash("Ukjent e-post.", "danger")
        else:
            from werkzeug.security import check_password_hash
            if user["password_hash"] and check_password_hash(user["password_hash"], password):
                session["user_id"] = user["id"]
                session["role"] = (user["role"] if "role" in user.keys() else None) or "user"
                flash("Innlogget.", "success")
                return redirect(url_for("dashboard"))
            else:
                flash("Feil passord.", "danger")
    return render_template("login.html")


@app.route("/adminlogin", methods=["GET", "POST"])
def admin_login():
    if request.method == "POST":
        email = (request.form.get("email") or "").strip().lower()
        password = request.form.get("password") or ""
        db = get_db()
        user = db.execute("SELECT * FROM users WHERE email=? AND role='admin'", (email,)).fetchone()
        if not user:
            flash("Ingen admin med den e-posten.", "danger")
        else:
            from werkzeug.security import check_password_hash
            if user["password_hash"] and check_password_hash(user["password_hash"], password):
                session["user_id"] = user["id"]
                session["role"] = "admin"
                flash("Innlogget som admin.", "success")
                return redirect(url_for("admin_dashboard"))
            else:
                flash("Feil passord.", "danger")
    return render_template("adminlogin.html")


@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        email = (request.form.get("email") or "").strip().lower()
        name = (request.form.get("name") or "").strip()
        password = request.form.get("password") or ""
        if not (email and password):
            flash("E-post og passord kreves.", "danger")
            return render_template("register.html")
        db = get_db()
        exists = db.execute("SELECT 1 FROM users WHERE email=?", (email,)).fetchone()
        if exists:
            flash("E-posten er allerede registrert.", "warning")
            return render_template("register.html")
        from werkzeug.security import generate_password_hash
        try:
            db.execute(
                "INSERT INTO users (email, name, password_hash, role) VALUES (?, ?, ?, 'user')",
                (email, name or None, generate_password_hash(password)),
            )
            db.commit()
            flash("Konto opprettet. Logg inn.", "success")
            return redirect(url_for("login"))
        except Exception as e:
            flash(f"Kunne ikke opprette konto: {e}", "danger")
    return render_template("register.html")


@app.get("/logout")
def logout():
    session.pop("user_id", None)
    session.pop("role", None)
    flash("Du er logget ut.", "info")
    return redirect(url_for("root"))


# --- Superuser ---
SUPERUSER_CODE = os.environ.get("SUPERUSER_CODE", "rigplan2025")
EMERGENCY_CODE = None
if os.environ.get("SUPERUSER_EMERGENCY") == "1":
    try:
        EMERGENCY_CODE = secrets.token_urlsafe(8)
    except Exception:
        EMERGENCY_CODE = f"superuser-{os.urandom(4).hex()}"
    try:
        print(f"[SUPERUSER] EMERGENCY_CODE: {EMERGENCY_CODE}")
    except Exception:
        pass


@app.route("/superuser", methods=["GET", "POST"])
def superuser_login():
    if request.method == "POST":
        code = request.form.get("code", "")
        if code == SUPERUSER_CODE or (EMERGENCY_CODE and code == EMERGENCY_CODE):
            session["superuser"] = True
            return redirect(url_for("superuser_panel"))
        flash("Feil kode. Prøv igjen.", "danger")
    if EMERGENCY_CODE:
        flash("Nødmodus aktiv: se server-konsollen for engangskode.", "info")
    return render_template("superuser_login.html")


@app.route("/superuser/logout")
def superuser_logout():
    session.pop("superuser", None)
    return redirect(url_for("superuser_login"))


@app.route("/superuser/panel", methods=["GET", "POST"])
@superuser_required
def superuser_panel():
    db = get_db()
    error = None
    if request.method == "POST" and request.form.get("new_rig_name"):
        rig_name = request.form.get("new_rig_name", "").strip()
        rig_desc = request.form.get("new_rig_desc", "").strip()
        if not rig_name:
            error = "Du må fylle ut rignavn."
        else:
            try:
                db.execute("INSERT INTO rigs(name, description) VALUES(?, ?)", (rig_name, rig_desc))
                db.commit()
                flash("Rigg er opprettet!", "success")
            except Exception as e:
                error = f"Kunne ikke opprette rigg: {e}"
    admins = db.execute("SELECT id, name, email, rig_id FROM users WHERE role='admin'").fetchall()
    rigs = db.execute("SELECT * FROM rigs ORDER BY id").fetchall()
    current_rig_id = None
    current_rig_name = None
    if session.get("user_id"):
        row = db.execute("SELECT rig_id FROM users WHERE id=?", (session["user_id"],)).fetchone()
        if row and row["rig_id"]:
            rig = db.execute("SELECT id, name FROM rigs WHERE id=?", (row["rig_id"],)).fetchone()
            if rig:
                current_rig_id = rig["id"]
                current_rig_name = rig["name"]
    return render_template("superuser_panel.html", admins=admins, rigs=rigs, error=error,
                           current_rig_id=current_rig_id, current_rig_name=current_rig_name)


@app.route("/superuser/rig/<int:rig_id>", methods=["GET", "POST"])
@superuser_required
def superuser_rig_detail(rig_id: int):
    db = get_db()
    error = None
    rig = db.execute("SELECT * FROM rigs WHERE id=?", (rig_id,)).fetchone()
    if not rig:
        flash("Rigg finnes ikke.", "danger")
        return redirect(url_for("superuser_panel"))
    if request.method == "POST" and request.form.get("delete_admin_id"):
        admin_id = request.form.get("delete_admin_id", type=int)
        if admin_id is None:
            error = "Ugyldig admin-id."
        else:
            try:
                db.execute("DELETE FROM users WHERE id=? AND role='admin'", (admin_id,))
                db.commit()
                flash("Admin er slettet!", "success")
                return redirect(url_for("superuser_rig_detail", rig_id=rig_id))
            except Exception as e:
                error = f"Kunne ikke slette admin: {e}"
    elif request.method == "POST" and all(k in request.form for k in ("name", "email", "password")):
        name = request.form.get("name", "").strip()
        email = request.form.get("email", "").strip().lower()
        password = request.form.get("password", "")
        if not (name and email and password):
            error = "Alle felt må fylles ut."
        else:
            from werkzeug.security import generate_password_hash
            try:
                db.execute(
                    "INSERT INTO users (email, name, password_hash, tenant_id, rig_id, role) VALUES (?, ?, ?, 0, ?, 'admin')",
                    (email, name, generate_password_hash(password), rig_id),
                )
                db.commit()
                flash("Admin er opprettet!", "success")
                return redirect(url_for("superuser_rig_detail", rig_id=rig_id))
            except Exception as e:
                error = f"Kunne ikke opprette admin: {e}"
    admins = db.execute("SELECT id, name, email FROM users WHERE role='admin' AND rig_id=? ORDER BY id", (rig_id,)).fetchall()
    return render_template("superuser_rig_detail.html", rig=rig, admins=admins, error=error)


# --- Root, health ---
@app.route("/", methods=["GET", "POST"])
def root():
    # Handle landing page POST choice (admin vs user)
    if request.method == "POST":
        lt = (request.form.get("login_type") or "").strip().lower()
        if lt == "admin":
            return redirect(url_for("admin_login"))
        if lt == "user":
            return redirect(url_for("login"))
        flash("Ugyldig valg.", "warning")
    # Default redirects for authenticated sessions
    if session.get("role") == "admin":
        return redirect(url_for("admin_dashboard_protected"))
    if session.get("user_id"):
        return redirect(url_for("dashboard"))
    return render_template("index.html")


@app.get("/health")
def health():
    try:
        db = get_db()
        row = db.execute("SELECT 1").fetchone()
        db_ok = True if row else False
    except Exception:
        db_ok = False
    return jsonify({"ok": True, "db": db_ok, "version": "v20250913"})


@app.get("/favicon.ico")
def favicon():
    # Serve our SVG as favicon to avoid 404; browsers will follow redirect.
    try:
        return redirect(url_for("static", filename="Yuplanlogo_offshore.svg"))
    except Exception:
        # Fallback empty response (no favicon)
        return "", 204


@app.get("/__routes")
def list_routes():
    try:
        rules = sorted([str(r) for r in app.url_map.iter_rules()])
        return jsonify({"ok": True, "routes": rules, "count": len(rules)})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500




if __name__ == "__main__":
    import os as _os
    _port = int(_os.environ.get("PORT") or _os.environ.get("YUPLAN_PORT") or 5000)
    try:
        interesting = {"/","/landing","/coming-soon","/ping","/login","/dashboard"}
        print("[yuplan] Starting on port", _port, "DEMO_MODE=", _os.environ.get("DEMO_MODE")=="1")
        for r in sorted(app.url_map.iter_rules(), key=lambda x: str(x)):
            if str(r) in interesting:
                print("   ", f"{r:15} -> {r.endpoint}")
    except Exception:
        pass
    app.run(host="127.0.0.1", port=_port, debug=False, use_reloader=False)
