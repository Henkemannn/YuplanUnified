import sys
from pathlib import Path

from docx import Document


def summarize(doc_path: Path):
    doc = Document(str(doc_path))
    summary = {
        "paragraphs": [],
        "tables": [],
        "headers": {"paragraphs": [], "tables": []},
        "footers": {"paragraphs": [], "tables": []},
    }

    # Body
    for i, p in enumerate(doc.paragraphs, 1):
        style = p.style.name if p.style else None
        text = p.text.strip()
        if text:
            summary["paragraphs"].append({"index": i, "style": style, "text": text})
    for ti, t in enumerate(doc.tables, 1):
        rows = []
        for ri, row in enumerate(t.rows):
            cells = []
            for ci, c in enumerate(row.cells):
                cells.append({
                    "r": ri,
                    "c": ci,
                    "text": c.text.strip().replace("\n", " "),
                })
            rows.append(cells)
        summary["tables"].append({"index": ti, "rows": rows})

    # Headers/Footers (first section only is common in templates)
    if doc.sections:
        sec = doc.sections[0]
        if sec.header:
            for i, p in enumerate(sec.header.paragraphs, 1):
                text = p.text.strip()
                if text:
                    summary["headers"]["paragraphs"].append({"index": i, "text": text})
            for ti, t in enumerate(sec.header.tables, 1):
                rows = []
                for ri, row in enumerate(t.rows):
                    cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                    rows.append(cells)
                summary["headers"]["tables"].append({"index": ti, "rows": rows})
        if sec.footer:
            for i, p in enumerate(sec.footer.paragraphs, 1):
                text = p.text.strip()
                if text:
                    summary["footers"]["paragraphs"].append({"index": i, "text": text})
            for ti, t in enumerate(sec.footer.tables, 1):
                rows = []
                for ri, row in enumerate(t.rows):
                    cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                    rows.append(cells)
                summary["footers"]["tables"].append({"index": ti, "rows": rows})
    return summary


def main():
    if len(sys.argv) < 2:
        print("Usage: inspect_docx.py <path-to-docx>")
        sys.exit(1)
    path = Path(sys.argv[1])
    s = summarize(path)
    # Print a compact preview
    print("Body paragraphs (non-empty):", len(s["paragraphs"]))
    for p in s["paragraphs"][:20]:
        print(f"P{p['index']:>3} [{p['style']}] => {p['text'][:100]}")
    if len(s["paragraphs"]) > 20:
        print("...")
    print("Body tables:", len(s["tables"]))
    for t in s["tables"]:
        max_cols = max((len(r) for r in t["rows"]), default=0)
        print(f"Table {t['index']}: {len(t['rows'])} rows, {max_cols} cols")
        for r in t["rows"][:8]:
            line = " | ".join(f"[{cell['r']},{cell['c']}] {cell['text'][:40]}" for cell in r[:8])
            print(line)
        if len(t["rows"]) > 8:
            print("...")

    # Header/Footer
    if s["headers"]["paragraphs"] or s["headers"]["tables"]:
        print("\nHeader content:")
        for p in s["headers"]["paragraphs"]:
            print(f"H-P{p['index']}: {p['text']}")
        for t in s["headers"]["tables"]:
            print(f"H-Table {t['index']}: {len(t['rows'])} rows")
            for r in t["rows"][:6]:
                print(" | ".join(r[:6]))
            if len(t["rows"]) > 6:
                print("...")
    if s["footers"]["paragraphs"] or s["footers"]["tables"]:
        print("\nFooter content:")
        for p in s["footers"]["paragraphs"]:
            print(f"F-P{p['index']}: {p['text']}")
        for t in s["footers"]["tables"]:
            print(f"F-Table {t['index']}: {len(t['rows'])} rows")
            for r in t["rows"][:6]:
                print(" | ".join(r[:6]))
            if len(t["rows"]) > 6:
                print("...")


if __name__ == "__main__":
    main()
